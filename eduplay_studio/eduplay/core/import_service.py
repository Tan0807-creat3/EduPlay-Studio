"""
Import Service - Handles file imports (DOCX, PDF, XLSX)
"""

import os
import re
from typing import List, Dict, Optional
from pathlib import Path

from eduplay.core.i18n import I18n
import unicodedata

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from openpyxl import load_workbook
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

class ImportService:
    """Service for importing questions from various file formats"""
    
    def __init__(self):
        """Initialize import service"""
        self.templates_dir = Path(__file__).parent.parent.parent / "assets_bundle" / "templates"
    
    def get_available_templates(self) -> List[Dict]:
        """Get list of available sample templates"""
        templates = []
        
        if self.templates_dir.exists():
            for file_path in self.templates_dir.glob("*.txt"):
                template_name = file_path.stem.replace("_", " ").title()
                templates.append({
                    "name": template_name,
                    "file": str(file_path),
                    "description": I18n.t("import.template_desc_prefix") + template_name.lower()
                })
        
        return templates
    
    def load_template(self, template_name: str) -> List[Dict]:
        """Load questions from a template file"""
        template_path = self.templates_dir / f"{template_name}.txt"
        
        if not template_path.exists():
            # Try with underscores
            template_path = self.templates_dir / f"{template_name.replace(' ', '_').lower()}.txt"
        
        if template_path.exists():
            return self.import_from_txt(str(template_path))
        else:
            return []
    
    def import_from_txt(self, file_path: str) -> List[Dict]:
        """Import questions from text file"""
        questions = []
        
        try:
            # Use utf-8-sig to handle BOM if present
            with open(file_path, 'r', encoding='utf-8-sig') as file:
                content = file.read()
            
            try:
                content = self._normalize_input(content)
            except Exception:
                pass
            
            # Parse questions based on format patterns
            # Try smart parse first as it handles blocks better
            smart_questions = self.parse_smart_format(content)
            if smart_questions:
                # Validate questions
                for q in smart_questions:
                    validated = self._validate_imported_question(q)
                    if validated:
                        questions.append(validated)
            
            # Legacy parsers are disabled to prevent duplicates and ensure consistent explanation extraction
            # questions.extend(self.parse_multiple_choice_questions(content))
            # questions.extend(self.parse_true_false_questions(content))
            # questions.extend(self.parse_fill_blank_questions(content))
            # questions.extend(self.parse_matching_questions(content))
            
            # If smart parser found nothing, we might try legacy? 
            # But we improved smart parser to handle everything.
            if not questions:
                 # Fallback just in case, but usually smart parser catches blocks
                 pass

            
        except Exception as e:
            print(f"Error importing from TXT: {e}")
        
        return questions
    
    def parse_smart_format(self, content: str) -> List[Dict]:
        """
        Smartly parse questions from text blocks.
        Handles:
        - Implicit options (A. B. - * or no prefix)
        - Implicit question types (MC, TF, FillBlank)
        - Implicit question text (first line(s))
        """
        try:
            content = self._normalize_input(content)
        except Exception:
            pass
        questions = []
        # Split by empty lines (2 or more newlines) OR by "Câu X:" headers OR by "X. Type" headers
        # This ensures we catch questions even if there are no empty lines between them.
        # Also supports headers like "2. Trắc nghiệm (Dạng khác)"
        type_keywords = r"(?:Trắc nghiệm|Đúng|Điền|Ghép|Tự luận|Multiple|True|Fill|Matching|Short)"
        pattern = r'(?:\n\s*\n)|(?=(?:^|\n)(?:Câu|Question|Pregunta|Frage)\s+\d+[:\.])|(?=(?:^|\n)\d+\.\s+.*' + type_keywords + r')'
        blocks = re.split(pattern, content, flags=re.MULTILINE)
        
        for block in blocks:
            lines = [line.strip() for line in block.split('\n') if line.strip()]
            if not lines:
                continue
            
            # Check for orphaned explanation (starts with "Giải thích:", "Explanation:", or "//")
            first_line_lower = lines[0].lower()
            is_explanation = (
                first_line_lower.startswith('giải thích:') or 
                first_line_lower.startswith('explanation:') or
                first_line_lower.startswith('//')
            )
            
            if is_explanation and questions:
                # Merge explanation to previous question
                expl_text = ""
                for line in lines:
                    if '//' in line:
                         parts = line.split('//', 1)
                         if len(parts) > 1:
                             content = parts[1].strip()
                             # Remove "Giải thích:" prefix if present in the comment
                             if content.lower().startswith('giải thích:'):
                                 content = content[11:].strip()
                             elif content.lower().startswith('explanation:'):
                                 content = content[12:].strip()
                             expl_text += content + "\n"
                    elif line.lower().startswith('giải thích:'):
                        expl_text += line[11:].strip() + "\n"
                    elif line.lower().startswith('explanation:'):
                        expl_text += line[12:].strip() + "\n"
                    else:
                        expl_text += line + "\n"
                
                prev_q = questions[-1]
                if prev_q.get("explanation"):
                    prev_q["explanation"] += "\n" + expl_text.strip()
                else:
                    prev_q["explanation"] = expl_text.strip()
                continue

            # Parse block
            q = self._parse_block(lines)
            if q:
                questions.append(q)
                
        return questions

    def _normalize_input(self, text: str) -> str:
        if not text:
            return ""
        t = text.replace('\r\n', '\n')
        t = t.replace('\u00A0', ' ')
        t = t.replace('\u200B', '')
        t = re.sub('[\u2000-\u200A\u202F\u205F]', ' ', t)
        return t

    def _remove_question_numbers(self, text: str) -> str:
        """Remove all question number prefixes from text"""
        if not text:
            return text
        # Remove patterns like "Câu 1:", "Question 1:", "câu 1", "questions 1", etc.
        text = re.sub(r'(?:Câu|Question|Pregunta|Frage|câu|question|pregunta|frage)\s+\d+[:\.]?\s*', '', text, flags=re.IGNORECASE)
        # Also remove standalone number prefixes like "1.", "2)", "1:" at start of lines
        text = re.sub(r'^\s*\d+[\.\):]\s*', '', text, flags=re.MULTILINE)
        return text
 
    def _clean_header(self, lines: List[str]) -> List[str]:
        """Remove type header and question number prefixes from question lines"""
        if not lines:
            return lines

        # Strip empty lines from start
        # work on a copy to avoid side effects if original list is used elsewhere?
        # _parse_block passes slices mostly.
        lines = list(lines)

        while lines and not lines[0].strip():
            lines.pop(0)

        if not lines:
            return lines

        first_line = lines[0].strip()

        # Remove question number prefixes like "Câu 1:", "Question 1:", "câu 1", "questions 1", etc.
        # Pattern matches: Câu/Question/Pregunta/Frage followed by number and optional punctuation
        question_prefix_pattern = r'^(?:Câu|Question|Pregunta|Frage|câu|question|pregunta|frage)\s+\d+[:\.]?\s*'
        if re.match(question_prefix_pattern, first_line, re.IGNORECASE):
            # Remove the prefix from the first line
            lines[0] = re.sub(question_prefix_pattern, '', first_line, flags=re.IGNORECASE).strip()
            # If the line becomes empty after removing prefix, remove it
            if not lines[0].strip():
                lines = lines[1:]

        # After removing question prefix, check if the remaining first line is a type header
        if lines:
            first_line = lines[0].strip()
            # Regex to match Type Headers
            type_keywords = [
                r"Trắc nghiệm", r"Đúng\s*[\/-]\s*Sai", r"Điền khuyết", r"Ghép đôi", r"Tự luận",
                r"Multiple Choice", r"True\s*[\/-]\s*False", r"Fill.*Blank", r"Matching", r"Short Answer"
            ]
            type_pattern = r'^(?:\d+[\.:]\s*)?.*(?:' + '|'.join(type_keywords) + r').*$'

            if re.match(type_pattern, first_line, re.IGNORECASE):
                lines = lines[1:]

        # Strip empty lines from start again
        while lines and not lines[0].strip():
            lines.pop(0)

        # Remove all question number prefixes from ALL lines (not just first line)
        # This handles cases where question numbers appear in the middle of text
        cleaned_lines = []
        for line in lines:
            # Remove patterns like "Câu 1:", "Question 1:", "câu 1", "questions 1", etc.
            cleaned_line = re.sub(r'(?:Câu|Question|Pregunta|Frage|câu|question|pregunta|frage)\s+\d+[:\.]?\s*', '', line, flags=re.IGNORECASE)
            # Also remove standalone number prefixes like "1.", "2)", "1:" at start of lines
            cleaned_line = re.sub(r'^\s*\d+[\.\):]\s*', '', cleaned_line)
            cleaned_lines.append(cleaned_line)

        return cleaned_lines

    def _parse_block(self, lines: List[str]) -> Optional[Dict]:
        """Parse a single block of lines into a question"""
        if not lines:
            return None
            
        # Strategy:
        # 1. Identify explicit structure (Fill Blank with ==, or MC with Answer: line)
        # 2. Identify Options vs Question
        
        explanation = ""
        # Extract explanation first (lines starting with // or "Giải thích:" or "Explanation:", or inline //)
        filtered_lines = []
        for line in lines:
            # Check for inline comments //
            if re.search(r'(^|\s)//', line):
                parts = line.split('//', 1)
                content_part = parts[0].strip()
                expl_part = parts[1].strip()
                
                # Check if expl_part starts with "Giải thích:" or "Explanation:" (optional cleanup)
                if expl_part.lower().startswith('giải thích:'):
                    expl_part = expl_part[11:].strip()
                elif expl_part.lower().startswith('explanation:'):
                    expl_part = expl_part[12:].strip()
                
                explanation += expl_part + " "
                
                if content_part:
                    filtered_lines.append(content_part)
            
            # Check for full line legacy explanation markers
            elif line.lower().startswith('giải thích:'):
                explanation += line[11:].strip() + " "
            elif line.lower().startswith('explanation:'):
                explanation += line[12:].strip() + " "
            elif re.fullmatch(r'[\s=\-\*_~]+', line):
                # Skip decoration/separator lines like "====", "---", "***"
                continue
            else:
                filtered_lines.append(line)
        lines = filtered_lines
        
        if not lines:
            return None
        
        first_line_lower = lines[0].strip().lower()
        if (
            first_line_lower.startswith('hướng dẫn:') or
            first_line_lower.startswith('instructions:') or
            first_line_lower.startswith('instrucciones:') or
            first_line_lower.startswith('anleitung:') or
            ('hướng dẫn' in first_line_lower) or
            ('guidelines' in first_line_lower)
        ):
            return None
        
        # Skip non-question header blocks like "Cấu trúc:" / "Structure:"
        if first_line_lower in ('cấu trúc:', 'structure:'):
            return None
        # Remove "Ví dụ:" / "Example:" headers but keep following lines for parsing
        if first_line_lower in ('ví dụ:', 'example:'):
            lines = lines[1:]
            if not lines:
                return None
        # Skip generic explanation/instruction sections
        if ('giải thích' in first_line_lower or 'explanation' in first_line_lower):
            return None

        # Check for Short Answer or Fill Blank (line starting with ==)
        has_q_header = any(re.match(r'^(?:Câu|Question|Pregunta|Frage)\s+\d+[:\.]', l, re.IGNORECASE) for l in lines)
        for i, line in enumerate(lines):
            if line.startswith('=='):
                tail = line[2:].strip()
                # Ignore lines that are pure decoration like "===="
                if not tail or re.fullmatch(r'[=\-\*_~]+', tail):
                    continue
                # If it looks like an option line (e.g. "A. ...", "1) ..."), treat it as MC correct marker, not short answer.
                try:
                    if re.match(r'^\s*([A-Fa-f]|\d{1,2})[\.\):]\s+.+', tail):
                        continue
                except Exception:
                    pass
                if has_q_header:
                    q_lines = self._clean_header(lines[:i])
                    question_text = "\n".join(q_lines).strip()
                    question_text = self._remove_question_numbers(question_text)
                    correct_answer = tail
                    
                    # Check if this is Fill Blank (has [blank] in question) or Short Answer (direct answer)
                    # If question has [blank] markers, it's fill_blank
                    # Otherwise, it's short_answer (direct answer expected)
                    if re.search(r'\[.*?\]', question_text):
                        # Replace [answer] with _____
                        question_text = re.sub(r'\[.*?\]', '_____', question_text)
                        return {
                            "type": "fill_blank",
                            "question": question_text,
                            "correct_answer": correct_answer,
                            "correct_answers": [correct_answer],
                            "answers": [correct_answer],
                            "explanation": explanation.strip()
                        }
                    else:
                        # This is a short answer question
                        return {
                            "type": "short_answer",
                            "question": question_text,
                            "answers": [correct_answer],
                            "explanation": explanation.strip()
                        }

        # Check for Inline Fill Blank (e.g. "Water boils at [100] degrees.")
        # Only if we don't have explicit answer line (which might override)
        
        # Clean header first
        cleaned_lines = self._clean_header(lines)
        full_text = "\n".join(cleaned_lines)

        # Look for [answer] pattern
        # Be careful not to match [citation] or [1] if it's not answer.
        # But user specified this format.
        fill_blank_matches = re.findall(r'\[(.*?)\]', full_text)
        if fill_blank_matches and has_q_header:
            # Assume the first one or all? User example has one.
            # If multiple, maybe it's multiple blanks? Current system supports 1 answer usually?
            # The system seems to support list of answers for fill blank (line 654: "answers": [answer_text]).
            # But line 372 uses "correct_answer": explicit_answer_char.
            # Let's support one for now or check data structure.
            # Line 192 uses "correct_answer". Line 462 uses "answers".
            # I should use "answers" (list) if possible, or "correct_answer" (string).
            # Let's use "answers" list.
            
            # Replace all [match] with _____
            question_text = full_text
            answers = []
            for match in fill_blank_matches:
                question_text = question_text.replace(f"[{match}]", "_____")
                answers.append(match.strip())

            question_text = self._remove_question_numbers(question_text)

            return {
                "type": "fill_blank",
                "question": question_text.strip(),
                "answers": answers, # Use list for consistency
                "correct_answers": answers, # For UI compatibility
                "correct_answer": answers[0] if answers else "", # Fallback
                "explanation": explanation.strip()
            }

        # Check for explicit "Answer:" line (Legacy support within block)
        explicit_answer_char = None
        legacy_answer_idx = -1
        
        for i, line in enumerate(lines):
            match = re.match(r'^\s*(?:[-*•]\s*)?(?:Đáp\s+án|Answer|Réponse|Respuesta|Antwort|Đáp\s+án\s+đúng|Correct\s+Answer)[:\s]\s*(.*)', line, re.IGNORECASE)
            if match:
                legacy_answer_idx = i
                explicit_answer_char = match.group(1).strip()
                try:
                    s_ans = explicit_answer_char
                    if ':' in s_ans:
                        s_ans = s_ans.split(':')[-1].strip()
                    m_letter = re.search(r'\b([A-Fa-f])\b', s_ans)
                    if m_letter:
                        explicit_answer_char = m_letter.group(1).upper()
                    else:
                        m_num = re.search(r'\b(\d{1,2})\b', s_ans)
                        if m_num:
                            explicit_answer_char = m_num.group(1)
                        else:
                            explicit_answer_char = s_ans.strip()
                except Exception:
                    pass
                # If explicit answer contains ==, strip it
                if explicit_answer_char.endswith('=='):
                    explicit_answer_char = explicit_answer_char[:-2].strip()
                elif explicit_answer_char.startswith('=='):
                     explicit_answer_char = explicit_answer_char[2:].strip()
                break
        
        # Remove Answer line from content processing if found
        if legacy_answer_idx != -1:
            lines.pop(legacy_answer_idx)

        # 2. Check for Answer Indicators
        # A valid question block MUST have either:
        # a) A line ending with '=='
        # b) An explicit "Answer:" line (explicit_answer_char)
        # c) Explicit True/False keywords as options (implied answer, though weak)
        # If none of these, it's likely instructions or headers -> Return None
        
        has_smart_answer = False
        smart_answer_line_idx = -1
        
        for i, line in enumerate(lines):
            if line.strip().endswith('=='):
                has_smart_answer = True
                smart_answer_line_idx = i
                break
        
        if not has_smart_answer and not explicit_answer_char:
            has_matching = False
            match_count = 0
            for line in lines[1:]:
                s = line.strip()
                # Check for matching separators
                # We accept lines without bullet points if they look like pairs
                if ('->' in s or ' = ' in s or ' - ' in s) and '==' not in s:
                    match_count += 1
            has_matching = match_count >= 2
            
            # Check for True/False keywords which might be implicit TF
            # Only if lines look like options
            has_tf = False
            for line in lines[1:]: # Skip Q text
                clean = line.strip().lower()
                if clean in ["đúng", "sai", "true", "false", "vrai", "faux", "falsch", "wahr", "verdadero", "falso"]:
                    has_tf = True
                    break
            has_mc_prefix = False
            for line in lines[1:]:
                cl = line.strip()
                if re.match(r'^[A-Fa-f][\.\)]\s', cl) or cl.startswith('- ') or cl.startswith('* ') or re.match(r'^[1-9]\.\s', cl):
                    has_mc_prefix = True
                    break
            if not has_tf and not has_matching and not has_mc_prefix:
                return None
        
        # Identify Options
        # Heuristic:
        # - If any line ends with ==, it's an option (and correct).
        # - If any line starts with prefix (A., -), it's an option.
        # - If we found explicit_answer_char, we expect options.
        
        option_prefixes = (
            'A.', 'B.', 'C.', 'D.', 'E.', 'F.',
            'a.', 'b.', 'c.', 'd.', 'e.', 'f.',
            'A)', 'B)', 'C)', 'D)', 'E)', 'F)',
            'a)', 'b)', 'c)', 'd)', 'e)', 'f)',
            'A:', 'B:', 'C:', 'D:', 'E:', 'F:',
            'a:', 'b:', 'c:', 'd:', 'e:', 'f:',
            '- ', '* ', '• ', '– ',
            '1.', '2.', '3.', '4.',
            '1)', '2)', '3)', '4)'
        )
        
        # Determine split point between Question and Options
        split_idx = -1
        
        # Prefix families for smart grouping
        prefix_families = [
            ['A.', 'B.', 'C.', 'D.', 'E.', 'F.'],
            ['a.', 'b.', 'c.', 'd.', 'e.', 'f.'],
            ['A)', 'B)', 'C)', 'D)', 'E)', 'F)'],
            ['a)', 'b)', 'c)', 'd)', 'e)', 'f)'],
            ['A:', 'B:', 'C:', 'D:', 'E:', 'F:'],
            ['a:', 'b:', 'c:', 'd:', 'e:', 'f:'],
            ['1.', '2.', '3.', '4.'],
            ['1)', '2)', '3)', '4)'],
            ['- ', '* ', '• ', '– ']
        ]

        # Strategy A: Use Smart Answer ('==') location to find Option Block Start
        if smart_answer_line_idx != -1:
            # We found '=='. This line IS an option.
            # Scan backwards to find the start of the option block.
            
            # Identify prefix of the answer line
            ans_line = lines[smart_answer_line_idx].strip()
            ans_prefix = None
            current_family = None
            
            for p in option_prefixes:
                if ans_line.startswith(p):
                    ans_prefix = p
                    break
            
            if ans_prefix:
                for family in prefix_families:
                    if ans_prefix in family:
                        current_family = family
                        break
            
            # Scan backwards
            curr_idx = smart_answer_line_idx
            first_option_idx = curr_idx
            
            for i in range(curr_idx - 1, -1, -1):
                line = lines[i].strip()
                if not line: continue
                
                is_option = False
                
                if re.match(r'^(?:Câu|Question|Pregunta|Frage)\s+\d+[:\.]', line, re.IGNORECASE):
                    break
                
                if ans_prefix:
                     # Check if line starts with a prefix from the SAME family
                     if current_family:
                         if any(line.startswith(p) for p in current_family):
                             is_option = True
                     else:
                         # Fallback if prefix not in known families (shouldn't happen with standard list)
                         if any(line.startswith(p) for p in option_prefixes):
                             is_option = True
                else:
                    first_option_idx = i
                    continue
                    
                if is_option:
                    first_option_idx = i
                elif ans_prefix:
                    # If we expect prefixes but this line has none or wrong family, it's likely the Question text
                    # E.g. "2. Header" (Family '1.') vs "A. Option" (Family 'A.')
                    break
                else:
                    break
            
            if ans_prefix is None and first_option_idx == 0:
                first_option_idx = 1
            split_idx = first_option_idx
            
        # Strategy B: Use First Prefix Match (Fallback if no smart answer or scanning failed)
        if split_idx == -1:
            # 1. Prioritize Letter Prefixes (A., B., etc.) to avoid matching numbered lists in question text
            letter_prefixes = [p for p in option_prefixes if p[0].lower() in 'abcdef']
            for i, line in enumerate(lines):
                if i == 0: continue
                clean_line = line.strip()
                if any(clean_line.startswith(p) for p in letter_prefixes):
                    split_idx = i
                    break
            
            # 2. If no letter prefixes, check other prefixes (Numbers, bullets)
            if split_idx == -1:
                other_prefixes = [p for p in option_prefixes if p[0].lower() not in 'abcdef']
                for i, line in enumerate(lines):
                    if i == 0: continue
                    clean_line = line.strip()
                    if any(clean_line.startswith(p) for p in other_prefixes):
                        split_idx = i
                        break
                
                # 3. Check for True/False keywords
                if split_idx == -1:
                    for i, line in enumerate(lines):
                        if i == 0: continue
                        clean_line = line.strip()
                        if clean_line.lower() in ["đúng", "sai", "true", "false", "vrai", "faux", "falsch", "wahr", "verdadero", "falso"]:
                            split_idx = i
                            break

            # 2. Check for Implicit Matching Pairs (lines with separators but no prefixes)
            if split_idx == -1:
                first_pair_idx = -1
                pair_count = 0
                for i, line in enumerate(lines):
                    if i == 0: continue # Skip question text
                    s = line.strip()
                    if ('->' in s or '=' in s or re.search(r'\s[-–—]\s', s)) and '==' not in s:
                        if first_pair_idx == -1:
                            first_pair_idx = i
                        pair_count += 1
                
                # If we found at least 2 pairs, assume this is a matching question
                if pair_count >= 2 and first_pair_idx != -1:
                    split_idx = first_pair_idx
        
        # 2. If no explicit marker found, but we have an explicit answer char (A, B..), 
        # try to find where options start based on A. B. (handled above) or just assume last N lines?
        # Or if we have implicit options (no prefix, no ==).
        if split_idx == -1:
            if explicit_answer_char:
                # Check if explicit answer is a TF keyword
                # If so, and we haven't found options, it's likely an Implicit True/False question
                # e.g. "Q: ... \n Answer: True"
                tf_keywords = ["đúng", "sai", "true", "false", "vrai", "faux", "falsch", "wahr", "verdadero", "falso"]
                if explicit_answer_char.lower() in tf_keywords:
                     # It is True/False!
                     is_true = explicit_answer_char.lower() in ["đúng", "true", "vrai", "verdadero", "wahr"]
                     question_text = self._remove_question_numbers("\n".join(lines).strip())
                     return {
                        "type": "true_false",
                        "question": question_text,
                        "correct_answer": is_true,
                        "explanation": explanation.strip()
                    }
                
                # If explicit answer is NOT TF, and we have no options, assume Fill Blank
                # But check if it looks like MC (A, B, C, D)
                if len(explicit_answer_char) == 1 and explicit_answer_char.upper() in "ABCD" and len(lines) > 1:
                     # Likely MC with implicit options?
                     pass
                else:
                     # Assume Fill Blank
                     question_text = self._remove_question_numbers("\n".join(lines).strip())
                     return {
                        "type": "fill_blank",
                        "question": question_text,
                        "correct_answers": [explicit_answer_char],
                        "answers": [explicit_answer_char],
                        "correct_answer": explicit_answer_char,
                        "explanation": explanation.strip()
                     }

            if len(lines) >= 3:
                # Heuristic: Line 0 is Question, Rest are Options
                split_idx = 1
            elif len(lines) == 2:
                # 2 lines. Q + Opt? Or Q + Answer (if explicit answer found)?
                # If explicit answer found, maybe lines are Q and Opts?
                # Case:
                # Q
                # True
                # Answer: True
                # -> split_idx should be 1.
                if explicit_answer_char:
                     split_idx = 1
        
        if split_idx == -1:
            # Can't find options. Maybe it's just a question text (Essay?) or incomplete.
            # If explicit_answer_char is present but no options found, maybe it's Short Answer/Fill Blank legacy?
            if explicit_answer_char:
                 question_text = self._remove_question_numbers("\n".join(lines).strip())
                 return {
                    "type": "fill_blank",
                    "question": question_text,
                    "correct_answer": explicit_answer_char,
                    "explanation": explanation.strip()
                }
            return None
            
        # Parse Options
        question_text_lines = lines[:split_idx]

        # Clean up Question Text: Remove "Type Headers"
        # Users often include "1. Multiple Choice", "2. True/False" etc. at the start.
        # We want to strip these lines so they don't appear in the actual question text.
        # But we must be careful not to strip valid question parts.
        # Heuristic: If the first line matches a known "Type Header" pattern, remove it.

        question_text_lines = self._clean_header(question_text_lines)

        # Strip empty lines from end of question text
        while question_text_lines and not question_text_lines[-1].strip():
            question_text_lines.pop()

        question_text = "\n".join(question_text_lines).strip()
        # Additional cleanup: remove any remaining question numbers from the final text
        question_text = self._remove_question_numbers(question_text)
        option_lines = lines[split_idx:]
        
        options = []
        for i, line in enumerate(option_lines):
            text = line.strip()
            is_correct = False
            
            # Leading correctness markers
            try:
                s0 = text.lstrip()
                if s0.startswith('=='):
                    is_correct = True
                    text = s0[2:].strip()
                elif s0.startswith('✓'):
                    is_correct = True
                    text = s0[1:].strip()
                elif s0.startswith('*') and len(s0) > 1:
                    is_correct = True
                    text = s0[1:].strip()
            except Exception:
                pass

            try:
                m_eq = re.search(r'\s*==\s*[.?!…]*\s*$', text)
                if m_eq:
                    is_correct = True
                    text = re.sub(r'\s*==\s*[.?!…]*\s*$', '', text).strip()
            except Exception:
                pass

            # Fallback: if '==' appears anywhere, treat it as a correctness marker
            try:
                if not is_correct and '==' in text:
                    is_correct = True
                    text = text.replace('==', ' ').strip()
            except Exception:
                pass
            
            matched_prefix = None
            for p in option_prefixes:
                if text.startswith(p):
                    matched_prefix = p
                    break
            
            if matched_prefix:
                text = text[len(matched_prefix):].strip()
                
                if explicit_answer_char:
                    clean_prefix = matched_prefix.strip(' .):')
                    if clean_prefix.lower() == explicit_answer_char.lower():
                        is_correct = True
            
            if explicit_answer_char and not is_correct:
                def _strip_accents(s: str) -> str:
                    try:
                        nf = unicodedata.normalize('NFD', s)
                        return ''.join(c for c in nf if not unicodedata.combining(c)).lower().strip()
                    except Exception:
                        return s.lower().strip()
                if text.lower().strip() == explicit_answer_char.lower().strip() or _strip_accents(text) == _strip_accents(explicit_answer_char):
                    is_correct = True
            
            try:
                m_annot = re.search(r'[\(\[\{]\s*(đúng|correct|true|vrai|wahr)\s*[\)\]\}]$', text, flags=re.IGNORECASE)
                m_tick = re.search(r'[\(\[\{]\s*✓\s*[\)\]\}]$', text)
                m_tick2 = re.search(r'✓\s*$', text)
                if m_annot or m_tick or m_tick2:
                    is_correct = True
                    text = re.sub(r'([\(\[\{]\s*(đúng|correct|true|vrai|wahr)\s*[\)\]\}])$', '', text, flags=re.IGNORECASE).strip()
                    text = re.sub(r'([\(\[\{]\s*✓\s*[\)\]\}])$', '', text).strip()
                    text = re.sub(r'✓\s*$', '', text).strip()
            except Exception:
                pass
            
            options.append({
                "text": text,
                "correct": is_correct
            })
            
        if not any(o['correct'] for o in options):
            expl_l = explanation.strip().lower()
            guess_idx = -1
            if expl_l:
                for i, o in enumerate(options):
                    t = o['text'].strip().lower()
                    if t and t in expl_l:
                        guess_idx = i
                        break
                if guess_idx == -1:
                    nums = []
                    for i, o in enumerate(options):
                        m = re.findall(r'\d+', o['text'])
                        if m:
                            try:
                                nums.append((i, int(m[-1])))
                            except:
                                pass
                    if nums:
                        target = None
                        m2 = re.findall(r'\d+', expl_l)
                        if m2:
                            try:
                                target = int(m2[-1])
                            except:
                                target = None
                        if target is not None:
                            for i, n in nums:
                                if n == target:
                                    guess_idx = i
                                    break
                        if guess_idx == -1:
                            nums.sort(key=lambda x: x[1], reverse=True)
                            guess_idx = nums[0][0]
            if guess_idx != -1 and 0 <= guess_idx < len(options):
                for i in range(len(options)):
                    options[i]['correct'] = (i == guess_idx)

        # Determine Type
        
        # Check if it is a Fill Blank Question (Smart Format)
        # Criteria: Only 1 option provided and it is marked correct (e.g. "Answer ==")
        # OR if it's explicitly marked as Fill Blank (via keywords) and has 1 option
        is_fill_blank_header = re.search(r'(?:Điền|Fill)', question_text, re.IGNORECASE)
        
        if (len(options) == 1 and options[0]['correct']) or (len(options) == 1 and is_fill_blank_header):
             return {
                "type": "fill_blank",
                "question": question_text,
                "correct_answers": [options[0]['text']],
                "answers": [options[0]['text']], # Legacy support
                "correct_answer": options[0]['text'], # UI compatibility
                "explanation": explanation.strip()
            }

        # Check if it is a Matching Question (Smart Format)
        # Criteria: Options contain '->' or '=' or ' - ', look like pairs
        
        # Relaxed matching criteria: If >= 50% of options have separators, treat as Matching
        # This handles cases where one line might be malformed or header-like
        arrow_count = sum(1 for o in options if '->' in o['text'])
        equal_count = sum(1 for o in options if ' = ' in o['text'])
        hyphen_count = sum(1 for o in options if re.search(r'\s[-–—]\s', o['text']))
        
        threshold = len(options) * 0.5
        has_arrow = arrow_count >= threshold
        has_equal = equal_count >= threshold
        has_hyphen = hyphen_count >= threshold
        
        if (has_arrow or has_equal or has_hyphen) and len(options) >= 2:
            pairs = []
            for o in options:
                parts = []
                if '->' in o['text']:
                    parts = o['text'].split('->', 1)
                elif '=' in o['text']:
                    parts = o['text'].split('=', 1)
                elif re.search(r'\s[-–—]\s', o['text']):
                    if ' - ' in o['text']:
                        parts = o['text'].split(' - ', 1)
                    elif ' – ' in o['text']:
                        parts = o['text'].split(' – ', 1)
                    elif ' — ' in o['text']:
                        parts = o['text'].split(' — ', 1)
                else:
                    continue
                    
                if len(parts) >= 2:
                    pairs.append({
                        "left": parts[0].strip(),
                        "right": parts[1].strip()
                    })
            
            if pairs:
                return {
                    "type": "matching",
                    "question": question_text,
                    "pairs": pairs,
                    "explanation": explanation.strip()
                }

        # Check if it is True/False
        # Criteria: 2 options, one True, one False (keywords)
        # Or explicit True/False keywords found
        tf_keywords = ["đúng", "sai", "true", "false", "vrai", "faux", "falsch", "wahr", "verdadero", "falso"]
        is_tf = False
        if len(options) == 2:
            opt_texts = [o['text'].lower() for o in options]
            if all(t in tf_keywords for t in opt_texts):
                is_tf = True
        
        if is_tf:
            # Determine correct answer (Boolean)
            correct_bool = True
            # Find which one is marked correct
            correct_opt = next((o for o in options if o['correct']), None)
            
            if correct_opt:
                if correct_opt['text'].lower() in ["sai", "false", "faux", "falsch", "falso"]:
                    correct_bool = False
            else:
                # If no marker, maybe check explicit_answer_char?
                if explicit_answer_char:
                    if explicit_answer_char.lower() in ["sai", "false", "faux", "falsch", "falso", "b"]:
                        correct_bool = False
            
            return {
                "type": "true_false",
                "question": question_text,
                "correct_answer": correct_bool,
                "explanation": explanation.strip()
            }

        # Default: Multiple Choice
        # We need to format options as Strings ("Content", "Content") - No prefix!
        # and set correct_answer to the Index (0, 1, 2, 3)
        
        formatted_options = []
        letters = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        correct_index = 0
        
        for i, opt in enumerate(options):
            # We used to add prefix here, but UI might prefer clean text
            # formatted_options.append(f"{letters[i] if i < len(letters) else str(i+1)}. {opt['text']}")
            formatted_options.append(opt['text'])
            
            if opt['correct']:
                correct_index = i
        
        # If explicit answer char was provided (e.g. "Answer: B"), trust it if it matches range
        if explicit_answer_char:
            clean_char = explicit_answer_char.strip(" .)")
            if len(clean_char) == 1 and clean_char.upper() in letters[:len(options)]:
                correct_letter = clean_char.upper()
                correct_index = letters.index(correct_letter)
            elif clean_char.isdigit():
                try:
                    num = int(clean_char)
                    if 1 <= num <= len(options):
                        correct_index = num - 1
                except Exception:
                    pass
        
        return {
            "type": "multiple_choice",
            "question": question_text,
            "options": formatted_options,
            "correct_answer": correct_index, # UI expects index
            "explanation": explanation.strip()
        }

    def parse_multiple_choice_questions(self, content: str) -> List[Dict]:
        """Parse multiple choice questions from text"""
        questions = []
        
        # Flexible pattern:
        # Capture from "Question X" to "Answer: Y"
        # We capture the body (text + options) and the answer
        pattern = r'(?:Câu|Question|Pregunta|Frage)\s+(\d+)[:\.]\s*(.+?)\s*(?:Đáp\s+án|Answer|Réponse|Respuesta|Antwort|Đáp\s+án\s+đúng|Correct\s+Answer):\s*([A-D])'
        
        # Split content into blocks roughly to avoid greedy matching issues across questions?
        # Actually, findall with non-greedy (.+?) usually works well for sequential blocks
        matches = re.findall(pattern, content, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            question_num, body, correct_answer = match
            
            # Now extract options from the body
            # Look for lines starting with A., B., C., D. (or A), B))
            # We assume options are at the end of the body
            
            # Regex for options: A. text
            opt_matches = list(re.finditer(r'(?:^|\n)\s*([A-D])[\.\)]\s*(.+?)(?=(?:\n\s*[A-D][\.\)])|$)', body, re.DOTALL))
            
            if not opt_matches:
                continue
                
            # Question text is everything before the first option
            first_opt_start = opt_matches[0].start()
            question_text = body[:first_opt_start].strip()
            
            options = []
            valid_options = True
            
            for m in opt_matches:
                label = m.group(1)
                text = m.group(2).strip()
                options.append({
                    "text": text,
                    "correct": label.upper() == correct_answer.upper()
                })
            
            if len(options) >= 2:
                # Extract explanation (look ahead after the match in original content)
                # This is hard with findall. 
                # Alternative: The "Answer" line was consumed. We check what's after.
                # But findall doesn't give us position.
                # We'll use a simpler explanation extraction or skip it for now.
                # Let's try to find explanation in the main loop if we iterate manually.
                
                # For now, let's just append the question.
                questions.append({
                    "type": "multiple_choice",
                    "question": question_text,
                    "options": options,
                    "explanation": "" 
                })
        
        return questions
    
    def parse_true_false_questions(self, content: str) -> List[Dict]:
        """Parse true/false questions from text"""
        questions = []
        
        # Pattern for true/false questions (supports VI/EN/FR/ES/DE)
        # Use negative lookahead to prevent matching across questions
        pattern = r'(?:Câu|Question|Pregunta|Frage)\s+(\d+)[:\.]\s*((?:(?!\n\s*(?:Câu|Question|Pregunta|Frage)\s+\d+).)+?)\n\s*(?:Đáp\s+án|Answer|Réponse|Respuesta|Antwort):\s*(Đúng|Sai|True|False|Vrai|Faux|Verdadero|Falso|Wahr|Falsch)'
        
        matches = re.findall(pattern, content, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            question_num, question_text, correct_answer = match
            
            is_true = correct_answer.lower() in ("đúng", "true", "vrai", "verdadero", "wahr")
            
            questions.append({
                "type": "true_false",
                "question": question_text.strip(),
                "correct_answer": is_true,
                "explanation": ""
            })
        
        return questions
    
    def parse_fill_blank_questions(self, content: str) -> List[Dict]:
        """Parse fill-in-the-blank questions from text"""
        questions = []
        
        # Pattern for fill blank questions (supports VI/EN/FR/ES/DE)
        # Use lookahead to stop at next question or end of string
        pattern = r'(?:Câu|Question|Pregunta|Frage)\s+(\d+)[:\.]\s*((?:(?!\n\s*(?:Câu|Question|Pregunta|Frage)\s+\d+).)+?)\n\s*(?:Đáp\s+án|Answer|Réponse|Respuesta|Antwort):\s*(.+?)(?=\n\s*(?:Câu|Question|Pregunta|Frage)|$)'
        
        matches = re.findall(pattern, content, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            question_num, question_text, correct_answer = match
            answer_text = correct_answer.strip()
            
            # Skip if it looks like True/False
            if answer_text.lower() in ("đúng", "true", "vrai", "verdadero", "wahr", "sai", "false", "faux", "falso", "falsch"):
                continue
                
            # Skip if it looks like Multiple Choice (Single letter A-D and body has options)
            if len(answer_text) == 1 and answer_text.upper() in "ABCD":
                # Check if body contains options like "A. " or "A) "
                if re.search(r'(?:^|\n)\s*[A-D][\.\)]', question_text):
                    continue
            
            questions.append({
                "type": "fill_blank",
                "question": question_text.strip(),
                "answers": [answer_text],
                "case_sensitive": False,
                "explanation": ""
            })
        
        return questions
    
    def parse_matching_questions(self, content: str) -> List[Dict]:
        """Parse matching questions from text"""
        questions = []
        
        # Look for matching blocks
        # Pattern: Question X: ... \n Column A - Column B \n ...
        pattern = r'(?:Câu|Question|Pregunta|Frage)\s+(\d+)[:\.]\s*((?:(?!\n\s*(?:Câu|Question|Pregunta|Frage)\s+\d+).)+?)\n\s*(?:Cột A|Column A|Colonne A|Columna A|Spalte A)\s*-\s*(?:Cột B|Column B|Colonne B|Columna B|Spalte B)\s*\n(.+?)(?=\n\s*(?:Câu|Question|Pregunta|Frage)|$)'
        
        matches = re.findall(pattern, content, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            question_num, question_text, pairs_text = match
            
            pairs = []
            for line in pairs_text.strip().split('\n'):
                line = line.strip()
                if not line: continue
                
                # Remove leading dash or bullet if present
                if line.startswith('- '):
                    line = line[2:]
                elif line.startswith('• '):
                    line = line[2:]
                elif line.startswith('-'): # Handle dash without space
                    line = line[1:].strip()
                
                # Split by -> or -
                if '->' in line:
                    parts = line.split('->', 1)
                elif '-' in line:
                    parts = line.split('-', 1)
                else:
                    continue

                if len(parts) == 2:
                    pairs.append({
                        "left": parts[0].strip(),
                        "right": parts[1].strip()
                    })
            
            if pairs:
                questions.append({
                    "type": "matching",
                    "question": question_text.strip(),
                    "pairs": pairs,
                    "explanation": ""
                })
        
        return questions
    
    def import_from_file(self, file_path: str) -> List[Dict]:
        """Import questions from file based on extension"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(I18n.t("import.error.file_not_found", path=file_path))
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.docx':
            return self.import_from_docx(file_path)
        elif file_ext == '.doc':
            return self.import_from_doc(file_path)
        elif file_ext == '.pdf':
            return self.import_from_pdf(file_path)
        elif file_ext == '.xlsx':
            return self.import_from_xlsx(file_path)
        else:
            raise ValueError(I18n.t("import.error.unsupported_format", ext=file_ext))

    def import_from_doc(self, file_path: str) -> List[Dict]:
        """Import questions from legacy DOC file (Windows COM fallback)"""
        # Try converting DOC -> DOCX via Word COM if available
        try:
            import win32com.client  # type: ignore
            import pythoncom  # type: ignore
        except ImportError:
            raise ImportError(I18n.t("import.error.doc_requires_win32com"))
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            tmp_docx = str(Path(file_path).with_suffix('.converted.docx'))
            wdFormatXMLDocument = 12
            doc.SaveAs(tmp_docx, FileFormat=wdFormatXMLDocument)
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()
            return self.import_from_docx(tmp_docx)
        except Exception as e:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            raise RuntimeError(I18n.t("import.error.doc_conversion_failed", error=str(e)))
    
    def import_from_docx(self, file_path: str) -> List[Dict]:
        """Import questions from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError(I18n.t("import.error.docx_missing"))
        
        questions = []
        
        try:
            doc = DocxDocument(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                full_text.append(text)
            
            # Join text and parse questions
            text_content = '\n'.join(full_text)
            raw_questions = self._parse_questions_from_text(text_content)
            
            # Validate and normalize questions
            for q in raw_questions:
                validated = self._validate_imported_question(q)
                if validated:
                    questions.append(validated)
            
        except Exception as e:
            raise RuntimeError(I18n.t("import.error.docx_read_failed", error=str(e)))
        
        return questions
    
    def import_from_pdf(self, file_path: str) -> List[Dict]:
        """Import questions from PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError(I18n.t("import.error.pdf_missing"))
        
        questions = []
        
        try:
            doc = fitz.open(file_path)
            full_text = []
            
            # Extract text from all pages
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    full_text.append(text.strip())
            
            doc.close()
            
            # Join text and parse questions
            text_content = '\n'.join(full_text)
            raw_questions = self._parse_questions_from_text(text_content)
            
            # Validate and normalize questions
            for q in raw_questions:
                validated = self._validate_imported_question(q)
                if validated:
                    questions.append(validated)
            
        except Exception as e:
            raise RuntimeError(I18n.t("import.error.pdf_read_failed", error=str(e)))
        
        return questions
    
    def import_from_xlsx(self, file_path: str) -> List[Dict]:
        """Import questions from XLSX file"""
        if not XLSX_AVAILABLE:
            raise ImportError(I18n.t("import.error.xlsx_missing"))
        
        questions = []
        
        try:
            wb = load_workbook(file_path)
            ws = wb.active
            
            # Expected column structure:
            # A: Question, B: Option1, C: Option2, D: Option3, E: Option4, 
            # F: Correct Answer, G: Explanation, H: Type
            
            for row in ws.iter_rows(min_row=2, values_only=True):  # Skip header
                if not row or not row[0]:  # Skip empty rows
                    continue
                
                question_data = {
                    "question": str(row[0]) if row[0] else "",
                    "options": [],
                    "correct_answer": str(row[5]) if len(row) > 5 and row[5] else "",
                    "explanation": str(row[6]) if len(row) > 6 and row[6] else "",
                    "type": str(row[7]).lower().replace(" ", "_") if len(row) > 7 and row[7] else "multiple_choice"
                }
                
                # Build options
                if len(row) > 4:
                    options = []
                    for i, option in enumerate(row[1:5], 1):
                        if option:
                            options.append(f"{chr(64 + i)}. {option}")
                    question_data["options"] = options
                
                # Validate and fix question
                validated_question = self._validate_imported_question(question_data)
                if validated_question:
                    questions.append(validated_question)
            
            wb.close()
            
        except Exception as e:
            raise RuntimeError(I18n.t("import.error.xlsx_read_failed", error=str(e)))
        
        return questions
    
    def _parse_questions_from_text(self, text: str) -> List[Dict]:
        """Parse questions from plain text"""
        questions = []
        
        try:
            text = self._normalize_input(text)
        except Exception:
            text = text.replace('\r\n', '\n').strip()
        
        # Parse using specific parsers
        # Try smart parse first as it handles blocks better
        smart_questions = self.parse_smart_format(text)
        if smart_questions:
            questions.extend(smart_questions)
        
        # Legacy parsers are disabled to prevent duplicates and ensure consistent explanation extraction
        # questions.extend(self.parse_multiple_choice_questions(text))
        # questions.extend(self.parse_true_false_questions(text))
        # questions.extend(self.parse_fill_blank_questions(text))
        # questions.extend(self.parse_matching_questions(text))
        
        return questions

    def _parse_custom_marker_blocks(self, text: str) -> List[Dict]:
        """Parse questions using custom markers:
        * starts a question
        Options are lines that look like A., a), 1), a., 1) ...
        Correct options are prefixed with '=='
        Explanation starts with '/' and continues until next '*'
        Supports multiple_choice, true_false, fill_blank, matching (left - right)
        """
        lines = [l.rstrip() for l in text.splitlines()]
        blocks: List[List[str]] = []
        current: Optional[List[str]] = None
        for ln in lines:
            stripped = ln.strip()
            if stripped.startswith('*'):
                if current:
                    blocks.append(current)
                current = [stripped[1:].strip()]
                continue
            if current is None:
                continue
            current.append(ln)
        if current:
            blocks.append(current)
        results: List[Dict] = []
        for blk in blocks:
            if not blk:
                continue
            question_text_parts: List[str] = []
            options: List[Dict] = []
            pairs: List[Dict] = []
            explanation_parts: List[str] = []
            in_explain = False
            seen_option = False
            for idx, raw in enumerate(blk):
                raw_stripped = str(raw).strip()
                if in_explain and re.match(r'^[A-D]\.\s+.+', raw_stripped):
                    break
                if raw_stripped.startswith('/'):
                    in_explain = True
                    explanation_parts.append(raw_stripped[1:].strip())
                    continue
                if in_explain:
                    explanation_parts.append(raw_stripped)
                    continue
                # Matching pair: left - right
                m_pair = re.match(r'^\s*(==)?\s*(.+?)\s*[-–>→]\s*(.+)\s*$', raw_stripped)
                if m_pair and not seen_option:
                    pairs.append({"left": m_pair.group(2).strip(), "right": m_pair.group(3).strip()})
                    continue
                # Option markers
                m_opt = re.match(r'^\s*(==)?\s*([A-Da-d]|\d{1,2})[\.\)]\s*(.+)\s*$', raw_stripped)
                if m_opt:
                    seen_option = True
                    is_correct = bool(m_opt.group(1))
                    opt_text = m_opt.group(3).strip()
                    options.append({"text": opt_text, "correct": is_correct})
                    continue
                # Continuation lines
                if seen_option and raw_stripped:
                    try:
                        options[-1]["text"] = f"{options[-1]['text']} {raw_stripped}".strip()
                    except Exception:
                        pass
                    continue
                
                # Ignore answer lines for fill_blank (starting with == but not option/pair)
                if raw_stripped.startswith('==') and idx > 0:
                    continue

                # Question text parts
                if raw_stripped:
                    question_text_parts.append(raw_stripped)
            q_text = ' '.join([p for p in question_text_parts if p]).strip()
            expl = ' '.join([p for p in explanation_parts if p]).strip()
            if pairs and not options:
                results.append({
                    "type": "matching",
                    "question": q_text or I18n.t("import.default.matching_question"),
                    "pairs": pairs,
                    "explanation": expl
                })
                continue
            if not options:
                # Fill blank: use any line prefixed with '==' as answer otherwise leave empty
                answer_lines = [re.sub(r'^\s*==\s*', '', l).strip() for l in blk if str(l).strip().startswith('==')]
                results.append({
                    "type": "fill_blank",
                    "question": q_text or I18n.t("import.default.fill_blank_question"),
                    "answers": answer_lines if answer_lines else [],
                    "case_sensitive": False,
                    "explanation": expl
                })
                continue
            # Determine type true/false if 2 options and match
            texts_lower = [str(o.get('text','')).strip().lower() for o in options]
            if len(options) == 2 and set(texts_lower) <= {"true","false","đúng","sai"}:
                correct_bool = options[0].get('correct', False)
                # if no marker, default first to True
                if not any(o.get('correct') for o in options):
                    correct_bool = texts_lower[0] in {"true","đúng"}
                results.append({
                    "type": "true_false",
                    "question": q_text,
                    "correct_answer": True if correct_bool else False,
                    "explanation": expl
                })
                continue
            # Multiple choice: build A-D options text list and correct letter
            texts = [o['text'] for o in options]
            letters = ['A','B','C','D']
            correct_letter = None
            for idx, o in enumerate(options):
                if o.get('correct'):
                    try:
                        correct_letter = letters[idx]
                    except Exception:
                        correct_letter = None
                    break
            results.append({
                "type": "multiple_choice",
                "question": q_text,
                "options": [f"{letters[i] if i < len(letters) else str(i+1)}. {texts[i]}" for i in range(len(texts))],
                "correct_answer": correct_letter or 'A',
                "explanation": expl
            })
        return results
    
    def _extract_question_from_match(self, match) -> Optional[Dict]:
        """Extract question data from regex match"""
        groups = match.groups()
        
        if len(groups) < 2:
            return None
        
        question_data = {
            "question": groups[1].strip(),
            "options": [],
            "correct_answer": "",
            "explanation": "",
            "type": "multiple_choice"
        }
        
        # Extract options (A, B, C, D)
        if len(groups) >= 6:
            for i in range(2, 6):
                if i < len(groups) and groups[i]:
                    option_text = groups[i].strip()
                    if option_text and not option_text.startswith('Answer:'):
                        question_data["options"].append(option_text)
        
        # Extract answer
        if len(groups) >= 7 and groups[6]:
            answer_text = groups[6].strip().upper()
            if len(answer_text) == 1 and answer_text in 'ABCD':
                question_data["correct_answer"] = answer_text
        
        # Extract explanation
        if len(groups) >= 8 and groups[7]:
            explanation = groups[7].strip()
            if explanation and not explanation.startswith('['):
                question_data["explanation"] = explanation
        
        return question_data
    
    def _validate_imported_question(self, question_data: Dict) -> Optional[Dict]:
        """Validate and fix imported question data"""
        if not question_data.get("question"):
            return None
        
        # Ensure required fields
        question_data["type"] = question_data.get("type", "multiple_choice")
        question_data["explanation"] = question_data.get("explanation", "")
        
        qt = (question_data.get("type") or "multiple_choice").lower().strip()
        
        # Map common type names to internal names
        if qt in ("fill_blank", "fill_in_the_blank", "dien_khuyet"):
            qt = "fill_blank"
        elif qt in ("matching", "ghep_cap"):
            qt = "matching"
        elif qt in ("true_false", "dung_sai"):
            qt = "true_false"
        elif qt in ("multiple_choice", "trac_nghiem"):
            qt = "multiple_choice"
            
        question_data["type"] = qt

        if qt == "multiple_choice":
            # Handle list of dicts for options (from smart parser)
            opts = question_data.get("options")
            if opts and isinstance(opts, list) and len(opts) > 0 and isinstance(opts[0], dict):
                # Extract text and find correct answer
                new_options = []
                correct_idx = -1
                for i, opt in enumerate(opts):
                    text = opt.get("text", "")
                    new_options.append(text)
                    if opt.get("correct"):
                        correct_idx = i
                
                question_data["options"] = new_options
                # If correct_answer wasn't set explicitly, use the one found in options
                if question_data.get("correct_answer") is None or question_data.get("correct_answer") == "":
                    if correct_idx != -1:
                        question_data["correct_answer"] = correct_idx

            if not question_data.get("options"):
                question_data["options"] = [
                    f"A. {I18n.t('import.default.option_1')}",
                    f"B. {I18n.t('import.default.option_2')}", 
                    f"C. {I18n.t('import.default.option_3')}",
                    f"D. {I18n.t('import.default.option_4')}"
                ]
            
            # Normalize correct_answer to index (int)
            ca = question_data.get("correct_answer")
            if ca is None or ca == "":
                question_data["correct_answer"] = 0
            elif isinstance(ca, str):
                # Convert "A", "B", "C", "D" to 0, 1, 2, 3
                clean_ca = ca.strip().upper()
                if len(clean_ca) == 1 and clean_ca in "ABCDEF":
                    question_data["correct_answer"] = ord(clean_ca) - ord('A')
                elif clean_ca.isdigit():
                    question_data["correct_answer"] = int(clean_ca)
            # If it's already an int (like 0), leave it alone

                
        elif qt == "true_false":
            # Ensure correct_answer is boolean or A/B
            ca = question_data.get("correct_answer")
            if isinstance(ca, str):
                ca = ca.strip().upper()
                if ca in ('A', 'TRUE', 'ĐÚNG', 'DUNG'):
                    question_data["correct_answer"] = True
                else:
                    question_data["correct_answer"] = False
            elif ca is None:
                 question_data["correct_answer"] = True
                 
        elif qt == "fill_blank":
            # Ensure correct_answers is a list
            ans = question_data.get("answers") or question_data.get("correct_answers")
            if not ans:
                ans = []
            if not isinstance(ans, list):
                ans = [str(ans)]
            question_data["correct_answers"] = ans
            # Remove legacy field if present
            if "answers" in question_data:
                del question_data["answers"]
                
        elif qt == "matching":
            # Ensure pairs is a list of dicts
            pairs = question_data.get("pairs") or []
            if not isinstance(pairs, list):
                pairs = []
            # Validate pair structure
            valid_pairs = []
            for p in pairs:
                if isinstance(p, dict) and 'left' in p and 'right' in p:
                    valid_pairs.append(p)
            question_data["pairs"] = valid_pairs
        
        # Clean up question text
        question_data["question"] = self._clean_text(question_data["question"])
        question_data["explanation"] = self._clean_text(question_data["explanation"])
        
        # Clean up options
        if "options" in question_data:
            question_data["options"] = [self._clean_text(opt) for opt in question_data["options"] if opt.strip()]
        
        return question_data
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""

        try:
            text = self._normalize_input(text)
        except Exception:
            pass

        text = re.sub(r'\s+', ' ', str(text).strip())
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        return text.strip()
    
    def create_sample_template(self, file_path: str, format_type: str = "docx", lang: str = "vi", template_type: str = "general"):
        """Create a sample template file for teachers"""
        if format_type == "docx" and DOCX_AVAILABLE:
            self._create_sample_docx(file_path, lang, template_type)
        elif format_type == "doc":
            self._create_sample_doc(file_path, lang, template_type)
        elif format_type == "xlsx" and XLSX_AVAILABLE:
            self._create_sample_xlsx(file_path, lang, template_type)
        elif format_type == "txt":
            self._create_sample_txt(file_path, lang, template_type)
        else:
            raise ValueError(I18n.t("import.error.template_creation_not_supported", format=format_type))

    def _create_sample_txt(self, file_path: str, lang: str, template_type: str = "general"):
        """Create sample TXT template by copying from assets"""
        # Determine source file name
        is_vi = lang == "vi"
        
        if template_type == "millionaire":
            src_name = "mau_cau_hoi_trieu_phu_vi.txt" if is_vi else "sample_questions_millionaire_en.txt"
        else:
            src_name = "mau_cau_hoi_tong_hop_vi.txt" if is_vi else "sample_questions_general_en.txt"
            
        src_path = self.templates_dir / src_name
        
        if not src_path.exists():
            # Fallback to English general if specific not found
            src_path = self.templates_dir / "sample_questions_general_en.txt"
            
        if not src_path.exists():
             # Last resort fallback content
             with open(file_path, "w", encoding="utf-8") as f:
                 f.write("Template file not found.")
             return

        # Copy content
        import shutil
        shutil.copy2(src_path, file_path)
    
    def _create_sample_docx(self, file_path: str, lang: str, template_type: str = "general"):
        """Create sample DOCX template"""
        doc = DocxDocument()
        l = lang or "vi"
        
        doc.add_heading(I18n.t("template.docx.title", l), 0)
        doc.add_paragraph(I18n.t("template.docx.subtitle", l))
        
        # Add notes about game types
        if template_type == "millionaire":
             doc.add_paragraph(I18n.t("template.docx.note_millionaire", l))
        else:
             doc.add_paragraph(I18n.t("template.docx.note_millionaire", l))
             doc.add_paragraph(I18n.t("template.docx.note_others", l))
        doc.add_paragraph("")
        
        doc.add_heading(I18n.t("template.docx.instructions_heading", l), level=1)
        doc.add_paragraph(I18n.t("template.docx.instructions_1", l))
        doc.add_paragraph(I18n.t("template.docx.instructions_2", l))
        doc.add_paragraph(I18n.t("template.docx.instructions_3", l))
        doc.add_paragraph(I18n.t("template.docx.instructions_4", l))
        doc.add_paragraph(I18n.t("template.docx.instructions_5", l))
        doc.add_paragraph("")
        
        doc.add_heading(I18n.t("template.docx.samples_heading", l), level=1)
        doc.add_paragraph("")

        # 1. Multiple Choice (Always included)
        doc.add_heading(I18n.t("template.docx.sample_mc_heading", l), level=2)
        doc.add_paragraph(I18n.t("template.docx.sample_mc_q", l))
        doc.add_paragraph(I18n.t("template.docx.sample_mc_a1", l))
        doc.add_paragraph(I18n.t("template.docx.sample_mc_a2", l))
        doc.add_paragraph(I18n.t("template.docx.sample_mc_a3", l))
        doc.add_paragraph(I18n.t("template.docx.sample_mc_a4", l))
        doc.add_paragraph(I18n.t("template.docx.sample_mc_expl", l))
        doc.add_paragraph("")

        if template_type != "millionaire":
            # 2. True/False
            doc.add_heading(I18n.t("template.docx.sample_tf_heading", l), level=2)
            doc.add_paragraph(I18n.t("template.docx.sample_tf_q", l))
            doc.add_paragraph(I18n.t("template.docx.sample_tf_true", l))
            doc.add_paragraph(I18n.t("template.docx.sample_tf_false", l))
            doc.add_paragraph(I18n.t("template.docx.sample_tf_expl", l))
            doc.add_paragraph("")

            # 3. Fill in the Blank
            doc.add_heading(I18n.t("template.docx.sample_fb_heading", l), level=2)
            doc.add_paragraph(I18n.t("template.docx.sample_fb_q", l))
            fb_ans = I18n.t("template.docx.sample_fb_answer", l)
            if fb_ans:
                doc.add_paragraph(fb_ans)
            doc.add_paragraph(I18n.t("template.docx.sample_fb_expl", l))
            doc.add_paragraph("")

            # 4. Matching
            doc.add_heading(I18n.t("template.docx.sample_match_heading", l), level=2)
            doc.add_paragraph(I18n.t("template.docx.sample_match_q", l))
            doc.add_paragraph(I18n.t("template.docx.sample_match_pair1", l))
            doc.add_paragraph(I18n.t("template.docx.sample_match_pair2", l))
            doc.add_paragraph(I18n.t("template.docx.sample_match_expl", l))
            doc.add_paragraph("")

            # 5. Short Answer
            doc.add_heading(I18n.t("template.docx.sample_sa_heading", l), level=2)
            doc.add_paragraph(I18n.t("template.docx.sample_sa_q", l))
            doc.add_paragraph(I18n.t("template.docx.sample_sa_answer", l))
            doc.add_paragraph(I18n.t("template.docx.sample_sa_expl", l))
            doc.add_paragraph("")
        
        doc.save(file_path)
    
    def _create_sample_xlsx(self, file_path: str, lang: str = "vi", template_type: str = "general"):
        """Create sample XLSX template"""
        from openpyxl import Workbook
        
        l = lang or "vi"
        wb = Workbook()
        ws = wb.active
        ws.title = I18n.t("import.xlsx.sheet_title", l)
        
        # Headers
        headers = [
            I18n.t("import.xlsx.header.question", l), 
            I18n.t("import.xlsx.header.option_a", l), 
            I18n.t("import.xlsx.header.option_b", l), 
            I18n.t("import.xlsx.header.option_c", l), 
            I18n.t("import.xlsx.header.option_d", l), 
            I18n.t("import.xlsx.header.correct_answer", l), 
            I18n.t("import.xlsx.header.explanation", l), 
            I18n.t("import.xlsx.header.type", l)
        ]
        
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Sample data
        sample_data = [
            [
                I18n.t("import.sample.q1", l), 
                I18n.t("import.sample.q1_opt1", l), 
                I18n.t("import.sample.q1_opt2", l), 
                I18n.t("import.sample.q1_opt3", l), 
                I18n.t("import.sample.q1_opt4", l), 
                "A", 
                I18n.t("import.sample.q1_expl", l), 
                "multiple_choice"
            ],
            [
                I18n.t("import.sample.q2", l), 
                "100", "50", "75", "25", 
                "A", 
                I18n.t("import.sample.q2_expl", l), 
                "multiple_choice"
            ]
        ]

        if template_type != "millionaire":
            sample_data.append([
                I18n.t("import.sample.q3", l), 
                I18n.t("import.default.true", l), 
                I18n.t("import.default.false", l), 
                "", "", 
                "A", 
                I18n.t("import.sample.q3_expl", l), 
                "true_false"
            ])
        
        for row_idx, row_data in enumerate(sample_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(file_path)

    def _create_sample_doc(self, file_path: str, lang: str = "vi", template_type: str = "general"):
        """Create sample DOC template by generating DOCX then converting via Word COM if available"""
        try:
            tmp_docx = str(Path(file_path).with_suffix('.tmp.docx'))
            self._create_sample_docx(tmp_docx, lang, template_type)
            try:
                import win32com.client  # type: ignore
                import pythoncom  # type: ignore
            except ImportError:
                # Fallback: just rename to .docx if COM not available
                Path(tmp_docx).replace(str(Path(file_path).with_suffix('.docx')))
                return
            try:
                pythoncom.CoInitialize()
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(tmp_docx)
                wdFormatDocument = 0
                doc.SaveAs(file_path, FileFormat=wdFormatDocument)
                doc.Close(False)
                word.Quit()
                pythoncom.CoUninitialize()
            finally:
                try:
                    Path(tmp_docx).unlink(missing_ok=True)
                except Exception:
                    pass
        except Exception as e:
            raise RuntimeError(I18n.t("import.error.doc_template_failed", error=str(e)))
"""
Nguyen-Thanh-Tan ¬_¬
"""
