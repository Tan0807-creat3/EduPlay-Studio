"""
Main Window - Central window for EduPlay Studio
"""

from PySide6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                               QPushButton, QLabel, QStackedWidget, QFrame, 
                               QDialog, QTextEdit, QTabWidget, QLineEdit, 
                               QComboBox, QSpinBox, QMessageBox, QSystemTrayIcon)
from PySide6.QtCore import Qt, Signal, QSize, QPropertyAnimation, QEasingCurve, QEvent, QPoint, QSequentialAnimationGroup, QParallelAnimationGroup, QObject, QThread, QPauseAnimation, QRect, QVariantAnimation, QTimer, QProcess
from PySide6.QtWidgets import QGraphicsOpacityEffect, QGraphicsBlurEffect, QGraphicsEffect, QGraphicsDropShadowEffect
from PySide6.QtGui import QFont, QIcon, QMovie, QFontMetrics, QPixmap, QPainter, QImage, QFontDatabase, QPainterPath, QLinearGradient, QColor, QPen, QCursor, QShortcut, QKeySequence
import os
import sys

from eduplay.core.asset_loader import load_asset_text, materialize_asset_file
from eduplay.core.command_palette import build_palette_items
from eduplay.core.path_resolver import PathResolver
from eduplay.ui.screens.home_screen import HomeScreen
from eduplay.ui.screens.new_project_screen import NewProjectScreen
from eduplay.ui.screens.browser_screen import BrowserScreen
from eduplay.ui.screens.player_screen import PlayerScreen
from eduplay.ui.screens.editor_screen import EditorScreen
from eduplay.ui.screens.editor_quiz_classic_screen import EditorQuizClassicScreen
from eduplay.ui.screens.editor_fishing_screen import EditorFishingScreen

from eduplay.ui.widgets.chat_widget import ChatWidget
from eduplay.ui.widgets.command_palette_dialog import CommandPaletteDialog
from eduplay.ui.widgets.settings_dialog import SettingsDialog
from eduplay.ui.widgets.left_nav_drawer import LeftNavDrawer, LeftEdgeHotZone
from eduplay.ui.icon_factory import build_standard_ui_icon, build_app_action_icon


class _BackgroundWorker(QObject):
    finished = Signal(object, object)

    def __init__(self, fn, *args, **kwargs):
        super().__init__()
        self._fn = fn
        self._args = args
        self._kwargs = kwargs

    def run(self):
        try:
            print("[Export] Background worker started")
        except Exception:
            pass
        result = None
        error = None
        try:
            result = self._fn(*self._args, **self._kwargs)
        except Exception as e:
            error = e
        try:
            print("[Export] Background worker finished, error:", error)
        except Exception:
            pass
        self.finished.emit(result, error)


class _StartupBoundaryWipeEffect(QGraphicsEffect):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._cut_y = 0

    def setCutY(self, cut_y: int):
        try:
            cut = int(cut_y)
        except Exception:
            cut = 0
        if cut < 0:
            cut = 0
        if cut != self._cut_y:
            self._cut_y = cut
            try:
                self.update()
            except Exception:
                pass

    def draw(self, painter: QPainter):
        offset = QPoint(0, 0)
        try:
            pixmap = self.sourcePixmap(
                Qt.LogicalCoordinates,
                offset,
                QGraphicsEffect.PadToEffectiveBoundingRect,
            )
        except Exception:
            try:
                pixmap = self.sourcePixmap(Qt.LogicalCoordinates, offset)
            except Exception:
                pixmap = self.sourcePixmap()
        if not pixmap or pixmap.isNull():
            return
        h = pixmap.height()
        cut = self._cut_y
        if cut <= 0:
            painter.drawPixmap(offset, pixmap)
            return
        if cut >= h:
            return
        painter.save()
        try:
            painter.setClipRect(offset.x(), offset.y() + cut, pixmap.width(), h - cut)
        except Exception:
            painter.restore()
            painter.drawPixmap(offset, pixmap)
            return
        painter.drawPixmap(offset, pixmap)
        painter.restore()


class _StartupLeftToRightWipeEffect(QGraphicsEffect):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._cut_x = 0

    def setCutX(self, cut_x: int):
        try:
            cut = int(cut_x)
        except Exception:
            cut = 0
        if cut < 0:
            cut = 0
        if cut != self._cut_x:
            self._cut_x = cut
            try:
                self.update()
            except Exception:
                pass

    def draw(self, painter: QPainter):
        offset = QPoint(0, 0)
        try:
            pixmap = self.sourcePixmap(
                Qt.LogicalCoordinates,
                offset,
                QGraphicsEffect.PadToEffectiveBoundingRect,
            )
        except Exception:
            try:
                pixmap = self.sourcePixmap(Qt.LogicalCoordinates, offset)
            except Exception:
                pixmap = self.sourcePixmap()
        if not pixmap or pixmap.isNull():
            return
        w = pixmap.width()
        cut = self._cut_x
        if cut <= 0:
            painter.drawPixmap(offset, pixmap)
            return
        if cut >= w:
            return
        painter.save()
        try:
            painter.setClipRect(offset.x() + cut, offset.y(), w - cut, pixmap.height())
        except Exception:
            painter.restore()
            painter.drawPixmap(offset, pixmap)
            return
        painter.drawPixmap(offset, pixmap)
        painter.restore()


class MainWindow(QMainWindow):
    """Main application window"""
    
    def __init__(self, project_manager=None, ai_service=None, asset_manager=None, settings_manager=None, import_service=None):
        super().__init__()
        self.project_manager = project_manager
        self.ai_service = ai_service
        self.asset_manager = asset_manager
        self.settings_manager = settings_manager
        self.import_service = import_service
        self._background_threads = []
        self._loading_overlay = None
        self._loading_label = None
        self._loading_movie = None
        self._loading_blur = None
        self._startup_overlay = None
        self._startup_logo_anim = None
        self._startup_logo_locked_rect = None
        self._startup_logo_hidden_after_dock = False
        self._startup_home_fade_done = False
        self._startup_home_fade_anim = None
        self._startup_title_timer = None
        self._startup_exit_delay = None
        self._startup_text_exit_timer = None
        self._startup_title_full = ""
        self._startup_title_index = 0
        self._startup_text_anim_group = None
        self._startup_char_labels = []
        self._startup_animation_requested = False
        self._home_header_title_anim = None
        self._chat_request_seq = 0
        self.setWindowTitle("EduPlay Studio")
        self.setMinimumSize(900, 600)
        try:
            from PySide6.QtWidgets import QApplication
            screen = QApplication.primaryScreen().geometry()
            w = max(1000, int(screen.width()*0.75))
            h = max(680, int(screen.height()*0.75))
            self.resize(w, h)
        except Exception:
            self.resize(1200, 800)
        icon_path = None
        try:
            icon_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../resources/icons"))
            candidates = [
                os.path.join(icon_dir, "icon.ico"),
                os.path.join(icon_dir, "icon.png"),
            ]
            for p in candidates:
                if os.path.exists(p):
                    icon_path = p
                    self.setWindowIcon(QIcon(p))
                    break
        except Exception:
            pass
        try:
            if QSystemTrayIcon.isSystemTrayAvailable():
                if not hasattr(self, "_tray_icon") or self._tray_icon is None:
                    self._tray_icon = QSystemTrayIcon(self)
                    ico = self.windowIcon()
                    if ico and not ico.isNull():
                        self._tray_icon.setIcon(ico)
                    else:
                        try:
                            from PySide6.QtGui import QIcon as _QIcon
                            if icon_path:
                                self._tray_icon.setIcon(_QIcon(icon_path))
                        except Exception:
                            pass
                    self._tray_icon.setToolTip("EduPlay Studio")
                    self._tray_icon.show()
        except Exception:
            self._tray_icon = None
        
        # Central widget and main layout
        self.central_widget = QWidget()
        self.central_widget.setObjectName("main-central")
        self.central_widget.setAttribute(Qt.WA_StyledBackground, True)
        self.setCentralWidget(self.central_widget)
        self.main_layout = QVBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        
        # Create header
        self.create_header()
        
        # Create stacked widget for screens
        self.stacked_widget = QStackedWidget()
        self.stacked_widget.setAttribute(Qt.WA_TranslucentBackground, True)
        self.stacked_widget.setStyleSheet("QStackedWidget { background: transparent; }")
        self.main_layout.addWidget(self.stacked_widget)
        
        # Create screens
        self.create_screens()

        try:
            self._init_left_nav_drawer()
        except Exception:
            self._left_nav_drawer = None
            self._left_nav_hotzone = None
            self._left_nav_hide_timer = None
        
        # Create floating chat widget
        self.chat_widget = ChatWidget()
        self.chat_widget.hide()
        try:
            if hasattr(self.chat_widget, 'set_language') and self.settings_manager:
                self.chat_widget.set_language(self.settings_manager.get_language())
        except Exception:
            pass
        try:
            self.chat_widget.dock_mode_changed.connect(self.on_chat_dock_changed)
            self.chat_widget.width_mode_changed.connect(self.on_chat_width_changed)
            self.chat_widget.detail_level_changed.connect(self.on_chat_detail_changed)
            self._chat_detail_level = 'chuẩn'
        except Exception:
            pass

        # Global floating chat toggle button
        self._chat_toggle_btn = QPushButton("", self)
        self._chat_toggle_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #7F56D9;
                border: none;
                border-radius: 24px;
                color: #FFFFFF;
                min-width: 48px;
                max-width: 48px;
                min-height: 48px;
                max-height: 48px;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #8B66E9;
            }
            QPushButton:pressed {
                background-color: #6B46C1;
            }
            """
        )
        try:
            self._chat_toggle_btn.setFixedSize(48, 48)
        except Exception:
            pass
        try:
            self._chat_toggle_btn.setIcon(build_app_action_icon("chat", self.style(), color_hex="#FFFFFF", size=22))
            self._chat_toggle_btn.setIconSize(QSize(22, 22))
            self._chat_toggle_btn.setToolTip("Edubot")
        except Exception:
            pass
        self._chat_toggle_btn.clicked.connect(self.show_chat)
        self._chat_toggle_btn.raise_()
        self._chat_toggle_btn.show()
        self._chat_drag_active = False
        self._chat_drag_offset = QPoint()
        self._chat_drag_moved = False
        self._chat_drag_start_pos = QPoint()
        self._chat_anchor_edge = 'right'
        self._chat_anchor_ratio = None
        try:
            if self.settings_manager:
                try:
                    edge = self.settings_manager.get("ui.chat_toggle.edge", None)
                except Exception:
                    edge = None
                try:
                    ratio_val = self.settings_manager.get("ui.chat_toggle.ratio", None)
                except Exception:
                    ratio_val = None
                if edge in ('left', 'right') and ratio_val is not None:
                    try:
                        self._chat_anchor_edge = edge
                        self._chat_anchor_ratio = float(ratio_val)
                    except Exception:
                        pass
        except Exception:
            pass
        try:
            btn = self._chat_toggle_btn
            size = btn.size()
            margin_x = 0
            margin_y = 20
            w = max(1, self.width())
            h = max(1, self.height())
            edge = getattr(self, '_chat_anchor_edge', 'right')
            ratio = getattr(self, '_chat_anchor_ratio', None)
            if ratio is None:
                x = w - size.width() - margin_x
                try:
                    initial_offset = 180
                except Exception:
                    initial_offset = 180
                y = h - size.height() - initial_offset
                y = max(margin_y, min(h - size.height() - margin_y, y))
            else:
                if edge == 'left':
                    x = margin_x
                else:
                    x = w - size.width() - margin_x
                center_y = int(h * float(ratio))
                y = center_y - size.height() // 2
                y = max(margin_y, min(h - size.height() - margin_y, y))
            btn.move(x, y)
        except Exception:
            pass
        self._chat_toggle_btn.installEventFilter(self)
        self._command_palette = None
        self._command_palette_shortcut = None
        
        # Connect signals
        self.connect_signals()
        self._init_command_palette()
        self._sync_home_quick_actions()
        
        # Show home screen by default
        self.show_screen("home")
        try:
            self._apply_theme_on_start()
        except Exception:
            pass

        try:
            from PySide6.QtCore import QTimer
            first_run = False
            try:
                if self.settings_manager:
                    if hasattr(self.settings_manager, "should_run_first_time_flow"):
                        first_run = bool(self.settings_manager.should_run_first_time_flow())
                    else:
                        # Nếu settings.json đã tồn tại trước khi mở app thì không còn là first_run nữa,
                        # kể cả khi key "first_run" trong file vẫn là True (ví dụ từ phiên bản cũ).
                        if getattr(self.settings_manager, "settings_file_existed", False):
                            first_run = False
                        else:
                            first_run = bool(self.settings_manager.get("first_run", True))
            except Exception:
                first_run = True
            try:
                self._pending_first_run_onboarding = False
            except Exception:
                self._pending_first_run_onboarding = False
            try:
                self._pending_whats_new_rc2 = False
            except Exception:
                self._pending_whats_new_rc2 = False
            try:
                if self.settings_manager:
                    lang0 = self.settings_manager.get_language() or "en"
                else:
                    lang0 = "en"
            except Exception:
                lang0 = "en"
            try:
                if self.settings_manager and hasattr(self.settings_manager, "needs_rc2_whats_new"):
                    dismissed = bool(self.settings_manager.get("whats_new.rc2.dismissed", False))
                    if (not dismissed) and bool(self.settings_manager.needs_rc2_whats_new()):
                        try:
                            self.settings_manager.ensure_rc2_header()
                        except Exception:
                            pass
                        self._pending_whats_new_rc2 = True
                        self._pending_whats_new_lang = lang0
            except Exception:
                pass
            if first_run:
                try:
                    onb_dismissed = bool(self.settings_manager.get("onboarding.dismissed", False)) if self.settings_manager else False
                except Exception:
                    onb_dismissed = False
                try:
                    wn_dismissed = bool(self.settings_manager.get("whats_new.rc2.dismissed", False)) if self.settings_manager else False
                except Exception:
                    wn_dismissed = False
                if not onb_dismissed:
                    self._pending_first_run_onboarding = True
                    self._pending_onboarding_lang = lang0
                # On a genuine first launch, also surface the "What's New"
                # window before the onboarding guide so users see what's new
                # together with the usage walkthrough.
                if not wn_dismissed and not getattr(self, "_pending_whats_new_rc2", False):
                    self._pending_whats_new_rc2 = True
                    if not getattr(self, "_pending_whats_new_lang", None):
                        self._pending_whats_new_lang = lang0
            try:
                QTimer.singleShot(1000, self._auto_install_ppt_addin_if_needed)
            except Exception:
                pass
            try:
                QTimer.singleShot(4000, self._start_powerpoint_mru_watcher_if_needed)
            except Exception:
                pass
        except Exception:
            pass

    def _auto_install_ppt_addin_if_needed(self):
        try:
            if not self.settings_manager:
                return
            enabled = bool(self.settings_manager.get("ppt_addin.auto_install_on_new_machine", True))
            installed_once = bool(self.settings_manager.get("ppt_addin.installed_once", False))
            if not enabled or installed_once:
                return
            from eduplay.core.ppt_vsto_addin_service import PptVstoAddinService

            def _do_install():
                return PptVstoAddinService(self.settings_manager).install_or_update()

            def _on_finished(result, error):
                try:
                    if error is not None:
                        try:
                            self.settings_manager.set("ppt_addin.installed_once", False)
                        except Exception:
                            pass
                        return
                    if not isinstance(result, dict):
                        try:
                            self.settings_manager.set("ppt_addin.installed_once", False)
                        except Exception:
                            pass
                        return
                    if result.get("installed"):
                        try:
                            from eduplay.core.i18n import I18n
                            lang = self.settings_manager.get_language() if self.settings_manager else "en"
                            title = I18n.t("settings.addin.autoinstall_notify_title", lang)
                            msg = I18n.t("settings.addin.autoinstall_notify_desc", lang)
                        except Exception:
                            title = "PowerPoint Add-in"
                            msg = "PowerPoint add-in installed."
                        try:
                            self.show_system_notification(title, msg, "success")
                        except Exception:
                            pass
                        try:
                            self._auto_seed_trusted_slides_if_needed()
                        except Exception:
                            pass
                    else:
                        try:
                            self.settings_manager.set("ppt_addin.installed_once", False)
                        except Exception:
                            pass
                except Exception:
                    pass

            try:
                self._run_background_task(_do_install, _on_finished)
            except Exception:
                pass
        except Exception:
            pass

    def _auto_seed_trusted_slides_if_needed(self):
        try:
            if not self.settings_manager:
                return
            enabled = bool(self.settings_manager.get("ppt_addin.auto_seed_trusted_slides", True))
            done = bool(self.settings_manager.get("ppt_addin.trusted_seed_done", False))
            if not enabled or done:
                return
            from eduplay.core.ppt_vsto_addin_service import PptVstoAddinService

            def _do_seed():
                return PptVstoAddinService(self.settings_manager).seed_trusted_slides_from_documents(max_files=1)

            def _on_finished(result, error):
                try:
                    if error is not None:
                        return
                    if isinstance(result, dict) and (result.get("copied") or result.get("skipped")):
                        try:
                            self.settings_manager.set("ppt_addin.trusted_seed_done", True)
                        except Exception:
                            pass
                except Exception:
                    pass

            try:
                self._run_background_task(_do_seed, _on_finished)
            except Exception:
                pass
        except Exception:
            pass

    def _start_powerpoint_mru_watcher_if_needed(self):
        try:
            if not self.settings_manager:
                return
            enabled = bool(self.settings_manager.get("ppt_addin.auto_seed_from_powerpoint_mru", True))
            if not enabled:
                return
            installed_once = bool(self.settings_manager.get("ppt_addin.installed_once", False))
            if not installed_once:
                return
            try:
                from PySide6.QtCore import QTimer
            except Exception:
                return
            try:
                if getattr(self, "_ppt_mru_timer", None) is not None:
                    return
            except Exception:
                pass
            try:
                sec = int(self.settings_manager.get("ppt_addin.powerpoint_mru_poll_sec", 15) or 15)
            except Exception:
                sec = 15
            sec = max(10, sec)
            t = QTimer(self)
            t.setInterval(int(sec * 1000))
            t.timeout.connect(self._poll_powerpoint_mru_and_seed_trusted)
            self._ppt_mru_timer = t
            try:
                t.start()
            except Exception:
                pass
            try:
                self._poll_powerpoint_mru_and_seed_trusted()
            except Exception:
                pass
        except Exception:
            pass

    def _poll_powerpoint_mru_and_seed_trusted(self):
        try:
            if not self.settings_manager:
                return
            enabled = bool(self.settings_manager.get("ppt_addin.auto_seed_from_powerpoint_mru", True))
            if not enabled:
                return
            installed_once = bool(self.settings_manager.get("ppt_addin.installed_once", False))
            if not installed_once:
                return
            from eduplay.core.ppt_vsto_addin_service import PptVstoAddinService
            svc = PptVstoAddinService(self.settings_manager)
            try:
                limit = int(self.settings_manager.get("ppt_addin.mru_open_set_limit", 3) or 3)
            except Exception:
                limit = 3
            limit = max(1, min(10, limit))
            mru_paths = svc.get_powerpoint_mru_paths(limit=limit)
            if not mru_paths:
                return
            last = str(self.settings_manager.get("ppt_addin.last_mru_source", "") or "").strip()
            sig = "|".join([str(p).strip().lower() for p in mru_paths if p])
            if last and sig and last.lower() == sig.lower():
                return

            def _do_copy():
                return svc.sync_trusted_slides_from_powerpoint_mru(limit=limit)

            def _on_finished(result, error):
                try:
                    if error is not None:
                        return
                    if not isinstance(result, dict):
                        return
                    if result.get("copied") or result.get("skipped") or result.get("deleted"):
                        try:
                            self.settings_manager.set("ppt_addin.last_mru_source", sig)
                        except Exception:
                            pass
                except Exception:
                    pass

            try:
                self._run_background_task(_do_copy, _on_finished)
            except Exception:
                pass
        except Exception:
            pass
    
    def create_header(self):
        """Create application header"""
        header = QFrame()
        header.setObjectName("header")
        header.setFixedHeight(60)
        
        # store header refs for responsive adjustments
        self.header_frame = header
        
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(20, 0, 20, 0)
        header_layout.setSpacing(10)
        self.header_layout = header_layout
        
        # Logo/Title
        title_label = QLabel("EduPlay Studio")
        self._title_label = title_label
        try:
            icon_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../resources/icons/icon.png"))
            if os.path.exists(icon_path):
                icon_lbl = QLabel()
                icon_lbl.setFixedSize(QSize(32, 32))
                icon_lbl.setStyleSheet("border-radius: 16px;")
                from PySide6.QtGui import QPixmap
                pm = QPixmap(icon_path).scaled(32, 32, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                icon_lbl.setPixmap(pm)
                self._header_icon = icon_lbl
                header_layout.addWidget(icon_lbl)
        except Exception:
            pass
        header_layout.addWidget(title_label)
        try:
            if hasattr(self, "_header_icon") and self._header_icon:
                header_layout.setAlignment(self._header_icon, Qt.AlignmentFlag.AlignVCenter)
            header_layout.setAlignment(title_label, Qt.AlignmentFlag.AlignVCenter)
        except Exception:
            pass
        
        header_layout.addStretch()
        
        try:
            theme = self.settings_manager.get_theme() or "dark"
        except Exception:
            theme = "dark"
        if theme == "dark":
            header_btn_bg = "#2D2F3A"
            header_btn_border = "#4A4E5A"
            header_btn_hover_bg = "#3A3C47"
            header_btn_hover_border = "#5A5E6A"
            header_btn_fg = "#FFFFFF"
        else:
            header_btn_bg = "#FFFFFF"
            header_btn_border = "#CBD5E1"
            header_btn_hover_bg = "#E5E7EB"
            header_btn_hover_border = "#94A3B8"
            header_btn_fg = "#0F1728"
        header_btn_style = f"""
            QPushButton {{
                background-color: {header_btn_bg};
                border: 1px solid {header_btn_border};
                border-radius: 10px;
                color: {header_btn_fg};
                font-size: 22px;
                padding: 0px;
            }}
            QPushButton:hover {{
                background-color: {header_btn_hover_bg};
                border-color: {header_btn_hover_border};
            }}
        """

        # Header buttons
        help_btn = QPushButton("")
        help_btn.setStyleSheet(header_btn_style)
        help_btn.setFixedSize(44, 44)
        try:
            help_btn.setCursor(Qt.PointingHandCursor)
            help_btn.setFocusPolicy(Qt.NoFocus)
            help_btn.setIcon(build_standard_ui_icon("help", self.style(), color_hex=header_btn_fg, size=18))
            help_btn.setIconSize(QSize(18, 18))
            help_btn.setToolTip("Help")
        except Exception:
            pass
        help_btn.clicked.connect(self.show_help)
        
        settings_btn = QPushButton("")
        settings_btn.setStyleSheet(header_btn_style)
        settings_btn.setFixedSize(44, 44)
        try:
            settings_btn.setCursor(Qt.PointingHandCursor)
            settings_btn.setFocusPolicy(Qt.NoFocus)
            settings_btn.setIcon(build_standard_ui_icon("settings", self.style(), color_hex=header_btn_fg, size=18))
            settings_btn.setIconSize(QSize(18, 18))
            settings_btn.setToolTip("Settings")
        except Exception:
            pass
        settings_btn.clicked.connect(self.show_settings)
        self._help_btn = help_btn
        self._settings_btn = settings_btn
        try:
            ai_ready = bool(self.ai_service and getattr(self.ai_service, "is_ready", False))
        except Exception:
            ai_ready = True
        if ai_ready:
            try:
                backend = ""
                if self.ai_service and hasattr(self.ai_service, "backend"):
                    backend = str(self.ai_service.backend() or "").strip().lower()
                if backend == "cli":
                    try:
                        lang = self.settings_manager.get_language() if self.settings_manager else 'en'
                    except Exception:
                        lang = 'en'
                    txt = "AI chậm (CLI)" if lang == 'vi' else "AI slow (CLI)"
                    btn = QPushButton(txt)
                    btn.setStyleSheet(
                        """
                        QPushButton {
                            background-color: #E5E7EB;
                            border: 1px solid #CBD5E1;
                            border-radius: 8px;
                            color: #0F1728;
                            font-weight: 700;
                            padding: 6px 12px;
                        }
                        QPushButton:hover { background-color: #DCE3EA; }
                        """
                    )
                    try:
                        btn.setCursor(Qt.PointingHandCursor)
                        btn.setFocusPolicy(Qt.NoFocus)
                    except Exception:
                        pass
                    btn.clicked.connect(self.show_ai_status)
                    self._ai_backend_badge_btn = btn
                    header_layout.addWidget(btn)
            except Exception:
                pass
        if not ai_ready:
            try:
                lang = self.settings_manager.get_language() if self.settings_manager else 'en'
            except Exception:
                lang = 'en'
            try:
                txt = "AI local chưa sẵn sàng" if lang == 'vi' else "Local AI not ready"
            except Exception:
                txt = "AI local chưa sẵn sàng" if lang == 'vi' else "Local AI not ready"
            ai_badge = QPushButton(txt)
            ai_badge.setStyleSheet(
                """
                QPushButton {
                    background-color: #F79009;
                    border: 1px solid #C46B07;
                    border-radius: 8px;
                    color: #1E1E24;
                    font-weight: 700;
                    padding: 6px 12px;
                }
                QPushButton:hover { background-color: #FF9F1A; }
                """
            )
            try:
                ai_badge.setCursor(Qt.PointingHandCursor)
                ai_badge.setFocusPolicy(Qt.NoFocus)
            except Exception:
                pass
            ai_badge.clicked.connect(self.show_ai_status)
            self._ai_badge_btn = ai_badge
            header_layout.addWidget(ai_badge)
        
        header_layout.addWidget(help_btn)
        header_layout.addWidget(settings_btn)
        
        self.main_layout.addWidget(header)

    def _apply_theme_on_start(self):
        try:
            theme = self.settings_manager.get_theme() if self.settings_manager else 'light'
            qss_name = "dark_theme.qss" if theme == 'dark' else "light_theme.qss"
            qss = load_asset_text(f"eduplay/resources/styles/{qss_name}")
            try:
                extra_qss = self._get_accessibility_qss()
                if extra_qss:
                    qss = f"{qss}\n{extra_qss}"
            except Exception:
                pass
            try:
                brand = self.settings_manager.get("brand_color", "#10B981") if self.settings_manager else "#10B981"
                if brand and brand.startswith('#'):
                    qss = qss.replace('#7F56D9', brand)
                self.setStyleSheet(qss)
            except Exception:
                self.setStyleSheet(qss)
            try:
                if hasattr(self, "central_widget") and self.central_widget:
                    if theme == "dark":
                        self.central_widget.setStyleSheet("#main-central { background-color: #0B0E14; }")
                    else:
                        self.central_widget.setStyleSheet(
                            "#main-central { background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #E0F2FE, stop:0.45 #ECFDF3, stop:1 #FEF3C7); }"
                        )
            except Exception:
                pass
        except Exception:
            pass
        try:
            self._apply_global_accessibility_font()
        except Exception:
            pass
        try:
            if hasattr(self, 'chat_widget') and hasattr(self.chat_widget, 'apply_theme'):
                self.chat_widget.apply_theme()
        except Exception:
            pass
    
    def create_screens(self):
        """Create all application screens"""
        # Home screen
        self.home_screen = HomeScreen()
        self.stacked_widget.addWidget(self.home_screen)
        
        # New project screen
        self.new_project_screen = NewProjectScreen()
        self.stacked_widget.addWidget(self.new_project_screen)
        
        # Browser screen
        self.browser_screen = BrowserScreen()
        self.stacked_widget.addWidget(self.browser_screen)
        
        # Player screen
        self.player_screen = PlayerScreen()
        self.stacked_widget.addWidget(self.player_screen)
        
        # Editor screens
        self.editor_screen = EditorScreen(self.project_manager, self.ai_service, self.import_service)
        self.stacked_widget.addWidget(self.editor_screen)
        self.editor_quiz_screen = EditorQuizClassicScreen(self.project_manager, self.ai_service, self.import_service)
        self.stacked_widget.addWidget(self.editor_quiz_screen)
        self.editor_fishing_screen = EditorFishingScreen(self.project_manager, self.ai_service, self.import_service)
        self.stacked_widget.addWidget(self.editor_fishing_screen)

        try:
            lang = self.settings_manager.get_language()
            if hasattr(self.home_screen, 'set_language'):
                self.home_screen.set_language(lang)
            if hasattr(self.new_project_screen, 'set_language'):
                self.new_project_screen.set_language(lang)
            if hasattr(self.browser_screen, 'set_language'):
                self.browser_screen.set_language(lang)
            if hasattr(self.player_screen, 'set_language'):
                self.player_screen.set_language(lang)
            if hasattr(self.editor_screen, 'set_language'):
                self.editor_screen.set_language(lang)
            if hasattr(self.editor_quiz_screen, 'set_language'):
                self.editor_quiz_screen.set_language(lang)
            if hasattr(self.editor_fishing_screen, 'set_language'):
                self.editor_fishing_screen.set_language(lang)

        except Exception:
            pass
    
    def connect_signals(self):
        """Connect signals between screens"""
        # Home screen connections
        self.home_screen.create_new_clicked.connect(
            lambda: self.show_screen("new_project")
        )
        self.home_screen.edit_project_clicked.connect(
            lambda: self.show_screen("browser")
        )
        self.home_screen.play_game_clicked.connect(
            lambda: self.show_screen("player")
        )
        self.home_screen.open_recent_clicked.connect(self._open_recent_project_from_home)
        self.home_screen.import_questions_clicked.connect(self._import_questions_from_home)
        self.home_screen.quick_publish_clicked.connect(self._quick_publish_from_home)
        self.home_screen.resume_clicked.connect(self._resume_project_from_home)
        self.home_screen.chat_toggled.connect(self.show_chat)
        self.home_screen.help_requested.connect(self.show_help)
        self.home_screen.settings_requested.connect(self.show_settings)
        self.chat_widget.message_sent.connect(self.on_chat_message)
        self.chat_widget.file_uploaded.connect(self.on_chat_file_uploaded)
        
        # New project screen connections
        self.new_project_screen.project_created.connect(self.on_project_created)
        self.new_project_screen.back_clicked.connect(
            lambda: self.show_screen("home")
        )
        
        # Browser screen connections
        self.browser_screen.project_selected.connect(self.on_project_selected)
        self.browser_screen.back_clicked.connect(
            lambda: self.show_screen("home")
        )
        self.browser_screen.create_new_clicked.connect(
            lambda: self.show_screen("new_project")
        )
        
        # Player screen connections
        self.player_screen.back_clicked.connect(
            lambda: self.show_screen("home")
        )
        
        # Editor screen connections
        self.editor_screen.back_to_home.connect(lambda: self.show_screen("home"))
        self.editor_screen.export_requested.connect(self.handle_export)
        self.editor_screen.publish_requested.connect(self.handle_publish_web)
        self.editor_quiz_screen.back_to_home.connect(lambda: self.show_screen("home"))
        self.editor_quiz_screen.export_requested.connect(self.handle_export)
        self.editor_quiz_screen.publish_requested.connect(self.handle_publish_web)
        self.editor_fishing_screen.back_to_home.connect(lambda: self.show_screen("home"))
        self.editor_fishing_screen.export_requested.connect(self.handle_export)
        self.editor_fishing_screen.publish_requested.connect(self.handle_publish_web)

    def _init_command_palette(self):
        try:
            self._command_palette_shortcut = QShortcut(QKeySequence("Ctrl+K"), self)
            self._command_palette_shortcut.setContext(Qt.ApplicationShortcut)
            self._command_palette_shortcut.activated.connect(self.show_command_palette)
        except Exception:
            self._command_palette_shortcut = None

    def _build_command_palette_items(self) -> list[dict]:
        try:
            projects = self.project_manager.get_all_projects() if self.project_manager else []
        except Exception:
            projects = []
        try:
            recent_projects = self.settings_manager.get_recent_projects() if self.settings_manager else []
        except Exception:
            recent_projects = []
        try:
            current_project = self.project_manager.get_current_project() if self.project_manager else {}
        except Exception:
            current_project = {}
        try:
            lang = self.settings_manager.get_language() if self.settings_manager else "en"
        except Exception:
            lang = "en"
        return build_palette_items(projects, recent_projects, (current_project or {}).get("id", ""), lang)

    def show_command_palette(self):
        try:
            lang = self.settings_manager.get_language() if self.settings_manager else "en"
        except Exception:
            lang = "en"
        try:
            items = self._build_command_palette_items()
            dlg = getattr(self, "_command_palette", None)
            if dlg is None or getattr(dlg, "_lang", "en") != lang:
                dlg = CommandPaletteDialog(items, self, lang=lang)
                dlg.item_triggered.connect(self._handle_command_palette_item)
                self._command_palette = dlg
            else:
                dlg.set_items(items)
            try:
                dlg.adjustSize()
            except Exception:
                pass
            x = max(24, int((self.width() - dlg.width()) / 2))
            y = max(24, int(self.height() * 0.12))
            dlg.move(self.mapToGlobal(QPoint(x, y)))
            dlg.open_and_focus()
        except Exception:
            pass

    def _handle_command_palette_item(self, item: dict):
        if not isinstance(item, dict):
            return
        kind = str(item.get("kind") or "")
        if kind == "project":
            project_id = str(item.get("project_id") or "").strip()
            if not project_id:
                return
            try:
                project_data = self.project_manager.load_project(project_id) if self.project_manager else None
            except Exception:
                project_data = None
            if project_data:
                try:
                    if self.settings_manager:
                        self.settings_manager.add_recent_project(project_data.get("id", ""), project_data.get("name", ""))
                except Exception:
                    pass
                self._open_editor_for_project(project_data)
            return

        action = str(item.get("action") or "")
        if action == "new_project":
            self.show_screen("new_project")
        elif action == "open_projects":
            self.show_screen("browser")
        elif action == "quick_preview":
            try:
                self.editor_screen.open_preview_window(preview_mode="quick")
            except Exception:
                pass
        elif action == "full_preview":
            try:
                self.editor_screen.open_preview_window(preview_mode="full")
            except Exception:
                pass
        elif action == "export_html":
            try:
                self.editor_screen.on_export_html()
            except Exception:
                pass

    def _sync_home_quick_actions(self):
        try:
            current_project = self.project_manager.get_current_project() if self.project_manager else None
        except Exception:
            current_project = None
        try:
            self.home_screen.set_quick_context(bool(current_project))
        except Exception:
            pass
        try:
            if getattr(self, "_left_nav_drawer", None):
                self._left_nav_drawer.set_quick_context(bool(current_project))
        except Exception:
            pass

    def _load_project_by_id(self, project_id: str):
        project_id = str(project_id or "").strip()
        if not project_id:
            return None
        try:
            project_data = self.project_manager.load_project(project_id) if self.project_manager else None
        except Exception:
            project_data = None
        if project_data:
            try:
                if self.settings_manager:
                    self.settings_manager.add_recent_project(project_data.get("id", ""), project_data.get("name", ""))
            except Exception:
                pass
        return project_data

    def _open_recent_project_from_home(self):
        try:
            recent = self.settings_manager.get_recent_projects() if self.settings_manager else []
        except Exception:
            recent = []
        project_data = None
        for item in recent or []:
            if not isinstance(item, dict):
                continue
            project_data = self._load_project_by_id(item.get("id", ""))
            if project_data:
                break
        if project_data:
            self._open_editor_for_project(project_data)
            return
        try:
            self.show_screen("browser")
        except Exception:
            pass

    def _resume_project_from_home(self):
        try:
            current_project = self.project_manager.get_current_project() if self.project_manager else None
        except Exception:
            current_project = None
        if current_project:
            self._open_editor_for_project(current_project)
            return
        self._open_recent_project_from_home()

    def _import_questions_from_home(self):
        try:
            current_project = self.project_manager.get_current_project() if self.project_manager else None
        except Exception:
            current_project = None
        if current_project:
            self._open_editor_for_project(current_project)
            try:
                self.editor_screen.on_import_questions()
            except Exception:
                pass
            return
        try:
            self.show_screen("browser")
        except Exception:
            pass

    def _quick_publish_from_home(self):
        try:
            current_project = self.project_manager.get_current_project() if self.project_manager else None
        except Exception:
            current_project = None
        if current_project:
            self._open_editor_for_project(current_project)
            try:
                self.editor_screen.on_export_html()
            except Exception:
                pass
            return
        self._open_recent_project_from_home()

    def _handle_left_nav_quick_action(self, action_key: str):
        key = str(action_key or "").strip().lower()
        if key == "recent":
            self._open_recent_project_from_home()
        elif key == "publish":
            self._quick_publish_from_home()
        elif key == "resume":
            self._resume_project_from_home()

    def _init_left_nav_drawer(self):
        self._left_nav_drawer = LeftNavDrawer(self)
        try:
            self._left_nav_drawer.setFixedWidth(126)
        except Exception:
            pass
        self._left_nav_drawer.hide()
        self._left_nav_hotzone = LeftEdgeHotZone(self)
        self._left_nav_hotzone.setGeometry(0, 0, 6, self.height())
        self._left_nav_hotzone.show()
        self._left_nav_show_timer = QTimer(self)
        self._left_nav_show_timer.setSingleShot(True)
        self._left_nav_show_timer.timeout.connect(self._show_left_nav_drawer)
        self._left_nav_hide_timer = QTimer(self)
        self._left_nav_hide_timer.setSingleShot(True)
        self._left_nav_hide_timer.timeout.connect(self._maybe_hide_left_nav_drawer)
        self._left_nav_anim = None
        self._left_nav_anim_active = False

        self._left_nav_hotzone.hover_entered.connect(self._schedule_show_left_nav_drawer)
        self._left_nav_hotzone.hover_left.connect(self._schedule_hide_left_nav_drawer)
        self._left_nav_drawer.hover_entered.connect(self._cancel_pending_left_nav_timers)
        self._left_nav_drawer.hover_left.connect(self._schedule_hide_left_nav_drawer)
        self._left_nav_drawer.navigate_requested.connect(self._handle_left_nav_navigate)
        self._left_nav_drawer.quick_action_requested.connect(self._handle_left_nav_quick_action)

        try:
            lang = self.settings_manager.get_language() if self.settings_manager else 'en'
            self._left_nav_drawer.set_language(lang)
        except Exception:
            pass
        try:
            self._sync_home_quick_actions()
        except Exception:
            pass
        self._update_left_nav_geometry(force=True)

    def _update_left_nav_geometry(self, force: bool = False):
        if not getattr(self, "_left_nav_drawer", None) or not getattr(self, "_left_nav_hotzone", None):
            return
        try:
            if (not force) and bool(getattr(self, "_left_nav_anim_active", False)):
                self._left_nav_hotzone.setGeometry(0, 0, 6, self.height())
                self._left_nav_hotzone.raise_()
                return
        except Exception:
            pass

        w = int(self._left_nav_drawer.width())
        h = int(self.height())
        self._left_nav_hotzone.setGeometry(0, 0, 6, h)
        self._left_nav_hotzone.raise_()
        if self._left_nav_drawer.isVisible():
            try:
                x = int(self._left_nav_drawer.pos().x())
            except Exception:
                x = 0
            self._left_nav_drawer.setGeometry(x, 0, w, h)
            self._left_nav_drawer.raise_()
        else:
            self._left_nav_drawer.setGeometry(-w, 0, w, h)

    def _cancel_hide_left_nav_drawer(self):
        try:
            if getattr(self, "_left_nav_hide_timer", None):
                self._left_nav_hide_timer.stop()
        except Exception:
            pass

    def _cancel_show_left_nav_drawer(self):
        try:
            if getattr(self, "_left_nav_show_timer", None):
                self._left_nav_show_timer.stop()
        except Exception:
            pass

    def _cancel_pending_left_nav_timers(self):
        self._cancel_show_left_nav_drawer()
        self._cancel_hide_left_nav_drawer()

    def _schedule_show_left_nav_drawer(self):
        self._cancel_hide_left_nav_drawer()
        try:
            if getattr(self, "_left_nav_show_timer", None):
                self._left_nav_show_timer.start(90)
        except Exception:
            pass

    def _schedule_hide_left_nav_drawer(self):
        self._cancel_show_left_nav_drawer()
        try:
            if getattr(self, "_left_nav_hide_timer", None):
                self._left_nav_hide_timer.start(220)
        except Exception:
            pass

    def _maybe_hide_left_nav_drawer(self):
        if not getattr(self, "_left_nav_drawer", None) or not getattr(self, "_left_nav_hotzone", None):
            return
        try:
            pos = self.mapFromGlobal(QCursor.pos())
        except Exception:
            pos = QPoint(-9999, -9999)
        try:
            if self._left_nav_drawer.geometry().contains(pos):
                return
        except Exception:
            pass
        try:
            if self._left_nav_hotzone.geometry().contains(pos):
                return
        except Exception:
            pass
        self._hide_left_nav_drawer()

    def _animate_left_nav(self, target_x: int, hide_after: bool):
        if not getattr(self, "_left_nav_drawer", None):
            return
        try:
            if getattr(self, "_left_nav_anim", None):
                self._left_nav_anim.stop()
        except Exception:
            pass
        try:
            if getattr(self, "_left_nav_anim_group", None):
                self._left_nav_anim_group.stop()
        except Exception:
            pass
        self._left_nav_anim_active = True
        try:
            effect = self._left_nav_drawer.graphicsEffect()
            if not isinstance(effect, QGraphicsOpacityEffect):
                effect = QGraphicsOpacityEffect(self._left_nav_drawer)
                effect.setOpacity(1.0 if not hide_after else 0.92)
                self._left_nav_drawer.setGraphicsEffect(effect)
        except Exception:
            effect = None
        anim = QPropertyAnimation(self._left_nav_drawer, b"pos", self)
        anim.setDuration(220)
        anim.setEasingCurve(QEasingCurve.InOutCubic)
        try:
            anim.setStartValue(self._left_nav_drawer.pos())
        except Exception:
            anim.setStartValue(QPoint(target_x, 0))
        anim.setEndValue(QPoint(int(target_x), 0))
        op_anim = None
        if effect is not None:
            op_anim = QPropertyAnimation(effect, b"opacity", self)
            op_anim.setDuration(180 if hide_after else 200)
            op_anim.setEasingCurve(QEasingCurve.OutCubic if not hide_after else QEasingCurve.InCubic)
            try:
                op_anim.setStartValue(float(effect.opacity()))
            except Exception:
                op_anim.setStartValue(1.0 if not hide_after else 0.92)
            op_anim.setEndValue(0.0 if hide_after else 1.0)
        group = QParallelAnimationGroup(self)
        group.addAnimation(anim)
        if op_anim is not None:
            group.addAnimation(op_anim)

        def _done():
            self._left_nav_anim_active = False
            if hide_after:
                try:
                    self._left_nav_drawer.hide()
                except Exception:
                    pass
                try:
                    self._left_nav_hotzone.setEnabled(True)
                except Exception:
                    pass
                try:
                    if isinstance(effect, QGraphicsOpacityEffect):
                        effect.setOpacity(1.0)
                except Exception:
                    pass
            else:
                try:
                    if isinstance(effect, QGraphicsOpacityEffect):
                        effect.setOpacity(1.0)
                except Exception:
                    pass
            self._update_left_nav_geometry(force=True)

        group.finished.connect(_done)
        self._left_nav_anim = anim
        self._left_nav_anim_group = group
        group.start()

    def _show_left_nav_drawer(self):
        if not getattr(self, "_left_nav_drawer", None) or not getattr(self, "_left_nav_hotzone", None):
            return
        self._cancel_pending_left_nav_timers()
        if not self._left_nav_drawer.isVisible():
            try:
                pos = self.mapFromGlobal(QCursor.pos())
                if not self._left_nav_hotzone.geometry().contains(pos):
                    return
            except Exception:
                pass
        try:
            self._left_nav_hotzone.setEnabled(False)
        except Exception:
            pass
        w = int(self._left_nav_drawer.width())
        h = int(self.height())
        if not self._left_nav_drawer.isVisible():
            self._left_nav_drawer.setGeometry(-w, 0, w, h)
            self._left_nav_drawer.show()
        try:
            self._left_nav_drawer.raise_()
        except Exception:
            pass
        self._animate_left_nav(0, hide_after=False)

    def _hide_left_nav_drawer(self):
        if not getattr(self, "_left_nav_drawer", None) or not getattr(self, "_left_nav_hotzone", None):
            return
        self._cancel_show_left_nav_drawer()
        w = int(self._left_nav_drawer.width())
        self._animate_left_nav(-w, hide_after=True)

    def _handle_left_nav_navigate(self, nav_key: str):
        key = str(nav_key or "")
        if key == "home":
            self.show_screen("home")
        elif key == "browser":
            self.show_screen("browser")
        elif key == "player":
            self.show_screen("player")
        self._schedule_hide_left_nav_drawer()

    def _current_window_scale(self) -> float:
        try:
            w = max(1, self.width())
            h = max(1, self.height())
            return max(0.55, min(1.75, min(w / 1200.0, h / 800.0)))
        except Exception:
            return 1.0

    def _update_window_title_for_screen(self, screen_name: str):
        try:
            from eduplay.core.i18n import I18n
            lang = self.settings_manager.get_language() if self.settings_manager else 'en'
            key_map = {"home": "title.home", "new_project": "title.new_project", "browser": "title.browser", "player": "title.player", "editor": "title.editor"}
            self.setWindowTitle(I18n.t(key_map.get(screen_name, "title.home"), lang))
        except Exception:
            self.setWindowTitle("EduPlay Studio")

    def _refresh_active_screen_polish(self):
        try:
            current = self.stacked_widget.currentWidget()
            if current is not None:
                current.updateGeometry()
                current.update()
        except Exception:
            pass
        try:
            self.stacked_widget.updateGeometry()
            self.stacked_widget.update()
        except Exception:
            pass
        try:
            scale = float(getattr(self, "_last_global_scale", self._current_window_scale()))
        except Exception:
            scale = self._current_window_scale()
        try:
            self._apply_global_scale(scale)
        except Exception:
            pass

    def _finalize_screen_switch(self, screen_name: str, defer_polish: bool = True):
        try:
            if screen_name in ("editor", "player") and self.chat_widget.isVisible() and not getattr(self.chat_widget, 'is_docked', False):
                self.chat_widget.hide()
        except Exception:
            pass
        self._update_window_title_for_screen(screen_name)
        try:
            if defer_polish:
                QTimer.singleShot(0, self._refresh_active_screen_polish)
            else:
                self._refresh_active_screen_polish()
        except Exception:
            self._refresh_active_screen_polish()

    def showEvent(self, event):
        try:
            # Ensure interactive state in case an overlay persisted
            self._hide_loading()
            self.setEnabled(True)
        except Exception:
            pass
        try:
            if not bool(getattr(self, "_startup_animation_requested", False)):
                self._startup_animation_requested = True
                self.start_startup_animation()
        except Exception:
            pass
        try:
            super().showEvent(event)
        except Exception:
            pass

    
    def show_screen(self, screen_name: str):
        """Show a specific screen"""
        try:
            self._current_screen_name = screen_name
        except Exception:
            pass
        try:
            if getattr(self, "_left_nav_drawer", None):
                nav_key = str(screen_name or "home")
                if nav_key.startswith("editor"):
                    nav_key = "browser"
                elif nav_key == "new_project":
                    nav_key = "browser"
                self._left_nav_drawer.set_active(nav_key)
                try:
                    hz = getattr(self, "_left_nav_hotzone", None)
                    if hz is not None and hasattr(hz, "set_active"):
                        marker = {}
                        try:
                            btn_home = getattr(self._left_nav_drawer, "btn_home", None)
                            btn_browser = getattr(self._left_nav_drawer, "btn_projects", None)
                            btn_player = getattr(self._left_nav_drawer, "btn_preview", None)
                            if btn_home is not None:
                                cy = btn_home.mapTo(self, QPoint(int(btn_home.width() / 2), int(btn_home.height() / 2))).y()
                                marker["home"] = hz.mapFrom(self, QPoint(0, cy)).y()
                            if btn_browser is not None:
                                cy = btn_browser.mapTo(self, QPoint(int(btn_browser.width() / 2), int(btn_browser.height() / 2))).y()
                                marker["browser"] = hz.mapFrom(self, QPoint(0, cy)).y()
                            if btn_player is not None:
                                cy = btn_player.mapTo(self, QPoint(int(btn_player.width() / 2), int(btn_player.height() / 2))).y()
                                marker["player"] = hz.mapFrom(self, QPoint(0, cy)).y()
                        except Exception:
                            marker = {}
                        try:
                            from eduplay.core.i18n import I18n

                            lang0 = self.settings_manager.get_language() if self.settings_manager else "en"
                            tmap = {"home": "nav.home", "browser": "nav.projects", "player": "nav.preview"}
                            tip = I18n.t(tmap.get(nav_key, "nav.home"), lang0)
                        except Exception:
                            tip = "Home" if nav_key == "home" else ("Projects" if nav_key == "browser" else "Preview")
                        hz.set_active(nav_key, marker_y=marker, tooltip_text=tip)
                except Exception:
                    pass
        except Exception:
            pass
        screen_map = {
            "home": 0,
            "new_project": 1,
            "browser": 2,
            "player": 3,
            "editor": 4,
            "editor_quiz": 5,
            "editor_fishing": 6,

        }
        
        index = screen_map.get(screen_name, 0)
        try:
            current_index = self.stacked_widget.currentIndex()
        except Exception:
            current_index = -1
        try:
            in_startup = not bool(getattr(self, "_startup_home_fade_done", False))
        except Exception:
            in_startup = False
        if index == current_index:
            self._finalize_screen_switch(screen_name)
            return
        if in_startup:
            try:
                self.stacked_widget.setGraphicsEffect(None)
            except Exception:
                pass
            self.stacked_widget.setCurrentIndex(index)
            self._finalize_screen_switch(screen_name)
            return
        # Avoid transition lag/jump on any transition involving Home.
        if current_index == 0 or index == 0:
            self.stacked_widget.setCurrentIndex(index)
            self._finalize_screen_switch(screen_name)
            return
        try:
            effect = QGraphicsOpacityEffect(self.stacked_widget)
            self.stacked_widget.setGraphicsEffect(effect)
            anim = QPropertyAnimation(effect, b"opacity", self)
            anim.setDuration(140)
            anim.setStartValue(1.0)
            anim.setEndValue(0.18)
            anim.setEasingCurve(QEasingCurve.InOutQuad)
            self._switch_effect = effect
            self._switch_anim = anim
            def _switch():
                if self.stacked_widget.graphicsEffect() is not effect:
                    return
                self.stacked_widget.setCurrentIndex(index)
                self._finalize_screen_switch(screen_name)
                anim2 = QPropertyAnimation(effect, b"opacity", self)
                anim2.setDuration(180)
                anim2.setStartValue(0.18)
                anim2.setEndValue(1.0)
                anim2.setEasingCurve(QEasingCurve.OutCubic)
                anim2.finished.connect(
                    lambda: self.stacked_widget.setGraphicsEffect(None)
                )
                anim2.start()
                self._switch_anim2 = anim2
            anim.finished.connect(_switch)
            anim.start()
        except Exception:
            self.stacked_widget.setCurrentIndex(index)
            self._finalize_screen_switch(screen_name)
    
    def on_project_created(self, project_data: dict):
        """Handle project creation"""
        # Persist project using ProjectManager and then load into appropriate editor
        try:
            from eduplay.core.project_manager import ProjectManager
            pm = self.project_manager or ProjectManager()
            name = project_data.get("name", "Untitled Project")
            description = project_data.get("description", "")
            game_type = project_data.get("game_type", "quiz_classic")
            created = pm.create_project(name, description, game_type)
            cfg = project_data.get("game_config")
            if cfg:
                # Normalize textual game type to internal value
                gt = (cfg.get("game_type") or "").lower()
                if "fish" in gt:
                    created["game_type"] = "fishing"
                    cfg["game_type"] = "Fishing Game"
                pm.update_game_config(cfg)
                created = pm.get_current_project() or created
            # Track recent
            try:
                if self.settings_manager:
                    self.settings_manager.add_recent_project(created.get("id",""), created.get("name",""))
            except Exception:
                pass
            try:
                self._sync_home_quick_actions()
            except Exception:
                pass
            self._open_editor_for_project(created)
        except Exception:
            # Fallback to previous behavior
            self.editor_screen.load_project(project_data)
            self.show_screen("editor")
    
    def on_project_selected(self, project_data: dict):
        """Handle project selection"""
        # Switch to appropriate editor screen with the selected project
        self._open_editor_for_project(project_data)

    def _open_editor_for_project(self, project_data: dict):
        try:
            from eduplay.core.i18n import I18n
            lang = self.settings_manager.get_language() if self.settings_manager else 'en'
        except Exception:
            I18n = None  # type: ignore
            lang = 'en'
        try:
            try:
                if I18n:
                    loading_msg = I18n.t('browser.loading_project', lang)
                else:
                    loading_msg = "Loading project..."
            except Exception:
                loading_msg = "Loading project..."
            self._show_loading(loading_msg, "default")
        except Exception:
            pass
        try:
            gt = (project_data or {}).get('game_type', '')
            gl = str(gt).lower()
            try:
                cfg = (project_data or {}).get('game_config', {}) or {}
            except Exception:
                cfg = {}
            try:
                marker0 = str((project_data or {}).get('variant_marker') or '').lower()
                cfg_marker0 = str((cfg.get('variant_marker') or '')).lower()
                fv0 = str((project_data or {}).get('force_variant') or '').lower()
                if (fv0 == 'fishing') or (marker0 == 'fishing') or (cfg_marker0 == 'fishing'):
                    gl = 'fishing'
                if ('bắt cá' in gl) or ('bat ca' in gl) or ('câu cá' in gl) or ('cau ca' in gl):
                    gl = 'fishing'
                cfg_gt0 = str((cfg.get('game_type') or '')).lower()
                if ('fishing' in cfg_gt0) or ('fish' in cfg_gt0) or ('bắt cá' in cfg_gt0) or ('bat ca' in cfg_gt0) or ('câu cá' in cfg_gt0) or ('cau ca' in cfg_gt0):
                    gl = 'fishing'
            except Exception:
                pass
            try:
                marker = str(project_data.get('variant_marker') or '').lower()
                cfg_marker = str((cfg.get('variant_marker') or '')).lower()
                if gl == 'fishing':
                    has_fish = bool(cfg.get('fish_objects')) or ('fishing_settings' in cfg)
                    cfg_gt = str((cfg.get('game_type') or '')).lower()
                    is_cfg_fishing = any(x in cfg_gt for x in ('fishing', 'fish', 'câu cá', 'cau ca', 'bắt cá', 'bat ca'))
                    if (marker == 'fishing') or (cfg_marker == 'fishing'):
                        has_fish = True
                        is_cfg_fishing = True
                    if not has_fish:
                        gl = 'quiz_classic'
                        project_data['game_type'] = 'quiz_classic'
                    elif not is_cfg_fishing:
                        for k in ['fish_objects', 'fish_speed', 'fish_count', 'base_speed', 'score_per_fish', 'fishing_settings']:
                            try:
                                cfg.pop(k, None)
                            except Exception:
                                pass
                        project_data['game_config'] = cfg
                        gl = 'quiz_classic'
                        project_data['game_type'] = 'quiz_classic'
            except Exception:
                pass
            if gl == 'fishing':
                self.editor_fishing_screen.load_project(project_data)
                self.show_screen('editor_fishing')
            else:
                self.editor_quiz_screen.load_project(project_data)
                self.show_screen('editor_quiz')
            try:
                self._sync_home_quick_actions()
            except Exception:
                pass
        finally:
            try:
                self._hide_loading()
            except Exception:
                pass
    
    def show_help(self):
        """Show help dialog — menu-driven: 3 section buttons first, then detail view with back button"""
        from PySide6.QtWidgets import QScrollArea, QStackedWidget
        from PySide6.QtCore import QTimer
        try:
            lang = self.settings_manager.get_language()
        except Exception:
            lang = 'vi'
        try:
            from eduplay.core.settings_manager import SettingsManager
            theme = SettingsManager().get_theme() or "dark"
        except Exception:
            theme = "dark"
        from eduplay.core.i18n import I18n

        help_dialog = QDialog(self)
        help_dialog.setWindowTitle(I18n.t('help.title', lang))
        try:
            help_dialog.setWindowFlags(
                Qt.Window | Qt.WindowTitleHint | Qt.WindowSystemMenuHint |
                Qt.WindowCloseButtonHint
            )
        except Exception:
            pass

        # Center on screen
        try:
            from PySide6.QtWidgets import QApplication
            screen = QApplication.primaryScreen()
            if screen:
                sg = screen.availableGeometry()
                dlg_w = min(820, max(640, int(sg.width() * 0.62)))
                dlg_h = min(680, max(520, int(sg.height() * 0.75)))
                help_dialog.resize(dlg_w, dlg_h)
                help_dialog.move(
                    sg.x() + (sg.width() - dlg_w) // 2,
                    sg.y() + (sg.height() - dlg_h) // 2,
                )
        except Exception:
            help_dialog.resize(800, 640)

        # Theme colors
        if theme == "dark":
            dlg_bg = "#0F172A"
            dlg_fg = "#E2E8F0"
            card_bg = "#1E293B"
            card_border = "#334155"
            muted = "#94A3B8"
            accent = "#7F56D9"
            btn_bg = "#1E293B"
            btn_border = "#334155"
            btn_hover = "#334155"
            content_bg = "#0F172A"
        else:
            dlg_bg = "#F8FAFC"
            dlg_fg = "#0F172A"
            card_bg = "#FFFFFF"
            card_border = "#E2E8F0"
            muted = "#64748B"
            accent = "#7F56D9"
            btn_bg = "#FFFFFF"
            btn_border = "#E2E8F0"
            btn_hover = "#EEF2FF"
            content_bg = "#FFFFFF"

        help_dialog.setStyleSheet(f"""
            QDialog {{ background-color: {dlg_bg}; color: {dlg_fg}; }}
            QLabel {{ background-color: transparent; color: {dlg_fg}; }}
            QLabel#muted {{ color: {muted}; }}
            QLabel#section-title {{ font-size: 18px; font-weight: 700; color: {dlg_fg}; }}
            QFrame#content-card {{
                background-color: {card_bg};
                border: 1px solid {card_border};
                border-radius: 16px;
            }}
            QPushButton {{
                background-color: {btn_bg};
                color: {dlg_fg};
                border: 1px solid {btn_border};
                border-radius: 12px;
                padding: 14px 20px;
                font-weight: 600;
                font-size: 14px;
                text-align: left;
            }}
            QPushButton:hover {{
                background-color: {btn_hover};
                border-color: {accent};
                color: {accent};
            }}
            QPushButton#back-btn {{
                background-color: transparent;
                border: none;
                color: {muted};
                font-size: 13px;
                padding: 4px 8px;
                text-align: left;
                font-weight: 500;
            }}
            QPushButton#back-btn:hover {{
                color: {accent};
                border: none;
            }}
            QPushButton#close-btn {{
                background-color: {accent};
                color: #FFFFFF;
                border: none;
                border-radius: 10px;
                padding: 10px 28px;
                font-weight: 600;
                font-size: 14px;
                text-align: center;
            }}
            QPushButton#close-btn:hover {{
                background-color: #8B66E9;
            }}
            QScrollArea {{ border: none; background: transparent; }}
            QScrollBar:vertical {{ background: transparent; width: 6px; border-radius: 3px; }}
            QScrollBar::handle:vertical {{ background: {card_border}; border-radius: 3px; min-height: 20px; }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0px; }}
        """)

        # Detailed content for each section
        def _t(key: str, fallback: str = "") -> str:
            try:
                return I18n.t(key, lang)
            except Exception:
                return fallback

        sections = [
            {
                "icon": "🧭",
                "title": _t("help.sections.overview.title", "Overview"),
                "subtitle": _t("help.sections.overview.subtitle", ""),
                "content": _t("help.sections.overview.content", ""),
            },
            {
                "icon": "🖥️",
                "title": _t("help.sections.interface.title", "Interface"),
                "subtitle": _t("help.sections.interface.subtitle", "Home, toolbar and navigation"),
                "content": _t("help.sections.interface.content", ""),
            },
            {
                "icon": "📁",
                "title": _t("help.sections.create_project.title", "Create Project"),
                "subtitle": _t("help.sections.create_project.subtitle", "Create and manage game projects"),
                "content": _t("help.sections.create_project.content", ""),
            },
            {
                "icon": "🤖",
                "title": _t("help.sections.edubot.title", "Edubot (AI)"),
                "subtitle": _t("help.sections.edubot.subtitle", "Auto-generate questions with AI"),
                "content": _t("help.sections.edubot.content", ""),
            },
            {
                "icon": "🚀",
                "title": _t("help.sections.preview_export.title", "Preview & Export"),
                "subtitle": _t("help.sections.preview_export.subtitle", "Test and share your game"),
                "content": _t("help.sections.preview_export.content", ""),
            },
            {
                "icon": "📊",
                "title": _t("help.sections.powerpoint_addin.title", "PowerPoint Add-in"),
                "subtitle": _t("help.sections.powerpoint_addin.subtitle", "Embed games into PowerPoint slides"),
                "content": _t("help.sections.powerpoint_addin.content", ""),
            },
        ]

        main_layout = QVBoxLayout(help_dialog)
        main_layout.setContentsMargins(24, 20, 24, 20)
        main_layout.setSpacing(0)

        # Header
        header_widget = QWidget()
        header_layout = QVBoxLayout(header_widget)
        header_layout.setContentsMargins(0, 0, 0, 16)
        header_layout.setSpacing(4)
        title_lbl = QLabel(_t("help.heading", "EduPlay Studio Help"))
        title_lbl.setObjectName("section-title")
        title_lbl.setStyleSheet(f"font-size: 22px; font-weight: 700; color: {dlg_fg};")
        sub_lbl = QLabel(_t("help.subheading", "Select a section to view detailed instructions."))
        sub_lbl.setObjectName("muted")
        sub_lbl.setStyleSheet(f"font-size: 13px; color: {muted};")
        header_layout.addWidget(title_lbl)
        header_layout.addWidget(sub_lbl)
        main_layout.addWidget(header_widget)

        # Stacked: page 0 = menu (3 buttons), page 1 = detail view
        stack = QStackedWidget()
        main_layout.addWidget(stack, 1)

        # --- Page 0: Menu ---
        menu_page = QWidget()
        menu_layout = QVBoxLayout(menu_page)
        menu_layout.setContentsMargins(0, 0, 0, 0)
        menu_layout.setSpacing(10)
        menu_layout.addStretch(1)
        for i, sec in enumerate(sections):
            btn = QPushButton(f"  {sec['icon']}  {sec['title']}")
            btn.setMinimumHeight(64)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {card_bg};
                    color: {dlg_fg};
                    border: 1px solid {card_border};
                    border-radius: 14px;
                    padding: 16px 20px;
                    font-size: 15px;
                    font-weight: 700;
                    text-align: left;
                }}
                QPushButton:hover {{
                    background-color: {btn_hover};
                    border-color: {accent};
                    color: {accent};
                }}
            """)
            sub = QLabel(f"   {sec['subtitle']}")
            sub.setStyleSheet(f"font-size: 12px; color: {muted}; margin-left: 4px;")
            container = QWidget()
            cl = QVBoxLayout(container)
            cl.setContentsMargins(0, 0, 0, 0)
            cl.setSpacing(0)
            cl.addWidget(btn)
            menu_layout.addWidget(container)

            def _mk_nav(idx):
                def _go():
                    stack.setCurrentIndex(1)
                    _load_detail(idx)
                return _go
            btn.clicked.connect(_mk_nav(i))

        menu_layout.addStretch(2)
        stack.addWidget(menu_page)  # index 0

        # --- Page 1: Detail view ---
        detail_page = QWidget()
        detail_layout = QVBoxLayout(detail_page)
        detail_layout.setContentsMargins(0, 0, 0, 0)
        detail_layout.setSpacing(10)

        # Back button row
        back_row = QWidget()
        back_row_layout = QHBoxLayout(back_row)
        back_row_layout.setContentsMargins(0, 0, 0, 0)
        back_btn = QPushButton("← Quay lại" if lang == "vi" else "← Back")
        back_btn.setObjectName("back-btn")
        back_btn.setMaximumWidth(120)
        back_btn.clicked.connect(lambda: stack.setCurrentIndex(0))
        back_row_layout.addWidget(back_btn)
        back_row_layout.addStretch()
        detail_layout.addWidget(back_row)

        # Section title in detail view
        detail_title = QLabel("")
        detail_title.setStyleSheet(f"font-size: 17px; font-weight: 700; color: {dlg_fg}; margin-bottom: 4px;")
        detail_layout.addWidget(detail_title)

        # Content area (scrollable)
        detail_scroll = QScrollArea()
        detail_scroll.setWidgetResizable(True)
        from PySide6.QtWidgets import QFrame
        try:
            detail_scroll.setFrameShape(QFrame.Shape.NoFrame)
        except Exception:
            try:
                detail_scroll.setFrameShape(QFrame.NoFrame)
            except Exception:
                pass
        detail_inner = QWidget()
        detail_inner.setStyleSheet(f"background: transparent;")
        detail_inner_layout = QVBoxLayout(detail_inner)
        detail_inner_layout.setContentsMargins(0, 0, 0, 0)
        from PySide6.QtWidgets import QTextBrowser
        detail_text = QTextBrowser()
        detail_text.setOpenExternalLinks(True)
        detail_text.setStyleSheet(f"""
            QTextBrowser {{
                background-color: {card_bg};
                color: {dlg_fg};
                border: 1px solid {card_border};
                border-radius: 14px;
                padding: 20px;
                font-size: 14px;
                line-height: 1.6;
            }}
            QScrollBar:vertical {{ background: transparent; width: 6px; }}
            QScrollBar::handle:vertical {{ background: {card_border}; border-radius: 3px; }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{ height: 0px; }}
        """)
        detail_inner_layout.addWidget(detail_text)
        detail_scroll.setWidget(detail_inner)
        detail_layout.addWidget(detail_scroll, 1)
        stack.addWidget(detail_page)  # index 1

        def _load_detail(idx: int):
            sec = sections[idx]
            detail_title.setText(f"{sec['icon']}  {sec['title']}")
            detail_text.setHtml(sec["content"])
            try:
                detail_scroll.verticalScrollBar().setValue(0)
            except Exception:
                pass

        # Start on menu page
        stack.setCurrentIndex(0)

        # Bottom action row
        action_row = QWidget()
        action_layout = QHBoxLayout(action_row)
        action_layout.setContentsMargins(0, 14, 0, 0)
        action_layout.setSpacing(10)
        # DOCX template download
        try:
            download_text = I18n.t('help.download_docx_btn', lang)
        except Exception:
            download_text = "⬇️ Mẫu DOCX" if lang == "vi" else "⬇️ DOCX Template"
        download_btn = QPushButton(download_text)
        download_btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {btn_bg};
                color: {dlg_fg};
                border: 1px solid {btn_border};
                border-radius: 10px;
                padding: 9px 18px;
                font-weight: 600;
                font-size: 13px;
            }}
            QPushButton:hover {{ background-color: {btn_hover}; border-color: {accent}; }}
        """)
        def _download_docx():
            try:
                from PySide6.QtWidgets import QFileDialog
                from eduplay.core.import_service import ImportService
                try:
                    from eduplay.ui.widgets.template_selection_dialog import TemplateSelectionDialog
                    dialog = TemplateSelectionDialog(self, lang)
                    if dialog.exec():
                        template_type = dialog.get_selected_type()
                    else:
                        return
                except Exception:
                    template_type = "general"
                dlg2 = QFileDialog(self)
                dlg2.setAcceptMode(QFileDialog.AcceptSave)
                dlg2.setFileMode(QFileDialog.AnyFile)
                try:
                    dlg2.setNameFilter("Word Document (*.docx);;Excel (*.xlsx);;Text (*.txt)")
                except Exception:
                    pass
                if dlg2.exec():
                    sp = dlg2.selectedFiles()[0] if dlg2.selectedFiles() else ""
                    if sp:
                        import os as _os
                        if not _os.path.splitext(sp)[1]:
                            sp += ".docx"
                        fmt = "xlsx" if sp.lower().endswith(".xlsx") else ("txt" if sp.lower().endswith(".txt") else "docx")
                        ImportService().create_sample_template(sp, fmt, lang, template_type)
                        try:
                            self._show_message(I18n.t('help.title', lang), I18n.t('help.download_done', lang), "success")
                        except Exception:
                            pass
            except Exception as e:
                try:
                    self._show_message("Error", str(e), "error")
                except Exception:
                    pass
        download_btn.clicked.connect(_download_docx)
        action_layout.addWidget(download_btn)
        action_layout.addStretch()
        close_btn = QPushButton(_t("help.close", "Close"))
        close_btn.setObjectName("close-btn")
        close_btn.clicked.connect(help_dialog.accept)
        action_layout.addWidget(close_btn)
        main_layout.addWidget(action_row)

        help_dialog.exec()

    def show_settings(self):
        """Show settings dialog"""
        try:
            from eduplay.core.i18n import I18n
            lang = self.settings_manager.get_language() if self.settings_manager else 'en'
        except Exception:
            I18n = None  # type: ignore
            lang = 'en'
        settings_dialog = SettingsDialog(self.settings_manager, self)
        restart_prompt_state = {"show": False}
        def _apply_settings():
            try:
                if I18n:
                    loading_msg = I18n.t('settings.loading', lang)
                else:
                    loading_msg = "Applying settings..."
            except Exception:
                loading_msg = "Applying settings..."
            try:
                self._show_loading(loading_msg, "settings")
            except Exception:
                pass
            try:
                self.on_settings_changed()
                restart_prompt_state["show"] = True
            finally:
                try:
                    self._hide_loading()
                except Exception:
                    pass
        try:
            settings_dialog.settings_changed.connect(_apply_settings)
        except Exception:
            try:
                settings_dialog.settings_changed.connect(self.on_settings_changed)
            except Exception:
                pass
        settings_dialog.exec()
        if restart_prompt_state["show"]:
            try:
                self._show_settings_restart_prompt()
            except Exception:
                pass

    def _settings_restart_prompt_strings(self):
        try:
            from eduplay.core.i18n import I18n
            lang = self.settings_manager.get_language() if self.settings_manager else "en"
            title = I18n.t("settings.restart_required_title", lang)
            message = I18n.t("settings.restart_required_message", lang)
            restart_now = I18n.t("settings.restart_now", lang)
            later = I18n.t("settings.restart_later", lang)
        except Exception:
            title = "Cần khởi động lại"
            message = "Một số thay đổi chỉ được áp dụng đầy đủ sau khi khởi động lại EduPlay Studio."
            restart_now = "Khởi động lại ngay"
            later = "Để sau"
        return title, message, restart_now, later

    def _show_settings_restart_prompt(self) -> bool:
        title, message, restart_now, later = MainWindow._settings_restart_prompt_strings(self)
        box = QMessageBox(self)
        try:
            theme = (self.settings_manager.get_theme() if self.settings_manager else "dark") or "dark"
        except Exception:
            theme = "dark"
        if str(theme).lower() == "dark":
            box.setStyleSheet(
                """
                QMessageBox { background-color: #0F172A; color: #E5E7EB; }
                QMessageBox QLabel { color: #E5E7EB; background: transparent; }
                QMessageBox QPushButton {
                    background-color: #111827;
                    color: #E5E7EB;
                    border: 1px solid #374151;
                    border-radius: 10px;
                    padding: 8px 14px;
                    min-width: 96px;
                }
                QMessageBox QPushButton:hover { background-color: #1F2937; color: #FFFFFF; border-color: #4B5563; }
                QMessageBox QPushButton:default { background-color: #7F56D9; border-color: #7F56D9; color: #FFFFFF; }
                """
            )
        else:
            box.setStyleSheet(
                """
                QMessageBox { background-color: #FFFFFF; color: #111827; }
                QMessageBox QLabel { color: #1F2937; background: transparent; }
                QMessageBox QPushButton {
                    background-color: #F3F4F6;
                    color: #111827;
                    border: 1px solid #D0D5DD;
                    border-radius: 10px;
                    padding: 8px 14px;
                    min-width: 96px;
                }
                QMessageBox QPushButton:hover { background-color: #E5E7EB; color: #111827; border-color: #9CA3AF; }
                QMessageBox QPushButton:default { background-color: #7F56D9; border-color: #7F56D9; color: #FFFFFF; }
                """
            )
        box.setIcon(QMessageBox.Warning)
        box.setWindowTitle(title)
        box.setText(message)
        restart_btn = box.addButton(restart_now, QMessageBox.AcceptRole)
        later_btn = box.addButton(later, QMessageBox.RejectRole)
        try:
            box.setDefaultButton(restart_btn)
        except Exception:
            pass
        box.exec()
        if box.clickedButton() is restart_btn:
            try:
                self._restart_application()
            except Exception:
                return False
            return True
        return False

    def _restart_application(self):
        try:
            if getattr(sys, "frozen", False):
                program = sys.executable
                args = list(sys.argv[1:])
            else:
                program = sys.executable
                args = list(sys.argv)
            started = QProcess.startDetached(program, args)
            if started:
                self.close()
        except Exception:
            pass

    def show_intro_credits(self):
        try:
            if getattr(self, "_startup_overlay", None):
                return
            self.start_startup_animation(force_credits=True)
        except Exception:
            pass
    
    def on_settings_changed(self):
        """Handle settings changes"""
        # Apply theme changes
        theme = self.settings_manager.get_theme() or "light"
        try:
            qss_name = "dark_theme.qss" if theme == 'dark' else "light_theme.qss"
            qss = load_asset_text(f"eduplay/resources/styles/{qss_name}")
            try:
                extra_qss = self._get_accessibility_qss()
                if extra_qss:
                    qss = f"{qss}\n{extra_qss}"
            except Exception:
                pass
            try:
                brand = self.settings_manager.get("brand_color", "#10B981")
                if brand and brand.startswith('#'):
                    qss = qss.replace('#7F56D9', brand)
                self.setStyleSheet(qss)
            except Exception:
                self.setStyleSheet(qss)
            try:
                if hasattr(self, "central_widget") and self.central_widget:
                    if theme == "dark":
                        self.central_widget.setStyleSheet("#main-central { background-color: #0B0E14; }")
                    else:
                        self.central_widget.setStyleSheet(
                            "#main-central { background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #E0F2FE, stop:0.45 #ECFDF3, stop:1 #FEF3C7); }"
                        )
            except Exception:
                pass
        except Exception:
            pass

        try:
            if hasattr(self.chat_widget, 'apply_theme'):
                self.chat_widget.apply_theme()
        except Exception:
            pass

        try:
            if hasattr(self, 'browser_screen') and hasattr(self.browser_screen, 'apply_theme'):
                self.browser_screen.apply_theme(theme)
        except Exception:
            pass

        try:
            if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'apply_theme'):
                self.editor_screen.apply_theme(theme)
            if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'refresh_autosave_settings'):
                self.editor_screen.refresh_autosave_settings()
        except Exception:
            pass

        try:
            if hasattr(self, 'editor_quiz_screen') and hasattr(self.editor_quiz_screen, 'apply_theme'):
                self.editor_quiz_screen.apply_theme(theme)
            if hasattr(self, 'editor_quiz_screen') and hasattr(self.editor_quiz_screen, 'refresh_autosave_settings'):
                self.editor_quiz_screen.refresh_autosave_settings()
        except Exception:
            pass

        try:
            if hasattr(self, 'editor_fishing_screen') and hasattr(self.editor_fishing_screen, 'apply_theme'):
                self.editor_fishing_screen.apply_theme(theme)
            if hasattr(self, 'editor_fishing_screen') and hasattr(self.editor_fishing_screen, 'refresh_autosave_settings'):
                self.editor_fishing_screen.refresh_autosave_settings()
        except Exception:
            pass

        try:
            if hasattr(self, 'home_screen') and hasattr(self.home_screen, 'refresh_theme'):
                self.home_screen.refresh_theme()
        except Exception:
            pass

        try:
            if hasattr(self, 'new_project_screen') and hasattr(self.new_project_screen, 'apply_theme'):
                self.new_project_screen.apply_theme(theme)
        except Exception:
            pass

        try:
            if hasattr(self, 'player_screen') and hasattr(self.player_screen, 'apply_theme'):
                self.player_screen.apply_theme(theme)
        except Exception:
            pass

        try:
            if hasattr(self, 'browser_screen') and hasattr(self.browser_screen, 'apply_theme'):
                self.browser_screen.apply_theme(theme)
        except Exception:
            pass

        # Apply global editor font settings
        try:
            self._apply_global_accessibility_font()
        except Exception:
            pass
        
        # Update language title
        from eduplay.core.i18n import I18n
        lang = self.settings_manager.get_language()
        try:
            I18n.set_locale(lang)
        except Exception:
            pass
        current_index = self.stacked_widget.currentIndex()
        index_map = {0: "home", 1: "new_project", 2: "browser", 3: "player", 4: "editor"}
        name = index_map.get(current_index, "home")
        key_map = {"home": "title.home", "new_project": "title.new_project", "browser": "title.browser", "player": "title.player", "editor": "title.editor"}

    def _get_accessibility_qss(self) -> str:
        try:
            cfg = self.settings_manager.get_accessibility_settings() if self.settings_manager else {}
        except Exception:
            cfg = {}
        if not isinstance(cfg, dict):
            cfg = {}
        if not bool(cfg.get("high_contrast", False)):
            return ""
        return """
QPushButton, QLineEdit, QTextEdit, QListWidget, QComboBox, QToolButton, QCheckBox, QSpinBox, QTabWidget::pane, QGroupBox {
    border-width: 2px;
}
QPushButton:focus, QLineEdit:focus, QTextEdit:focus, QListWidget:focus, QComboBox:focus, QToolButton:focus, QCheckBox:focus, QSpinBox:focus {
    outline: none;
    border: 2px solid #FFD60A;
}
QLabel {
    color: palette(window-text);
}
"""

    def _apply_global_accessibility_font(self):
        try:
            from PySide6.QtWidgets import QApplication
            editor = self.settings_manager.get_editor_settings() if self.settings_manager else {}
            accessibility = self.settings_manager.get_accessibility_settings() if self.settings_manager else {}
            family = editor.get("font_family") or "Segoe UI"
            base_size = int(editor.get("font_size") or 12)
            scale_pct = int(accessibility.get("ui_scale", 100) or 100)
            size = max(9, int(round(base_size * (scale_pct / 100.0))))
            app_font = QFont(family, size)
            QApplication.setFont(app_font)
        except Exception:
            pass
        self.setWindowTitle(I18n.t(key_map.get(name, "title.home"), lang))
        
        try:
            if hasattr(self.chat_widget, 'set_language'):
                self.chat_widget.set_language(lang)
            if hasattr(self.home_screen, 'set_language'):
                self.home_screen.set_language(lang)
            if hasattr(self.browser_screen, 'set_language'):
                self.browser_screen.set_language(lang)
            if hasattr(self.player_screen, 'set_language'):
                self.player_screen.set_language(lang)
            if hasattr(self.editor_screen, 'set_language'):
                self.editor_screen.set_language(lang)
            if hasattr(self.editor_quiz_screen, 'set_language'):
                self.editor_quiz_screen.set_language(lang)
            if hasattr(self.editor_fishing_screen, 'set_language'):
                self.editor_fishing_screen.set_language(lang)
            if hasattr(self.new_project_screen, 'set_language'):
                self.new_project_screen.set_language(lang)
            if getattr(self, "_left_nav_drawer", None):
                self._left_nav_drawer.set_language(lang)
        except Exception:
            pass
        
        try:
            ai_ready = bool(self.ai_service and getattr(self.ai_service, "is_ready", False))
        except Exception:
            ai_ready = True
        try:
            if hasattr(self, '_ai_badge_btn'):
                self._ai_badge_btn.setVisible(not ai_ready)
            else:
                if not ai_ready:
                    try:
                        lang = self.settings_manager.get_language() if self.settings_manager else 'en'
                    except Exception:
                        lang = 'en'
                    try:
                        txt = "AI local chưa sẵn sàng" if lang == 'vi' else "Local AI not ready"
                    except Exception:
                        txt = "AI local chưa sẵn sàng" if lang == 'vi' else "Local AI not ready"
                    btn = QPushButton(txt)
                    btn.setStyleSheet(
                        """
                        QPushButton {
                            background-color: #F79009;
                            border: 1px solid #C46B07;
                            border-radius: 8px;
                            color: #1E1E24;
                            font-weight: 700;
                            padding: 6px 12px;
                        }
                        QPushButton:hover { background-color: #FF9F1A; }
                        """
                    )
                    btn.clicked.connect(self.show_ai_status)
                    self._ai_badge_btn = btn
                    try:
                        if hasattr(self, 'main_layout') and self.main_layout and self.main_layout.count() > 0:
                            header = self.main_layout.itemAt(0).widget()
                            if header:
                                header.layout().addWidget(btn)
                    except Exception:
                        pass
        except Exception:
            pass
        try:
            self._hide_loading()
        except Exception:
            pass

    def show_ai_status(self):
        try:
            lang = self.settings_manager.get_language() if self.settings_manager else "en"
        except Exception:
            lang = "en"
        err = ""
        try:
            err = str(getattr(self.ai_service, "error_message", "") or "").strip()
        except Exception:
            err = ""
        if lang == "vi":
            title = "AI local"
            msg = "AI local chưa sẵn sàng."
            if err:
                msg = msg + "\n\nLỗi: " + err
            msg = msg + "\n\nCách xử lý:\n- Cài dependency: pip install -r requirements.txt\n- Lần đầu cần internet để tải model GGUF và llama.cpp (CLI)"
        else:
            title = "Local AI"
            msg = "Local AI is not ready."
            if err:
                msg = msg + "\n\nError: " + err
            msg = msg + "\n\nFix:\n- Install dependencies: pip install -r requirements.txt\n- First run needs internet to download the GGUF model and llama.cpp (CLI)"
        try:
            self._show_message(title, msg, "warning")
        except Exception:
            try:
                QMessageBox.information(self, title, msg)
            except Exception:
                pass
    
    def show_chat(self):
        """Show/hide chat widget"""
        if self.chat_widget.isVisible():
            self.chat_widget.hide()
        else:
            self._position_chat_widget()
            self.chat_widget.show()
            self.chat_widget.raise_()
            self.chat_widget.activateWindow()
        try:
            self._update_chat_layout_for_dock()
        except Exception:
            pass

    def _position_chat_widget(self):
        try:
            screen_rect = self.geometry()
            if getattr(self.chat_widget, 'is_docked', False):
                frac = 0.32 if getattr(self.chat_widget, 'width_mode', 'narrow') == 'narrow' else 0.42
                chat_width = max(320, int(screen_rect.width() * frac))
                header_h = 60
                chat_height = max(360, int(screen_rect.height() - header_h))
                margin = 0
                x = screen_rect.x() + screen_rect.width() - chat_width - margin
                y = screen_rect.y() + header_h
            else:
                if getattr(self.chat_widget, 'width_mode', 'narrow') == 'narrow':
                    frac = 0.32
                else:
                    frac = 0.50
                chat_width = max(320, int(screen_rect.width() * frac))
                chat_height = max(360, int(screen_rect.height() * 0.6))
                margin = 20
                x = screen_rect.x() + screen_rect.width() - chat_width - margin
                y = screen_rect.y() + screen_rect.height() - chat_height - margin - 60
            try:
                setattr(self.chat_widget, "_manual_positioned", True)
            except Exception:
                pass
            self.chat_widget.setGeometry(x, y, chat_width, chat_height)
        except Exception:
            pass

    def _update_chat_layout_for_dock(self):
        try:
            left, top, _, bottom = self.main_layout.getContentsMargins()
            self.main_layout.setContentsMargins(left, top, 0, bottom)
        except Exception:
            pass

    def resizeEvent(self, event):
        super().resizeEvent(event)
        try:
            # header responsive adjustments based on window width
            w = self.width()
            if hasattr(self, 'header_frame') and self.header_frame:
                self.header_frame.setFixedHeight(52 if w < 900 else 60)
            if hasattr(self, 'header_layout') and self.header_layout:
                m = 12 if w < 900 else 20
                try:
                    self.header_layout.setContentsMargins(m, 0, m, 0)
                except Exception:
                    pass
            if hasattr(self, '_ai_badge_btn'):
                try:
                    self._ai_badge_btn.setVisible(w >= 900)
                except Exception:
                    pass
            if hasattr(self, '_help_btn'):
                try:
                    self._help_btn.setVisible(w >= 800)
                except Exception:
                    pass
            if hasattr(self, '_title_label') and self._title_label:
                try:
                    f = self._title_label.font()
                    f.setPointSize(18 if w < 900 else (22 if w < 1200 else 24))
                    self._title_label.setFont(f)
                    self._title_label.setText("EduPlay" if w < 800 else "EduPlay Studio")
                except Exception:
                    pass
            if hasattr(self, '_header_icon'):
                try:
                    self._header_icon.setVisible(w >= 900)
                except Exception:
                    pass

            btn = max(36, min(64, int(self.height() * 0.05)))
            if self.width() < 900:
                btn = 36
            else:
                btn = max(btn, 44)
            if hasattr(self, '_help_btn'):
                self._help_btn.setFixedSize(btn, btn)
            if hasattr(self, '_settings_btn'):
                self._settings_btn.setFixedSize(btn, btn)
            if hasattr(self, '_chat_toggle_btn'):
                size = 48
                btn = self._chat_toggle_btn
                btn.setFixedSize(size, size)
                margin_x = 0
                margin_y = 20
                edge = getattr(self, '_chat_anchor_edge', 'right')
                ratio = getattr(self, '_chat_anchor_ratio', None)
                if ratio is None:
                    x = self.width() - size - margin_x
                    y = self.height() - size - margin_y
                else:
                    if edge == 'left':
                        x = margin_x
                    else:
                        x = self.width() - size - margin_x
                    center_y = int(self.height() * float(ratio))
                    y = center_y - size // 2
                    y = max(margin_y, min(self.height() - size - margin_y, y))
                btn.move(x, y)
            if hasattr(self, 'chat_widget') and self.chat_widget.isVisible():
                try:
                    self._position_chat_widget()
                except Exception:
                    pass
        except Exception:
            pass
        try:
            scale = self._current_window_scale()
            has_startup = bool(getattr(self, "_startup_overlay", None)) and bool(getattr(self, "_startup_overlay").isVisible() if getattr(self, "_startup_overlay", None) else False)
            if has_startup:
                try:
                    self._startup_pending_scale = float(scale)
                except Exception:
                    self._startup_pending_scale = scale
            else:
                self._apply_global_scale(scale)
                try:
                    self._startup_pending_scale = None
                except Exception:
                    pass
        except Exception:
            pass
        try:
            if hasattr(self, 'chat_widget'):
                self._update_chat_layout_for_dock()
        except Exception:
            pass
        try:
            if getattr(self, "_left_nav_drawer", None):
                self._update_left_nav_geometry()
        except Exception:
            pass
        try:
            if hasattr(self, '_loading_overlay') and self._loading_overlay and self._loading_overlay.isVisible():
                try:
                    self._loading_overlay.setGeometry(self.rect())
                except Exception:
                    pass
        except Exception:
            pass

        try:
            if hasattr(self, '_startup_overlay') and self._startup_overlay and self._startup_overlay.isVisible():
                try:
                    self._startup_overlay.setGeometry(self.rect())
                except Exception:
                    pass
                try:
                    logo_locked = getattr(self, "_startup_logo_locked_rect", None)
                    logo = getattr(self, "_startup_logo_label", None)
                    if logo_locked and logo:
                        logo.setGeometry(logo_locked)
                        try:
                            if getattr(self, "_startup_logo_hidden_after_dock", False):
                                logo.hide()
                        except Exception:
                            pass
                    elif not getattr(self, "_startup_logo_anim", None):
                        self._update_startup_overlay_layout()
                except Exception:
                    pass
        except Exception:
            pass

    def _run_background_task(self, func, on_finished, loading_message=None, icon_variant="default"):
        try:
            if loading_message is not None:
                try:
                    self._show_loading(loading_message, icon_variant)
                except Exception:
                    pass
        except Exception:
            pass
        try:
            thread = QThread()
        except Exception:
            thread = None
        if not thread:
            try:
                result = None
                error = None
                try:
                    result = func()
                except Exception as e:
                    error = e
                if on_finished:
                    try:
                        on_finished(result, error)
                    except Exception:
                        pass
            finally:
                try:
                    self._hide_loading()
                except Exception:
                    pass
            return
        try:
            worker = _BackgroundWorker(func)
        except Exception:
            worker = None
        if not worker:
            try:
                self._hide_loading()
            except Exception:
                pass
            return
        try:
            worker.moveToThread(thread)
        except Exception:
            pass
        try:
            if self._background_threads is None:
                self._background_threads = []
        except Exception:
            self._background_threads = []
        try:
            if getattr(self, "_background_workers", None) is None:
                self._background_workers = []
        except Exception:
            self._background_workers = []
        try:
            self._background_threads.append(thread)
        except Exception:
            pass
        try:
            self._background_workers.append(worker)
        except Exception:
            pass

        bg_owner = None
        try:
            bg_owner = f"bg_{id(thread)}"
        except Exception:
            bg_owner = "bg"
        try:
            l0 = "en"
            try:
                if self.settings_manager:
                    l0 = str(self.settings_manager.get_language() or "en")
            except Exception:
                l0 = "en"
            status_text = "Edubot đang xử lý..." if str(l0).lower().startswith("vi") else "Edubot is processing..."
            cw = getattr(self, "chat_widget", None)
            if cw is not None and cw.isVisible():
                cw.set_status(status_text, owner=bg_owner)
        except Exception:
            pass

        def _finalize(result, error):
            try:
                cb_error = None
                cb_tb = None
                try:
                    self._hide_loading()
                except Exception:
                    pass
                try:
                    cw = getattr(self, "chat_widget", None)
                    if cw is not None:
                        cw.clear_status(owner=bg_owner)
                except Exception:
                    pass
                try:
                    if on_finished:
                        try:
                            on_finished(result, error)
                        except Exception as e:
                            cb_error = e
                            try:
                                import traceback

                                cb_tb = traceback.format_exc()
                            except Exception:
                                cb_tb = None
                except Exception:
                    pass
                try:
                    if worker in self._background_workers:
                        self._background_workers.remove(worker)
                except Exception:
                    pass
                if cb_error is not None:
                    try:
                        print("[Export] on_finished crashed:", cb_error)
                        try:
                            if cb_tb:
                                print(cb_tb)
                        except Exception:
                            pass
                    except Exception:
                        pass
                    try:
                        cw = getattr(self, "chat_widget", None)
                        if cw is not None and bool(getattr(cw, "_busy", False)):
                            try:
                                cw.set_ai_response("Lỗi nội bộ khi cập nhật kết quả AI. Bạn thử gửi lại." if hasattr(self, "settings_manager") and (self.settings_manager.get_language() == "vi") else "Internal error while updating AI result. Please retry.")
                            except Exception:
                                pass
                            try:
                                cw.set_busy(False)
                            except Exception:
                                pass
                    except Exception:
                        pass
            except Exception:
                pass

        def _on_worker_finished(result, error):
            try:
                from PySide6.QtCore import QTimer
                try:
                    QTimer.singleShot(0, self, lambda: _finalize(result, error))
                except Exception:
                    QTimer.singleShot(0, lambda: _finalize(result, error))
            except Exception:
                try:
                    _finalize(result, error)
                except Exception:
                    pass

        def _on_thread_finished():
            try:
                if thread in self._background_threads:
                    self._background_threads.remove(thread)
            except Exception:
                pass

        try:
            worker.finished.connect(_on_worker_finished)
        except Exception:
            pass
        try:
            worker.finished.connect(thread.quit)
        except Exception:
            pass
        try:
            worker.finished.connect(worker.deleteLater)
        except Exception:
            pass
        try:
            thread.finished.connect(_on_thread_finished)
        except Exception:
            pass
        try:
            thread.finished.connect(thread.deleteLater)
        except Exception:
            pass
        try:
            thread.started.connect(worker.run)
        except Exception:
            pass
        try:
            thread.start()
        except Exception as start_error:
            try:
                self._hide_loading()
            except Exception:
                pass
            try:
                from PySide6.QtCore import QTimer
                try:
                    QTimer.singleShot(0, self, lambda se=start_error: _finalize(None, se))
                except Exception:
                    QTimer.singleShot(0, lambda se=start_error: _finalize(None, se))
            except Exception:
                try:
                    _finalize(None, start_error)
                except Exception:
                    pass

    def _update_startup_overlay_layout(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            if not overlay:
                return
            logo = getattr(self, "_startup_logo_label", None)
            title = getattr(self, "_startup_title_label", None)
            if not logo:
                return
            size = logo.width() or logo.height() or 200
            w = max(1, overlay.width())
            h = max(1, overlay.height())
            target_x = int((w - size) / 2)
            target_y = int(h / 2) - int(size * 0.8)
            logo.setGeometry(target_x, target_y, size, size)
            if title:
                title_height = title.sizeHint().height()
                title_width = max(1, title.sizeHint().width())
                try:
                    tpm = title.pixmap()
                    if tpm and not tpm.isNull():
                        title_width = max(1, tpm.width())
                        title_height = max(title_height, tpm.height())
                except Exception:
                    pass
                title_y = target_y + size + 24
                title_x = max(0, int((w - title_width) / 2))
                title.setGeometry(title_x, title_y, title_width, title_height)
                try:
                    title.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                except Exception:
                    pass
        except Exception:
            pass

    def _snap_chat_toggle_to_edge(self):
        try:
            btn = self._chat_toggle_btn
        except AttributeError:
            return
        margin_x = 0
        margin_y = 20
        size = btn.size()
        x = btn.x()
        y = btn.y()
        w = max(1, self.width())
        h = max(1, self.height())
        dist_left = x
        dist_right = w - (x + size.width())
        if dist_left <= dist_right:
            self._chat_anchor_edge = 'left'
            new_x = margin_x
        else:
            self._chat_anchor_edge = 'right'
            new_x = w - size.width() - margin_x
        center_y = y + size.height() / 2.0
        self._chat_anchor_ratio = max(0.0, min(1.0, center_y / float(h)))
        new_y_center = h * self._chat_anchor_ratio
        new_y = int(new_y_center - size.height() / 2.0)
        new_y = max(margin_y, min(h - size.height() - margin_y, new_y))
        btn.move(new_x, new_y)
        try:
            if self.settings_manager:
                try:
                    self.settings_manager.set("ui.chat_toggle.edge", self._chat_anchor_edge)
                except Exception:
                    pass
                try:
                    self.settings_manager.set("ui.chat_toggle.ratio", float(self._chat_anchor_ratio or 0.0))
                except Exception:
                    pass
        except Exception:
            pass

    def _create_gradient_title_pixmap(self, text: str, theme: str = "dark") -> QPixmap:
        try:
            font_path = str(materialize_asset_file("eduplay/resources/fonts/FC-Ethnocentric-Rg.otf"))
            family = None
            if os.path.exists(font_path):
                fid = QFontDatabase.addApplicationFont(font_path)
                fams = QFontDatabase.applicationFontFamilies(fid)
                if fams:
                    family = fams[0]
            f = QFont(family or "Arial")
            f.setPointSize(54)
            f.setBold(True)
            try:
                f.setLetterSpacing(QFont.AbsoluteSpacing, -1.6)
            except Exception:
                pass
            fm = QFontMetrics(f)
            margin = 18
            tw = fm.horizontalAdvance(text)
            th = fm.height()
            w = max(1, tw + margin * 2)
            h = max(1, th + margin * 2)
            img = QImage(w, h, QImage.Format.Format_ARGB32)
            img.fill(0)
            p = QPainter(img)
            p.setRenderHint(QPainter.Antialiasing, True)
            p.setRenderHint(QPainter.TextAntialiasing, True)
            path = QPainterPath()
            baseline = margin + fm.ascent()
            path.addText(margin, baseline, f, text)
            outline = QPen(QColor(255, 255, 255, 230))
            outline.setWidth(8)
            outline.setJoinStyle(Qt.MiterJoin)
            p.setPen(outline)
            p.setBrush(Qt.NoBrush)
            p.drawPath(path)
            grad = QLinearGradient(0, margin, w, h - margin)
            if theme == "dark":
                grad.setColorAt(0.0, QColor("#FFD166"))
                grad.setColorAt(0.4, QColor("#FF5EA3"))
                grad.setColorAt(1.0, QColor("#7C83FF"))
            else:
                grad.setColorAt(0.0, QColor("#F59E0B"))
                grad.setColorAt(0.5, QColor("#EF4444"))
                grad.setColorAt(1.0, QColor("#8B5CF6"))
            p.setPen(Qt.NoPen)
            p.setBrush(grad)
            p.drawPath(path)
            p.end()
            pm = QPixmap.fromImage(self._trim_transparent_image(img))
            return pm
        except Exception:
            pm = QPixmap()
            return pm

    def _create_gradient_char_pixmap(self, ch: str, theme: str = "dark") -> QPixmap:
        try:
            font_path = str(materialize_asset_file("eduplay/resources/fonts/FC-Ethnocentric-Rg.otf"))
            family = None
            if os.path.exists(font_path):
                fid = QFontDatabase.addApplicationFont(font_path)
                fams = QFontDatabase.applicationFontFamilies(fid)
                if fams:
                    family = fams[0]
            f = QFont(family or "Arial")
            f.setPointSize(54)
            f.setBold(True)
            try:
                f.setLetterSpacing(QFont.AbsoluteSpacing, -1.6)
            except Exception:
                pass
            fm = QFontMetrics(f)
            # Keep enough horizontal padding for thick outline to avoid clipping first/last glyph.
            margin_x = 12
            margin_y = 10
            adv = max(1, fm.horizontalAdvance(ch))
            h = fm.height()
            w = adv + margin_x * 2
            img = QImage(w, h + margin_y * 2, QImage.Format.Format_ARGB32)
            img.fill(0)
            p = QPainter(img)
            p.setRenderHint(QPainter.Antialiasing, True)
            p.setRenderHint(QPainter.TextAntialiasing, True)
            path = QPainterPath()
            baseline = margin_y + fm.ascent()
            path.addText(margin_x, baseline, f, ch)
            outline = QPen(QColor(255, 255, 255, 230))
            outline.setWidth(8)
            outline.setJoinStyle(Qt.MiterJoin)
            p.setPen(outline)
            p.setBrush(Qt.NoBrush)
            p.drawPath(path)
            grad = QLinearGradient(0, margin_y, w, h + margin_y)
            if theme == "dark":
                grad.setColorAt(0.0, QColor("#FFD166"))
                grad.setColorAt(0.4, QColor("#FF5EA3"))
                grad.setColorAt(1.0, QColor("#7C83FF"))
            else:
                grad.setColorAt(0.0, QColor("#F59E0B"))
                grad.setColorAt(0.5, QColor("#EF4444"))
                grad.setColorAt(1.0, QColor("#8B5CF6"))
            p.setPen(Qt.NoPen)
            p.setBrush(grad)
            p.drawPath(path)
            p.end()
            return QPixmap.fromImage(self._trim_transparent_image(img))
        except Exception:
            return QPixmap()

    def _trim_transparent_image(self, image: QImage) -> QImage:
        try:
            if not image or image.isNull():
                return image
            w = image.width()
            h = image.height()
            min_x, min_y = w, h
            max_x, max_y = -1, -1
            for y in range(h):
                for x in range(w):
                    if QColor(image.pixel(x, y)).alpha() > 0:
                        if x < min_x:
                            min_x = x
                        if y < min_y:
                            min_y = y
                        if x > max_x:
                            max_x = x
                        if y > max_y:
                            max_y = y
            if max_x < min_x or max_y < min_y:
                return image
            pad = 2
            left = max(0, min_x - pad)
            top = max(0, min_y - pad)
            right = min(w - 1, max_x + pad)
            bottom = min(h - 1, max_y + pad)
            return image.copy(left, top, right - left + 1, bottom - top + 1)
        except Exception:
            return image

    def _cleanup_char_typing(self, fade_out: bool = False):
        try:
            labels = getattr(self, "_startup_char_labels", []) or []
        except Exception:
            labels = []
        if not labels:
            return
        try:
            if fade_out:
                for lbl in labels:
                    try:
                        eff = lbl.graphicsEffect()
                        if not isinstance(eff, QGraphicsOpacityEffect):
                            eff = QGraphicsOpacityEffect(lbl)
                            lbl.setGraphicsEffect(eff)
                        anim = QPropertyAnimation(eff, b"opacity", self)
                        anim.setDuration(180)
                        anim.setStartValue(1.0)
                        anim.setEndValue(0.0)
                        anim.setEasingCurve(QEasingCurve.OutCubic)
                        anim.start(QPropertyAnimation.DeleteWhenStopped)
                    except Exception:
                        pass
        except Exception:
            pass
        try:
            for lbl in labels:
                try:
                    lbl.hide()
                    lbl.deleteLater()
                except Exception:
                    pass
        except Exception:
            pass
        try:
            self._startup_char_labels = []
        except Exception:
            pass

    def start_startup_animation(self, force_credits: bool = False):
        try:
            try:
                accessibility = self.settings_manager.get_accessibility_settings() if self.settings_manager else {}
            except Exception:
                accessibility = {}
            if bool((accessibility or {}).get("reduce_motion", False)) and not force_credits:
                # Skip animation but still trigger post-startup prompts (onboarding/whats new)
                try:
                    self._finish_startup_animation()
                except Exception:
                    pass
                return
            if getattr(self, "_startup_overlay", None):
                return
            try:
                self._startup_home_fade_done = False
            except Exception:
                pass
            overlay = QWidget(self)
            overlay.setObjectName("startup-overlay")
            try:
                theme = self.settings_manager.get_theme() if self.settings_manager else 'light'
            except Exception:
                theme = 'dark'
            if theme == 'dark':
                overlay_style = "background-color: #020617;"
                title_color = "#E5E7EB"
            else:
                overlay_style = (
                    "background: qlineargradient(x1:0, y1:0, x2:1, y2:1, "
                    "stop:0 #E0F2FE, stop:0.45 #ECFDF3, stop:1 #FEF3C7);"
                )
                title_color = "#0F1728"
            overlay.setStyleSheet(overlay_style)
            overlay.setGeometry(self.rect())
            overlay.raise_()
            self._startup_overlay = overlay
            self._startup_force_credits = bool(force_credits)
            self._startup_theme = theme
            try:
                from PySide6.QtGui import QPixmap
                icon_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../resources/icons/icon.png"))
                pm = None
                if os.path.exists(icon_path):
                    pm = QPixmap(icon_path).scaled(240, 240, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            except Exception:
                pm = None
            logo_label = QLabel(overlay)
            logo_label.setAlignment(Qt.AlignCenter)
            size = 200
            if pm and not pm.isNull():
                try:
                    logo_label.setStyleSheet("background: transparent;")
                except Exception:
                    pass
                logo_label.setPixmap(pm)
                size = pm.width()
            w = max(1, overlay.width())
            h = max(1, overlay.height())
            target_x = int((w - size) / 2)
            target_y = int(h / 2) - int(size * 0.8)
            logo_label.setGeometry(target_x, target_y, size, size)
            self._startup_logo_label = logo_label
            title_label = QLabel(overlay)
            title_label.setAlignment(Qt.AlignCenter)
            text_pm = self._create_gradient_title_pixmap("EduPlay - Studio", theme=theme)
            if text_pm and not text_pm.isNull():
                try:
                    title_label.setStyleSheet("background: transparent;")
                except Exception:
                    pass
                title_label.setPixmap(text_pm)
                try:
                    shadow = QGraphicsDropShadowEffect(title_label)
                    shadow.setBlurRadius(28)
                    shadow.setOffset(0, 8)
                    shadow.setColor(QColor(0, 0, 0, 160))
                    title_label.setGraphicsEffect(shadow)
                except Exception:
                    pass
            else:
                title_label.setStyleSheet(
                    f"color: {title_color}; font-size: 34px; font-weight: 800; letter-spacing: 0.5px;"
                )
                title_label.setText("EduPlay - Studio")
            title_height = title_label.sizeHint().height()
            title_width = max(1, title_label.sizeHint().width())
            try:
                tpm0 = title_label.pixmap()
                if tpm0 and not tpm0.isNull():
                    title_width = max(1, tpm0.width())
                    title_height = max(title_height, tpm0.height())
            except Exception:
                pass
            title_y = target_y + size + 24
            title_label.setGeometry(0, title_y, w, title_height)
            self._startup_title_label = title_label
            try:
                eff = QGraphicsOpacityEffect(logo_label)
                logo_label.setGraphicsEffect(eff)
                fade = QPropertyAnimation(eff, b"opacity", self)
                fade.setDuration(900)
                fade.setStartValue(0.0)
                fade.setEndValue(1.0)
                fade.start(QPropertyAnimation.DeleteWhenStopped)
            except Exception:
                pass
            overlay.show()
            try:
                self._update_startup_overlay_layout()
            except Exception:
                pass
            try:
                self._start_startup_title_typing()
            except Exception:
                self._finish_startup_animation()
        except Exception:
            try:
                self._finish_startup_animation()
            except Exception:
                pass

    def _start_startup_title_typing(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            title = getattr(self, "_startup_title_label", None)
            from PySide6.QtCore import QTimer
            try:
                self._cleanup_char_typing(fade_out=False)
            except Exception:
                pass
            if title:
                try:
                    title.show()
                    title.raise_()
                except Exception:
                    pass
                try:
                    rect = title.geometry()
                    try:
                        pm = title.pixmap()
                    except Exception:
                        pm = None
                    text_w = max(1, pm.width()) if (pm and not pm.isNull()) else max(1, title.sizeHint().width())
                    text_h = max(1, pm.height()) if (pm and not pm.isNull()) else max(1, title.sizeHint().height())
                    if rect.width() > text_w + 2:
                        text_x = rect.x() + int((rect.width() - text_w) / 2)
                        text_y = rect.y() + int((rect.height() - text_h) / 2)
                        title.setGeometry(text_x, text_y, text_w, text_h)
                        rect = title.geometry()
                    title.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
                    wipe_in = _StartupLeftToRightWipeEffect(title)
                    # Start fully clipped, then reveal to the right while text stays fixed.
                    wipe_in.setCutX(max(1, rect.width()) + 2)
                    title.setGraphicsEffect(wipe_in)
                    reveal = QVariantAnimation(self)
                    reveal.setDuration(700)
                    reveal.setStartValue(max(1, rect.width()) + 2)
                    reveal.setEndValue(0)
                    reveal.setEasingCurve(QEasingCurve.InOutCubic)
                    reveal.valueChanged.connect(lambda v: wipe_in.setCutX(int(v)))
                    self._startup_text_anim_group = reveal
                    reveal.start()
                except Exception:
                    pass
            timer = QTimer(self)
            timer.setSingleShot(True)
            def _after_intro_phase():
                try:
                    if bool(getattr(self, "_startup_force_credits", False)):
                        self._start_startup_credits_sequence(force_credits=True)
                    else:
                        self._start_startup_text_exit()
                except Exception:
                    self._finish_startup_animation()
            timer.timeout.connect(_after_intro_phase)
            self._startup_title_timer = timer
            # Include reveal time + hold time so intro is readable.
            timer.start(3400 if (overlay and title) else 1800)
        except Exception:
            try:
                self._finish_startup_animation()
            except Exception:
                pass

    def _finish_startup_animation(self):
        try:
            if getattr(self, "_startup_logo_anim", None):
                try:
                    self._startup_logo_anim.stop()
                except Exception:
                    pass
                self._startup_logo_anim = None
            try:
                self._startup_logo_locked_rect = None
            except Exception:
                pass
            try:
                self._startup_logo_hidden_after_dock = False
            except Exception:
                pass
            if getattr(self, "_startup_text_anim_group", None):
                try:
                    self._startup_text_anim_group.stop()
                except Exception:
                    pass
                self._startup_text_anim_group = None
            if getattr(self, "_startup_char_labels", None):
                try:
                    for lbl in self._startup_char_labels:
                        try:
                            lbl.hide()
                            lbl.deleteLater()
                        except Exception:
                            pass
                except Exception:
                    pass
                self._startup_char_labels = []
            if getattr(self, "_startup_exit_delay", None):
                try:
                    self._startup_exit_delay.stop()
                except Exception:
                    pass
                self._startup_exit_delay = None
            if getattr(self, "_startup_text_exit_timer", None):
                try:
                    self._startup_text_exit_timer.stop()
                except Exception:
                    pass
                self._startup_text_exit_timer = None
            if getattr(self, "_startup_title_timer", None):
                try:
                    self._startup_title_timer.stop()
                except Exception:
                    pass
                self._startup_title_timer = None
            if getattr(self, "_startup_overlay", None):
                try:
                    self._startup_overlay.hide()
                    self._startup_overlay.deleteLater()
                except Exception:
                    pass
                self._startup_overlay = None
            try:
                self.stacked_widget.setGraphicsEffect(None)
            except Exception:
                pass
            try:
                self._startup_home_fade_done = True
            except Exception:
                pass
            try:
                pending = getattr(self, "_startup_pending_scale", None)
                if pending:
                    self._apply_global_scale(float(pending))
            except Exception:
                try:
                    w = max(1, self.width())
                    h = max(1, self.height())
                    scale = max(0.55, min(1.75, min(w/1200.0, h/800.0)))
                    self._apply_global_scale(scale)
                except Exception:
                    pass
            try:
                self._startup_pending_scale = None
            except Exception:
                pass
            try:
                self._show_post_startup_prompts()
            except Exception:
                pass
            try:
                self._animate_home_header_title_reveal()
            except Exception:
                pass
            try:
                home = getattr(self, "home_screen", None)
                if home and hasattr(home, "animate_cards_entrance"):
                    from PySide6.QtCore import QTimer
                    QTimer.singleShot(30, home.animate_cards_entrance)
            except Exception:
                pass
        except Exception:
            pass

    def _animate_home_header_title_reveal(self):
        """Reveal header title from left to right while keeping its position fixed."""
        try:
            title = getattr(self, "_title_label", None)
            if not title or not title.isVisible():
                return
            try:
                prev = getattr(self, "_home_header_title_anim", None)
                if prev:
                    prev.stop()
            except Exception:
                pass
            try:
                title.setGraphicsEffect(None)
            except Exception:
                pass
            wipe = _StartupLeftToRightWipeEffect(title)
            title.setGraphicsEffect(wipe)
            try:
                text_w = max(1, title.fontMetrics().horizontalAdvance(title.text()) + 2)
            except Exception:
                text_w = max(1, title.width())
            clip_w = max(1, min(max(1, title.width()), text_w))
            wipe.setCutX(clip_w + 2)
            anim = QVariantAnimation(self)
            anim.setDuration(460)
            anim.setStartValue(clip_w + 2)
            anim.setEndValue(0)
            anim.setEasingCurve(QEasingCurve.InOutCubic)
            anim.valueChanged.connect(lambda v: wipe.setCutX(int(v)))
            # Ensure no full-text frame appears before reveal starts.
            try:
                title.hide()
            except Exception:
                pass
            def _done():
                try:
                    title.setGraphicsEffect(None)
                except Exception:
                    pass
                self._home_header_title_anim = None
            anim.finished.connect(_done)
            self._home_header_title_anim = anim
            from PySide6.QtCore import QTimer
            def _start():
                try:
                    title.show()
                    title.raise_()
                except Exception:
                    pass
                try:
                    anim.start()
                except Exception:
                    pass
            QTimer.singleShot(0, _start)
        except Exception:
            pass

    def _show_post_startup_prompts(self):
        try:
            if bool(getattr(self, "_post_startup_prompt_running", False)):
                return
        except Exception:
            pass
        try:
            self._post_startup_prompt_running = True
        except Exception:
            pass

        try:
            from PySide6.QtCore import QTimer
        except Exception:
            QTimer = None  # type: ignore

        def _continue_after_whats_new():
            try:
                self._pending_whats_new_rc2 = False
            except Exception:
                pass
            try:
                self._maybe_show_onboarding()
            except Exception:
                pass
            try:
                self._post_startup_prompt_running = False
            except Exception:
                pass

        try:
            pending_whats_new = bool(getattr(self, "_pending_whats_new_rc2", False))
        except Exception:
            pending_whats_new = False
        if pending_whats_new:
            try:
                if QTimer:
                    QTimer.singleShot(120, lambda: self._show_whats_new_rc2(_continue_after_whats_new))
                else:
                    self._show_whats_new_rc2(_continue_after_whats_new)
            except Exception:
                _continue_after_whats_new()
            return

        try:
            if QTimer:
                QTimer.singleShot(120, self._maybe_show_onboarding)
            else:
                self._maybe_show_onboarding()
        except Exception:
            pass
        try:
            self._post_startup_prompt_running = False
        except Exception:
            pass

    def _show_whats_new_rc2(self, done_cb=None):
        try:
            if not self.settings_manager:
                if done_cb:
                    done_cb()
                return
        except Exception:
            if done_cb:
                done_cb()
            return
        try:
            lang = getattr(self, "_pending_whats_new_lang", None) or self.settings_manager.get_language() or "en"
        except Exception:
            lang = "en"
        try:
            from eduplay.ui.widgets.whats_new_dialog import WhatsNewDialog
        except Exception:
            if done_cb:
                done_cb()
            return
        dlg = WhatsNewDialog(self, lang=lang)

        def _closed(_dont_show: bool):
            try:
                if self.settings_manager:
                    self.settings_manager.set("whats_new.rc2.dismissed", True)
            except Exception:
                pass
            try:
                if done_cb:
                    done_cb()
            except Exception:
                pass

        try:
            dlg.closed.connect(_closed)
        except Exception:
            pass
        try:
            dlg.exec()
        except Exception:
            try:
                dlg.show()
            except Exception:
                pass
            try:
                if done_cb:
                    done_cb()
            except Exception:
                pass

    def _maybe_show_onboarding(self):
        try:
            pending = bool(getattr(self, "_pending_first_run_onboarding", False))
        except Exception:
            pending = False
        if not pending:
            return
        try:
            lang = getattr(self, "_pending_onboarding_lang", None) or (self.settings_manager.get_language() if self.settings_manager else "en") or "en"
        except Exception:
            lang = "en"

        def _build_overlay():
            try:
                home = getattr(self, "home_screen", None)
                marks: list[tuple[QWidget, str]] = []
                try:
                    from eduplay.core.i18n import I18n

                    step_create_card = I18n.t("onboarding.step_create_card", lang)
                    step_edit_card = I18n.t("onboarding.step_edit_card", lang)
                    step_play_card = I18n.t("onboarding.step_play_card", lang)
                except Exception:
                    step_create_card = "Tạo mới: bắt đầu dự án học tập mới" if lang == "vi" else "Create: start a new learning project"
                    step_edit_card = "Chỉnh sửa: mở dự án đã lưu" if lang == "vi" else "Edit: open an existing project"
                    step_play_card = "Chơi: thử hoặc xuất bản game đã tạo" if lang == "vi" else "Play: test or publish your game"

                # Ensure home screen is fully laid out before inspecting widgets
                if home is not None:
                    try:
                        home.updateGeometry()
                        home.layout().update()
                        home.layout().activate()
                    except Exception:
                        pass

                # Guide the 3 main home cards (not the smaller inner buttons)
                # so the highlight borders and text bubbles align with the
                # large TẠO MỚI / CHỈNH SỬA / CHƠI tiles.
                if home is not None:
                    c1 = getattr(home, "create_card", None)
                    c2 = getattr(home, "edit_card", None)
                    c3 = getattr(home, "play_card", None)
                    if c1 is not None and c1.isVisible():
                        marks.append((c1, step_create_card))
                    if c2 is not None and c2.isVisible():
                        marks.append((c2, step_edit_card))
                    if c3 is not None and c3.isVisible():
                        marks.append((c3, step_play_card))

                if not marks:
                    _show_fallback_welcome()
                    return

                from eduplay.ui.widgets.coach_marks_overlay import CoachMarksOverlay
                overlay = CoachMarksOverlay(self, marks, lang=lang)

                def _done(_dont_show: bool):
                    try:
                        self._pending_first_run_onboarding = False
                    except Exception:
                        pass
                    try:
                        if self.settings_manager:
                            self.settings_manager.set("first_run", False)
                            self.settings_manager.set("onboarding.dismissed", True if _dont_show else False)
                    except Exception:
                        pass
                    # After the quick onboarding overlay, automatically open
                    # the detailed help guide so first-time users get a full
                    # walkthrough of the current interface.
                    try:
                        if QTimer:
                            QTimer.singleShot(0, self.show_help)
                    except Exception:
                        pass

                overlay.finished.connect(_done)
                overlay.show()
                overlay.raise_()
                overlay.activateWindow()
            except Exception:
                _show_fallback_welcome()

        def _show_fallback_welcome():
            try:
                from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QLabel, QPushButton
                from PySide6.QtCore import Qt
                from eduplay.core.i18n import I18n

                fallback = QFrame(self)
                fallback.setObjectName("onboarding-fallback")
                fallback.setAttribute(Qt.WA_StyledBackground, True)
                fallback.setFocusPolicy(Qt.NoFocus)

                title = I18n.t("onboarding.fallback_title", lang)
                body = I18n.t("onboarding.fallback_body", lang)
                done_text = I18n.t("onboarding.done", lang)

                lay = QVBoxLayout(fallback)
                lay.setContentsMargins(32, 32, 32, 32)
                lay.setSpacing(16)

                t = QLabel(title, fallback)
                t.setAlignment(Qt.AlignCenter)
                t.setStyleSheet("color:#EDEEF3;font-size:22px;font-weight:800;background:transparent;")

                b = QLabel(body, fallback)
                b.setAlignment(Qt.AlignCenter)
                b.setWordWrap(True)
                b.setStyleSheet("color:#EDEEF3;font-size:14px;font-weight:400;background:transparent;")

                btn = QPushButton(done_text, fallback)
                btn.setCursor(Qt.PointingHandCursor)
                btn.setStyleSheet(
                    "QPushButton{background-color:#7F56D9;color:#FFFFFF;border:none;border-radius:12px;padding:10px 24px;font-weight:800;} QPushButton:hover{background-color:#6941C6;}"
                )

                lay.addStretch()
                lay.addWidget(t)
                lay.addWidget(b)
                lay.addWidget(btn, 0, Qt.AlignCenter)
                lay.addStretch()

                fallback.setStyleSheet(
                    "#onboarding-fallback{background-color:rgba(15,23,42,0.92);border:2px solid rgba(127,86,217,0.9);border-radius:20px;}"
                )
                fallback.setGeometry(self.rect().adjusted(80, 80, -80, -80))

                def _close():
                    try:
                        self._pending_first_run_onboarding = False
                    except Exception:
                        pass
                    try:
                        if self.settings_manager:
                            self.settings_manager.set("first_run", False)
                            self.settings_manager.set("onboarding.dismissed", True)
                    except Exception:
                        pass
                    try:
                        fallback.hide()
                        fallback.deleteLater()
                    except Exception:
                        pass

                btn.clicked.connect(_close)
                fallback.show()
                fallback.raise_()
                fallback.activateWindow()
            except Exception:
                try:
                    self._pending_first_run_onboarding = False
                except Exception:
                    pass

        try:
            QTimer.singleShot(800, _build_overlay)
        except Exception:
            _build_overlay()

    def _start_startup_exit_sequence(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            if not overlay:
                self._finish_startup_animation()
                return
            try:
                # Hide moving startup elements before fading overlay to avoid ghost flash.
                logo = getattr(self, "_startup_logo_label", None)
                title = getattr(self, "_startup_title_label", None)
                # Keep docked logo visible until home fade completes.
                if logo:
                    try:
                        # Freeze exactly at the docked rect to avoid last-moment reposition jumps.
                        self._startup_logo_locked_rect = QRect(logo.geometry())
                    except Exception:
                        pass
                    try:
                        locked = getattr(self, "_startup_logo_locked_rect", None)
                        if locked:
                            logo.setGeometry(locked)
                    except Exception:
                        pass
                    logo.show()
                    logo.raise_()
                if title:
                    title.hide()
            except Exception:
                pass
            try:
                # Avoid overlay opacity fade to prevent one-frame geometry drift at the end.
                # Hold docked logo briefly, then remove startup overlay immediately.
                from PySide6.QtCore import QTimer
                timer = QTimer(self)
                timer.setSingleShot(True)
                try:
                    locked = getattr(self, "_startup_logo_locked_rect", None)
                    if logo and locked:
                        logo.setGeometry(locked)
                        logo.raise_()
                except Exception:
                    pass
                timer.timeout.connect(self._finish_startup_animation)
                self._startup_exit_delay = timer
                timer.start(220)
            except Exception:
                self._finish_startup_animation()
        except Exception:
            self._finish_startup_animation()

    def _start_startup_text_exit(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            logo = getattr(self, "_startup_logo_label", None)
            if not (overlay and logo):
                self._finish_startup_animation()
                return
            title = getattr(self, "_startup_title_label", None)
            try:
                self._cleanup_char_typing(fade_out=False)
            except Exception:
                pass
            title = getattr(self, "_startup_title_label", None)
            if not title:
                self._animate_startup_logo_to_header()
                return
            rect = title.geometry()
            try:
                pm = title.pixmap()
            except Exception:
                pm = None
            text_w = 0
            text_h = rect.height()
            if pm and not pm.isNull():
                text_w = max(1, pm.width())
                text_h = max(1, pm.height())
            else:
                text_w = max(1, title.sizeHint().width())
            if rect.width() > text_w + 2:
                text_x = rect.x() + int((rect.width() - text_w) / 2)
                text_y = rect.y() + int((rect.height() - text_h) / 2)
                title.setGeometry(text_x, text_y, text_w, text_h)
                rect = title.geometry()
            try:
                title.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)
            except Exception:
                pass
            wipe = _StartupLeftToRightWipeEffect(title)
            title.setGraphicsEffect(wipe)
            wipe_anim = QVariantAnimation(self)
            wipe_anim.setDuration(700)
            wipe_anim.setStartValue(0)
            wipe_anim.setEndValue(max(1, rect.width()) + 2)
            wipe_anim.setEasingCurve(QEasingCurve.InOutCubic)
            wipe_anim.valueChanged.connect(lambda v: wipe.setCutX(int(v)))
            def _after():
                try:
                    title.hide()
                    title.setGraphicsEffect(None)
                except Exception:
                    pass
                try:
                    self._animate_startup_logo_to_header()
                except Exception:
                    self._start_startup_exit_sequence()
            wipe_anim.finished.connect(_after)
            self._startup_text_anim_group = wipe_anim
            wipe_anim.start()
        except Exception:
            self._finish_startup_animation()

    def _header_icon_target_rect_on_overlay(self, overlay: QWidget, header_icon: QWidget) -> QRect:
        try:
            if not (overlay and header_icon):
                return QRect()
            try:
                if header_icon.isVisible():
                    top_left_on_main = header_icon.mapTo(self, QPoint(0, 0))
                    top_left_on_overlay = overlay.mapFrom(self, top_left_on_main)
                    tw = max(1, header_icon.width())
                    th = max(1, header_icon.height())
                    return QRect(top_left_on_overlay.x(), top_left_on_overlay.y(), tw, th)
            except Exception:
                pass
            # Compact header fallback: anchor to real title position first.
            title = getattr(self, "_title_label", None)
            target_size = 32
            if title:
                top_left_on_main = title.mapTo(self, QPoint(0, 0))
                top_left_on_overlay = overlay.mapFrom(self, top_left_on_main)
                base_h = max(1, title.height())
                y = top_left_on_overlay.y() + int((base_h - target_size) / 2)
                x = max(0, top_left_on_overlay.x() - target_size - 10)
                return QRect(x, y, target_size, target_size)
            # Last fallback uses header margin if title is unavailable.
            header = getattr(self, "header_frame", None)
            if header:
                try:
                    lm = self.header_layout.contentsMargins().left()
                except Exception:
                    lm = 20
                header_top_left = header.mapTo(self, QPoint(0, 0))
                header_on_overlay = overlay.mapFrom(self, header_top_left)
                x = max(0, int(header_on_overlay.x() + lm))
                y = max(0, int(header_on_overlay.y() + (header.height() - target_size) / 2))
                return QRect(x, y, target_size, target_size)
            x = 20
            y = 14
            return QRect(x, y, target_size, target_size)
        except Exception:
            return QRect()

    def _animate_startup_logo_to_header(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            logo = getattr(self, "_startup_logo_label", None)
            header_icon = getattr(self, "_header_icon", None)
            if not (overlay and logo and header_icon):
                self._start_startup_exit_sequence()
                return
            start_rect = logo.geometry()
            try:
                logo.setScaledContents(True)
                try:
                    # Flush layout before mapping to reduce drift from late geometry updates.
                    if self.layout():
                        self.layout().activate()
                    parent = header_icon.parentWidget()
                    if parent and parent.layout():
                        parent.layout().activate()
                    QApplication.processEvents()
                except Exception:
                    pass
                end_rect = self._header_icon_target_rect_on_overlay(overlay, header_icon)
                if end_rect.isNull() or end_rect.width() <= 0 or end_rect.height() <= 0:
                    raise RuntimeError("invalid header icon target")
            except Exception:
                self._start_startup_exit_sequence()
                return
            anim = QPropertyAnimation(logo, b"geometry", self)
            anim.setDuration(1400)
            anim.setStartValue(start_rect)
            anim.setEndValue(end_rect)
            anim.setEasingCurve(QEasingCurve.InOutCubic)
            def _done():
                try:
                    # Keep final position stable; do not recompute fallback target here.
                    logo.setGeometry(end_rect)
                except Exception:
                    pass
                try:
                    self._startup_logo_locked_rect = QRect(end_rect)
                except Exception:
                    self._startup_logo_locked_rect = end_rect
                try:
                    self._startup_logo_anim = None
                except Exception:
                    pass
                try:
                    # Keep overlay logo visible at docked position until exit fade is done.
                    self._startup_logo_hidden_after_dock = False
                    if logo:
                        logo.show()
                except Exception:
                    pass
                try:
                    self._start_startup_exit_sequence()
                except Exception:
                    self._finish_startup_animation()
            anim.finished.connect(_done)
            self._startup_logo_anim = anim
            anim.start()
        except Exception:
            self._start_startup_exit_sequence()

    def _start_startup_credits_sequence(self, force_credits: bool = False):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            logo = getattr(self, "_startup_logo_label", None)
            title = getattr(self, "_startup_title_label", None)
            if not overlay or not logo or not title:
                self._start_startup_text_exit()
                return
            try:
                title.show()
            except Exception:
                pass
            try:
                self._cleanup_char_typing(fade_out=True)
            except Exception:
                pass
            if not force_credits:
                try:
                    from eduplay.core.settings_manager import SettingsManager
                    sm = self.settings_manager or SettingsManager()
                    sm.set("intro_credits_shown", True)
                except Exception:
                    pass
            w = max(1, overlay.width())
            h = max(1, overlay.height())
            logo_rect = logo.geometry()
            title_rect = title.geometry()
            try:
                logo.setScaledContents(True)
            except Exception:
                pass
            try:
                title.setScaledContents(True)
            except Exception:
                pass
            logo_scale = 0.44
            title_scale = 0.72
            logo_w = max(1, int(logo_rect.width() * logo_scale))
            logo_h = max(1, int(logo_rect.height() * logo_scale))
            title_pm = title.pixmap()
            if title_pm and not title_pm.isNull():
                base_tw = title_pm.width()
                base_th = title_pm.height()
                title_w = max(1, int(base_tw * title_scale))
                title_h = max(1, int(base_th * title_scale))
            else:
                title_w = max(1, int(title_rect.width() * title_scale))
                title_h = max(1, int(title_rect.height() * title_scale))
            top_y = max(12, int(h * 0.04))
            total_width = logo_w + 16 + title_w
            start_x = int((w - total_width) / 2)
            logo_target_x = start_x
            logo_target_y = top_y
            title_target_x = logo_target_x + logo_w + 16
            title_target_y = logo_target_y + int((logo_h - title_h) / 2)
            logo_target = QRect(logo_target_x, logo_target_y, logo_w, logo_h)
            title_target = QRect(title_target_x, title_target_y, title_w, title_h)
            move_group = QParallelAnimationGroup(self)
            logo_anim = QPropertyAnimation(logo, b"geometry", self)
            logo_anim.setDuration(800)
            logo_anim.setStartValue(logo_rect)
            logo_anim.setEndValue(logo_target)
            logo_anim.setEasingCurve(QEasingCurve.InOutCubic)
            title_anim = QPropertyAnimation(title, b"geometry", self)
            title_anim.setDuration(800)
            title_anim.setStartValue(title_rect)
            title_anim.setEndValue(title_target)
            title_anim.setEasingCurve(QEasingCurve.InOutCubic)
            move_group.addAnimation(logo_anim)
            move_group.addAnimation(title_anim)
            def _after_move():
                try:
                    self._show_startup_credits_content()
                except Exception:
                    self._start_startup_text_exit()
            move_group.finished.connect(_after_move)
            self._startup_logo_anim = move_group
            move_group.start()
        except Exception:
            self._start_startup_text_exit()

    def _show_startup_credits_content(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            if not overlay:
                self._start_startup_text_exit()
                return
            logo = getattr(self, "_startup_logo_label", None)
            theme = getattr(self, "_startup_theme", "light")
            credits = QWidget(overlay)
            try:
                credits.setStyleSheet("background: transparent;")
            except Exception:
                pass
            credits_layout = QVBoxLayout(credits)
            credits_layout.setContentsMargins(40, 40, 40, 40)
            credits_layout.setSpacing(24)
            name_color = "#E5E7EB" if theme == "dark" else "#0F1728"
            role_color = "#CBD5E1" if theme == "dark" else "#64748B"
            email_color = "#E5E7EB" if theme == "dark" else "#0F1728"
            try:
                from eduplay.core.i18n import I18n
                from eduplay.core.settings_manager import SettingsManager
                lang = SettingsManager().get_language() or "en"
            except Exception:
                I18n = None
                lang = "en"
            row1 = QHBoxLayout()
            row1.setSpacing(24)
            left1 = QVBoxLayout()
            left1.setSpacing(4)
            if I18n:
                dev_name_text = I18n.t("credits.dev_name", lang)
                dev_role_text = I18n.t("credits.dev_role", lang)
                dev_email_text = I18n.t("credits.dev_email", lang)
            else:
                dev_name_text = "Nguyễn Thanh Tân"
                dev_role_text = "Developer"
                dev_email_text = "- Email: ttmeomeo807@gmail.com – tannguyenthanh498@gmail.com"
            name1 = QLabel(dev_name_text, credits)
            name1.setStyleSheet(f"color: {name_color}; font-size: 22px; font-weight: 800;")
            role1 = QLabel(dev_role_text, credits)
            role1.setStyleSheet(f"color: {role_color}; font-size: 16px;")
            email1 = QLabel(dev_email_text, credits)
            email1.setStyleSheet(f"color: {email_color}; font-size: 14px;")
            left1.addWidget(name1)
            left1.addWidget(role1)
            left1.addWidget(email1)
            row1.addLayout(left1)
            qr1 = QLabel(credits)
            qr1.setAlignment(Qt.AlignCenter)
            try:
                from PySide6.QtGui import QPixmap
                qr1_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../resources/icons/ORCID (meomeo).png"))
                if os.path.exists(qr1_path):
                    pm1 = QPixmap(qr1_path).scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    qr1.setPixmap(pm1)
            except Exception:
                pass
            row1.addWidget(qr1, 0, Qt.AlignRight | Qt.AlignVCenter)
            credits_layout.addLayout(row1)
            row2 = QHBoxLayout()
            row2.setSpacing(24)
            qr2 = QLabel(credits)
            qr2.setAlignment(Qt.AlignCenter)
            try:
                from PySide6.QtGui import QPixmap
                qr2_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../resources/icons/ORCID (maomao).png"))
                if os.path.exists(qr2_path):
                    pm2 = QPixmap(qr2_path).scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                    qr2.setPixmap(pm2)
            except Exception:
                pass
            row2.addWidget(qr2, 0, Qt.AlignLeft | Qt.AlignVCenter)
            right2 = QVBoxLayout()
            right2.setSpacing(4)
            if I18n:
                editor_name_text = I18n.t("credits.editor_name", lang)
                editor_role_text = I18n.t("credits.editor_role", lang)
                editor_email_text = I18n.t("credits.editor_email", lang)
            else:
                editor_name_text = "Bùi Ngô Anh Thư"
                editor_role_text = "Editor"
                editor_email_text = "- Email: btmaomao262@gmail.com"
            name2 = QLabel(editor_name_text, credits)
            name2.setStyleSheet(f"color: {name_color}; font-size: 22px; font-weight: 800;")
            role2 = QLabel(editor_role_text, credits)
            role2.setStyleSheet(f"color: {role_color}; font-size: 16px;")
            email2 = QLabel(editor_email_text, credits)
            email2.setStyleSheet(f"color: {email_color}; font-size: 14px;")
            right2.addWidget(name2)
            right2.addWidget(role2)
            right2.addWidget(email2)
            row2.addLayout(right2)
            credits_layout.addLayout(row2)
            try:
                if I18n:
                    ack_title_text = I18n.t("credits.ack.title", lang)
                    ack_body_text = I18n.t("credits.ack.body", lang)
                else:
                    ack_title_text = "Acknowledgements"
                    ack_body_text = (
                        "EduPlay Studio sincerely expresses its gratitude to the individuals, communities, and organizations "
                        "whose work, ideas, and contributions have directly or indirectly supported the development of this project."
                    )
            except Exception:
                ack_title_text = "Acknowledgements"
                ack_body_text = ""
            ack_title = QLabel(ack_title_text, credits)
            ack_title.setObjectName("startup-ack-title")
            ack_title.setStyleSheet(f"color: {name_color}; font-size: 18px; font-weight: 900; margin-top: 8px;")
            ack_body = QLabel(ack_body_text, credits)
            ack_body.setObjectName("startup-ack-body")
            ack_body.setWordWrap(True)
            ack_body.setStyleSheet(f"color: {role_color}; font-size: 14px; line-height: 1.35;")
            credits_layout.addWidget(ack_title)
            credits_layout.addWidget(ack_body)
            credits.adjustSize()
            w = max(1, overlay.width())
            h = max(1, overlay.height())
            cw = min(credits.sizeHint().width(), w - 80)
            ch = credits.sizeHint().height()
            start_x = int((w - cw) / 2)
            credits.setGeometry(start_x, h, cw, ch)
            credits.show()
            self._startup_credits_widget = credits
            boundary_y = None
            try:
                title = getattr(self, "_startup_title_label", None)
                if title:
                    pm = title.pixmap()
                    if pm and not pm.isNull():
                        img = pm.toImage()
                        try:
                            img = img.convertToFormat(QImage.Format.Format_ARGB32)
                        except Exception:
                            pass
                        width = img.width()
                        height = img.height()
                        x_step = 2
                        best_y = int(height * 0.75)
                        best_count = -1
                        samples_per_row = max(1, int(width / x_step))
                        threshold = int(samples_per_row * 0.35)
                        for yy in range(height):
                            cnt = 0
                            for xx in range(0, width, x_step):
                                c = img.pixelColor(xx, yy)
                                if c.alpha() < 12:
                                    continue
                                if (c.red() + c.green() + c.blue()) <= 180:
                                    cnt += 1
                            if cnt > best_count:
                                best_count = cnt
                                best_y = yy
                        if best_count < threshold:
                            best_y = int(height * 0.75)
                        boundary_y = title.geometry().y() + int(best_y)
                    else:
                        boundary_y = title.geometry().y() + int(title.height() * 0.75)
            except Exception:
                boundary_y = None
            if boundary_y is None:
                boundary_y = int(h * 0.25)
            try:
                boundary_y = max(0, boundary_y + 48)
            except Exception:
                pass
            try:
                wipe = _StartupBoundaryWipeEffect(credits)
                credits.setGraphicsEffect(wipe)
                self._startup_credits_wipe_effect = wipe
            except Exception:
                wipe = None
            def _update_wipe(v):
                if not wipe:
                    return
                try:
                    y = int(v.y())
                except Exception:
                    try:
                        y = int(credits.y())
                    except Exception:
                        y = 0
                wipe.setCutY(boundary_y - y)
            _update_wipe(credits.pos())
            start_pos = QPoint(start_x, h)
            end_pos = QPoint(start_x, int(boundary_y - ch))
            scroll_distance = max(1, start_pos.y() - end_pos.y())
            # Keep the credits roll at a constant speed so it feels cinematic
            # instead of easing, pausing, then jumping into the next segment.
            pixels_per_second = max(56.0, min(72.0, float(h) / 14.0))
            scroll_duration = int((float(scroll_distance) / pixels_per_second) * 1000.0)
            scroll_duration = max(7000, min(16000, scroll_duration))
            seq = QSequentialAnimationGroup(self)
            move_roll = QPropertyAnimation(credits, b"pos", self)
            move_roll.setDuration(scroll_duration)
            move_roll.setStartValue(start_pos)
            move_roll.setEndValue(end_pos)
            move_roll.setEasingCurve(QEasingCurve.Linear)
            try:
                move_roll.valueChanged.connect(_update_wipe)
            except Exception:
                pass
            seq.addAnimation(move_roll)
            def _done():
                try:
                    credits.hide()
                    credits.deleteLater()
                except Exception:
                    pass
                try:
                    self._startup_credits_widget = None
                except Exception:
                    pass
                try:
                    self._start_startup_credits_exit_to_center()
                except Exception:
                    self._start_startup_text_exit()
            seq.finished.connect(_done)
            self._startup_text_anim_group = seq
            seq.start()
        except Exception:
            self._start_startup_text_exit()

    def _start_startup_credits_exit_to_center(self):
        try:
            overlay = getattr(self, "_startup_overlay", None)
            logo = getattr(self, "_startup_logo_label", None)
            title = getattr(self, "_startup_title_label", None)
            if not overlay or not logo or not title:
                self._start_startup_text_exit()
                return
            w = max(1, overlay.width())
            h = max(1, overlay.height())
            size = logo.width() or logo.height() or 200
            center_x = int((w - size) / 2)
            center_y = int(h / 2) - int(size * 0.8)
            logo_target = QRect(center_x, center_y, logo.width(), logo.height())
            title_height = title.height()
            title_target = QRect(0, center_y + logo.height() + 24, w, title_height)
            move_group = QParallelAnimationGroup(self)
            logo_anim = QPropertyAnimation(logo, b"geometry", self)
            logo_anim.setDuration(800)
            logo_anim.setStartValue(logo.geometry())
            logo_anim.setEndValue(logo_target)
            logo_anim.setEasingCurve(QEasingCurve.InOutCubic)
            title_anim = QPropertyAnimation(title, b"geometry", self)
            title_anim.setDuration(800)
            title_anim.setStartValue(title.geometry())
            title_anim.setEndValue(title_target)
            title_anim.setEasingCurve(QEasingCurve.InOutCubic)
            move_group.addAnimation(logo_anim)
            move_group.addAnimation(title_anim)
            def _after_move():
                try:
                    self._start_startup_text_exit()
                except Exception:
                    self._finish_startup_animation()
            move_group.finished.connect(_after_move)
            self._startup_logo_anim = move_group
            move_group.start()
        except Exception:
            self._start_startup_text_exit()

    def eventFilter(self, obj, event):
        if hasattr(self, '_chat_toggle_btn') and obj is self._chat_toggle_btn:
            if event.type() == QEvent.MouseButtonPress and event.button() == Qt.LeftButton:
                self._chat_drag_active = True
                self._chat_drag_moved = False
                try:
                    self._chat_drag_start_pos = self._chat_toggle_btn.pos()
                except Exception:
                    self._chat_drag_start_pos = QPoint()
                try:
                    self._chat_drag_offset = event.position().toPoint()
                except AttributeError:
                    self._chat_drag_offset = event.pos()
                return False
            if event.type() == QEvent.MouseMove and self._chat_drag_active and (event.buttons() & Qt.LeftButton):
                try:
                    global_pos = event.globalPosition().toPoint()
                except AttributeError:
                    global_pos = event.globalPos()
                local = self.mapFromGlobal(global_pos)
                new_x = local.x() - self._chat_drag_offset.x()
                new_y = local.y() - self._chat_drag_offset.y()
                w = max(1, self.width())
                h = max(1, self.height())
                bw = self._chat_toggle_btn.width()
                bh = self._chat_toggle_btn.height()
                new_x = max(0, min(w - bw, new_x))
                new_y = max(0, min(h - bh, new_y))
                try:
                    start = getattr(self, "_chat_drag_start_pos", self._chat_toggle_btn.pos())
                    dx = new_x - start.x()
                    dy = new_y - start.y()
                    if abs(dx) > 3 or abs(dy) > 3:
                        self._chat_drag_moved = True
                except Exception:
                    self._chat_drag_moved = True
                self._chat_toggle_btn.move(new_x, new_y)
                return True
            if event.type() == QEvent.MouseButtonRelease and self._chat_drag_active:
                self._chat_drag_active = False
                if getattr(self, "_chat_drag_moved", False):
                    self._snap_chat_toggle_to_edge()
                    return True
                return False
        return super().eventFilter(obj, event)

    def _apply_global_scale(self, scale: float):
        try:
            self._last_global_scale = float(scale)
        except Exception:
            self._last_global_scale = scale
        try:
            from PySide6.QtWidgets import QApplication, QPushButton
            base = 12
            f = QApplication.font()
            if f.pointSize() > 0:
                f.setPointSize(max(1, int(base * scale)))
            QApplication.setFont(f)
            if hasattr(self, '_help_btn') and self._help_btn:
                s = max(36, int(44 * scale))
                self._help_btn.setFixedSize(s, s)
            if hasattr(self, '_settings_btn') and self._settings_btn:
                s = max(36, int(44 * scale))
                self._settings_btn.setFixedSize(s, s)
            if hasattr(self, '_chat_toggle_btn') and self._chat_toggle_btn:
                self._chat_toggle_btn.setFixedSize(48, 48)
                try:
                    bf = self._chat_toggle_btn.font()
                    bf.setPointSizeF(max(18.0, min(26.0, 22.0 * float(scale))))
                    self._chat_toggle_btn.setFont(bf)
                except Exception:
                    pass
            if hasattr(self, 'chat_widget') and self.chat_widget:
                try:
                    if hasattr(self.chat_widget, 'set_scale'):
                        self.chat_widget.set_scale(scale)
                except Exception:
                    pass
            if hasattr(self, '_title_label') and self._title_label:
                tf = self._title_label.font()
                tf.setPointSize(max(18, int(22 * scale)))
                self._title_label.setFont(tf)
            current = self.stacked_widget.currentWidget()
            if current is self.home_screen and hasattr(self.home_screen, 'apply_style'):
                self.home_screen.apply_style(scale)
            if current is self.browser_screen and hasattr(self.browser_screen, 'set_scale'):
                self.browser_screen.set_scale(scale)
            if current is self.new_project_screen and hasattr(self.new_project_screen, 'set_scale'):
                try:
                    self.new_project_screen.set_scale(scale)
                except Exception:
                    pass
            if current is self.player_screen and hasattr(self.player_screen, 'set_scale'):
                try:
                    self.player_screen.set_scale(scale)
                except Exception:
                    pass
            if current is self.editor_screen and hasattr(self.editor_screen, 'set_scale'):
                try:
                    self.editor_screen.set_scale(scale)
                except Exception:
                    pass
        except Exception:
            pass

    def on_chat_dock_changed(self, docked: bool):
        try:
            if self.chat_widget.isVisible():
                self._position_chat_widget()
        except Exception:
            pass
        try:
            self._update_chat_layout_for_dock()
        except Exception:
            pass

    def on_chat_width_changed(self, mode: str):
        try:
            if self.chat_widget.isVisible():
                self._position_chat_widget()
        except Exception:
            pass
        try:
            self._update_chat_layout_for_dock()
        except Exception:
            pass

    def on_chat_detail_changed(self, level: str):
        try:
            self._chat_detail_level = level
        except Exception:
            pass

    def on_chat_message(self, message: str):
        try:
            try:
                language = self.settings_manager.get_language()
            except Exception:
                language = 'en'
            try:
                ui_lang = str(getattr(self.chat_widget, "language", "") or "").strip()
                if ui_lang:
                    language = ui_lang
            except Exception:
                pass
            raw_msg = ""
            try:
                raw_msg = str(message or "")
            except Exception:
                raw_msg = ""
            msg = raw_msg.strip()
            attachments = []
            try:
                if "\n[ATTACHMENTS]" in msg or msg.strip().endswith("[ATTACHMENTS]") or "[ATTACHMENTS]" in msg:
                    parts = msg.split("[ATTACHMENTS]", 1)
                    msg = str(parts[0] or "").strip()
                    tail = str(parts[1] or "")
                    files = []
                    for ln in tail.splitlines():
                        s = str(ln or "").strip()
                        if not s:
                            continue
                        files.append(s)
                    attachments = [f for f in files if f]
            except Exception:
                attachments = []
            lower = msg.lower()
            try:
                self._last_chat_user_msg = msg
                self._last_chat_language = language
            except Exception:
                pass
            try:
                print(f"[AI-STEP] received_message lang={language} attachments_hint={('[ATTACHMENTS]' in raw_msg)}")
            except Exception:
                pass

            try:
                import time as _time

                now = float(_time.monotonic() or 0.0)
                guard = getattr(self, "_chat_quota_guard", None)
                if not isinstance(guard, dict):
                    guard = {
                        "banned_until": 0.0,
                        "strike": 0,
                        "last_strike_ts": 0.0,
                        "recent": [],
                        "last_norm": "",
                        "repeat_streak": 0,
                    }
                    self._chat_quota_guard = guard

                banned_until = float(guard.get("banned_until", 0.0) or 0.0)
                if banned_until and now < banned_until:
                    rem = int(max(1.0, banned_until - now))
                    if language == "vi":
                        text = f"Bạn đang bị tạm khóa {rem}s vì gửi quá nhanh/lặp lại. Quota không tự nhiên mà có đâu bạn. Nếu tiếp tục spam, thời gian khóa sẽ tăng gấp đôi."
                    else:
                        text = f"You are temporarily blocked for {rem}s due to rapid/repeated messages. Quota isn't free. Repeated abuse will double the cooldown."
                    try:
                        self.chat_widget.set_busy(False)
                    except Exception:
                        pass
                    try:
                        self.chat_widget.add_message("System", text, "ai")
                    except Exception:
                        try:
                            self.chat_widget.set_ai_response(text)
                        except Exception:
                            pass
                    return

                def _norm_text(s: str) -> str:
                    try:
                        import re as _re

                        t = str(s or "").strip().lower()
                        t = _re.sub(r"\s+", " ", t)
                        return t
                    except Exception:
                        try:
                            return str(s or "").strip().lower()
                        except Exception:
                            return ""

                norm = _norm_text(msg)
                recent = guard.get("recent", []) or []
                try:
                    recent = [(float(t), str(n)) for t, n in recent if (now - float(t)) <= 15.0]
                except Exception:
                    recent = []
                recent.append((now, norm))
                guard["recent"] = recent

                last_norm = str(guard.get("last_norm", "") or "")
                if norm and norm == last_norm:
                    guard["repeat_streak"] = int(guard.get("repeat_streak", 0) or 0) + 1
                else:
                    guard["repeat_streak"] = 1 if norm else 0
                    guard["last_norm"] = norm

                too_fast = False
                try:
                    if len(recent) >= 6:
                        t0 = float(recent[0][0])
                        too_fast = (now - t0) <= 10.0
                except Exception:
                    too_fast = False
                too_repeat = int(guard.get("repeat_streak", 0) or 0) >= 3

                if (too_fast or too_repeat) and norm:
                    last_strike_ts = float(guard.get("last_strike_ts", 0.0) or 0.0)
                    if last_strike_ts and (now - last_strike_ts) > 900.0:
                        guard["strike"] = 0
                    strike = int(guard.get("strike", 0) or 0) + 1
                    guard["strike"] = strike
                    guard["last_strike_ts"] = now

                    dur = 120.0 * (2.0 ** float(max(0, strike - 1)))
                    if dur > 3600.0:
                        dur = 3600.0
                    guard["banned_until"] = now + dur

                    rem = int(max(1.0, dur))
                    if language == "vi":
                        text = f"Bạn đang gửi quá nhanh/lặp lại. Quota không tự nhiên mà có đâu bạn. Bạn bị cấm dùng Edubot trong {rem}s."
                    else:
                        text = f"You're sending too fast/repeating. Quota isn't free. You are blocked from using Edubot for {rem}s."
                    try:
                        self.chat_widget.set_busy(False)
                    except Exception:
                        pass
                    try:
                        self.chat_widget.add_message("System", text, "ai")
                    except Exception:
                        try:
                            self.chat_widget.set_ai_response(text)
                        except Exception:
                            pass
                    return
            except Exception:
                pass

            try:
                ml = str(msg or "").strip().lower()
            except Exception:
                ml = ""
            try:
                is_greeting = (not attachments) and (ml in ("chào", "xin chào", "hello", "hi", "hey"))
            except Exception:
                is_greeting = False
            if is_greeting:
                try:
                    self.chat_widget.set_busy(False)
                except Exception:
                    pass
                try:
                    self.chat_widget.set_ai_response("Xin chào! Bạn cần tôi giúp gì hôm nay?" if language == "vi" else "Hello! How can I help you today?")
                except Exception:
                    pass
                return

            try:
                is_capability = (not attachments) and (ml in ("bạn làm được gì", "ban lam duoc gi", "edubot làm được gì", "edubot lam duoc gi", "help", "giúp gì", "giup gi"))
            except Exception:
                is_capability = False
            if is_capability:
                try:
                    from eduplay.core.i18n import I18n
                    if language == "vi":
                        text = (
                            "Tôi có thể giúp bạn:\n"
                            "• Tạo câu hỏi từ nội dung bài giảng/tài liệu\n"
                            "• Tạo trắc nghiệm, đúng/sai, điền khuyết, trả lời ngắn\n"
                            "• Thêm câu hỏi vào đúng dự án bạn chỉ định\n"
                            "• Gợi ý đáp án và giải thích ngắn gọn\n\n"
                            "Bạn thử gửi: “Tạo 5 câu trắc nghiệm về … và thêm vào dự án Thử Thách”"
                        )
                    else:
                        text = (
                            "I can help you:\n"
                            "• Create questions from your documents/lesson content\n"
                            "• Generate multiple-choice, true/false, fill-in, short answers\n"
                            "• Add questions into the project you specify\n"
                            "• Provide brief explanations\n\n"
                            "Try: “Create 5 quiz questions about … and add to project Thử Thách”"
                        )
                    self.chat_widget.set_ai_response(text)
                except Exception:
                    self.chat_widget.set_ai_response("Tôi có thể tạo câu hỏi từ tài liệu và thêm vào dự án bạn chỉ định." if language == "vi" else "I can create questions from documents and add them to your project.")
                return

            focus_title = ""
            try:
                import re
                patterns = [
                    r"(?:bài|bai)\s+(.+?)\s+(?:trong|tu|từ)\s+file",
                    r"(?:bài|bai)\s+(.+?)(?:\s+và|\s+va|\s*,|\s*;|\s*\.|\s*$)",
                    r"(?:tiết|tiet)\s+(.+?)(?:\s+và|\s+va|\s*,|\s*;|\s*\.|\s*$)",
                    r"(?:mục|muc|phần|phan|chủ đề|chu de)\s+(.+?)(?:\s+và|\s+va|\s*,|\s*;|\s*\.|\s*$)",
                ]
                for pat in patterns:
                    m_title = re.search(pat, msg, flags=re.IGNORECASE)
                    if m_title:
                        focus_title = str(m_title.group(1) or "").strip().strip('"').strip("'")
                        if focus_title:
                            break
            except Exception:
                focus_title = ""

            try:
                import re
                m_proj = re.search(r"(?:dự án|du an|project)\s+([^\n\r,.;]+)", msg, flags=re.IGNORECASE)
                target_project_name = str(m_proj.group(1) or "").strip().strip('"').strip("'") if m_proj else ""
            except Exception:
                target_project_name = ""
            current_project = None
            try:
                active_editor = None
                for _nm in ("editor_quiz_screen", "editor_fishing_screen", "editor_screen"):
                    try:
                        _scr = getattr(self, _nm, None)
                    except Exception:
                        _scr = None
                    if not _scr:
                        continue
                    try:
                        _p = _scr.get_current_project() if hasattr(_scr, "get_current_project") else getattr(_scr, "current_project", None)
                    except Exception:
                        _p = None
                    if isinstance(_p, dict) and str(_p.get("id") or "").strip():
                        active_editor = _scr
                        current_project = _p
                        break
                if active_editor is None:
                    current_project = None
            except Exception:
                current_project = None
            try:
                if (not current_project) and target_project_name:
                    projects = self.project_manager.get_all_projects()
                    found = None
                    for p in projects:
                        try:
                            if str(p.get("name") or "").strip().lower() == target_project_name.lower():
                                found = p
                                break
                        except Exception:
                            continue
                    if found:
                        proj_data = self.project_manager.load_project(found.get("id"))
                        if proj_data:
                            self._open_editor_for_project(proj_data)
                            current_project = proj_data
            except Exception:
                pass
            try:
                if not (isinstance(current_project, dict) and str(current_project.get("id") or "").strip()):
                    pm_proj = None
                    try:
                        pm_proj = self.project_manager.get_current_project()
                    except Exception:
                        pm_proj = None
                    if isinstance(pm_proj, dict) and str(pm_proj.get("id") or "").strip():
                        current_project = pm_proj
                if isinstance(current_project, dict) and str(current_project.get("id") or "").strip():
                    try:
                        cur_pm = self.project_manager.get_current_project() or {}
                    except Exception:
                        cur_pm = {}
                    try:
                        cur_pid = str((cur_pm or {}).get("id") or "").strip()
                    except Exception:
                        cur_pid = ""
                    try:
                        pid2 = str(current_project.get("id") or "").strip()
                    except Exception:
                        pid2 = ""
                    if pid2 and (not cur_pid or cur_pid != pid2):
                        try:
                            self.project_manager.current_project = current_project
                        except Exception:
                            pass
            except Exception:
                pass
            try:
                from eduplay.core.i18n import I18n
                try:
                    loading_msg = I18n.t('ai.loading', language)
                except Exception:
                    loading_msg = "Đang suy nghĩ..." if language == 'vi' else "Thinking..."
            except Exception:
                loading_msg = "Đang suy nghĩ..." if language == 'vi' else "Thinking..."
            game_type = (current_project or {}).get('game_type', 'quiz')
            topic = (current_project or {}).get('name', 'chủ đề hiện tại')
            app_context = {}
            try:
                app_context["current_screen"] = getattr(self, "_current_screen_name", "home")
            except Exception:
                pass
            try:
                if current_project:
                    app_context["current_project"] = {
                        "name": current_project.get("name"),
                        "game_type": current_project.get("game_type"),
                    }
            except Exception:
                pass
            try:
                idx = getattr(active_editor, "current_question_index", -1) if active_editor is not None else -1
                app_context["selected_question_index"] = idx
            except Exception:
                pass
            try:
                qs = (current_project or {}).get("questions", [])
                app_context["question_count"] = len(qs) if isinstance(qs, list) else 0
            except Exception:
                pass
            try:
                msgs = getattr(self.chat_widget, "messages", []) or []
                user_hist = []
                for m in reversed(msgs):
                    try:
                        t = str((m or {}).get("type") or "")
                        sender = str((m or {}).get("sender") or "")
                        content = str((m or {}).get("message") or "").strip()
                    except Exception:
                        continue
                    if not content:
                        continue
                    if sender.lower() == "system":
                        continue
                    if t != "user":
                        continue
                    try:
                        if len(content) > 600:
                            content = content[:600].rstrip() + "..."
                    except Exception:
                        pass
                    user_hist.append({"role": "user", "content": content})
                    if len(user_hist) >= 5:
                        break
                user_hist.reverse()
                app_context["chat_history"] = user_hist
            except Exception:
                pass

            attachments_for_ai = list(attachments or [])
            wants_questions_request = False
            requested_q_count = None
            try:
                import re
                ml2 = str(msg or "").lower()
                wants_questions_request = any(k in ml2 for k in ("câu hỏi", "trắc nghiệm", "quiz")) and any(k in ml2 for k in ("tạo", "generate", "trích xuất", "extract", "soạn", "làm"))
                m_cnt = re.search(r"(?:tạo|tao|generate|create|soạn|soan|viết|viet)\s+(\d+)\s+(?:câu|cau|questions?)", msg, flags=re.IGNORECASE)
                if not m_cnt:
                    m_cnt = re.search(r"(\d+)\s+(?:câu|cau|questions?)", msg, flags=re.IGNORECASE)
                if m_cnt:
                    try:
                        requested_q_count = int(m_cnt.group(1))
                    except Exception:
                        requested_q_count = None
                    if requested_q_count is not None and (requested_q_count <= 0 or requested_q_count > 100):
                        requested_q_count = None
            except Exception:
                wants_questions_request = False
                requested_q_count = None

            wants_edit_request = False
            wants_delete_request = False
            target_question_indices = []
            try:
                import re

                ml3 = str(msg or "").lower()
                wants_edit_request = any(k in ml3 for k in ("sửa", "sua", "chỉnh", "chinh", "chỉnh sửa", "chinh sua", "update", "edit")) and any(
                    k in ml3
                    for k in (
                        "giải thích",
                        "giai thich",
                        "explanation",
                        "phân tích",
                        "phan tich",
                        "đáp án",
                        "dap an",
                        "lựa chọn",
                        "lua chon",
                        "options",
                        "option",
                        "câu hỏi",
                        "cau hoi",
                    )
                )
                wants_delete_request = any(k in ml3 for k in ("xóa", "xoa", "delete", "remove")) and any(
                    k in ml3 for k in ("câu hỏi", "cau hoi", "question", "quiz")
                )

                try:
                    from eduplay.core.ai_tool_payloads import extract_question_numbers
                except Exception:
                    extract_question_numbers = None
                uniq = extract_question_numbers(msg) if extract_question_numbers else []
                try:
                    qs_all = (current_project or {}).get("questions", [])
                    total_q = len(qs_all) if isinstance(qs_all, list) else 0
                except Exception:
                    total_q = 0
                if uniq and total_q:
                    for n in uniq[:12]:
                        idx0 = int(n) - 1
                        if 0 <= idx0 < total_q:
                            target_question_indices.append(idx0)
            except Exception:
                wants_edit_request = False
                wants_delete_request = False
                target_question_indices = []

            try:
                if target_question_indices:
                    qs_all = (current_project or {}).get("questions", [])
                    targets = []
                    if isinstance(qs_all, list):
                        for idx0 in target_question_indices[:12]:
                            try:
                                q = qs_all[idx0]
                            except Exception:
                                q = None
                            if not isinstance(q, dict):
                                continue
                            try:
                                qid = str(q.get("id") or "").strip()
                            except Exception:
                                qid = ""
                            if not qid:
                                continue
                            try:
                                qt = str(q.get("question") or "").strip()
                            except Exception:
                                qt = ""
                            try:
                                exp = str(q.get("explanation") or "").strip()
                            except Exception:
                                exp = ""
                            try:
                                qtype = str(q.get("type") or "").strip()
                            except Exception:
                                qtype = ""
                            try:
                                opts = q.get("options")
                            except Exception:
                                opts = None
                            try:
                                ca = q.get("correct_answer")
                            except Exception:
                                ca = None
                            targets.append(
                                {
                                    "index": int(idx0) + 1,
                                    "question_id": qid,
                                    "type": qtype,
                                    "question": qt[:260],
                                    "options": opts,
                                    "correct_answer": ca,
                                    "explanation": exp[:260],
                                }
                            )
                    if targets:
                        app_context["target_questions"] = targets
                        try:
                            ids2 = [str(t.get("question_id") or "").strip() for t in targets if isinstance(t, dict) and str(t.get("question_id") or "").strip()]
                            self._last_target_question_ids = ids2
                            self._pending_update_question_ids = list(ids2)
                            self._pending_delete_question_ids = list(ids2)
                        except Exception:
                            self._last_target_question_ids = []
                            self._pending_update_question_ids = []
                            self._pending_delete_question_ids = []
            except Exception:
                pass

            selected_model = None
            chat_fallback_model = None
            use_task_model_flag = False
            try:
                import os as _os

                def _resolve_chat_model() -> str | None:
                    env_m = str(_os.getenv("GROQ_CHAT_MODEL") or "").strip()
                    if env_m:
                        return env_m
                    try:
                        s0 = str(self.settings_manager.get("ai_settings.chat_model", "") or "").strip()
                    except Exception:
                        s0 = ""
                    if s0:
                        return s0
                    return "llama-3.1-8b-instant"

                def _resolve_task_model() -> str:
                    env_m = str(_os.getenv("GROQ_TASK_MODEL") or "").strip()
                    if env_m:
                        return env_m
                    try:
                        s = str(self.settings_manager.get("ai_settings.task_model", "") or "").strip()
                    except Exception:
                        s = ""
                    return s or "qwen/qwen3-32b"

                def _looks_like_app_task(text: str) -> bool:
                    try:
                        t = str(text or "").strip().lower()
                    except Exception:
                        return False
                    if not t:
                        return False
                    if "[attachments]" in t:
                        return True
                    tool_words = ("add_question", "update_question", "create_project", "open_project", "set_question_image", "update_game_config", "search_image")
                    if any(w in t for w in tool_words):
                        return True
                    verbs = ("tạo", "tao", "thêm", "them", "sửa", "sua", "chỉnh", "chinh", "cập nhật", "cap nhat", "xóa", "xoa", "mở", "mo", "đổi", "doi", "add", "create", "update", "edit", "open", "set", "bật", "bat", "tắt", "tat", "enable", "disable")
                    objs = ("câu hỏi", "cau hoi", "trắc nghiệm", "trac nghiem", "quiz", "dự án", "du an", "project", "cấu hình", "cau hinh", "game", "hình", "hinh", "ảnh", "anh", "image", "trộn", "tron", "xáo trộn", "xao tron", "shuffle", "random")
                    if any(v in t for v in verbs) and any(o in t for o in objs):
                        return True
                    return False

                chat_model = _resolve_chat_model()
                task_model = _resolve_task_model()
                use_task_model = bool(attachments_for_ai) or bool(wants_questions_request) or _looks_like_app_task(msg)
                selected_model = task_model if use_task_model else chat_model
                chat_fallback_model = chat_model
                use_task_model_flag = bool(use_task_model)
            except Exception:
                selected_model = None
                chat_fallback_model = None
                use_task_model_flag = False

            def _build_attachment_suffix(att_list: list, focus: str = "", teach_only: bool = False) -> str:
                if not att_list:
                    return ""
                try:
                    import os
                    import shutil
                    import tempfile
                    from pathlib import Path
                    import re
                    try:
                        print(f"[AI-STEP] ingest_attachments count={len(att_list)}")
                    except Exception:
                        pass

                    proj = None
                    try:
                        proj = self.project_manager.get_current_project()
                    except Exception:
                        proj = None

                    media_dir = None
                    if isinstance(proj, dict) and str(proj.get("id") or "").strip():
                        pid = str(proj.get("id") or "").strip()
                        media_dir = self.project_manager.projects_dir / pid / "media"
                        try:
                            media_dir.mkdir(parents=True, exist_ok=True)
                        except Exception:
                            pass

                    def _unique_name(base: str, ext: str) -> str:
                        safe_base = re.sub(r"[^a-zA-Z0-9]+", "_", str(base or "")).strip("_")[:48] or "file"
                        safe_ext = str(ext or "").strip()
                        if safe_ext and not safe_ext.startswith("."):
                            safe_ext = "." + safe_ext
                        if not safe_ext:
                            safe_ext = ".bin"
                        cand = (media_dir / f"{safe_base}{safe_ext}") if media_dir else Path(f"{safe_base}{safe_ext}")
                        idx = 1
                        while media_dir and cand.exists():
                            cand = media_dir / f"{safe_base}_{idx}{safe_ext}"
                            idx += 1
                        return cand.name

                    def _extract_docx_text(d) -> str:
                        parts = []
                        seen = set()
                        try:
                            for para in getattr(d, "paragraphs", []) or []:
                                try:
                                    t = str(getattr(para, "text", "") or "").strip()
                                except Exception:
                                    t = ""
                                if t and t not in seen:
                                    seen.add(t)
                                    parts.append(t)
                        except Exception:
                            pass
                        try:
                            tables = getattr(d, "tables", []) or []
                        except Exception:
                            tables = []
                        try:
                            for tb in tables:
                                try:
                                    rows = getattr(tb, "rows", []) or []
                                except Exception:
                                    rows = []
                                for row in rows:
                                    try:
                                        cells = getattr(row, "cells", []) or []
                                    except Exception:
                                        cells = []
                                    vals = []
                                    for cell in cells:
                                        try:
                                            ct = str(getattr(cell, "text", "") or "").strip()
                                        except Exception:
                                            ct = ""
                                        if ct:
                                            ct = " ".join(ct.replace("\r", "").split())
                                            vals.append(ct)
                                    line = " | ".join([v for v in vals if v])
                                    if line and line not in seen:
                                        seen.add(line)
                                        parts.append(line)
                        except Exception:
                            pass
                        return "\n".join(parts).strip()

                    def _read_doc_via_com(doc_path: Path) -> str:
                        try:
                            import win32com.client  # type: ignore
                            import pythoncom  # type: ignore
                        except Exception:
                            return ""
                        tmp_docx = ""
                        try:
                            pythoncom.CoInitialize()
                            word = win32com.client.Dispatch("Word.Application")
                            word.Visible = False
                            doc = word.Documents.Open(str(doc_path))
                            tmp_docx = os.path.join(tempfile.gettempdir(), f"eduplay_{doc_path.stem}.converted.docx")
                            wdFormatXMLDocument = 12
                            doc.SaveAs(tmp_docx, FileFormat=wdFormatXMLDocument)
                            doc.Close(False)
                            word.Quit()
                            pythoncom.CoUninitialize()
                        except Exception:
                            try:
                                pythoncom.CoUninitialize()
                            except Exception:
                                pass
                            tmp_docx = ""
                        if not tmp_docx or not os.path.exists(tmp_docx):
                            return ""
                        try:
                            from docx import Document
                            d = Document(str(tmp_docx))
                            return _extract_docx_text(d)
                        except Exception:
                            return ""

                    blocks = []

                    def _extract_focus_excerpt(text: str, focus_title_in: str) -> tuple[str, bool]:
                        try:
                            focus_s = str(focus_title_in or "").strip()
                        except Exception:
                            focus_s = ""
                        if not focus_s:
                            return str(text or "").strip(), True
                        try:
                            lines = [ln.rstrip() for ln in str(text or "").splitlines()]
                        except Exception:
                            lines = []
                        if not lines:
                            return str(text or "").strip(), False
                        focus_l = focus_s.lower()
                        hit = -1
                        for idx, ln in enumerate(lines):
                            try:
                                if focus_l in str(ln).lower():
                                    hit = idx
                                    break
                            except Exception:
                                continue
                        if hit < 0:
                            return "", False
                        start = max(0, hit - 5)
                        end = min(len(lines), hit + 140)
                        heading_re = None
                        try:
                            heading_re = re.compile(r"^\s*(?:BÀI|BAI|TIẾT|TIET|CHỦ\s*ĐỀ|CHU\s*DE|PHẦN|PHAN)\b", flags=re.IGNORECASE)
                        except Exception:
                            heading_re = None
                        if heading_re is not None:
                            j = hit + 5
                            while j < min(len(lines), hit + 220):
                                t = str(lines[j] or "").strip()
                                if t and heading_re.search(t) and (focus_l not in t.lower()):
                                    end = j
                                    break
                                j += 1
                        excerpt = "\n".join([ln for ln in lines[start:end] if str(ln or "").strip()]).strip()
                        if len(excerpt) > 2400:
                            excerpt = excerpt[:2400].rstrip() + "..."
                        return excerpt, True

                    def _extract_taught_scope(text: str) -> str:
                        try:
                            raw = str(text or "")
                        except Exception:
                            raw = ""
                        if not raw.strip():
                            return ""
                        try:
                            lines = [ln.strip() for ln in raw.replace("\r", "\n").split("\n") if ln.strip()]
                        except Exception:
                            lines = []
                        if not lines:
                            return ""
                        try:
                            import re as _re
                        except Exception:
                            _re = None
                        exclude_prefix = ("môn", "lop", "lớp", "trường", "giáo viên", "giao vien", "ngày", "tuần", "tiết", "tiet", "thời gian", "thoi gian")
                        include_terms = (
                            "nội dung",
                            "noi dung",
                            "kiến thức",
                            "kien thuc",
                            "ghi nhớ",
                            "ghi nho",
                            "ví dụ",
                            "vi du",
                            "củng cố",
                            "cung co",
                            "luyện tập",
                            "luyen tap",
                            "bài tập",
                            "bai tap",
                            "hoạt động",
                            "hoat dong",
                        )
                        picked = []
                        seen = set()
                        for ln in lines:
                            l = ln.lower()
                            if any(l.startswith(p + ":") for p in exclude_prefix):
                                continue
                            if ("|" in ln) and ("hoạt động" in l or "hoat dong" in l or "nội dung" in l or "noi dung" in l):
                                parts = [p.strip() for p in ln.split("|") if p.strip()]
                                if parts:
                                    tail = parts[-1]
                                    if tail and tail.lower() not in seen:
                                        seen.add(tail.lower())
                                        picked.append(tail)
                                continue
                            if any(t in l for t in include_terms):
                                if ln and l not in seen:
                                    seen.add(l)
                                    picked.append(ln)
                                continue
                            if _re is not None and _re.search(r"\b\d{1,4}\s*[\-\+×x\*:/]\s*\d{1,4}\b", ln):
                                if ln and l not in seen:
                                    seen.add(l)
                                    picked.append(ln)
                                continue
                        if not picked:
                            return ""
                        out = "\n".join(picked).strip()
                        if len(out) > 2400:
                            out = out[:2400].rstrip() + "..."
                        return out

                    def _extract_headings(text: str) -> list:
                        try:
                            import re as _re2
                        except Exception:
                            _re2 = None
                        if _re2 is None:
                            return []
                        lines = [ln.strip() for ln in str(text or "").splitlines() if str(ln or "").strip()]
                        head_re = _re2.compile(r"^\s*(?:BÀI|BAI|TIẾT|TIET|CHỦ\s*ĐỀ|CHU\s*DE|PHẦN|PHAN)\b.*$", flags=_re2.IGNORECASE)
                        out = []
                        seen = set()
                        for ln in lines:
                            if head_re.search(ln):
                                key = ln.lower()
                                if key not in seen:
                                    seen.add(key)
                                    out.append(ln)
                            if len(out) >= 12:
                                break
                        return out

                    def _make_lesson_description(taught_text: str, focus_in: str = "") -> str:
                        taught_s = str(taught_text or "").strip()
                        if not taught_s:
                            return ""
                        low = taught_s.lower()
                        parts = []
                        if focus_in:
                            parts.append(f"Chủ đề: {str(focus_in).strip()}")
                        if ("số bị trừ" in low) or ("so bi tru" in low) or ("số trừ" in low) or ("so tru" in low) or ("hiệu" in low) or ("hieu" in low):
                            parts.append("Kiến thức trọng tâm: phép trừ và các thành phần (số bị trừ, số trừ, hiệu).")
                        if ("đặt tính" in low) or ("dat tinh" in low):
                            parts.append("Kĩ năng: đặt tính rồi tính.")
                        if ("giải" in low) or ("giai" in low) or ("lời văn" in low) or ("loi van" in low):
                            parts.append("Kĩ năng: giải bài toán có lời văn liên quan.")
                        if not parts:
                            lines = [ln.strip() for ln in taught_s.splitlines() if ln.strip()]
                            pick = lines[:6]
                            parts = ["Nội dung đã dạy:", *pick]
                        desc = "\n".join(parts).strip()
                        if len(desc) > 900:
                            desc = desc[:900].rstrip() + "..."
                        return desc

                    for i, fp in enumerate(att_list[:6]):
                        p = Path(str(fp or "").strip())
                        if not p.exists() or not p.is_file():
                            continue
                        try:
                            print(f"[AI-STEP] extract_attachment file={p.name}")
                        except Exception:
                            pass
                        try:
                            size = int(p.stat().st_size)
                        except Exception:
                            size = 0
                        if size and size > 25 * 1024 * 1024:
                            continue

                        saved_rel = ""
                        if media_dir:
                            try:
                                name = _unique_name(p.stem, p.suffix or ".bin")
                                dest = media_dir / name
                                try:
                                    src_res = p.resolve()
                                    media_res = media_dir.resolve()
                                    in_media = (src_res == media_res) or (media_res in src_res.parents)
                                except Exception:
                                    in_media = False
                                if not in_media:
                                    shutil.copy2(str(p), str(dest))
                                    saved_rel = f"media/{dest.name}"
                                else:
                                    saved_rel = f"media/{p.name}"
                            except Exception:
                                saved_rel = ""

                        lower_name = p.name.lower()
                        is_pdf = lower_name.endswith(".pdf")
                        is_docx = lower_name.endswith(".docx")
                        is_doc = lower_name.endswith(".doc")
                        content = ""
                        if is_pdf:
                            try:
                                print(f"[AI-STEP] extract_pdf file={p.name}")
                            except Exception:
                                pass
                            try:
                                import fitz
                                doc = fitz.open(str(p))
                                chunks = []
                                for pi in range(min(6, int(getattr(doc, "page_count", 0) or 0))):
                                    try:
                                        chunks.append(doc.load_page(pi).get_text("text"))
                                    except Exception:
                                        pass
                                content = "\n".join([c for c in chunks if str(c or "").strip()]).strip()
                            except Exception:
                                content = ""
                        elif is_docx:
                            try:
                                print(f"[AI-STEP] extract_docx file={p.name}")
                            except Exception:
                                pass
                            try:
                                from docx import Document
                                d = Document(str(p))
                                content = _extract_docx_text(d)
                            except Exception:
                                content = ""
                        elif is_doc:
                            try:
                                print(f"[AI-STEP] extract_doc file={p.name}")
                            except Exception:
                                pass
                            content = _read_doc_via_com(p)
                        else:
                            try:
                                content = p.read_text(encoding="utf-8", errors="ignore")
                            except Exception:
                                content = ""

                        try:
                            content = str(content or "").strip()
                        except Exception:
                            content = ""
                        focus_found = True
                        focus_excerpt = ""
                        if content and focus:
                            focus_excerpt, focus_found = _extract_focus_excerpt(content, focus)
                            if focus_found and focus_excerpt:
                                content = focus_excerpt
                        if focus and (not focus_found) and wants_questions_request:
                            heads = _extract_headings(content)
                            msg_heads = ""
                            if heads:
                                msg_heads = "\n" + "\n".join([f"- {h}" for h in heads])
                            self.chat_widget.set_ai_response(
                                "Tôi không tìm thấy đúng phần/bài bạn chỉ định trong giáo án. Bạn có thể gửi lại tên bài đúng như trong file, hoặc chọn 1 dòng trong danh sách sau:"
                                + msg_heads
                            )
                            return ""
                        if len(content) > 3000:
                            content = content[:3000].rstrip() + "..."
                        taught = _extract_taught_scope(content)
                        lesson_desc = _make_lesson_description(taught, focus)
                        grade = ""
                        try:
                            import re as _re3
                            m_grade = _re3.search(r"\bl(?:ớp|op)\s*[:\-]?\s*(\d{1,2})\b", content, flags=_re3.IGNORECASE)
                            if m_grade:
                                grade = str(m_grade.group(1)).strip()
                        except Exception:
                            grade = ""
                        keywords = []
                        try:
                            import re as _re4
                            src_kw = taught or content or ""
                            tokens = _re4.findall(r"[A-Za-zÀ-ỹ]+", str(src_kw).lower())
                            stop = {
                                "và", "là", "của", "cho", "các", "một", "những", "trong", "với", "khi", "để", "về", "theo", "có", "đã", "sẽ",
                                "này", "đó", "bài", "tiết", "phần", "hoạt", "động", "nội", "dung", "kiến", "thức", "luyện", "tập", "củng", "cố",
                                "môn", "lớp", "giao", "viên", "học", "sinh", "giáo", "dạy"
                            }
                            freq = {}
                            for t in tokens:
                                if len(t) < 3:
                                    continue
                                if t in stop:
                                    continue
                                freq[t] = int(freq.get(t, 0) or 0) + 1
                            ordered = sorted(freq.items(), key=lambda x: (-int(x[1] or 0), x[0]))
                            keywords = [k for k, _ in ordered[:24]]
                        except Exception:
                            keywords = []

                        header = f"FILE_{i+1}: {p.name}"
                        if saved_rel:
                            header += f"\nSAVED_IN_MEDIA: {saved_rel}"
                        if focus:
                            header += f"\nFOCUS: {focus}"
                        if grade:
                            header += f"\nTARGET_GRADE: {grade}"
                        if keywords:
                            header += "\nKEYWORDS_ALLOWED (chỉ dùng thuật ngữ trong danh sách; được dùng số liệu; tránh khái niệm ngoài danh sách):\n" + ", ".join(keywords)
                        if taught:
                            header += "\nPHẦN GIÁO VIÊN ĐÃ DẠY (chỉ dùng phần này để tạo bài luyện tập/câu hỏi):\n" + taught
                        if lesson_desc:
                            header += "\nTÓM TẮT BÀI DẠY (dùng để định hướng tạo luyện tập):\n" + lesson_desc
                        if teach_only:
                            src_for_ai = taught or content
                            if src_for_ai:
                                header += "\nCONTENT:\n" + src_for_ai
                        else:
                            if content:
                                header += "\nCONTENT:\n" + content
                        blocks.append(header.strip())

                    if not blocks:
                        return ""
                    try:
                        from eduplay.core.i18n import I18n
                        header_txt = I18n.t("ai.attachments_header", language)
                    except Exception:
                        header_txt = "TÀI LIỆU ĐÍNH KÈM (dùng để trả lời chính xác hơn):" if language == "vi" else "ATTACHMENTS (use these as context):"
                    return "\n\n" + str(header_txt) + "\n" + "\n\n".join(blocks)
                except Exception:
                    return ""

            def _run_ai(prompt_text: str, att_list: list | None = None):
                suffix_cache = {"suffix": ""}
                try:
                    self._chat_request_seq = int(getattr(self, "_chat_request_seq", 0)) + 1
                except Exception:
                    self._chat_request_seq = 1
                req_id = int(getattr(self, "_chat_request_seq", 1))
                try:
                    self._chat_active_req_id = req_id
                except Exception:
                    pass
                try:
                    self.chat_widget.set_busy(True)
                except Exception:
                    pass

                try:
                    import psutil, os as _os
                    if not hasattr(self, "_chat_prev_priority"):
                        self._chat_prev_priority = None
                    proc = psutil.Process(_os.getpid())
                    prev = proc.nice()
                    self._chat_prev_priority = prev
                    proc.nice(psutil.ABOVE_NORMAL_PRIORITY_CLASS)
                except Exception:
                    pass

                def _progress(step: str):
                    try:
                        try:
                            print(f"[AI-STEP] {str(step)}")
                        except Exception:
                            pass
                        from PySide6.QtCore import QTimer

                        def _append():
                            try:
                                if int(getattr(self, "_chat_active_req_id", 0)) != int(req_id):
                                    return
                            except Exception:
                                return
                            try:
                                self.chat_widget.set_status(str(step), owner=req_id)
                            except Exception:
                                try:
                                    self.chat_widget.set_ai_response(str(step))
                                except Exception:
                                    pass

                        try:
                            QTimer.singleShot(0, self, _append)
                        except Exception:
                            QTimer.singleShot(0, _append)
                    except Exception:
                        pass

                def _task():
                    try:
                        try:
                            print("[AI-STEP] call_ai")
                        except Exception:
                            pass
                        suffix = ""
                        if att_list:
                            try:
                                from eduplay.core.i18n import I18n
                                _progress(I18n.t("ai.reading_attachments", language))
                            except Exception:
                                try:
                                    _progress("Đang đọc tài liệu đính kèm..." if language == "vi" else "Reading attachments...")
                                except Exception:
                                    pass
                            suffix = _build_attachment_suffix(att_list, focus_title, teach_only=bool(wants_questions_request))
                            try:
                                suffix_cache["suffix"] = suffix
                            except Exception:
                                pass
                        full_prompt = prompt_text
                        if suffix:
                            full_prompt = (full_prompt or "") + suffix
                        try:
                            from eduplay.core.i18n import I18n
                            _progress(I18n.t("ai.loading", language))
                        except Exception:
                            try:
                                _progress("Đang suy nghĩ..." if language == "vi" else "Thinking...")
                            except Exception:
                                pass
                        try:
                            res = self.ai_service.chat_with_ai(
                                full_prompt,
                                language,
                                app_context=app_context,
                                model=selected_model,
                                progress_cb=_progress,
                            )
                        except TypeError:
                            res = self.ai_service.chat_with_ai(
                                full_prompt,
                                language,
                                app_context=app_context,
                                progress_cb=_progress,
                            )
                        try:
                            rs = str(res or "")
                        except Exception:
                            rs = ""
                        try:
                            rsl = rs.lower()
                        except Exception:
                            rsl = ""
                        try:
                            invalid_model = ("error communicating with ai (400" in rsl) and ("model" in rsl)
                            invalid_model = invalid_model or (("invalid model" in rsl) or ("model not found" in rsl) or ("does not exist" in rsl) or ("not supported" in rsl))
                        except Exception:
                            invalid_model = False
                        if invalid_model and use_task_model_flag and chat_fallback_model and (str(chat_fallback_model) != str(selected_model)):
                            try:
                                return self.ai_service.chat_with_ai(
                                    full_prompt,
                                    language,
                                    app_context=app_context,
                                    model=chat_fallback_model,
                                    progress_cb=_progress,
                                )
                            except Exception:
                                return rs
                        return rs
                    except TypeError:
                        try:
                            res2 = self.ai_service.chat_with_ai(
                                prompt_text,
                                language,
                                app_context=app_context,
                                model=selected_model,
                                progress_cb=_progress,
                            )
                        except TypeError:
                            res2 = self.ai_service.chat_with_ai(
                                prompt_text,
                                language,
                                app_context=app_context,
                                progress_cb=_progress,
                            )
                        try:
                            rs2 = str(res2 or "")
                        except Exception:
                            rs2 = ""
                        try:
                            rs2l = rs2.lower()
                        except Exception:
                            rs2l = ""
                        try:
                            invalid_model2 = ("error communicating with ai (400" in rs2l) and ("model" in rs2l)
                            invalid_model2 = invalid_model2 or (("invalid model" in rs2l) or ("model not found" in rs2l) or ("does not exist" in rs2l) or ("not supported" in rs2l))
                        except Exception:
                            invalid_model2 = False
                        if invalid_model2 and use_task_model_flag and chat_fallback_model and (str(chat_fallback_model) != str(selected_model)):
                            try:
                                return self.ai_service.chat_with_ai(
                                    prompt_text,
                                    language,
                                    app_context=app_context,
                                    model=chat_fallback_model,
                                    progress_cb=_progress,
                                )
                            except Exception:
                                return rs2
                        return rs2
                def _done(result, error):
                    try:
                        if int(getattr(self, "_chat_active_req_id", 0)) != int(req_id):
                            return
                    except Exception:
                        pass
                    try:
                        self.chat_widget.clear_status(owner=req_id)
                    except Exception:
                        pass
                    try:
                        self.chat_widget.set_busy(True)
                    except Exception:
                        pass
                    if error:
                        self.chat_widget.set_ai_response(str(error))
                        try:
                            self.chat_widget.set_busy(False)
                        except Exception:
                            pass
                        try:
                            import psutil, os as _os
                            prev = getattr(self, "_chat_prev_priority", None)
                            if prev is not None:
                                psutil.Process(_os.getpid()).nice(prev)
                        except Exception:
                            pass
                    else:
                        try:
                            result_s = str(result or "")
                        except Exception:
                            result_s = ""

                        def _extract_tools(text: str):
                            cmds = []
                            keep_lines = []
                            try:
                                import re
                            except Exception:
                                re = None
                            def _first_balanced_json_chunk(s: str) -> str:
                                try:
                                    t = str(s or "")
                                except Exception:
                                    return ""
                                if not t:
                                    return ""
                                a = -1
                                for ch in ("{", "["):
                                    try:
                                        i0 = t.find(ch)
                                    except Exception:
                                        i0 = -1
                                    if i0 != -1 and (a == -1 or i0 < a):
                                        a = i0
                                if a == -1:
                                    return t.strip()
                                curly = 0
                                square = 0
                                in_str = False
                                esc = False
                                for i1 in range(a, len(t)):
                                    c = t[i1]
                                    if in_str:
                                        if esc:
                                            esc = False
                                            continue
                                        if c == "\\":
                                            esc = True
                                            continue
                                        if c == '"':
                                            in_str = False
                                        continue
                                    if c == '"':
                                        in_str = True
                                        continue
                                    if c == "{":
                                        curly += 1
                                    elif c == "}":
                                        curly -= 1
                                    elif c == "[":
                                        square += 1
                                    elif c == "]":
                                        square -= 1
                                    if (curly <= 0) and (square <= 0) and (i1 > a):
                                        return t[a : i1 + 1].strip()
                                return t[a:].strip()
                            tool_names = (
                                "CREATE_PROJECT",
                                "NEW_PROJECT",
                                "NEWPROJECT",
                                "PROJECT_CREATE",
                                "CREATE_NEW_PROJECT",
                                "OPEN_PROJECT",
                                "PROJECT_OPEN",
                                "ADD_QUESTION",
                                "UPDATE_QUESTION",
                                "DELETE_QUESTION",
                                "REMOVE_QUESTION",
                                "REMOVE_QUESTIONS",
                                "DEL_QUESTION",
                                "DEL_QUESTIONS",
                                "SET_QUESTION_IMAGE",
                                "SET_QUESTION_IMAGE_URL",
                                "SEARCH_IMAGE",
                                "UPDATE_GAME_CONFIG",
                                "READ_PROJECT_DETAILS",
                                "WEB_SEARCH",
                                "WEB_FETCH",
                            )
                            norm = str(text or "")
                            try:
                                if re is not None:
                                    pat = "|".join([str(x) for x in tool_names])
                                    norm = re.sub(r"\s+(?=(?:" + pat + r")\s*:)", "\n", norm, flags=re.IGNORECASE)
                                    norm = re.sub(r"\s+(?=\[(?:" + pat + r")\s*:)", "\n", norm, flags=re.IGNORECASE)
                            except Exception:
                                norm = str(text or "")
                            lines = str(norm or "").splitlines()
                            i = 0
                            while i < len(lines):
                                raw = lines[i]
                                s = str(raw or "").strip()
                                if s.startswith("[") and ("]" in s) and (":" in s[1:]):
                                    try:
                                        close_idx = s.find("]")
                                        inside = s[1:close_idx].strip()
                                        tail = s[close_idx + 1 :].strip()
                                        il = inside.upper()
                                        if il.startswith("TOOL_NAME:") or il.startswith("TOOL:"):
                                            cmd2 = inside.split(":", 1)[1].strip().upper()
                                            if cmd2:
                                                cmds.append((cmd2, tail))
                                                i += 1
                                                continue
                                    except Exception:
                                        pass
                                if s.startswith("[") and s.endswith("]") and (":" in s[1:]):
                                    try:
                                        colon = s.find(":")
                                        cmd = s[1:colon].strip().upper()
                                        rest = s[colon + 1 : -1].strip()
                                        if cmd:
                                            cmds.append((cmd, rest))
                                        i += 1
                                        continue
                                    except Exception:
                                        pass

                                cmd_name = ""
                                inline_rest = ""
                                if s and (":" not in s) and (not s.startswith("[")) and (str(s).strip().upper() in tool_names):
                                    cmd_name = str(s).strip().upper()
                                    inline_rest = ""
                                elif re is not None:
                                    m1 = re.match(r"^(?:[\-\*\u2022]?\s*)?([A-Za-z_][A-Za-z0-9_]*)\s*:\s*(.*)$", s)
                                    if m1:
                                        cmd_name = str(m1.group(1) or "").strip().upper()
                                        inline_rest = str(m1.group(2) or "").strip()
                                    else:
                                        m2 = re.match(r"^(?:[\-\*\u2022]?\s*)?([A-Za-z_][A-Za-z0-9_]*)\s+(.+)$", s)
                                        if m2:
                                            cmd_name = str(m2.group(1) or "").strip().upper()
                                            inline_rest = str(m2.group(2) or "").strip()

                                if cmd_name in tool_names:
                                    buf = []
                                    if inline_rest:
                                        buf.append(inline_rest)
                                    j = i + 1
                                    depth = 0
                                    started = bool(inline_rest)
                                    saw_json = False
                                    if inline_rest:
                                        depth += inline_rest.count("{") - inline_rest.count("}")
                                        depth += inline_rest.count("[") - inline_rest.count("]")
                                        if ("{" in inline_rest) or ("[" in inline_rest):
                                            saw_json = True
                                    if saw_json and (depth <= 0) and inline_rest:
                                        arg_one = " ".join(str(inline_rest or "").replace("\r", "").split())
                                        arg_one = _first_balanced_json_chunk(arg_one)
                                        if arg_one:
                                            cmds.append((cmd_name, arg_one))
                                            i += 1
                                            continue
                                    while j < len(lines):
                                        ln = str(lines[j] or "")
                                        t = ln.strip()
                                        if (not started) and (not t):
                                            j += 1
                                            continue
                                        if (not started) and t:
                                            started = True
                                        if started:
                                            buf.append(ln)
                                            depth += t.count("{") - t.count("}")
                                            depth += t.count("[") - t.count("]")
                                            if ("{" in t) or ("[" in t):
                                                saw_json = True
                                            if saw_json and (depth <= 0):
                                                break
                                        j += 1
                                    arg = "\n".join(buf).strip()
                                    if arg.startswith("```"):
                                        arg = arg.strip("`").strip()
                                    arg_one = " ".join(arg.replace("\r", "").split())
                                    arg_one = _first_balanced_json_chunk(arg_one)
                                    if arg_one:
                                        cmds.append((cmd_name, arg_one))
                                        i = j + 1
                                        continue

                                if (not keep_lines) and str(s).strip().lower() in ("assistant", "assistant:", "assistant："):
                                    i += 1
                                    continue
                                keep_lines.append(str(raw or ""))
                                i += 1
                            return cmds, "\n".join(keep_lines).strip()

                        def _process_text(text_in: str):
                            try:
                                from eduplay.core.i18n import I18n
                            except Exception:
                                I18n = None
                            info_cmds = {"READ_PROJECT_DETAILS", "WEB_SEARCH", "WEB_FETCH"}
                            state = {"tool_retry": 0}
                            try:
                                debug_ai = bool(os.environ.get("EDUPLAY_DEBUG_AI"))
                            except Exception:
                                debug_ai = False

                            def _detect_requested_question_count(user_text: str) -> int | None:
                                try:
                                    import re
                                    t = str(user_text or "")
                                except Exception:
                                    t = ""
                                if not t:
                                    return None
                                m = re.search(r"(?:tạo|tao|generate|create|soạn|soan|viết|viet)\s+(\d+)\s+(?:câu|cau|questions?)", t, flags=re.IGNORECASE)
                                if not m:
                                    m = re.search(r"(\d+)\s+(?:câu|cau|questions?)", t, flags=re.IGNORECASE)
                                if not m:
                                    return None
                                try:
                                    n = int(m.group(1))
                                except Exception:
                                    return None
                                if n <= 0 or n > 100:
                                    return None
                                return n

                            def _detect_intent(user_text: str):
                                try:
                                    ml = str(user_text or "").lower().strip()
                                except Exception:
                                    ml = ""
                                try:
                                    short = len(ml) <= 12
                                except Exception:
                                    short = False

                                greet_terms = ("xin chào", "chào", "hello", "hi", "hey")
                                is_greeting = short and any(t == ml or ml.startswith(t + " ") for t in greet_terms)

                                wants_open_project = ("mở dự án" in ml) or ("mở project" in ml) or ("open project" in ml)
                                open_ctx = (("mở" in ml) or ("open" in ml)) and (("dự án" in ml) or ("du an" in ml) or ("project" in ml))
                                create_strict = (
                                    ("tạo dự án" in ml)
                                    or ("tao du an" in ml)
                                    or ("tạo project" in ml)
                                    or ("create project" in ml)
                                )
                                create_loose = (("dự án mới" in ml) or ("du an moi" in ml) or ("new project" in ml))
                                wants_project_create = bool(create_strict or (create_loose and (not open_ctx)))
                                if wants_open_project and not create_strict:
                                    wants_project_create = False
                                wants_delete = any(k in ml for k in ("xóa", "xoá", "delete", "remove"))
                                wants_questions = any(
                                    k in ml
                                    for k in (
                                        "câu hỏi",
                                        "trắc nghiệm",
                                        "quiz",
                                        "câu trả lời",
                                        "tra loi",
                                        "đáp án",
                                        "dap an",
                                        "trả lời ngắn",
                                        "tra loi ngan",
                                        "tự luận",
                                        "tu luan",
                                        "đúng sai",
                                        "dung sai",
                                        "đúng/sai",
                                        "dung/sai",
                                        "đúng-sai",
                                        "dung-sai",
                                        "true false",
                                        "true/false",
                                        "true-false",
                                    )
                                )
                                wants_add_into_project = any(k in ml for k in ("thêm vào", "đưa vào", "add vào", "add to", "thêm thẳng", "import vào"))
                                wants_generate = any(k in ml for k in ("tạo", "generate", "trích xuất", "extract", "làm", "soạn", "thêm", "them", "bổ sung", "bo sung"))
                                wants_edit = any(k in ml for k in ("sửa", "sua", "chỉnh sửa", "chinh sua", "update", "edit", "chỉnh", "chinh")) and any(
                                    k in ml
                                    for k in (
                                        "câu hỏi",
                                        "cau hoi",
                                        "giải thích",
                                        "giai thich",
                                        "explanation",
                                        "phân tích",
                                        "phan tich",
                                        "đáp án",
                                        "dap an",
                                        "lựa chọn",
                                        "lua chon",
                                        "options",
                                        "option",
                                    )
                                )
                                wants_image = any(k in ml for k in ("ảnh", "anh", "hình", "hinh", "image", "picture")) and any(
                                    k in ml for k in ("chèn", "chen", "gắn", "gan", "đặt", "dat", "set", "thêm", "them", "add", "tải", "tai", "download")
                                )
                                wants_config = (
                                    any(k in ml for k in ("thời gian", "thoi gian", "time", "timer", "cấu hình", "cau hinh", "config",
                                                          "trộn", "tron", "xáo trộn", "xao tron", "shuffle", "randomize",
                                                          "ngẫu nhiên", "ngau nhien", "random"))
                                    and any(k in ml for k in ("đổi", "doi", "cập nhật", "cap nhat", "update", "set", "chỉnh", "chinh",
                                                              "bật", "bat", "tắt", "tat", "bật lên", "tắt đi", "on", "off",
                                                              "cho phép", "cho phep", "không cho", "ko cho", "disable", "enable"))
                                )
                                wants_questions = wants_questions and (wants_add_into_project or wants_generate or wants_edit)

                                wants_action = wants_questions or wants_edit or wants_delete or wants_project_create or wants_open_project
                                return {
                                    "ml": ml,
                                    "is_greeting": is_greeting,
                                    "wants_questions": wants_questions,
                                    "wants_edit": wants_edit,
                                    "wants_image": wants_image,
                                    "wants_config": wants_config,
                                    "wants_project_create": wants_project_create,
                                    "wants_open_project": wants_open_project,
                                    "wants_delete": wants_delete,
                                    "wants_action": wants_action,
                                }

                            def _filter_tools(matches: list, intent: dict) -> list:
                                filtered = []
                                try:
                                    proj = self.project_manager.get_current_project()
                                    has_project = bool(proj and isinstance(proj, dict) and str(proj.get("id") or "").strip())
                                except Exception:
                                    has_project = False

                                for cmd, args in matches:
                                    c = str(cmd or "").upper().strip()
                                    a = str(args or "").strip()
                                    al = a.lower()
                                    if c in ("NEW_PROJECT", "NEWPROJECT", "PROJECT_CREATE", "CREATE_NEW_PROJECT"):
                                        c = "CREATE_PROJECT"
                                    if c == "CREATE_PROJECT":
                                        if ("project name" in al) or ("game_type" in al) or ("mô tả" in al) or ("mo ta" in al):
                                            continue
                                        if intent.get("wants_project_create"):
                                            filtered.append((c, a))
                                            continue
                                        continue
                                    if c == "OPEN_PROJECT":
                                        if intent.get("wants_open_project"):
                                            filtered.append((c, a))
                                        continue
                                    if c == "PROJECT_OPEN":
                                        if intent.get("wants_open_project"):
                                            filtered.append(("OPEN_PROJECT", a))
                                        continue
                                    if c == "ADD_QUESTION":
                                        if intent.get("wants_questions"):
                                            filtered.append((c, a))
                                        continue
                                    if c == "UPDATE_QUESTION":
                                        if intent.get("wants_edit"):
                                            filtered.append((c, a))
                                        continue
                                    if c == "DELETE_QUESTION":
                                        if intent.get("wants_delete"):
                                            filtered.append((c, a))
                                        continue
                                    if c in ("SET_QUESTION_IMAGE", "SET_QUESTION_IMAGE_URL", "SEARCH_IMAGE"):
                                        if intent.get("wants_image") or intent.get("wants_questions") or intent.get("wants_edit"):
                                            filtered.append((c, a))
                                        continue
                                    if c == "UPDATE_GAME_CONFIG":
                                        if intent.get("wants_config"):
                                            filtered.append((c, a))
                                        continue
                                    if c in info_cmds:
                                        filtered.append((c, a))
                                        continue
                                return filtered

                            def _run_follow_up(prompt_text: str):
                                def _task2():
                                    try:
                                        suf = str((suffix_cache or {}).get("suffix") or "")
                                    except Exception:
                                        suf = ""
                                    if suf:
                                        prompt_text2 = str(prompt_text or "") + "\n\n" + suf
                                    else:
                                        prompt_text2 = str(prompt_text or "")
                                    try:
                                        model2 = str(chat_fallback_model or chat_model or "llama-3.1-8b-instant").strip() or None
                                    except Exception:
                                        model2 = None
                                    try:
                                        try:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                model=model2,
                                                progress_cb=_progress,
                                            )
                                        except TypeError:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                progress_cb=_progress,
                                            )
                                    except TypeError:
                                        try:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                model=model2,
                                                progress_cb=_progress,
                                            )
                                        except TypeError:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                progress_cb=_progress,
                                            )

                                def _done2(result2, error2):
                                    try:
                                        if debug_ai:
                                            try:
                                                r2 = str(result2 or "")
                                            except Exception:
                                                r2 = ""
                                            print(f"[AI-DEBUG] follow_up_done error={error2} len={len(r2)} preview={r2[:160]!r}")
                                    except Exception:
                                        pass
                                    if error2:
                                        self.chat_widget.set_ai_response(str(error2))
                                        try:
                                            self.chat_widget.set_busy(False)
                                        except Exception:
                                            pass
                                        try:
                                            import psutil, os as _os
                                            prev = getattr(self, "_chat_prev_priority", None)
                                            if prev is not None:
                                                psutil.Process(_os.getpid()).nice(prev)
                                        except Exception:
                                            pass
                                        return
                                    try:
                                        nxt = str(result2 or "")
                                    except Exception:
                                        nxt = ""
                                    _handle_text(nxt)

                                try:
                                    self._run_background_task(_task2, _done2, None, "default")
                                except Exception as e:
                                    try:
                                        self.chat_widget.set_ai_response(str(e))
                                    except Exception:
                                        pass
                                    try:
                                        self.chat_widget.set_busy(False)
                                    except Exception:
                                        pass
                                    try:
                                        import psutil, os as _os
                                        prev = getattr(self, "_chat_prev_priority", None)
                                        if prev is not None:
                                            psutil.Process(_os.getpid()).nice(prev)
                                    except Exception:
                                        pass

                            def _run_follow_up_chat(prompt_text: str):
                                def _task2():
                                    try:
                                        suf = str((suffix_cache or {}).get("suffix") or "")
                                    except Exception:
                                        suf = ""
                                    if suf:
                                        prompt_text2 = str(prompt_text or "") + "\n\n" + suf
                                    else:
                                        prompt_text2 = str(prompt_text or "")
                                    try:
                                        model2 = str(chat_fallback_model or chat_model or "llama-3.1-8b-instant").strip() or None
                                    except Exception:
                                        model2 = None
                                    try:
                                        try:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                model=model2,
                                                progress_cb=_progress,
                                            )
                                        except TypeError:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                progress_cb=_progress,
                                            )
                                    except TypeError:
                                        try:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                model=model2,
                                                progress_cb=_progress,
                                            )
                                        except TypeError:
                                            return self.ai_service.chat_with_ai(
                                                prompt_text2,
                                                language,
                                                app_context=app_context,
                                                progress_cb=_progress,
                                            )

                                def _done2(result2, error2):
                                    if error2:
                                        try:
                                            self.chat_widget.set_ai_response(str(error2))
                                        except Exception:
                                            pass
                                    else:
                                        try:
                                            txt = str(result2 or "").strip()
                                            if not txt:
                                                try:
                                                    from eduplay.core.i18n import I18n
                                                    txt = I18n.t("ai.done_short", language)
                                                except Exception:
                                                    txt = "Đã xử lý xong." if language == "vi" else "Done."
                                            self.chat_widget.set_ai_response(txt)
                                        except Exception:
                                            pass
                                    try:
                                        self.chat_widget.set_busy(False)
                                    except Exception:
                                        pass
                                    try:
                                        import psutil, os as _os
                                        prev = getattr(self, "_chat_prev_priority", None)
                                        if prev is not None:
                                            psutil.Process(_os.getpid()).nice(prev)
                                    except Exception:
                                        pass

                                try:
                                    self._run_background_task(_task2, _done2, None, "default")
                                except Exception as e:
                                    try:
                                        self.chat_widget.set_ai_response(str(e))
                                    except Exception:
                                        pass
                                    try:
                                        self.chat_widget.set_busy(False)
                                    except Exception:
                                        pass
                                    try:
                                        import psutil, os as _os
                                        prev = getattr(self, "_chat_prev_priority", None)
                                        if prev is not None:
                                            psutil.Process(_os.getpid()).nice(prev)
                                    except Exception:
                                        pass

                            def _handle_text(cur_text: str):
                                try:
                                    import re

                                    t0 = str(cur_text or "")
                                    t0 = re.sub(r"<think>[\s\S]*?</think>", "", t0, flags=re.IGNORECASE).strip()
                                    cur_text = t0
                                except Exception:
                                    pass
                                intent = _detect_intent(msg)
                                requested_count = _detect_requested_question_count(msg) if intent.get("wants_questions") else None
                                matches, display_text = _extract_tools(cur_text)
                                matches = _filter_tools(matches, intent)
                                try:
                                    if debug_ai:
                                        mnames = [str(c or "") for c, _a in (matches or [])][:8]
                                        print(f"[AI-DEBUG] handle_text len={len(str(cur_text or ''))} matches={len(matches or [])} retry={int(state.get('tool_retry', 0) or 0)} cmds={mnames} display_len={len(str(display_text or ''))}")
                                except Exception:
                                    pass
                                try:
                                    self.chat_widget.set_busy(True)
                                except Exception:
                                    pass

                                def _looks_like_provider_error(text: str) -> bool:
                                    try:
                                        t = str(text or "").strip()
                                    except Exception:
                                        return False
                                    if not t:
                                        return False
                                    tl = t.lower()
                                    bad_terms = (
                                        "ai service not ready",
                                        "thiếu groq_api_key",
                                        "groq api key",
                                        "không kết nối được groq",
                                        "lỗi ssl",
                                        "lỗi proxy",
                                        "timeout",
                                        "error communicating with ai",
                                        "giới hạn tốc độ (429",
                                    )
                                    return any(x in tl for x in bad_terms)

                                def _looks_like_clarification(text: str) -> bool:
                                    try:
                                        import re
                                        t = str(text or "").strip()
                                    except Exception:
                                        return False
                                    if not t:
                                        return False
                                    if len(t) > 400:
                                        return False
                                    if "?" not in t:
                                        return False
                                    pat = r"(bạn\s+(?:có\s+thể|vui\s+lòng|cho\s+mình|cho\s+tôi|nói\s+rõ|xác\s+nhận)|mình\s+(?:cần|chưa\s+rõ|không\s+rõ)|vui\s+lòng\s+cho\s+biết|bạn\s+muốn\s+|bạn\s+cần\s+)"
                                    try:
                                        return re.search(pat, t, flags=re.IGNORECASE) is not None
                                    except Exception:
                                        return False

                                def _finish():
                                    try:
                                        self.chat_widget.set_busy(False)
                                    except Exception:
                                        pass
                                    try:
                                        import psutil, os as _os
                                        prev = getattr(self, "_chat_prev_priority", None)
                                        if prev is not None:
                                            psutil.Process(_os.getpid()).nice(prev)
                                    except Exception:
                                        pass

                                if intent.get("is_greeting"):
                                    display = str(display_text or "").strip()
                                    if not display:
                                        display = str(cur_text or "").strip()
                                    if display:
                                        self.chat_widget.set_ai_response(display)
                                    else:
                                        self.chat_widget.set_ai_response("Xin chào!" if language == "vi" else "Hello!")
                                    _finish()
                                    return

                                if _looks_like_provider_error(cur_text):
                                    self.chat_widget.set_ai_response(str(cur_text or "").strip())
                                    _finish()
                                    return

                                try:
                                    sent = str(cur_text or "").strip().lower()
                                except Exception:
                                    sent = ""
                                if sent in ("ko có trong danh sách lệnh", "không có trong danh sách lệnh"):
                                    if language == "vi":
                                        self.chat_widget.set_ai_response("Hiện tại phần mềm chưa hỗ trợ công cụ này vì EduPlay - Studio đang trong giai đoạn phát triển. Xin thứ lỗi vì sự bất tiện này, hãy quay lại sớm vì EduPlay - Studio sẽ phát triển liên tục!!!")
                                    else:
                                        self.chat_widget.set_ai_response("This tool is not supported yet because EduPlay Studio is still under development. Sorry for the inconvenience—please check back soon as EduPlay Studio is improving continuously.")
                                    _finish()
                                    return

                                if not matches:
                                    not_supported_vi = "Hiện tại phần mềm chưa hỗ trợ công cụ này vì EduPlay - Studio đang trong giai đoạn phát triển. Xin thứ lỗi vì sự bất tiện này, hãy quay lại sớm vì EduPlay - Studio sẽ phát triển liên tục!!!"
                                    not_supported_en = "Our app is still under development and doesn't support all requests yet. Sorry—please check back soon as we're improving it."
                                    if intent.get("wants_action") and _looks_like_clarification(cur_text):
                                        try:
                                            if debug_ai:
                                                print("[AI-DEBUG] branch=clarification")
                                        except Exception:
                                            pass
                                        self.chat_widget.set_ai_response(str(cur_text or "").strip())
                                        _finish()
                                        return
                                    if intent.get("wants_action") and int(state.get("tool_retry", 0) or 0) < 3:
                                        state["tool_retry"] = int(state.get("tool_retry", 0) or 0) + 1
                                        try:
                                            if debug_ai:
                                                print(f"[AI-DEBUG] branch=force_tools retry={int(state.get('tool_retry', 0) or 0)}")
                                        except Exception:
                                            pass
                                        lang_rule = "TIẾNG VIỆT" if language == "vi" else "ENGLISH"
                                        req_cnt_line = ""
                                        try:
                                            if intent.get("wants_questions") and requested_count:
                                                req_cnt_line = f"- Người dùng yêu cầu {int(requested_count)} câu hỏi: phải xuất đúng {int(requested_count)} dòng ADD_QUESTION (không thừa/không thiếu).\n"
                                        except Exception:
                                            req_cnt_line = ""
                                        qa_type_line = ""
                                        knowledge_line = ""
                                        try:
                                            mlx = str(intent.get("ml") or "")
                                            if "trả lời ngắn" in mlx or "tra loi ngan" in mlx or "tự luận" in mlx or "tu luan" in mlx:
                                                qa_type_line = "- Ưu tiên tạo câu hỏi type short_answer.\n"
                                        except Exception:
                                            qa_type_line = ""
                                        try:
                                            if intent.get("wants_questions") and requested_count and int(requested_count) >= 4:
                                                knowledge_line = "- Trong bộ câu hỏi, thêm 1–2 câu kiểm tra KIẾN THỨC đúng bài; còn lại là luyện tập.\n"
                                        except Exception:
                                            knowledge_line = ""
                                        prompt_force = (
                                            "SYSTEM:\n"
                                            "Tôi đang chạy trong EduPlay Studio.\n"
                                            "Tôi chỉ được phép xuất các lệnh tool sau (đúng y hệt tên lệnh):\n"
                                            "- CREATE_PROJECT, OPEN_PROJECT, ADD_QUESTION, UPDATE_QUESTION, DELETE_QUESTION,\n"
                                            "- SET_QUESTION_IMAGE, SET_QUESTION_IMAGE_URL, SEARCH_IMAGE,\n"
                                            "- UPDATE_GAME_CONFIG, READ_PROJECT_DETAILS, WEB_SEARCH, WEB_FETCH.\n"
                                            "\n"
                                            "QUY TẮC BẮT BUỘC:\n"
                                            "- Nếu yêu cầu của người dùng KHÔNG ánh xạ được sang một lệnh trong danh sách trên: chỉ trả về đúng 1 dòng: ko có trong danh sách lệnh\n"
                                            "- Nếu yêu cầu ánh xạ được: chỉ xuất các dòng tool hợp lệ (mỗi lệnh 1 dòng). Không xuất chữ giải thích.\n"
                                            "- Không tự ý thêm câu hỏi nếu người dùng không yêu cầu tạo/thêm câu hỏi.\n"
                                            + req_cnt_line +
                                            f"- Nếu tạo câu hỏi: nội dung phải bằng {lang_rule}.\n"
                                            + qa_type_line +
                                            "\n"
                                            "AI_OUTPUT_TEXT:\n"
                                            + str(cur_text or "").strip()
                                            + "\n\n"
                                            "User request:\n"
                                            + str(msg or "")
                                        )
                                        _run_follow_up(prompt_force)
                                        return
                                    display = str(display_text or "").strip()
                                    if display:
                                        self.chat_widget.set_ai_response(display)
                                        _finish()
                                        return
                                    if intent.get("wants_action"):
                                        if intent.get("wants_questions"):
                                            if language == "vi":
                                                self.chat_widget.set_ai_response(
                                                    "Tôi chưa đủ thông tin để tạo câu hỏi đúng. Bạn cho tôi biết: (1) lớp mấy, (2) bài/tiết hoặc nội dung đã dạy, (3) muốn bao nhiêu câu và dạng nào (trắc nghiệm/trả lời ngắn)?"
                                                )
                                            else:
                                                self.chat_widget.set_ai_response(
                                                    "I need a bit more info to generate the questions correctly: (1) grade, (2) lesson/topic taught, (3) how many questions and which format?"
                                                )
                                        else:
                                            self.chat_widget.set_ai_response(not_supported_vi if language == "vi" else not_supported_en)
                                        _finish()
                                        return
                                    if language == 'vi':
                                        self.chat_widget.set_ai_response("AI không trả lời (kết quả rỗng). Bạn thử lại hoặc kiểm tra model.")
                                    else:
                                        self.chat_widget.set_ai_response("AI returned an empty response. Please retry or check your model.")
                                    _finish()
                                    return

                                try:
                                    self._show_loading(loading_msg, "default")
                                except Exception:
                                    pass

                                results = []
                                try:
                                    try:
                                        if I18n is not None:
                                            self.chat_widget.set_status(I18n.t("ai.executing_tools", language), owner=req_id)
                                        else:
                                            raise Exception("no i18n")
                                    except Exception:
                                        self.chat_widget.set_status("Đang thực thi lệnh trong EduPlay..." if language == "vi" else "Executing commands...", owner=req_id)
                                    try:
                                        pending_ids = []
                                        delete_pending_ids = []
                                        for c, _a in (matches or []):
                                            if str(c or "").strip().upper() in ("UPDATE_QUESTION", "DELETE_QUESTION"):
                                                tqs = (app_context or {}).get("target_questions") or []
                                                if isinstance(tqs, list):
                                                    for t in tqs:
                                                        if not isinstance(t, dict):
                                                            continue
                                                        qid = str(t.get("question_id") or "").strip()
                                                        if qid:
                                                            pending_ids.append(qid)
                                                if not pending_ids:
                                                    try:
                                                        prev_ids = getattr(self, "_last_target_question_ids", []) or []
                                                        if isinstance(prev_ids, list):
                                                            pending_ids = [str(x or "").strip() for x in prev_ids if str(x or "").strip()]
                                                    except Exception:
                                                        pending_ids = []
                                                if not pending_ids:
                                                    try:
                                                        raw_user = str(getattr(self, "_last_chat_user_msg", "") or "")
                                                    except Exception:
                                                        raw_user = ""
                                                    try:
                                                        proj = self.project_manager.get_current_project() or {}
                                                        if not proj:
                                                            for _nm2 in ("editor_quiz_screen", "editor_fishing_screen", "editor_screen"):
                                                                try:
                                                                    _scr2 = getattr(self, _nm2, None)
                                                                except Exception:
                                                                    _scr2 = None
                                                                if not _scr2:
                                                                    continue
                                                                try:
                                                                    proj = _scr2.get_current_project() if hasattr(_scr2, "get_current_project") else getattr(_scr2, "current_project", None)
                                                                except Exception:
                                                                    proj = {}
                                                                if isinstance(proj, dict) and str(proj.get("id") or "").strip():
                                                                    break
                                                    except Exception:
                                                        proj = {}
                                                    qs = proj.get("questions", []) if isinstance(proj, dict) else []
                                                    if qs and isinstance(qs, list):
                                                        try:
                                                            from eduplay.core.ai_tool_payloads import extract_question_numbers
                                                        except Exception:
                                                            extract_question_numbers = None
                                                        nums = extract_question_numbers(raw_user) if extract_question_numbers else []
                                                        for n in nums[:12]:
                                                            idx0 = int(n) - 1
                                                            if 0 <= idx0 < len(qs):
                                                                cand = qs[idx0]
                                                                if isinstance(cand, dict):
                                                                    qid = str(cand.get("id") or "").strip()
                                                                    if qid:
                                                                        pending_ids.append(qid)
                                                break
                                        # --- Bulk delete: detect "xoá N câu cuối" ---
                                        try:
                                            import re as _re2
                                            _num_words_del = {
                                                "một": 1, "mot": 1, "hai": 2, "ba": 3,
                                                "bốn": 4, "bon": 4, "tư": 4, "tu": 4,
                                                "năm": 5, "nam": 5, "sáu": 6, "sau": 6,
                                                "bảy": 7, "bay": 7, "tám": 8, "tam": 8,
                                                "chín": 9, "chin": 9, "mười": 10, "muoi": 10,
                                            }
                                            _raw_del = str(getattr(self, "_last_chat_user_msg", "") or "").lower()
                                            _m_del = _re2.search(
                                                r"(?:xóa|xoa|delete|remove)\s+(\d+|[a-zàáạảãâầấậẩẫăằắặẳẵđèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ]+)\s+(?:câu|cau)[^\n\r]*\s+cuối",
                                                _raw_del, flags=_re2.IGNORECASE
                                            )
                                            if not _m_del:
                                                _m_del = _re2.search(
                                                    r"(?:xóa|xoa|delete|remove)\s+(\d+|[a-zàáạảãâầấậẩẫăằắặẳẵđèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ]+)\s+(?:câu|cau)\s+(?:hỏi\s+)?cuối",
                                                    _raw_del, flags=_re2.IGNORECASE
                                                )
                                            if _m_del:
                                                _g = str(_m_del.group(1) or "").strip()
                                                _del_n = int(_g) if _g.isdigit() else int(_num_words_del.get(_g, 0))
                                                if _del_n >= 1:
                                                    _del_proj = self.project_manager.get_current_project() or {}
                                                    _del_qs = _del_proj.get("questions", []) if isinstance(_del_proj, dict) else []
                                                    if isinstance(_del_qs, list) and _del_qs:
                                                        _take = min(_del_n, len(_del_qs))
                                                        # Take from the end (last N questions)
                                                        for _dq in _del_qs[-_take:]:
                                                            if isinstance(_dq, dict):
                                                                _dqid = str(_dq.get("id") or "").strip()
                                                                if _dqid:
                                                                    delete_pending_ids.append(_dqid)
                                        except Exception:
                                            delete_pending_ids = []
                                        self._pending_update_question_ids = list(pending_ids)
                                        self._pending_delete_question_ids = list(delete_pending_ids) if delete_pending_ids else list(pending_ids)
                                    except Exception:
                                        self._pending_update_question_ids = []
                                        self._pending_delete_question_ids = []
                                    for cmd, args in matches:
                                        results.append(self._execute_ai_tool(cmd, args.strip()))
                                finally:
                                    try:
                                        self._hide_loading()
                                    except Exception:
                                        pass
                                    try:
                                        self.chat_widget.clear_status(owner=req_id)
                                    except Exception:
                                        pass

                                ok_all = True
                                follow_up_chunks = []
                                action_only = True
                                executed_cmds = []
                                add_ok_count = 0
                                for r in results:
                                    try:
                                        executed_cmds.append(str((r or {}).get("cmd") or "").upper())
                                        if not (r or {}).get("ok"):
                                            ok_all = False
                                        if str((r or {}).get("cmd") or "").upper() == "ADD_QUESTION" and (r or {}).get("ok"):
                                            add_ok_count += 1
                                        if (r or {}).get("follow_up"):
                                            follow_up_chunks.append(str((r or {}).get("follow_up") or ""))
                                        if str((r or {}).get("cmd") or "").upper() in info_cmds:
                                            action_only = False
                                    except Exception:
                                        pass

                                if follow_up_chunks:
                                    merged = "\n\n".join([c for c in follow_up_chunks if c.strip()])
                                    if merged.strip():
                                        try:
                                            m2 = str(merged or "")
                                        except Exception:
                                            m2 = ""
                                        m2l = m2.lower()
                                        if ("tool_error" in m2l) or ("invalid" in m2l) or ("args must be" in m2l) or ("must be valid" in m2l):
                                            prompt2 = (
                                                "SYSTEM_TOOL_OUTPUT:\n"
                                                + m2
                                                + "\n\n"
                                                + "User request:\n"
                                                + str(msg or "")
                                            )
                                            _run_follow_up(prompt2)
                                            return

                                        prompt2 = (
                                            "SYSTEM:\n"
                                            "Tôi là Edubot trong EduPlay Studio.\n"
                                            "Hãy phản hồi người dùng như chat bình thường sau khi hệ thống đã thực hiện xong yêu cầu.\n"
                                            "- Xưng 'Tôi' và gọi người dùng là 'bạn'.\n"
                                            "- Nêu rõ bạn đã làm gì và kết quả (đã mở gì, đã thêm bao nhiêu câu, ảnh đã lưu vào thư mục media...).\n"
                                            "- Không xuất bất kỳ lệnh tool nào.\n"
                                            "- Không dùng từ 'tài liệu'.\n"
                                            "\n"
                                            "SYSTEM_TOOL_OUTPUT:\n"
                                            + merged
                                            + "\n\n"
                                            + "User request:\n"
                                            + str(msg or "")
                                        )
                                        _run_follow_up_chat(prompt2)
                                        return
                                    _finish()
                                    return

                                if ok_all and action_only:
                                    try:
                                        if intent.get("wants_questions") and requested_count and add_ok_count < requested_count:
                                            need = int(requested_count - add_ok_count)
                                            proj = None
                                            try:
                                                proj = self.project_manager.get_current_project()
                                            except Exception:
                                                proj = None
                                            existing_titles = []
                                            try:
                                                qs = (proj or {}).get("questions", [])
                                                if isinstance(qs, list) and qs:
                                                    for q in qs[-min(12, len(qs)):]:
                                                        try:
                                                            qt = str((q or {}).get("question") or "").strip()
                                                            if qt:
                                                                existing_titles.append(qt[:140])
                                                        except Exception:
                                                            continue
                                            except Exception:
                                                existing_titles = []
                                            lang_rule = "TIẾNG VIỆT" if language == "vi" else "ENGLISH"
                                            prompt_more = (
                                                "SYSTEM:\n"
                                                "Tôi đang chạy trong EduPlay Studio.\n"
                                                "Tôi là TRỢ LÝ GIÁO VIÊN. Mọi câu hỏi tạo ra là bài luyện tập cho học sinh.\n"
                                                "Hãy tạo thêm câu hỏi để đủ đúng số lượng người dùng yêu cầu.\n"
                                                "YÊU CẦU:\n"
                                                "- Chỉ xuất các dòng tool ADD_QUESTION.\n"
                                                f"- Nội dung câu hỏi/giải thích phải bằng {lang_rule}.\n"
                                                f"- Tạo đúng {need} câu nữa (không thừa/không thiếu).\n"
                                                "- Chỉ dùng nội dung trong file đính kèm; không bịa.\n"
                                                "- Không trùng ý/câu với danh sách đã có.\n"
                                                "- Trong bộ câu hỏi, có thể thêm 1–2 câu kiểm tra KIẾN THỨC đúng bài; còn lại là luyện tập.\n"
                                                "- Không xuất thêm bất kỳ chữ nào ngoài các dòng ADD_QUESTION.\n"
                                            )
                                            if focus_title:
                                                prompt_more += f"- Chỉ tạo câu hỏi trong phạm vi: {focus_title}.\n"
                                            if existing_titles:
                                                prompt_more += "\nCÂU ĐÃ CÓ (tránh trùng):\n" + "\n".join(f"- {t}" for t in existing_titles) + "\n"
                                            prompt_more += "\nUser request:\n" + str(msg or "")
                                            _run_follow_up(prompt_more)
                                            return
                                        if intent.get("wants_questions") and ("ADD_QUESTION" not in executed_cmds) and (not state.get("asked_to_add")):
                                            state["asked_to_add"] = True
                                            proj = None
                                            try:
                                                proj = self.project_manager.get_current_project()
                                            except Exception:
                                                proj = None
                                            try:
                                                pid = str((proj or {}).get("id") or "")
                                                pname = str((proj or {}).get("name") or "")
                                            except Exception:
                                                pid = ""
                                                pname = ""
                                            lang_rule = "TIẾNG VIỆT" if language == "vi" else "ENGLISH"
                                            req_cnt_line = ""
                                            try:
                                                if intent.get("wants_questions") and requested_count:
                                                    req_cnt_line = f"Tạo đúng {int(requested_count)} câu hỏi (không thừa/không thiếu).\n"
                                            except Exception:
                                                req_cnt_line = ""
                                            prompt3 = (
                                                "SYSTEM:\n"
                                                "Tôi đang chạy trong EduPlay Studio.\n"
                                                "Tôi là TRỢ LÝ GIÁO VIÊN. Mọi câu hỏi tạo ra là bài luyện tập cho học sinh.\n"
                                                f"Project hiện tại đã sẵn sàng: name={pname}.\n"
                                                "Bây giờ hãy tạo câu hỏi theo yêu cầu người dùng và thêm thẳng vào project.\n"
                                                "Chỉ xuất các dòng tool, ưu tiên ADD_QUESTION (có thể nhiều dòng). Không xuất thêm chữ giải thích.\n"
                                                + req_cnt_line +
                                                f"Nội dung phải bằng {lang_rule}.\n"
                                                "Chỉ tạo câu hỏi dựa trên PHẦN GIÁO VIÊN ĐÃ DẠY trong nội dung file đính kèm. Không lan man, không nâng cao.\n"
                                                "Trong bộ câu hỏi, có thể thêm 1–2 câu kiểm tra KIẾN THỨC đúng bài; còn lại là luyện tập.\n"
                                                "\n"
                                                "User request:\n"
                                                + str(msg or "")
                                            )
                                            _run_follow_up(prompt3)
                                            return
                                    except Exception:
                                        pass
                                    try:
                                        import json
                                        proj = None
                                        try:
                                            proj = self.project_manager.get_current_project()
                                        except Exception:
                                            proj = None
                                        summary = {
                                            "executed_commands": executed_cmds,
                                            "added_questions": int(add_ok_count),
                                            "project_id": (proj or {}).get("id") if isinstance(proj, dict) else None,
                                            "project_name": (proj or {}).get("name") if isinstance(proj, dict) else None,
                                        }
                                        prompt_ok = (
                                            "SYSTEM:\n"
                                            "Tôi là Edubot trong EduPlay Studio.\n"
                                            "Hãy phản hồi người dùng như chat bình thường sau khi hệ thống đã thực hiện xong yêu cầu.\n"
                                            "- Xưng 'Tôi' và gọi người dùng là 'bạn'.\n"
                                            "- Nêu rõ bạn đã làm gì và kết quả.\n"
                                            "- Không xuất bất kỳ lệnh tool nào.\n"
                                            "- Không dùng từ 'tài liệu'.\n"
                                            "\n"
                                            "SYSTEM_TOOL_OUTPUT:\n"
                                            + json.dumps(summary, ensure_ascii=False)
                                            + "\n\n"
                                            + "User request:\n"
                                            + str(msg or "")
                                        )
                                        _run_follow_up_chat(prompt_ok)
                                        return
                                    except Exception:
                                        try:
                                            from eduplay.core.i18n import I18n
                                            self.chat_widget.set_ai_response(I18n.t("ai.done_short", language))
                                        except Exception:
                                            self.chat_widget.set_ai_response("Đã xử lý xong." if language == "vi" else "Done.")
                                        _finish()
                                        return

                                if ok_all:
                                    try:
                                        if debug_ai:
                                            print("[AI-DEBUG] branch=tools_ok_all")
                                    except Exception:
                                        pass
                                    display = str(display_text or "").strip()
                                    if display:
                                        self.chat_widget.set_ai_response(display)
                                    else:
                                        try:
                                            from eduplay.core.i18n import I18n
                                            self.chat_widget.set_ai_response(I18n.t("ai.done_short", language))
                                        except Exception:
                                            self.chat_widget.set_ai_response("OK" if language != "vi" else "Đã xử lý xong.")
                                    _finish()
                                    return

                                bad = []
                                for r in results:
                                    try:
                                        if not (r or {}).get("ok"):
                                            bad.append(str((r or {}).get("cmd") or ""))
                                    except Exception:
                                        pass
                                if bad:
                                    try:
                                        if debug_ai:
                                            print(f"[AI-DEBUG] branch=tools_failed bad={bad}")
                                    except Exception:
                                        pass
                                    try:
                                        bad_u = [str(x or "").upper().strip() for x in bad if str(x or "").strip()]
                                    except Exception:
                                        bad_u = []
                                    if bad_u and all(x == "OPEN_PROJECT" for x in bad_u):
                                        _finish()
                                        return
                                    if language == "vi":
                                        self.chat_widget.set_ai_response("Tôi chưa thực hiện được yêu cầu vừa rồi. Bạn thử nói rõ hơn (ví dụ: tên dự án chính xác / câu số mấy) hoặc thử lại nhé.")
                                    else:
                                        self.chat_widget.set_ai_response("I couldn't complete the last request. Please clarify (exact project name / which question number) and try again.")
                                _finish()
                                return

                            _handle_text(str(text_in or ""))

                        _process_text(result_s)
                            
                try:
                    _progress("Đang xử lý...")
                except Exception:
                    pass
                self._run_background_task(_task, _done, None, "default")

            def _run_translate(text: str, target: str):
                def _task():
                    return self.ai_service.translate_content(text, target)
                def _done(result, error):
                    if error:
                        self.chat_widget.set_ai_response(str(error))
                    else:
                        self.chat_widget.set_ai_response(result)
                self._run_background_task(_task, _done, None, "default")

            # Check for hardcoded shortcuts first
            if lower.startswith('tạo 5 câu hỏi mức dễ'):
                lang_rule = "TIẾNG VIỆT" if language == "vi" else "ENGLISH"
                prompt = (
                    f"Tạo đúng 5 câu hỏi mức Đơn giản cho {game_type} về chủ đề '{topic}'.\n"
                    "YÊU CẦU:\n"
                    "- Chỉ xuất đúng 5 dòng tool ADD_QUESTION.\n"
                    "- Mỗi dòng là: ADD_QUESTION: {JSON_1_DÒNG}\n"
                    "- Bộ 5 câu phải đủ 5 loại (mỗi loại đúng 1 câu): multiple_choice, true_false, fill_blank, matching, short_answer.\n"
                    f"- Nội dung câu hỏi/giải thích phải bằng {lang_rule}.\n"
                    "- Với multiple_choice: options phải là mảng, correct_answer là chỉ số (0-based).\n"
                    "- Với true_false: correct_answer là true/false.\n"
                    "- Với fill_blank/short_answer: answers là mảng string.\n"
                    "- Với matching: pairs là mảng {left,right}.\n"
                    "- Không xuất thêm bất kỳ chữ nào ngoài các dòng ADD_QUESTION.\n"
                )
                _run_ai(prompt, attachments_for_ai)
                return
            if lower.startswith('tạo 5 câu hỏi mức trung'):
                lang_rule = "TIẾNG VIỆT" if language == "vi" else "ENGLISH"
                prompt = (
                    f"Tạo đúng 5 câu hỏi mức Trung bình cho {game_type} về chủ đề '{topic}'.\n"
                    "YÊU CẦU:\n"
                    "- Chỉ xuất đúng 5 dòng tool ADD_QUESTION.\n"
                    "- Mỗi dòng là: ADD_QUESTION: {JSON_1_DÒNG}\n"
                    "- Bộ 5 câu phải đủ 5 loại (mỗi loại đúng 1 câu): multiple_choice, true_false, fill_blank, matching, short_answer.\n"
                    f"- Nội dung câu hỏi/giải thích phải bằng {lang_rule}.\n"
                    "- Explanation phải rõ ràng, có bước/luận giải ngắn gọn.\n"
                    "- Không xuất thêm bất kỳ chữ nào ngoài các dòng ADD_QUESTION.\n"
                )
                _run_ai(prompt, attachments_for_ai)
                return
            if lower.startswith('tạo 5 câu hỏi mức khó'):
                lang_rule = "TIẾNG VIỆT" if language == "vi" else "ENGLISH"
                prompt = (
                    f"Tạo đúng 5 câu hỏi mức Khó cho {game_type} về chủ đề '{topic}', có bẫy hợp lý và so sánh các phương án.\n"
                    "YÊU CẦU:\n"
                    "- Chỉ xuất đúng 5 dòng tool ADD_QUESTION.\n"
                    "- Mỗi dòng là: ADD_QUESTION: {JSON_1_DÒNG}\n"
                    "- Bộ 5 câu phải đủ 5 loại (mỗi loại đúng 1 câu): multiple_choice, true_false, fill_blank, matching, short_answer.\n"
                    f"- Nội dung câu hỏi/giải thích phải bằng {lang_rule}.\n"
                    "- Explanation phải chỉ ra vì sao các đáp án sai dễ nhầm.\n"
                    "- Không xuất thêm bất kỳ chữ nào ngoài các dòng ADD_QUESTION.\n"
                )
                _run_ai(prompt, attachments_for_ai)
                return
            if lower.startswith('tổng hợp đáp án'):
                try:
                    qs = (current_project or {}).get('questions', [])
                    sample = qs[:10]
                    import json
                    prompt = "Tổng hợp và phân tích đáp án của bộ câu hỏi sau, nêu điểm mạnh/yếu và gợi ý cải thiện.\n\n" + json.dumps(sample, ensure_ascii=False, indent=2)
                    _run_ai(prompt, attachments_for_ai)
                    return
                except Exception as e:
                    self.chat_widget.set_ai_response(str(e))
                    return

            # Translation handling
            target_map = {
                'english': 'en', 'en': 'en',
                'vietnamese': 'vi', 'vi': 'vi', 'tiếng việt': 'vi',
                'french': 'fr', 'fr': 'fr', 'tiếng pháp': 'fr',
                'spanish': 'es', 'es': 'es', 'tiếng tây ban nha': 'es',
                'german': 'de', 'de': 'de', 'tiếng đức': 'de',
                'chinese': 'zh', 'zh': 'zh', 'tiếng trung': 'zh',
                'japanese': 'ja', 'ja': 'ja', 'tiếng nhật': 'ja',
                'korean': 'ko', 'ko': 'ko', 'tiếng hàn': 'ko',
                'russian': 'ru', 'ru': 'ru', 'tiếng nga': 'ru',
                'italian': 'it', 'it': 'it', 'tiếng ý': 'it',
                'portuguese': 'pt', 'pt': 'pt', 'tiếng bồ đào nha': 'pt'
            }
            translated = False
            import re
            m = re.match(r"translate\s+(to\s+)?([a-zA-Z\s]+):\s*(.+)", lower)
            mv = re.match(r"dịch\s+(sang\s+)?([a-zA-Z\s]+):\s*(.+)", lower)
            if m or mv:
                grp = m or mv
                lang_name = grp.group(2).strip()
                text = msg[msg.find(':')+1:].strip()
                target = target_map.get(lang_name, 'en')
                _run_translate(text, target)
                translated = True

            if not translated:
                try:
                    import re
                    simple = str(msg or "").strip()
                    if simple and ("\n" not in simple) and (len(simple) <= 140) and (not attachments_for_ai):
                        def _clean_name(s: str) -> str:
                            try:
                                t = str(s or "").strip().strip('"').strip("'").strip()
                            except Exception:
                                t = ""
                            if not t:
                                return ""
                            try:
                                t = t[:80].strip()
                            except Exception:
                                pass
                            return t

                        def _norm_name(s: str) -> str:
                            try:
                                import unicodedata
                                t = unicodedata.normalize("NFD", str(s or "").lower())
                                t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
                            except Exception:
                                t = str(s or "").lower()
                            try:
                                t = re.sub(r"[^a-z0-9\s]+", " ", t)
                                t = re.sub(r"\s+", " ", t).strip()
                            except Exception:
                                t = " ".join(t.split())
                            return t

                        m_create = re.match(
                            r"^(?:tạo|tao|create)\s+(?:một\s+)?(?:dự\s*án|du\s*an|project)(?:\s+mới|\s+new)?(?:\s+tên\s*(?:là)?|\s+ten\s*(?:la)?|\s+named)?\s+(.+)$",
                            simple,
                            flags=re.IGNORECASE,
                        )
                        m_open = re.match(
                            r"^(?:mở|mo|open)\s+(?:dự\s*án|du\s*an|project)\s+(.+)$",
                            simple,
                            flags=re.IGNORECASE,
                        )
                        m_count = re.search(r"\b(?:bao\s+nhieu|m[ấa]y)\s+(?:dự\s*án|du\s*an|projects?)\b", simple, flags=re.IGNORECASE)
                        m_count2 = re.search(r"\b(?:danh\s+s[aá]ch|list)\s+(?:dự\s*án|du\s*an|projects?)\b", simple, flags=re.IGNORECASE)
                        if m_count or m_count2:
                            try:
                                projects = self.project_manager.get_all_projects() or []
                            except Exception:
                                projects = []
                            names = []
                            for p in projects:
                                if not isinstance(p, dict):
                                    continue
                                try:
                                    nm = str(p.get("name") or "").strip()
                                except Exception:
                                    nm = ""
                                if nm:
                                    names.append(nm)
                            try:
                                names = names[:12]
                            except Exception:
                                pass
                            if language == "vi":
                                if names:
                                    text = f"Tôi thấy bạn đang có {len(projects)} dự án:\n- " + "\n- ".join(names)
                                else:
                                    text = "Tôi thấy bạn chưa có dự án nào."
                            else:
                                if names:
                                    text = f"You currently have {len(projects)} projects:\n- " + "\n- ".join(names)
                                else:
                                    text = "You don't have any projects yet."
                            try:
                                self.chat_widget.set_busy(False)
                            except Exception:
                                pass
                            try:
                                self.chat_widget.add_message("System", text, "ai")
                            except Exception:
                                pass
                            return
                        if m_create:
                            name = _clean_name(m_create.group(1))
                            if name:
                                ml = simple.lower()
                                gtype = "quiz_classic"
                                if any(x in ml for x in ("bắt cá", "bat ca", "câu cá", "cau ca", "fishing", "fish")):
                                    gtype = "fishing"
                                if any(x in ml for x in ("triệu phú", "trieu phu", "millionaire")):
                                    gtype = "quiz_millionaire"
                                new_proj = self.project_manager.create_project(name, description="", game_type=gtype)
                                try:
                                    if self.settings_manager:
                                        self.settings_manager.add_recent_project(new_proj.get("id", ""), new_proj.get("name", ""))
                                except Exception:
                                    pass
                                self._open_editor_for_project(new_proj)
                                try:
                                    self.chat_widget.set_busy(False)
                                except Exception:
                                    pass
                                try:
                                    self.chat_widget.add_message("System", str(new_proj.get("name") or name), "ai")
                                except Exception:
                                    pass
                                return
                        if m_open:
                            target = _clean_name(m_open.group(1))
                            if target:
                                projects = self.project_manager.get_all_projects()
                                best = None
                                best_score = 0.0
                                nt = _norm_name(target)
                                try:
                                    from difflib import SequenceMatcher
                                except Exception:
                                    SequenceMatcher = None
                                for p in projects or []:
                                    try:
                                        nm = str((p or {}).get("name") or "").strip()
                                    except Exception:
                                        nm = ""
                                    if not nm:
                                        continue
                                    nn = _norm_name(nm)
                                    if not nn:
                                        continue
                                    try:
                                        r0 = float(SequenceMatcher(None, nt, nn).ratio()) if (SequenceMatcher is not None and nt) else 0.0
                                    except Exception:
                                        r0 = 0.0
                                    try:
                                        if nt and (nt in nn or nn in nt):
                                            r0 = min(1.0, r0 + 0.12)
                                    except Exception:
                                        pass
                                    if r0 > best_score:
                                        best_score = r0
                                        best = p
                                if best and best_score >= 0.75:
                                    proj_data = self.project_manager.load_project(best.get("id"))
                                    if proj_data:
                                        self._open_editor_for_project(proj_data)
                                        try:
                                            self.chat_widget.set_busy(False)
                                        except Exception:
                                            pass
                                        try:
                                            self.chat_widget.add_message("System", str(proj_data.get("name") or ""), "ai")
                                        except Exception:
                                            pass
                                        return
                except Exception:
                    pass
                prompt = msg
                if focus_title and attachments_for_ai:
                    if language == "vi":
                        prompt = prompt + f"\n\nChỉ dùng nội dung đúng phần có tiêu đề: {focus_title}."
                    else:
                        prompt = prompt + f"\n\nUse only the section titled: {focus_title}."
                try:
                    if wants_edit_request:
                        prompt = (
                            str(prompt or "").rstrip()
                            + "\n\n"
                            + (
                                "YÊU CẦU SỬA:\n"
                                "- Chỉ dùng UPDATE_QUESTION.\n"
                                "- Mỗi câu cần sửa: 1 dòng UPDATE_QUESTION.\n"
                                "- BẮT BUỘC có question_id đúng theo target_questions trong ngữ cảnh.\n"
                                "- Output chỉ gồm các dòng bắt đầu bằng 'UPDATE_QUESTION: ' (không có dấu [], không có gạch đầu dòng, không có markdown).\n"
                                "- Args BẮT BUỘC là JSON hợp lệ 1 dòng, dùng dấu nháy kép \". Không dùng nháy đơn '. Không xuống dòng trong JSON.\n"
                                "- Schema chuẩn: {\"question_id\":\"<id>\",\"patch\":{\"explanation\":\"<giải thích mới>\"}}\n"
                                "- Không dùng các key khác ngoài question_id và patch.explanation.\n"
                                "- Chỉ cập nhật explanation (không đổi question/options/correct_answer).\n"
                                "- Không xuất chữ giải thích ngoài các dòng tool.\n"
                                "- Nếu thiếu target_questions hoặc không chắc question_id: KHÔNG xuất UPDATE_QUESTION; hãy hỏi lại người dùng để xác nhận số câu cần sửa.\n"
                            )
                        )
                except Exception:
                    pass
                try:
                    if wants_delete_request:
                        # Build question list context for AI to understand positions
                        _del_ctx_lines = ""
                        try:
                            _del_proj = self.project_manager.get_current_project() or {}
                            _del_qs = _del_proj.get("questions", []) if isinstance(_del_proj, dict) else []
                            if isinstance(_del_qs, list) and _del_qs:
                                _del_total = len(_del_qs)
                                _del_ctx_lines = f"Danh sách câu hỏi hiện tại ({_del_total} câu):\n"
                                for _di, _dq in enumerate(_del_qs):
                                    _dqt = str((_dq or {}).get("question") or "")[:60] if isinstance(_dq, dict) else ""
                                    _del_ctx_lines += f"- Câu {_di+1}: {_dqt}\n"
                        except Exception:
                            _del_ctx_lines = ""
                        prompt = (
                            str(prompt or "").rstrip()
                            + "\n\n"
                            + (
                                "YÊU CẦU XÓA:\n"
                                "- Chỉ dùng DELETE_QUESTION.\n"
                                "- Phân tích ý nghĩa: 'xoá N câu cuối' = xoá N câu ở CUỐI danh sách. 'xoá câu thứ N' = xoá đúng câu số N.\n"
                                "- Mỗi câu cần xóa: 1 dòng DELETE_QUESTION với {\"question_number\": <số thứ tự 1-based>}.\n"
                                "- Nếu 'xoá N câu cuối': output N dòng DELETE_QUESTION với question_number của N câu cuối (câu cuối nhất xoá trước).\n"
                                "- Output chỉ gồm các dòng DELETE_QUESTION. Không có chữ giải thích.\n"
                                "- Args là JSON 1 dòng, dùng dấu nháy kép.\n"
                                + (_del_ctx_lines)
                            )
                        )
                except Exception:
                    pass
                _run_ai(prompt, attachments_for_ai)

        except Exception as e:
            self.chat_widget.set_ai_response(str(e))

    def _execute_ai_tool(self, cmd: str, args: str):
        out = {"cmd": str(cmd or "").upper(), "ok": False, "follow_up": ""}
        try:
            cmd = str(cmd or "").strip().upper()
            if cmd == "PROJECT_OPEN":
                cmd = "OPEN_PROJECT"
            if cmd in ("REMOVE_QUESTION", "REMOVE_QUESTIONS", "DEL_QUESTION", "DEL_QUESTIONS"):
                cmd = "DELETE_QUESTION"
            out["cmd"] = cmd

            def _parse_json_lenient(s: str):
                import json
                import ast
                try:
                    raw = str(s or "").strip()
                except Exception:
                    raw = ""
                if raw.startswith("```"):
                    raw = raw.strip("`").strip()
                candidates = []
                if raw:
                    candidates.append(raw)
                try:
                    i = raw.find("{")
                    j = raw.rfind("}")
                    if i != -1 and j != -1 and j > i:
                        candidates.append(raw[i : j + 1])
                except Exception:
                    pass
                try:
                    i = raw.find("[")
                    j = raw.rfind("]")
                    if i != -1 and j != -1 and j > i:
                        candidates.append(raw[i : j + 1])
                except Exception:
                    pass
                try:
                    raw2 = raw.replace("\r", "").replace("\n", " ").strip()
                    if raw2 and raw2 not in candidates:
                        candidates.append(raw2)
                    try:
                        i = raw2.find("{")
                        j = raw2.rfind("}")
                        if i != -1 and j != -1 and j > i:
                            cand = raw2[i : j + 1]
                            if cand not in candidates:
                                candidates.append(cand)
                    except Exception:
                        pass
                    try:
                        i = raw2.find("[")
                        j = raw2.rfind("]")
                        if i != -1 and j != -1 and j > i:
                            cand = raw2[i : j + 1]
                            if cand not in candidates:
                                candidates.append(cand)
                    except Exception:
                        pass
                except Exception:
                    pass
                for cand in candidates:
                    try:
                        return json.loads(cand)
                    except Exception:
                        pass
                    try:
                        obj = ast.literal_eval(cand)
                        if isinstance(obj, (dict, list)):
                            return obj
                    except Exception:
                        pass
                raise ValueError("Invalid JSON payload")

            def _get_lang():
                try:
                    return self.settings_manager.get_language()
                except Exception:
                    return "en"

            def _t(key: str, fallback: str):
                try:
                    from eduplay.core.i18n import I18n
                    val = I18n.t(key, _get_lang())
                    if not val or val == key:
                        return fallback
                    return val
                except Exception:
                    return fallback

            def _safe_media_target() -> str:
                proj = self.project_manager.get_current_project()
                if not proj:
                    raise ValueError("No project open")
                pid = str(proj.get("id") or "").strip()
                if not pid:
                    raise ValueError("Invalid project id")
                media_dir = self.project_manager.projects_dir / pid / "media"
                try:
                    media_dir.mkdir(parents=True, exist_ok=True)
                except Exception:
                    pass
                return str(media_dir)

            if cmd == "OPEN_PROJECT":
                target = str(args or "").strip()
                try:
                    cand_obj = _parse_json_lenient(target)
                except Exception:
                    cand_obj = None
                if isinstance(cand_obj, dict):
                    try:
                        target = str(
                            cand_obj.get("project_name")
                            or cand_obj.get("projectName")
                            or cand_obj.get("name")
                            or cand_obj.get("id")
                            or target
                        ).strip()
                    except Exception:
                        target = str(args or "").strip()
                if not target:
                    self.chat_widget.add_message(
                        "System",
                        ("Bạn muốn mở dự án nào? Ví dụ: “mở dự án Bạo Lực Học Đường”." if _get_lang() == "vi" else "Which project do you want to open? Example: “open project Bạo Lực Học Đường”."),
                        False,
                    )
                    return out
                projects = self.project_manager.get_all_projects()
                found = None
                def _safe_strip_lower(v: str) -> str:
                    try:
                        return str(v or "").strip().lower()
                    except Exception:
                        return ""
                for p in projects:
                    try:
                        pid = str(p.get("id") or "").strip()
                        pname = str(p.get("name") or "").strip()
                        if pid and pid == target:
                            found = p
                            break
                        if pname and _safe_strip_lower(pname) == _safe_strip_lower(target):
                            found = p
                            break
                    except Exception:
                        continue
                if not found:
                    try:
                        import unicodedata
                        import re
                        from difflib import SequenceMatcher
                    except Exception:
                        unicodedata = None
                        re = None
                        SequenceMatcher = None

                    def _norm_name(s: str) -> str:
                        try:
                            t = str(s or "").strip().lower()
                        except Exception:
                            t = ""
                        if not t:
                            return ""
                        try:
                            if unicodedata is not None:
                                t = unicodedata.normalize("NFD", t)
                                t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
                        except Exception:
                            pass
                        try:
                            if re is not None:
                                t = re.sub(r"[^a-z0-9\s]+", " ", t)
                                t = re.sub(r"\s+", " ", t).strip()
                        except Exception:
                            t = " ".join(t.split())
                        return t

                    nt = _norm_name(target)
                    if nt:
                        for p in projects:
                            try:
                                nm = str(p.get("name") or "")
                            except Exception:
                                nm = ""
                            if not nm:
                                continue
                            if _norm_name(nm) == nt:
                                found = p
                                break
                    if found:
                        proj_data = self.project_manager.load_project(found.get('id'))
                        if not proj_data:
                            self.chat_widget.add_message("System", _t("ai.tool.project_load_failed", f"Failed to load project: {found.get('name')}"), False)
                            return out
                        self._open_editor_for_project(proj_data)
                        out["ok"] = True
                        try:
                            import json
                            out["follow_up"] = json.dumps(
                                {
                                    "opened_project_id": proj_data.get("id"),
                                    "opened_project_name": proj_data.get("name"),
                                    "question_count": len(proj_data.get("questions", []) or []),
                                },
                                ensure_ascii=False,
                            )
                        except Exception:
                            pass
                        return out
                    scored = []
                    for p in projects:
                        try:
                            nm = str(p.get("name") or "")
                        except Exception:
                            nm = ""
                        if not nm:
                            continue
                        nn = _norm_name(nm)
                        if not nn:
                            continue
                        try:
                            if SequenceMatcher is not None and nt:
                                r0 = float(SequenceMatcher(None, nt, nn).ratio())
                            else:
                                r0 = 0.0
                        except Exception:
                            r0 = 0.0
                        try:
                            if nt and (nt in nn or nn in nt):
                                r0 = min(1.0, r0 + 0.12)
                        except Exception:
                            pass
                        scored.append((r0, p))
                    scored.sort(key=lambda x: float(x[0] or 0.0), reverse=True)

                    best = scored[0][1] if scored else None
                    best_score = float(scored[0][0]) if scored else 0.0
                    if best and best_score >= 0.87:
                        proj_data = self.project_manager.load_project(best.get("id"))
                        if proj_data:
                            self._open_editor_for_project(proj_data)
                            out["ok"] = True
                            try:
                                bname = str(best.get("name") or "").strip()
                            except Exception:
                                bname = ""
                            try:
                                same_name = bool(_norm_name(bname) and _norm_name(target) and (_norm_name(bname) == _norm_name(target)))
                            except Exception:
                                same_name = False
                            if _get_lang() == "vi":
                                if same_name or best_score >= 0.97:
                                    self.chat_widget.add_message("System", f"{bname}", False)
                                else:
                                    self.chat_widget.add_message("System", f"{bname}", False)
                            else:
                                if same_name or best_score >= 0.97:
                                    self.chat_widget.add_message("System", f"{bname}", False)
                                else:
                                    self.chat_widget.add_message("System", f"{bname}", False)
                            return out

                    sugg = []
                    for r0, p in scored[:3]:
                        try:
                            nm = str(p.get("name") or "").strip()
                        except Exception:
                            nm = ""
                        if nm:
                            sugg.append(nm)
                    if _get_lang() == "vi":
                        if sugg:
                            self.chat_widget.add_message(
                                "System",
                                "Không tìm thấy đúng tên dự án. Có phải bạn muốn mở một trong các dự án này không?\n- " + "\n- ".join(sugg),
                                False,
                            )
                        else:
                            self.chat_widget.add_message("System", "Không tìm thấy dự án bạn yêu cầu. Bạn bấm “Duyệt dự án” để chọn đúng tên nhé.", False)
                    else:
                        if sugg:
                            self.chat_widget.add_message(
                                "System",
                                "Project not found. Did you mean one of these?\n- " + "\n- ".join(sugg),
                                False,
                            )
                        else:
                            self.chat_widget.add_message("System", "Project not found. Please use “Browse projects” to pick the correct one.", False)
                    return out
                proj_data = self.project_manager.load_project(found.get('id'))
                if not proj_data:
                    self.chat_widget.add_message("System", _t("ai.tool.project_load_failed", f"Failed to load project: {found.get('name')}"), False)
                    return out
                self._open_editor_for_project(proj_data)
                out["ok"] = True
                try:
                    import json
                    out["follow_up"] = json.dumps(
                        {
                            "opened_project_id": proj_data.get("id"),
                            "opened_project_name": proj_data.get("name"),
                            "question_count": len(proj_data.get("questions", []) or []),
                        },
                        ensure_ascii=False,
                    )
                except Exception:
                    pass
                return out

            if cmd == "CREATE_PROJECT":
                payload = None
                try:
                    payload = _parse_json_lenient(args)
                except Exception:
                    payload = None
                name = ""
                gtype = ""
                desc = ""
                if isinstance(payload, dict):
                    try:
                        name = str(payload.get("project_name") or payload.get("name") or payload.get("title") or "").strip()
                    except Exception:
                        name = ""
                    try:
                        gtype = str(payload.get("game_type") or payload.get("gameType") or payload.get("type") or "").strip()
                    except Exception:
                        gtype = ""
                    try:
                        desc = str(payload.get("description") or payload.get("desc") or payload.get("note") or "").strip()
                    except Exception:
                        desc = ""
                if not name:
                    parts = str(args or "").split('|')
                    if len(parts) < 2:
                        self.chat_widget.add_message("System", _t("ai.tool.create_project_format", "Invalid CREATE_PROJECT format. Use: name|type|desc"), False)
                        return out
                    name = parts[0].strip()
                    gtype = parts[1].strip()
                    desc = parts[2].strip() if len(parts) > 2 else ""
                # Clean up name: strip filler phrases so only the actual project title remains
                try:
                    import re as _re
                    _filler = (
                        r"^(?:tạo|tao|tôi sẽ tạo|sẽ tạo|hãy tạo|create|tên là|tên|ten la|ten|"
                        r"dự án mới|du an moi|dự án|du an|project|new project|mới|new)\s+"
                    )
                    name = _re.sub(_filler, "", name, flags=_re.IGNORECASE).strip().strip('"').strip("'")
                    # Remove trailing descriptive suffixes like "về ...", "thuộc loại ...", "loại game ..."
                    name = _re.sub(
                        r"\s+(?:về|ve|cho|thuộc|thuoc|loại|loai|game|dành cho|danh cho|với|voi)\s+.*$",
                        "", name, flags=_re.IGNORECASE
                    ).strip()
                except Exception:
                    pass
                try:
                    desc = desc.strip().strip('"').strip("'").strip()
                    while desc.endswith(")") and desc.count("(") < desc.count(")"):
                        desc = desc[:-1].rstrip()
                    while desc.startswith("(") and desc.count("(") > desc.count(")"):
                        desc = desc[1:].lstrip()
                except Exception:
                    pass
                if not gtype:
                    gtype = "quiz_classic"
                new_proj = self.project_manager.create_project(name, description=desc, game_type=gtype)
                self._open_editor_for_project(new_proj)
                out["ok"] = True
                try:
                    import json
                    out["follow_up"] = json.dumps(
                        {
                            "created_project_id": new_proj.get("id"),
                            "created_project_name": new_proj.get("name"),
                            "game_type": new_proj.get("game_type"),
                        },
                        ensure_ascii=False,
                    )
                except Exception:
                    pass
                return out

            if cmd == "ADD_QUESTION":
                try:
                    if not self.project_manager.get_current_project():
                        import re

                        raw_user = str(getattr(self, "_last_chat_user_msg", "") or "")
                        ml = raw_user.lower()
                        m = re.search(
                            r"(?:dự án|du an|project)(?:\s+mới|\s+new)?(?:\s+tên\s*là|\s+ten\s*la)?\s+([^\n\r,.;]+)",
                            raw_user,
                            flags=re.IGNORECASE,
                        )
                        target_name = ""
                        if m:
                            target_name = str(m.group(1) or "").strip().strip('"').strip("'")
                        if target_name:
                            try:
                                target_name = target_name[:80].strip()
                            except Exception:
                                pass
                        if target_name:
                            projects = self.project_manager.get_all_projects()
                            found = None
                            for p in projects:
                                try:
                                    if str(p.get("name") or "").strip().lower() == target_name.lower():
                                        found = p
                                        break
                                except Exception:
                                    continue
                            if found:
                                proj_data = self.project_manager.load_project(found.get("id"))
                                if proj_data:
                                    self._open_editor_for_project(proj_data)
                            else:
                                create_requested = any(
                                    x in ml
                                    for x in (
                                        "tạo dự án",
                                        "tao du an",
                                        "create project",
                                        "dự án mới",
                                        "du an moi",
                                        "new project",
                                    )
                                )
                                if create_requested:
                                    try:
                                        gtype = "quiz_classic"
                                        if any(x in ml for x in ("bắt cá", "bat ca", "câu cá", "cau ca", "fishing", "fish")):
                                            gtype = "fishing"
                                        if any(x in ml for x in ("triệu phú", "trieu phu", "millionaire")):
                                            gtype = "quiz_millionaire"
                                        new_proj = self.project_manager.create_project(target_name, description="", game_type=gtype)
                                        try:
                                            if self.settings_manager:
                                                self.settings_manager.add_recent_project(new_proj.get("id", ""), new_proj.get("name", ""))
                                        except Exception:
                                            pass
                                        self._open_editor_for_project(new_proj)
                                    except Exception:
                                        pass
                                else:
                                    self.chat_widget.add_message(
                                        "System",
                                        (
                                            f"Không tìm thấy dự án “{target_name}”. Nếu muốn tạo mới, hãy gõ: “tạo dự án {target_name}”. Nếu muốn mở dự án khác, hãy gõ: “mở dự án <tên>”."
                                            if _get_lang() == "vi"
                                            else f"Project “{target_name}” not found. To create it, reply: “create project {target_name}”. To open another project, reply: “open project <name>”."
                                        ),
                                        False,
                                    )
                                    return out

                    if not self.project_manager.get_current_project():
                        self.chat_widget.add_message(
                            "System",
                            ("Bạn muốn thêm vào dự án nào? Ví dụ: 'thêm vào dự án Thử Thách'." if _get_lang() == "vi" else "Which project should I add to? Example: 'add to project Thử Thách'."),
                            False,
                        )
                        return out
                except Exception:
                    pass

                try:
                    payload = _parse_json_lenient(args)
                except Exception:
                    out["follow_up"] = (
                        "TOOL_ERROR: ADD_QUESTION args must be valid one-line JSON.\n"
                        "Please output only tool lines, one per line.\n"
                        "Example:\n"
                        "[ADD_QUESTION:{\"type\":\"multiple_choice\",\"question\":\"...\",\"options\":[\"A\",\"B\",\"C\",\"D\"],\"correct_answer\":0,\"explanation\":\"...\"}]"
                    )
                    return out
                items = []
                if isinstance(payload, list):
                    items = payload
                elif isinstance(payload, dict) and isinstance(payload.get("questions"), list):
                    items = payload.get("questions") or []
                elif isinstance(payload, dict):
                    items = [payload]
                else:
                    self.chat_widget.add_message("System", _t("ai.tool.invalid_question_payload", "Invalid question payload."), False)
                    return out

                def _is_image_question(text: str) -> bool:
                    try:
                        t = str(text or "").strip().lower()
                    except Exception:
                        t = ""
                    if not t:
                        return False
                    try:
                        import re
                        if re.search(r"\bhình\s*\d+\b", t, flags=re.IGNORECASE):
                            return True
                    except Exception:
                        pass
                    triggers = (
                        "hình ảnh",
                        "bức ảnh",
                        "trong ảnh",
                        "trong hình",
                        "ảnh dưới",
                        "hình dưới",
                        "xem ảnh",
                        "xem hình",
                        "quan sát ảnh",
                        "quan sát hình",
                        "nhìn vào ảnh",
                        "nhìn vào hình",
                        "in the image",
                        "in the picture",
                        "shown in the image",
                        "shown in the picture",
                        "look at the image",
                        "look at the picture",
                    )
                    return any(x in t for x in triggers)

                def _prepare_local_image_path(image_value: str) -> str:
                    try:
                        s = str(image_value or "").strip()
                    except Exception:
                        s = ""
                    if not s:
                        return ""
                    sl = s.lower()
                    if sl.startswith("data:"):
                        return ""
                    try:
                        ai_settings = self.settings_manager.get_ai_settings() or {}
                        allow_net = bool(ai_settings.get("allow_internet", False))
                        allow_img = bool(ai_settings.get("allow_image_download", False))
                    except Exception:
                        allow_net = False
                        allow_img = False

                    from pathlib import Path
                    import shutil
                    import uuid
                    import re

                    media_dir = _safe_media_target()
                    media_root = Path(media_dir)

                    def _unique_dest(stem: str, suffix: str) -> Path:
                        base = re.sub(r"[^a-zA-Z0-9]+", "_", str(stem or "").strip())[:40].strip("_") or "img"
                        suf = str(suffix or "").strip()
                        if not suf.startswith("."):
                            suf = "." + suf if suf else ".png"
                        candidate = media_root / f"{base}{suf}"
                        idx = 1
                        while candidate.exists():
                            candidate = media_root / f"{base}_{idx}{suf}"
                            idx += 1
                        return candidate

                    if sl.startswith("http://") or sl.startswith("https://"):
                        if not (allow_net and allow_img):
                            return ""
                        try:
                            from urllib.parse import urlparse
                            import requests
                            try:
                                host = str(urlparse(s).netloc or "").strip().lower()
                            except Exception:
                                host = ""
                            blocked = ("via.placeholder.com", "placeholder.com", "placehold.co", "picsum.photos")
                            if host and any(b in host for b in blocked):
                                return ""
                            resp = requests.get(s, timeout=8)
                            if not (200 <= resp.status_code < 300):
                                return ""
                            ctype = str(resp.headers.get("content-type") or "").lower()
                            if "image/" not in ctype:
                                return ""
                            if "png" in ctype:
                                ext = ".png"
                            elif "jpeg" in ctype or "jpg" in ctype:
                                ext = ".jpg"
                            elif "gif" in ctype:
                                ext = ".gif"
                            elif "webp" in ctype:
                                ext = ".webp"
                            else:
                                ext = ".png"
                            dest = _unique_dest("img_" + uuid.uuid4().hex[:10], ext)
                            dest.write_bytes(resp.content)
                            return str(dest)
                        except Exception:
                            return ""

                    p = Path(s)
                    if not p.is_absolute():
                        pp = s.replace("\\", "/").strip()
                        if pp.startswith("media/"):
                            pp = pp.split("/", 1)[1]
                        p = media_root / pp
                    if (not p.exists()) or (not p.is_file()):
                        return ""
                    try:
                        p_res = p.resolve()
                        media_res = media_root.resolve()
                        in_media = (p_res == media_res) or (media_res in p_res.parents)
                    except Exception:
                        in_media = False
                    if in_media:
                        return str(p)
                    dest = _unique_dest(p.stem, p.suffix or ".png")
                    try:
                        shutil.copy2(str(p), str(dest))
                        return str(dest)
                    except Exception:
                        return ""

                def _normalize_question(q: dict) -> dict | None:
                    if not isinstance(q, dict):
                        return None
                    q2 = dict(q)
                    try:
                        t_raw = str(q2.get("type") or "").strip().lower()
                    except Exception:
                        t_raw = ""
                    aliases = {
                        "mcq": "multiple_choice",
                        "multiple-choice": "multiple_choice",
                        "multiple choice": "multiple_choice",
                        "trac_nghiem": "multiple_choice",
                        "trắc nghiệm": "multiple_choice",
                        "truefalse": "true_false",
                        "true/false": "true_false",
                        "true-false": "true_false",
                        "dung_sai": "true_false",
                        "đúng sai": "true_false",
                        "đúng/sai": "true_false",
                        "fillblank": "fill_blank",
                        "fill-blank": "fill_blank",
                        "fill blank": "fill_blank",
                        "dien_khuyet": "fill_blank",
                        "điền khuyết": "fill_blank",
                        "matching": "matching",
                        "noi": "matching",
                        "nối": "matching",
                        "ghep": "matching",
                        "ghép": "matching",
                        "short": "short_answer",
                        "short-answer": "short_answer",
                        "short answer": "short_answer",
                        "tra_loi_ngan": "short_answer",
                        "trả lời ngắn": "short_answer",
                        "tu_luan": "short_answer",
                        "tự luận": "short_answer",
                    }
                    qtype = aliases.get(t_raw, t_raw) if t_raw else "multiple_choice"
                    if qtype not in ("multiple_choice", "true_false", "fill_blank", "matching", "short_answer"):
                        qtype = "multiple_choice"
                    q2["type"] = qtype

                    if not str(q2.get("question") or "").strip():
                        for k in ("prompt", "content", "text", "title", "cau_hoi", "cauhoi", "question_text"):
                            v = q2.get(k)
                            if v is not None and str(v).strip():
                                q2["question"] = str(v).strip()
                                break
                    if not str(q2.get("question") or "").strip():
                        return None

                    if q2.get("explanation") is None:
                        q2["explanation"] = ""
                    try:
                        exp = str(q2.get("explanation") or "").strip()
                    except Exception:
                        exp = ""
                    if exp:
                        try:
                            import re
                            exp2 = exp
                            exp2 = re.sub(r"^\s*(?:theo\s+tài\s+liệu|theo\s+tai\s+lieu)\s*[:,\-]?\s*", "", exp2, flags=re.IGNORECASE)
                            exp2 = re.sub(r"^\s*(?:dựa\s+trên|dua\s+tren|dựa\s+vào|dua\s+vao|căn\s+cứ|can\s+cu)\s+(?:tài\s+liệu|tai\s+lieu|nội\s+dung|noi\s+dung)[^,:;\-]{0,60}\s*[:,\-]\s*", "", exp2, flags=re.IGNORECASE)
                            exp2 = re.sub(r"^\s*(?:theo\s+nội\s+dung|theo\s+noi\s+dung)\s*[:,\-]?\s*", "", exp2, flags=re.IGNORECASE)
                            exp2 = re.sub(r"^\s*(?:tài\s+liệu|tai\s+lieu)\s+(?:nêu\s+rõ|noi\s+ro|cho\s+biết|chi\s+ra|chỉ\s+ra|đề\s+cập|de\s+cap)\s*(?:rằng|la)?\s*[:,\-]?\s*", "", exp2, flags=re.IGNORECASE)
                            exp2 = re.sub(r"\b(?:trong\s+)?(?:tài\s+liệu|tai\s+lieu)\b", "", exp2, flags=re.IGNORECASE)
                            exp2 = re.sub(r"\s{2,}", " ", exp2, flags=re.IGNORECASE).strip()
                            q2["explanation"] = exp2.strip()
                        except Exception:
                            q2["explanation"] = exp

                    def _coerce_str_list(v) -> list[str]:
                        if v is None:
                            return []
                        if isinstance(v, list):
                            out0 = []
                            for x in v:
                                try:
                                    s0 = str(x or "").strip()
                                except Exception:
                                    s0 = ""
                                if s0:
                                    out0.append(s0)
                            return out0
                        try:
                            s = str(v or "").strip()
                        except Exception:
                            s = ""
                        if not s:
                            return []
                        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
                        if lines:
                            return lines
                        return [s] if s else []

                    def _coerce_bool(v) -> bool | None:
                        if isinstance(v, bool):
                            return bool(v)
                        if isinstance(v, int):
                            return bool(int(v))
                        try:
                            s = str(v or "").strip().lower()
                        except Exception:
                            s = ""
                        if not s:
                            return None
                        if s in ("true", "t", "yes", "y", "1", "đúng", "dung"):
                            return True
                        if s in ("false", "f", "no", "n", "0", "sai"):
                            return False
                        return None

                    if qtype == "multiple_choice":
                        extracted_opts = []
                        try:
                            import re
                            qtxt = str(q2.get("question") or "")
                            lines = [ln.strip() for ln in qtxt.splitlines() if ln.strip()]
                            opt_re = re.compile(r"^\s*([A-F])[\)\.\:\-]\s*(.+?)\s*$", flags=re.IGNORECASE)
                            opt_hits = []
                            keep = []
                            for ln in lines:
                                m = opt_re.match(ln)
                                if m:
                                    opt_hits.append((m.group(1).upper(), m.group(2).strip()))
                                else:
                                    keep.append(ln)
                            if len(opt_hits) >= 2:
                                opt_hits.sort(key=lambda x: x[0])
                                extracted_opts = [t for _, t in opt_hits if t]
                                q2["question"] = "\n".join(keep).strip() if keep else q2.get("question")
                        except Exception:
                            extracted_opts = []

                        opts = q2.get("options")
                        if opts is None:
                            for k in ("choices", "dap_an", "lua_chon", "choices_text"):
                                if k in q2:
                                    opts = q2.get(k)
                                    break
                        opts_list: list[str] = []
                        correct_idx: int | None = None
                        if isinstance(opts, dict):
                            try:
                                ordered = []
                                for kk in ["A", "B", "C", "D", "E", "F"]:
                                    if kk in opts:
                                        ordered.append(opts.get(kk))
                                if not ordered:
                                    ordered = list(opts.values())
                                opts_list = [str(x).strip() for x in ordered if str(x).strip()]
                            except Exception:
                                opts_list = []
                        elif isinstance(opts, str):
                            opts_list = [s.strip() for s in str(opts).splitlines() if s.strip()]
                            if opts_list and all(len(s) >= 2 and s[0].upper() in "ABCDEF" and s[1] in ").:" for s in opts_list[: min(6, len(opts_list))]):
                                opts_list = [s[2:].strip() for s in opts_list]
                        elif isinstance(opts, list):
                            if opts and isinstance(opts[0], dict):
                                tmp = []
                                for it in opts:
                                    try:
                                        tmp.append(str(it.get("text", "")).strip())
                                        if bool(it.get("correct")):
                                            correct_idx = len(tmp) - 1
                                    except Exception:
                                        continue
                                opts_list = [x for x in tmp if x]
                            else:
                                opts_list = [str(x).strip() for x in opts if str(x).strip()]

                        if extracted_opts:
                            if (not opts_list) or all(str(x).strip().upper() in ("A", "B", "C", "D", "E", "F") for x in (opts_list or [])):
                                opts_list = extracted_opts

                        correct_raw = q2.get("correct_answer")
                        if correct_raw is None:
                            for k in ("correct", "answer", "correct_index", "correctIndex", "correct_option"):
                                if k in q2:
                                    correct_raw = q2.get(k)
                                    break
                        if correct_idx is None:
                            if isinstance(correct_raw, int):
                                correct_idx = int(correct_raw)
                            elif correct_raw is not None:
                                s = str(correct_raw).strip()
                                su = s.upper()
                                if len(su) == 1 and su in "ABCDEF":
                                    correct_idx = ord(su) - ord("A")
                                elif s.isdigit():
                                    try:
                                        correct_idx = int(s)
                                    except Exception:
                                        correct_idx = 0
                                else:
                                    try:
                                        for i, o in enumerate(opts_list):
                                            if str(o).strip().lower() == s.strip().lower():
                                                correct_idx = i
                                                break
                                    except Exception:
                                        pass
                        if correct_idx is None:
                            correct_idx = 0
                        if correct_idx < 0:
                            correct_idx = 0
                        if opts_list and correct_idx >= len(opts_list):
                            correct_idx = max(0, len(opts_list) - 1)
                        if len(opts_list) < 2:
                            return None
                        q2["options"] = opts_list
                        q2["correct_answer"] = int(correct_idx)
                        try:
                            q2.pop("answers", None)
                        except Exception:
                            pass
                        try:
                            q2.pop("pairs", None)
                        except Exception:
                            pass

                    elif qtype == "true_false":
                        ca = q2.get("correct_answer")
                        if ca is None:
                            for k in ("correct", "answer"):
                                if k in q2:
                                    ca = q2.get(k)
                                    break
                        b = _coerce_bool(ca)
                        if b is None:
                            return None
                        q2["correct_answer"] = bool(b)
                        for k in ("options", "answers", "pairs", "case_sensitive", "correct_index", "correctIndex"):
                            try:
                                q2.pop(k, None)
                            except Exception:
                                pass

                    elif qtype == "fill_blank":
                        ans = q2.get("answers")
                        if ans is None:
                            for k in ("correct_answers", "correctAnswers", "blanks", "solutions", "solution"):
                                if k in q2:
                                    ans = q2.get(k)
                                    break
                        answers = _coerce_str_list(ans)
                        if not answers:
                            return None
                        q2["answers"] = answers
                        cs = q2.get("case_sensitive")
                        q2["case_sensitive"] = bool(cs) if isinstance(cs, bool) else False
                        for k in ("options", "correct_answer", "correct", "answer", "pairs"):
                            try:
                                q2.pop(k, None)
                            except Exception:
                                pass

                    elif qtype == "short_answer":
                        ans = q2.get("answers")
                        if ans is None:
                            for k in ("correct_answers", "correctAnswers", "solutions", "solution", "answer"):
                                if k in q2:
                                    ans = q2.get(k)
                                    break
                        answers = _coerce_str_list(ans)
                        if not answers:
                            return None
                        q2["answers"] = answers
                        for k in ("options", "correct_answer", "correct", "pairs", "case_sensitive"):
                            try:
                                q2.pop(k, None)
                            except Exception:
                                pass

                    elif qtype == "matching":
                        pairs_raw = q2.get("pairs")
                        if pairs_raw is None:
                            for k in ("matching_pairs", "matchingPairs", "match_pairs", "matchPairs", "items"):
                                if k in q2:
                                    pairs_raw = q2.get(k)
                                    break
                        pairs_out = []
                        if isinstance(pairs_raw, dict):
                            for lk, rv in pairs_raw.items():
                                try:
                                    l0 = str(lk or "").strip()
                                except Exception:
                                    l0 = ""
                                try:
                                    r0 = str(rv or "").strip()
                                except Exception:
                                    r0 = ""
                                if l0 and r0:
                                    pairs_out.append({"left": l0, "right": r0})
                        elif isinstance(pairs_raw, list):
                            for it in pairs_raw:
                                if isinstance(it, dict):
                                    try:
                                        l0 = str(it.get("left") or it.get("a") or it.get("term") or "").strip()
                                    except Exception:
                                        l0 = ""
                                    try:
                                        r0 = str(it.get("right") or it.get("b") or it.get("definition") or "").strip()
                                    except Exception:
                                        r0 = ""
                                    if l0 and r0:
                                        pairs_out.append({"left": l0, "right": r0})
                                elif isinstance(it, (list, tuple)) and len(it) >= 2:
                                    try:
                                        l0 = str(it[0] or "").strip()
                                    except Exception:
                                        l0 = ""
                                    try:
                                        r0 = str(it[1] or "").strip()
                                    except Exception:
                                        r0 = ""
                                    if l0 and r0:
                                        pairs_out.append({"left": l0, "right": r0})
                                else:
                                    try:
                                        s0 = str(it or "").strip()
                                    except Exception:
                                        s0 = ""
                                    if s0 and ("-" in s0 or "–" in s0 or ":" in s0):
                                        sep = "–" if "–" in s0 else ("-" if "-" in s0 else ":")
                                        a0, b0 = (s0.split(sep, 1) + [""])[:2]
                                        a0 = str(a0 or "").strip()
                                        b0 = str(b0 or "").strip()
                                        if a0 and b0:
                                            pairs_out.append({"left": a0, "right": b0})
                        elif pairs_raw is None:
                            try:
                                import re
                                qtxt = str(q2.get("question") or "")
                                lines = [ln.strip() for ln in qtxt.splitlines() if ln.strip()]
                                hits = []
                                for ln in lines:
                                    if ("-" in ln) or ("–" in ln) or (":" in ln):
                                        sep = "–" if "–" in ln else ("-" if "-" in ln else ":")
                                        a0, b0 = (ln.split(sep, 1) + [""])[:2]
                                        a0 = str(a0 or "").strip()
                                        b0 = str(b0 or "").strip()
                                        if a0 and b0 and len(a0) <= 60 and len(b0) <= 120:
                                            hits.append({"left": a0, "right": b0})
                                if len(hits) >= 2:
                                    pairs_out = hits
                                    q2["question"] = re.sub(r"[\r\n]+", "\n", qtxt).strip().splitlines()[0].strip() if qtxt else q2.get("question")
                            except Exception:
                                pass
                        if len(pairs_out) < 2:
                            return None
                        q2["pairs"] = pairs_out
                        for k in ("options", "correct_answer", "correct", "answer", "answers", "case_sensitive"):
                            try:
                                q2.pop(k, None)
                            except Exception:
                                pass

                    img_val = None
                    for k in ("image", "image_path", "imagePath", "image_url", "imageUrl"):
                        if k in q2 and q2.get(k):
                            img_val = q2.get(k)
                            break
                    qtxt_for_img = str(q2.get("question") or "")
                    if not _is_image_question(qtxt_for_img):
                        for k in ("image", "image_path", "imagePath", "image_url", "imageUrl"):
                            try:
                                q2.pop(k, None)
                            except Exception:
                                pass
                    else:
                        local_img = _prepare_local_image_path(str(img_val)) if img_val is not None else ""
                        if local_img:
                            q2["image"] = local_img
                        else:
                            try:
                                q2.pop("image", None)
                            except Exception:
                                pass

                    if not str(q2.get("question") or "").strip():
                        return None
                    return q2

                ok_any = False
                created_ids = []
                for q_data in items:
                    nq = _normalize_question(q_data) if isinstance(q_data, dict) else None
                    if nq and self.project_manager.add_question(nq):
                        ok_any = True
                        try:
                            qid = str(nq.get("id") or "").strip()
                            if qid:
                                created_ids.append(qid)
                        except Exception:
                            pass
                if ok_any:
                    current = self.project_manager.get_current_project()
                    if current:
                        self._open_editor_for_project(current)
                    out["ok"] = True
                    try:
                        if created_ids and current:
                            import json
                            created = []
                            try:
                                qs = current.get("questions", []) if isinstance(current, dict) else []
                            except Exception:
                                qs = []
                            try:
                                idset = set(created_ids)
                            except Exception:
                                idset = set()
                            if isinstance(qs, list) and idset:
                                for idx2, q in enumerate(qs):
                                    if not isinstance(q, dict):
                                        continue
                                    qid2 = str(q.get("id") or "").strip()
                                    if qid2 and qid2 in idset:
                                        qtype2 = str(q.get("type") or "").strip()
                                        qtext2 = str(q.get("question") or "").strip()
                                        exp2 = str(q.get("explanation") or "").strip()
                                        img2 = str(q.get("image") or "").strip()
                                        opts2 = q.get("options")
                                        ca2 = q.get("correct_answer")
                                        ans2 = q.get("answers")
                                        correct_text = ""
                                        try:
                                            if isinstance(opts2, list) and isinstance(ca2, int) and 0 <= int(ca2) < len(opts2):
                                                correct_text = str(opts2[int(ca2)]).strip()
                                        except Exception:
                                            correct_text = ""
                                        created.append(
                                            {
                                                "index": idx2 + 1,
                                                "question_id": qid2,
                                                "type": qtype2,
                                                "question": qtext2[:260],
                                                "options": opts2,
                                                "correct_answer": ca2,
                                                "correct_text": correct_text[:160],
                                                "answers": ans2,
                                                "explanation": exp2[:260],
                                                "image": img2,
                                            }
                                        )
                            out["follow_up"] = json.dumps(
                                {
                                    "created_question_ids": created_ids,
                                    "created_questions": created,
                                    "project_id": current.get("id"),
                                    "project_name": current.get("name"),
                                },
                                ensure_ascii=False,
                            )
                    except Exception:
                        pass
                else:
                    self.chat_widget.add_message("System", _t("ai.tool.add_question_failed", "Failed to add question."), False)
                return out

            if cmd == "UPDATE_QUESTION":
                try:
                    from eduplay.core.ai_tool_payloads import parse_update_question_payload, build_updated_question
                except Exception:
                    parse_update_question_payload = None
                    build_updated_question = None
                try:
                    payload = parse_update_question_payload(args) if parse_update_question_payload else _parse_json_lenient(args)
                except Exception:
                    payload = _parse_json_lenient(args)
                qid_from_queue = False
                qid = str(
                    (payload or {}).get("question_id")
                    or (payload or {}).get("questionId")
                    or (payload or {}).get("qid")
                    or (payload or {}).get("id")
                    or ""
                ).strip()
                patch = (payload or {}).get("patch")
                full = (payload or {}).get("question")
                proj = self.project_manager.get_current_project() or {}
                if not proj:
                    try:
                        proj = {}
                        for _nm3 in ("editor_quiz_screen", "editor_fishing_screen", "editor_screen"):
                            try:
                                _scr3 = getattr(self, _nm3, None)
                            except Exception:
                                _scr3 = None
                            if not _scr3:
                                continue
                            try:
                                proj = _scr3.get_current_project() if hasattr(_scr3, "get_current_project") else getattr(_scr3, "current_project", None)
                            except Exception:
                                proj = {}
                            if isinstance(proj, dict) and str(proj.get("id") or "").strip():
                                break
                    except Exception:
                        proj = {}
                try:
                    if proj and isinstance(proj, dict):
                        pid = str(proj.get("id") or "").strip()
                    else:
                        pid = ""
                except Exception:
                    pid = ""
                if pid:
                    try:
                        cur_pm = self.project_manager.get_current_project() or {}
                    except Exception:
                        cur_pm = {}
                    try:
                        cur_pid = str((cur_pm or {}).get("id") or "").strip()
                    except Exception:
                        cur_pid = ""
                    if not cur_pid or cur_pid != pid:
                        try:
                            self.project_manager.current_project = proj
                        except Exception:
                            pass
                questions = proj.get("questions", []) if isinstance(proj, dict) else []
                if (not isinstance(questions, list)) or (not questions):
                    try:
                        lp = None
                        for _nm4 in ("editor_quiz_screen", "editor_fishing_screen", "editor_screen"):
                            try:
                                _scr4 = getattr(self, _nm4, None)
                            except Exception:
                                _scr4 = None
                            if not _scr4:
                                continue
                            lp = getattr(_scr4, "left_panel", None)
                            if lp is not None:
                                break
                        qlist = getattr(lp, "questions_list", None) if lp is not None else None
                        tmp = []
                        if qlist is not None:
                            for ii in range(int(qlist.count())):
                                try:
                                    it = qlist.item(ii)
                                except Exception:
                                    it = None
                                if not it:
                                    continue
                                try:
                                    qobj = it.data(Qt.UserRole)
                                except Exception:
                                    qobj = None
                                if isinstance(qobj, dict):
                                    tmp.append(qobj)
                        if tmp:
                            questions = tmp
                    except Exception:
                        pass
                if (not qid) and isinstance(questions, list) and questions:
                    idx_raw = None
                    try:
                        for k in ("question_index", "questionIndex", "index", "question_number", "questionNumber", "number", "q", "stt", "no", "seq"):
                            if k in (payload or {}):
                                idx_raw = (payload or {}).get(k)
                                break
                    except Exception:
                        idx_raw = None
                    idx0 = None
                    try:
                        if isinstance(idx_raw, int):
                            idx0 = int(idx_raw)
                        elif idx_raw is not None:
                            sidx = str(idx_raw).strip()
                            if sidx.isdigit():
                                idx0 = int(sidx)
                            else:
                                try:
                                    import re

                                    m = re.search(r"\d{1,3}", sidx)
                                    if m:
                                        idx0 = int(m.group(0))
                                except Exception:
                                    idx0 = None
                    except Exception:
                        idx0 = None
                    try:
                        if idx0 is not None:
                            if idx0 >= 1:
                                idx0 = int(idx0) - 1
                            if 0 <= int(idx0) < len(questions):
                                cand = questions[int(idx0)]
                                if isinstance(cand, dict):
                                    qid = str(cand.get("id") or "").strip()
                    except Exception:
                        pass
                if not qid:
                    try:
                        pending = getattr(self, "_pending_update_question_ids", []) or []
                        if isinstance(pending, list) and pending:
                            qid = str(pending.pop(0) or "").strip()
                            self._pending_update_question_ids = pending
                            if qid:
                                qid_from_queue = True
                    except Exception:
                        pass
                if not qid:
                    try:
                        pending2 = getattr(self, "_last_target_question_ids", []) or []
                        if isinstance(pending2, list) and pending2:
                            qid = str(pending2.pop(0) or "").strip()
                            self._last_target_question_ids = pending2
                            if qid:
                                qid_from_queue = True
                    except Exception:
                        pass
                if not qid:
                    try:
                        import re

                        raw_user = str(getattr(self, "_last_chat_user_msg", "") or "")
                        try:
                            from eduplay.core.ai_tool_payloads import extract_question_numbers
                        except Exception:
                            extract_question_numbers = None
                        hits = extract_question_numbers(raw_user) if extract_question_numbers else []
                        hits = [x for x in hits if 1 <= int(x) <= len(questions)]
                        hits = list(dict.fromkeys(hits))
                        if hits:
                            ids = []
                            for h in hits[:12]:
                                idx0 = int(h) - 1
                                if 0 <= idx0 < len(questions):
                                    cand = questions[idx0]
                                    if isinstance(cand, dict):
                                        qx = str(cand.get("id") or "").strip()
                                        if qx:
                                            ids.append(qx)
                            if ids:
                                qid = str(ids[0] or "").strip()
                                if len(ids) > 1:
                                    self._pending_update_question_ids = list(ids[1:])
                    except Exception:
                        pass
                if not qid:
                    self.chat_widget.add_message("System", _t("ai.tool.missing_question_id", "Missing question_id."), False)
                    try:
                        out["follow_up"] = "TOOL_ERROR: Missing question_id. Expected: UPDATE_QUESTION: {\"question_id\":\"...\",\"patch\":{\"explanation\":\"...\"}}"
                    except Exception:
                        pass
                    return out
                current_q = None
                if isinstance(questions, list):
                    for q in questions:
                        if isinstance(q, dict) and q.get("id") == qid:
                            current_q = q
                            break
                if (current_q is None) and pid:
                    try:
                        loaded = self.project_manager.load_project(pid)
                    except Exception:
                        loaded = None
                    if isinstance(loaded, dict):
                        try:
                            proj = loaded
                            questions = proj.get("questions", []) if isinstance(proj, dict) else []
                            try:
                                self.editor_screen.load_project(proj)
                            except Exception:
                                pass
                        except Exception:
                            pass
                        if isinstance(questions, list):
                            for q in questions:
                                if isinstance(q, dict) and q.get("id") == qid:
                                    current_q = q
                                    break
                new_q = None
                if build_updated_question:
                    try:
                        new_q = build_updated_question(payload or {}, current_q)
                    except Exception:
                        new_q = None
                if not isinstance(new_q, dict):
                    try:
                        if qid and qid_from_queue:
                            pending_back = getattr(self, "_pending_update_question_ids", []) or []
                            if not isinstance(pending_back, list):
                                pending_back = []
                            if qid not in pending_back:
                                pending_back.insert(0, qid)
                            self._pending_update_question_ids = pending_back
                    except Exception:
                        pass
                    self.chat_widget.add_message("System", _t("ai.tool.update_question_payload", "Edubot nhận lệnh sửa không hợp lệ. Tôi đang thử lại..."), False)
                    try:
                        out["follow_up"] = (
                            "TOOL_ERROR: Invalid UPDATE_QUESTION payload.\n"
                            "Output UPDATE_QUESTION as ONE LINE JSON.\n"
                            "If you are updating explanation only, use:\n"
                            "UPDATE_QUESTION: {\"patch\":{\"explanation\":\"...\"}}\n"
                            "Do NOT output empty {}."
                        )
                    except Exception:
                        pass
                    return out
                if self.project_manager.update_question(qid, new_q):
                    current = self.project_manager.get_current_project()
                    if current:
                        self._open_editor_for_project(current)
                    out["ok"] = True
                else:
                    self.chat_widget.add_message("System", _t("ai.tool.update_question_failed", f"Failed to update question: {qid}"), False)
                return out

            if cmd == "DELETE_QUESTION":
                try:
                    payload = _parse_json_lenient(args)
                except Exception:
                    payload = None

                proj = self.project_manager.get_current_project() or {}
                questions = proj.get("questions", []) if isinstance(proj, dict) else []

                qid = ""
                idx0 = None
                try:
                    if isinstance(payload, dict):
                        qid = str(
                            payload.get("question_id")
                            or payload.get("questionId")
                            or payload.get("qid")
                            or payload.get("id")
                            or ""
                        ).strip()
                        for k in ("question_index", "questionIndex", "index", "question_number", "questionNumber", "number", "q", "stt", "no", "seq"):
                            if k in payload and payload.get(k) is not None:
                                idx0 = payload.get(k)
                                break
                except Exception:
                    qid = ""
                    idx0 = None

                if not qid:
                    try:
                        sraw = str(args or "").strip()
                    except Exception:
                        sraw = ""
                    if sraw:
                        if sraw.isdigit():
                            idx0 = int(sraw)
                        elif "=" in sraw:
                            try:
                                import re
                                m = re.search(r"(?:^|[,\s])(?:index|question_number|questionNumber|number|no|stt)\s*=\s*(\d+)", sraw, flags=re.IGNORECASE)
                                if m:
                                    idx0 = int(m.group(1))
                            except Exception:
                                pass
                        else:
                            qid = sraw if sraw.startswith("q_") else qid

                # --- Priority 1: use _pending_delete_question_ids (populated for bulk/last-N deletes) ---
                if not qid:
                    try:
                        pending = getattr(self, "_pending_delete_question_ids", []) or []
                        if isinstance(pending, list) and pending:
                            # If there are multiple pending IDs, do a bulk delete of ALL of them at once
                            if len(pending) > 1:
                                import json as _json_del
                                deleted_bulk = []
                                for _bid in list(pending):
                                    if self.project_manager.delete_question(_bid):
                                        deleted_bulk.append({"question_id": _bid})
                                self._pending_delete_question_ids = []
                                current = self.project_manager.get_current_project()
                                if current:
                                    self._open_editor_for_project(current)
                                out["ok"] = bool(deleted_bulk)
                                try:
                                    out["follow_up"] = _json_del.dumps(
                                        {
                                            "deleted": deleted_bulk,
                                            "deleted_count": len(deleted_bulk),
                                            "remaining_count": len((current or {}).get("questions", []) or []) if isinstance(current, dict) else None,
                                        },
                                        ensure_ascii=False,
                                    )
                                except Exception:
                                    pass
                                return out
                            else:
                                qid = str(pending.pop(0) or "").strip()
                                self._pending_delete_question_ids = pending
                    except Exception:
                        pass

                if (not qid) and (idx0 is not None) and isinstance(questions, list) and questions:
                    try:
                        n = int(idx0)
                        if n >= 1:
                            n = n - 1
                        if 0 <= n < len(questions):
                            cand = questions[n]
                            if isinstance(cand, dict):
                                qid = str(cand.get("id") or "").strip()
                    except Exception:
                        pass

                if (not qid) and isinstance(questions, list) and questions:
                    try:
                        raw_user = str(getattr(self, "_last_chat_user_msg", "") or "")
                        try:
                            from eduplay.core.ai_tool_payloads import extract_question_numbers
                        except Exception:
                            extract_question_numbers = None
                        hits = extract_question_numbers(raw_user) if extract_question_numbers else []
                        hits = [x for x in hits if 1 <= int(x) <= len(questions)]
                        hits = list(dict.fromkeys(hits))
                        if hits:
                            n = int(hits[0]) - 1
                            cand = questions[n]
                            if isinstance(cand, dict):
                                qid = str(cand.get("id") or "").strip()
                    except Exception:
                        pass

                if not qid:
                    delete_count = 1
                    try:
                        if isinstance(payload, dict):
                            for k in ("count", "n", "num", "quantity", "how_many", "so_cau", "so_cau_hoi"):
                                if k in payload and payload.get(k) is not None:
                                    delete_count = int(payload.get(k))
                                    break
                    except Exception:
                        delete_count = 1
                    if delete_count <= 1:
                        try:
                            import re
                            raw_user = str(getattr(self, "_last_chat_user_msg", "") or "")
                            raw0 = str(raw_user or "").lower()
                            num_words = {
                                "một": 1,
                                "mot": 1,
                                "hai": 2,
                                "ba": 3,
                                "bốn": 4,
                                "bon": 4,
                                "tư": 4,
                                "tu": 4,
                                "năm": 5,
                                "nam": 5,
                                "sáu": 6,
                                "sau": 6,
                                "bảy": 7,
                                "bay": 7,
                                "tám": 8,
                                "tam": 8,
                                "chín": 9,
                                "chin": 9,
                                "mười": 10,
                                "muoi": 10,
                            }
                            m = re.search(r"(?:xóa|xoá|delete|remove)\s+(\d+)\s+(?:câu|cau)", raw0, flags=re.IGNORECASE)
                            if not m:
                                m = re.search(r"(?:xóa|xoá|delete|remove)\s+([a-zàáạảãâầấậẩẫăằắặẳẵđèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ]+)\s+(?:câu|cau)", raw0, flags=re.IGNORECASE)
                            if m and m.group(1):
                                g = str(m.group(1) or "").strip()
                                if g.isdigit():
                                    delete_count = int(g)
                                else:
                                    delete_count = int(num_words.get(g, 1))
                            if delete_count <= 1:
                                m2 = re.search(r"(?:xóa|xoá|delete|remove)\s+(?:\d+|[a-zàáạảãâầấậẩẫăằắặẳẵđèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ]+)\s+(?:câu|cau)[^\n\r]*\s+cuối", raw0, flags=re.IGNORECASE)
                                if m2 and ("hai" in raw0 or "2" in raw0):
                                    delete_count = 2
                        except Exception:
                            delete_count = 1
                    if delete_count > 1 and isinstance(questions, list) and questions:
                        deleted = []
                        try:
                            import json
                        except Exception:
                            json = None
                        max_del = min(int(delete_count), len(questions))
                        for _ in range(max_del):
                            cur = self.project_manager.get_current_project() or {}
                            qs = cur.get("questions", []) if isinstance(cur, dict) else []
                            if not isinstance(qs, list) or not qs:
                                break
                            last = qs[-1]
                            qid2 = str((last or {}).get("id") or "").strip() if isinstance(last, dict) else ""
                            if not qid2:
                                break
                            dnum = len(qs)
                            dtext = str((last or {}).get("question") or "").strip()[:200] if isinstance(last, dict) else ""
                            if not self.project_manager.delete_question(qid2):
                                break
                            deleted.append({"question_id": qid2, "question_number": dnum, "question": dtext})
                        current = self.project_manager.get_current_project()
                        if current:
                            self._open_editor_for_project(current)
                        out["ok"] = bool(deleted)
                        if json is not None:
                            try:
                                out["follow_up"] = json.dumps(
                                    {
                                        "deleted": deleted,
                                        "deleted_count": len(deleted),
                                        "remaining_count": len((current or {}).get("questions", []) or []) if isinstance(current, dict) else None,
                                    },
                                    ensure_ascii=False,
                                )
                            except Exception:
                                pass
                        return out
                    if isinstance(questions, list) and questions:
                        try:
                            last = questions[-1]
                            if isinstance(last, dict):
                                qid = str(last.get("id") or "").strip()
                                idx0 = len(questions)
                        except Exception:
                            qid = ""
                            idx0 = None
                    if not qid:
                        self.chat_widget.add_message("System", _t("ai.tool.missing_question_id", "Missing question_id."), False)
                        try:
                            out["follow_up"] = "TOOL_ERROR: Missing question_id. Expected: DELETE_QUESTION: {\"question_number\":7} or {\"question_id\":\"q_...\"}"
                        except Exception:
                            pass
                        return out

                deleted_text = ""
                deleted_index = None
                try:
                    if isinstance(questions, list):
                        for ii, q in enumerate(questions):
                            if isinstance(q, dict) and str(q.get("id") or "").strip() == qid:
                                deleted_index = ii + 1
                                deleted_text = str(q.get("question") or "").strip()[:200]
                                break
                except Exception:
                    deleted_text = ""
                    deleted_index = None

                if self.project_manager.delete_question(qid):
                    current = self.project_manager.get_current_project()
                    if current:
                        self._open_editor_for_project(current)
                    out["ok"] = True
                    try:
                        import json
                        out["follow_up"] = json.dumps(
                            {
                                "deleted_question_id": qid,
                                "deleted_question_number": deleted_index,
                                "deleted_question": deleted_text,
                                "remaining_count": len((current or {}).get("questions", []) or []) if isinstance(current, dict) else None,
                            },
                            ensure_ascii=False,
                        )
                    except Exception:
                        pass
                else:
                    self.chat_widget.add_message("System", _t("ai.tool.delete_question_failed", f"Failed to delete question: {qid}"), False)
                return out

            if cmd in ("SET_QUESTION_IMAGE", "SET_QUESTION_IMAGE_URL", "SEARCH_IMAGE"):
                payload = _parse_json_lenient(args) if cmd != "SEARCH_IMAGE" else _parse_json_lenient(args)
                qid = str((payload or {}).get("question_id") or "").strip()
                if not qid:
                    idx_raw = None
                    try:
                        if isinstance(payload, dict):
                            for k in ("question_index", "questionIndex", "index", "question_number", "questionNumber", "number", "q", "stt", "no", "seq"):
                                if k in payload and payload.get(k) is not None:
                                    idx_raw = payload.get(k)
                                    break
                    except Exception:
                        idx_raw = None
                    idx0 = None
                    try:
                        if isinstance(idx_raw, int):
                            idx0 = int(idx_raw)
                        elif idx_raw is not None:
                            sidx = str(idx_raw).strip()
                            if sidx.isdigit():
                                idx0 = int(sidx)
                    except Exception:
                        idx0 = None
                    if idx0 is not None:
                        try:
                            proj = self.project_manager.get_current_project() or {}
                        except Exception:
                            proj = {}
                        questions = proj.get("questions", []) if isinstance(proj, dict) else []
                        try:
                            n = int(idx0)
                            if n >= 1:
                                n = n - 1
                            if isinstance(questions, list) and 0 <= n < len(questions):
                                cand = questions[n]
                                if isinstance(cand, dict):
                                    qid = str(cand.get("id") or "").strip()
                        except Exception:
                            pass
                if not qid:
                    self.chat_widget.add_message("System", _t("ai.tool.missing_question_id", "Missing question_id."), False)
                    return out

                if cmd == "SET_QUESTION_IMAGE":
                    image_path = str((payload or {}).get("image_path") or "").strip()
                    if not image_path:
                        self.chat_widget.add_message("System", _t("ai.tool.set_image_format", "SET_QUESTION_IMAGE requires {question_id, image_path}."), False)
                        return out
                    from pathlib import Path
                    import shutil
                    media_dir = _safe_media_target()
                    src = Path(image_path)
                    if not src.is_absolute():
                        p = str(image_path or "").replace("\\", "/").strip()
                        if p.startswith("media/"):
                            p = p.split("/", 1)[1]
                        cand = Path(media_dir) / p
                        if cand.exists():
                            src = cand
                    if not src.exists() or not src.is_file():
                        self.chat_widget.add_message("System", _t("ai.tool.image_not_found", f"Image not found: {image_path}"), False)
                        return out
                    try:
                        media_root = Path(media_dir).resolve()
                        src_res = src.resolve()
                        in_media = (src_res == media_root) or (media_root in src_res.parents)
                    except Exception:
                        in_media = False
                    if in_media:
                        abs_path = str(src)
                    else:
                        stem = src.stem
                        suffix = src.suffix
                        candidate = Path(media_dir) / f"{stem}{suffix}"
                        idx = 1
                        while candidate.exists():
                            candidate = Path(media_dir) / f"{stem}_{idx}{suffix}"
                            idx += 1
                        shutil.copy2(str(src), str(candidate))
                        abs_path = str(candidate)
                else:
                    try:
                        ai_settings = self.settings_manager.get_ai_settings() or {}
                        allow_net = bool(ai_settings.get("allow_internet", False))
                        allow_img = bool(ai_settings.get("allow_image_download", False))
                    except Exception:
                        allow_net = False
                        allow_img = False
                    if not (allow_net and allow_img):
                        self.chat_widget.add_message("System", _t("ai.tool.internet_disabled", "Internet tools are disabled in Settings."), False)
                        return out

                    url = ""
                    if cmd == "SET_QUESTION_IMAGE_URL":
                        url = str((payload or {}).get("url") or "").strip()
                    if cmd == "SEARCH_IMAGE":
                        query = str((payload or {}).get("query") or "").strip()
                        if not query:
                            self.chat_widget.add_message("System", _t("ai.tool.search_image_format", "SEARCH_IMAGE requires {question_id, query}."), False)
                            return out
                        try:
                            import requests
                            api = "https://commons.wikimedia.org/w/api.php"
                            r = requests.get(
                                api,
                                params={
                                    "action": "query",
                                    "format": "json",
                                    "generator": "search",
                                    "gsrsearch": query,
                                    "gsrlimit": 1,
                                    "prop": "imageinfo",
                                    "iiprop": "url",
                                    "iiurlwidth": 1024,
                                },
                                timeout=15,
                            )
                            data = r.json() if (200 <= r.status_code < 300) else {}
                            pages = (((data or {}).get("query") or {}).get("pages") or {})
                            first = None
                            for _, v in (pages.items() if isinstance(pages, dict) else []):
                                first = v
                                break
                            if first:
                                ii = (first.get("imageinfo") or [])
                                if isinstance(ii, list) and ii:
                                    url = str(ii[0].get("thumburl") or ii[0].get("url") or "").strip()
                        except Exception:
                            url = ""
                    if not url:
                        self.chat_widget.add_message("System", _t("ai.tool.image_url_not_found", "Could not find a suitable image URL."), False)
                        return out
                    try:
                        import requests
                        from pathlib import Path
                        import re
                        import tempfile
                        resp = requests.get(url, timeout=20)
                        if not (200 <= resp.status_code < 300):
                            raise ValueError(f"HTTP {resp.status_code}")
                        ctype = str(resp.headers.get("content-type") or "").lower()
                        if "image/" not in ctype:
                            raise ValueError("Not an image")
                        ext = ""
                        if "png" in ctype:
                            ext = ".png"
                        elif "jpeg" in ctype or "jpg" in ctype:
                            ext = ".jpg"
                        elif "gif" in ctype:
                            ext = ".gif"
                        elif "webp" in ctype:
                            ext = ".webp"
                        if not ext:
                            ext = ".png"
                        tmp = Path(tempfile.gettempdir()) / ("eduplay_img_" + re.sub(r"[^a-zA-Z0-9]+", "_", qid)[:30] + ext)
                        tmp.write_bytes(resp.content)
                        media_dir = _safe_media_target()
                        dest = Path(media_dir) / tmp.name
                        idx = 1
                        while dest.exists():
                            dest = Path(media_dir) / (tmp.stem + f"_{idx}" + tmp.suffix)
                            idx += 1
                        dest.write_bytes(tmp.read_bytes())
                        abs_path = str(dest)
                    except Exception as e:
                        self.chat_widget.add_message("System", _t("ai.tool.image_download_failed", f"Failed to download image: {e}"), False)
                        return out

                proj = self.project_manager.get_current_project()
                questions = proj.get("questions", []) if isinstance(proj, dict) else []
                target_q = None
                if isinstance(questions, list):
                    for q in questions:
                        if isinstance(q, dict) and q.get("id") == qid:
                            target_q = q
                            break
                if not isinstance(target_q, dict):
                    self.chat_widget.add_message("System", _t("ai.tool.question_not_found", f"Question not found: {qid}"), False)
                    return out
                new_q = dict(target_q)
                new_q["image"] = abs_path
                if self.project_manager.update_question(qid, new_q):
                    current = self.project_manager.get_current_project()
                    if current:
                        self._open_editor_for_project(current)
                    out["ok"] = True
                else:
                    self.chat_widget.add_message("System", _t("ai.tool.set_image_failed", f"Failed to set image for {qid}"), False)
                return out

            if cmd == "WEB_SEARCH":
                try:
                    ai_settings = self.settings_manager.get_ai_settings() or {}
                    allow_net = bool(ai_settings.get("allow_internet", False))
                except Exception:
                    allow_net = False
                if not allow_net:
                    self.chat_widget.add_message("System", _t("ai.tool.internet_disabled", "Internet tools are disabled in Settings."), False)
                    return out
                query = str(args or "").strip()
                if not query:
                    self.chat_widget.add_message("System", _t("ai.tool.web_search_format", "WEB_SEARCH requires a query string."), False)
                    return out
                try:
                    import requests
                    r = requests.get(
                        "https://api.duckduckgo.com/",
                        params={"q": query, "format": "json", "no_redirect": 1, "no_html": 1},
                        timeout=15,
                    )
                    data = r.json() if (200 <= r.status_code < 300) else {}
                    abstract = str((data or {}).get("AbstractText") or "").strip()
                    heading = str((data or {}).get("Heading") or "").strip()
                    src = str((data or {}).get("AbstractURL") or "").strip()
                    lines = []
                    if heading:
                        lines.append(f"Query: {query}\nTitle: {heading}")
                    else:
                        lines.append(f"Query: {query}")
                    if abstract:
                        lines.append(abstract)
                    if src:
                        lines.append(f"Source: {src}")
                    out["follow_up"] = "\n".join(lines).strip()
                    out["ok"] = True
                    return out
                except Exception as e:
                    self.chat_widget.add_message("System", _t("ai.tool.web_search_failed", f"Web search failed: {e}"), False)
                    return out

            if cmd == "WEB_FETCH":
                try:
                    ai_settings = self.settings_manager.get_ai_settings() or {}
                    allow_net = bool(ai_settings.get("allow_internet", False))
                except Exception:
                    allow_net = False
                if not allow_net:
                    self.chat_widget.add_message("System", _t("ai.tool.internet_disabled", "Internet tools are disabled in Settings."), False)
                    return out
                url = str(args or "").strip()
                try:
                    url = url.strip("`").strip().strip('"').strip("'").strip()
                    url = url.strip("<>").strip()
                except Exception:
                    pass
                try:
                    if url.lower().startswith("www."):
                        url = "https://" + url
                except Exception:
                    pass
                if not url:
                    self.chat_widget.add_message("System", _t("ai.tool.web_fetch_format", "WEB_FETCH requires a URL."), False)
                    return out
                if not (url.lower().startswith("https://") or url.lower().startswith("http://")):
                    self.chat_widget.add_message("System", _t("ai.tool.web_fetch_format", "WEB_FETCH requires a URL."), False)
                    return out
                try:
                    import re
                    m = re.search(r"drive\.google\.com\/file\/d\/([^\/]+)", url)
                    if m:
                        fid = m.group(1)
                        url = f"https://drive.google.com/uc?export=download&id={fid}"
                    else:
                        m2 = re.search(r"drive\.google\.com\/open\\?id=([^&]+)", url)
                        if m2:
                            fid = m2.group(1)
                            url = f"https://drive.google.com/uc?export=download&id={fid}"
                except Exception:
                    pass
                try:
                    import requests
                    import io
                    headers = {"User-Agent": "EduPlayStudio/1.0"}
                    r = requests.get(url, headers=headers, timeout=30)
                    if not (200 <= r.status_code < 300):
                        self.chat_widget.add_message("System", _t("ai.tool.web_fetch_failed", f"WEB_FETCH failed ({r.status_code})."), False)
                        return out
                    ctype = str(r.headers.get("content-type") or "").lower()
                    fname = ""
                    try:
                        cd = str(r.headers.get("content-disposition") or "")
                        import re
                        mfn = re.search(r"filename\\*=UTF-8''([^;]+)", cd, flags=re.IGNORECASE)
                        if mfn:
                            fname = mfn.group(1)
                        else:
                            mfn2 = re.search(r'filename="([^"]+)"', cd, flags=re.IGNORECASE)
                            if mfn2:
                                fname = mfn2.group(1)
                            else:
                                mfn3 = re.search(r"filename=([^;]+)", cd, flags=re.IGNORECASE)
                                if mfn3:
                                    fname = mfn3.group(1).strip().strip('"').strip("'")
                        try:
                            from urllib.parse import unquote
                            fname = unquote(fname)
                        except Exception:
                            pass
                    except Exception:
                        fname = ""
                    raw_bytes = b""
                    try:
                        raw_bytes = r.content or b""
                    except Exception:
                        raw_bytes = b""
                    try:
                        max_bytes = 12 * 1024 * 1024
                        if len(raw_bytes) > max_bytes:
                            raw_bytes = raw_bytes[:max_bytes]
                    except Exception:
                        pass

                    text = ""
                    is_pdf = ("application/pdf" in ctype) or (raw_bytes[:4] == b"%PDF")
                    is_docx = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ctype) or url.lower().endswith(".docx")
                    is_img = ctype.startswith("image/") or raw_bytes[:8] == b"\x89PNG\r\n\x1a\n" or raw_bytes[:3] == b"\xff\xd8\xff"
                    extracted_rel = []
                    media_dir = None
                    try:
                        media_dir = _safe_media_target()
                    except Exception:
                        media_dir = None

                    if is_pdf:
                        try:
                            import fitz
                            doc = fitz.open(stream=raw_bytes, filetype="pdf")
                            chunks = []
                            try:
                                page_count = int(doc.page_count)
                            except Exception:
                                page_count = 0
                            limit_pages = min(8, page_count) if page_count else 8
                            for i in range(limit_pages):
                                try:
                                    chunks.append(doc.load_page(i).get_text("text"))
                                except Exception:
                                    pass
                            text = "\n".join([c for c in chunks if str(c or "").strip()]).strip()
                            if media_dir:
                                try:
                                    from pathlib import Path
                                    import re
                                    media_root = Path(media_dir)
                                    saved = 0
                                    for page_index in range(min(6, int(doc.page_count or 0))):
                                        try:
                                            page = doc.load_page(page_index)
                                            imgs = page.get_images(full=True) or []
                                        except Exception:
                                            imgs = []
                                        for img_i, img in enumerate(imgs):
                                            if saved >= 8:
                                                break
                                            try:
                                                xref = img[0]
                                                data = doc.extract_image(xref) or {}
                                                blob = data.get("image") or b""
                                                ext = str(data.get("ext") or "png").lower().strip().lstrip(".")
                                                if not blob:
                                                    continue
                                                base = f"webfetch_{re.sub(r'[^a-zA-Z0-9]+', '_', str(url)[:32])}_p{page_index+1}_{img_i+1}"
                                                base = re.sub(r"[^a-zA-Z0-9]+", "_", base).strip("_")[:48] or "img"
                                                name = f"{base}.{ext or 'png'}"
                                                dest = media_root / name
                                                idx = 1
                                                while dest.exists():
                                                    dest = media_root / f"{base}_{idx}.{ext or 'png'}"
                                                    idx += 1
                                                dest.write_bytes(blob)
                                                extracted_rel.append(f"media/{dest.name}")
                                                saved += 1
                                            except Exception:
                                                continue
                                        if saved >= 8:
                                            break
                                except Exception:
                                    pass
                        except Exception as e:
                            self.chat_widget.add_message("System", _t("ai.tool.web_fetch_failed", f"WEB_FETCH failed: {e}"), False)
                            return out
                    elif is_docx:
                        try:
                            from docx import Document
                            doc = Document(io.BytesIO(raw_bytes))
                            paras = []
                            for p in doc.paragraphs:
                                try:
                                    t = str(p.text or "").strip()
                                except Exception:
                                    t = ""
                                if t:
                                    paras.append(t)
                            text = "\n".join(paras).strip()
                            if media_dir:
                                try:
                                    from pathlib import Path
                                    import re
                                    rels = getattr(getattr(doc, "part", None), "rels", {}) or {}
                                    media_root = Path(media_dir)
                                    saved = 0
                                    for r2 in rels.values():
                                        if saved >= 8:
                                            break
                                        try:
                                            reltype = str(getattr(r2, "reltype", "") or "")
                                            if "image" not in reltype:
                                                continue
                                            part = getattr(r2, "target_part", None)
                                            blob = getattr(part, "blob", None)
                                            if not blob:
                                                continue
                                            ref = str(getattr(r2, "target_ref", "") or "")
                                            ext = ""
                                            try:
                                                ext = Path(ref).suffix.lower()
                                            except Exception:
                                                ext = ""
                                            if not ext:
                                                ext = ".png"
                                            base = f"webfetch_{re.sub(r'[^a-zA-Z0-9]+', '_', str(url)[:32])}_{saved+1}"
                                            base = re.sub(r"[^a-zA-Z0-9]+", "_", base).strip("_")[:48] or "img"
                                            dest = media_root / f"{base}{ext}"
                                            idx = 1
                                            while dest.exists():
                                                dest = media_root / f"{base}_{idx}{ext}"
                                                idx += 1
                                            dest.write_bytes(blob)
                                            extracted_rel.append(f"media/{dest.name}")
                                            saved += 1
                                        except Exception:
                                            continue
                                except Exception:
                                    pass
                        except Exception as e:
                            self.chat_widget.add_message("System", _t("ai.tool.web_fetch_failed", f"WEB_FETCH failed: {e}"), False)
                            return out
                    elif is_img:
                        try:
                            if media_dir:
                                try:
                                    from pathlib import Path
                                    import re
                                    base = ""
                                    try:
                                        base = Path(fname).stem if fname else ""
                                    except Exception:
                                        base = ""
                                    if not base:
                                        base = re.sub(r"[^a-zA-Z0-9]+", "_", str(url).split("/")[-1]).strip("_")
                                    if not base:
                                        base = "webfetch_img"
                                    base = re.sub(r"[^a-zA-Z0-9]+", "_", base).strip("_")[:48] or "img"
                                    ext = ""
                                    try:
                                        ext = Path(fname).suffix.lower() if fname else ""
                                    except Exception:
                                        ext = ""
                                    if not ext:
                                        ext = ".png"
                                    dest = Path(media_dir) / f"{base}{ext}"
                                    idx = 1
                                    while dest.exists():
                                        dest = Path(media_dir) / f"{base}_{idx}{ext}"
                                        idx += 1
                                    dest.write_bytes(raw_bytes)
                                    extracted_rel.append(f"media/{dest.name}")
                                except Exception:
                                    pass
                            from eduplay.core.ocr_service import OcrService
                            ocr = OcrService()
                            text = ocr.ocr_image_bytes(raw_bytes, lang="vie+eng")
                        except Exception:
                            text = ""
                    else:
                        if "text" in ctype or "json" in ctype or "xml" in ctype or "html" in ctype or not ctype:
                            try:
                                r.encoding = r.encoding or "utf-8"
                            except Exception:
                                pass
                            try:
                                text = str(r.text or "")
                            except Exception:
                                text = ""

                    if is_pdf and (not text or len(text.strip()) < 200):
                        try:
                            from eduplay.core.ocr_service import OcrService
                            ocr = OcrService()
                            ocr_text = ocr.ocr_pdf_bytes(raw_bytes, lang="vie+eng", max_pages=3, dpi=220)
                            if ocr_text and len(ocr_text.strip()) > len((text or "").strip()):
                                text = ocr_text
                        except Exception:
                            pass

                    if not text:
                        msg2 = "WEB_FETCH returned empty content (maybe unsupported file type or OCR not available)."
                        self.chat_widget.add_message("System", _t("ai.tool.web_fetch_failed", msg2), False)
                        return out
                    try:
                        import re
                        cleaned = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", text)
                        cleaned = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", cleaned)
                        cleaned = re.sub(r"(?is)<noscript[^>]*>.*?</noscript>", " ", cleaned)
                        cleaned = re.sub(r"(?is)<[^>]+>", " ", cleaned)
                        cleaned = re.sub(r"[ \t\r\f\v]+", " ", cleaned)
                        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
                        cleaned = cleaned.strip()
                    except Exception:
                        cleaned = text.strip()
                    try:
                        max_len = 7000
                        if len(cleaned) > max_len:
                            cleaned = cleaned[:max_len].rstrip() + "..."
                    except Exception:
                        pass
                    note = ""
                    try:
                        if is_pdf and len(cleaned) < 200:
                            note = "\n\nNOTE: Không trích được nhiều text từ PDF (có thể là scan). Nếu OCR chưa cài được (tesseract.exe), kết quả có thể thiếu."
                    except Exception:
                        note = ""
                    img_note = ""
                    if extracted_rel:
                        img_note = "\n\nEXTRACTED_IMAGES:\n" + "\n".join(extracted_rel[:8])
                    if fname:
                        out["follow_up"] = f"URL: {url}\nFilename: {fname}\nContent-Type: {ctype}\n\n{cleaned}{note}{img_note}"
                    else:
                        out["follow_up"] = f"URL: {url}\nContent-Type: {ctype}\n\n{cleaned}{note}{img_note}"
                    out["ok"] = True
                    return out
                except Exception as e:
                    self.chat_widget.add_message("System", _t("ai.tool.web_fetch_failed", f"WEB_FETCH failed: {e}"), False)
                    return out

            elif cmd == "READ_PROJECT_DETAILS":
                import json
                proj = None
                try:
                    proj = self.project_manager.get_current_project()
                except Exception:
                    proj = None
                if proj and isinstance(proj, dict):
                    out["follow_up"] = json.dumps({"current_project": proj}, ensure_ascii=False)
                else:
                    try:
                        projects = self.project_manager.get_all_projects() or []
                    except Exception:
                        projects = []
                    try:
                        items = [{"name": (p or {}).get("name")} for p in (projects or []) if isinstance(p, dict)]
                    except Exception:
                        items = []
                    out["follow_up"] = json.dumps({"project_count": len(items), "projects": items[:50]}, ensure_ascii=False)
                out["ok"] = True
                return out
            
            if cmd == "UPDATE_GAME_CONFIG":
                cfg = _parse_json_lenient(args)
                try:
                    cur = self.project_manager.get_current_project() or {}
                except Exception:
                    cur = {}
                try:
                    base_cfg = (cur or {}).get("game_config") or {}
                except Exception:
                    base_cfg = {}
                merged = dict(base_cfg) if isinstance(base_cfg, dict) else {}
                if isinstance(cfg, dict):
                    merged.update(cfg)
                if self.project_manager.update_game_config(merged):
                    current = self.project_manager.get_current_project()
                    if current:
                        self._open_editor_for_project(current)
                    out["ok"] = True
                else:
                    self.chat_widget.add_message("System", _t("ai.tool.update_config_failed", "Failed to update config."), False)
                return out

            self.chat_widget.add_message("System", _t("ai.tool.unknown", f"Unknown tool: {cmd}"), False)
            return out

        except Exception as e:
            try:
                print(f"[AI-TOOL-ERROR] {str(e)}")
            except Exception:
                pass
            try:
                out["follow_up"] = f"TOOL_ERROR: {str(e)}"
            except Exception:
                pass
            self.chat_widget.add_message(
                "System",
                ("Tôi gặp lỗi khi xử lý yêu cầu. Bạn bấm gửi lại giúp tôi nhé." if _get_lang() == "vi" else "Something went wrong while processing your request. Please try again."),
                False,
            )
            return out


    def on_chat_file_uploaded(self, content: str):
        try:
            try:
                language = self.settings_manager.get_language()
            except Exception:
                language = 'en'
            try:
                from eduplay.core.i18n import I18n
                try:
                    loading_msg = I18n.t('ai.loading', language)
                except Exception:
                    loading_msg = "Đang suy nghĩ..." if language == 'vi' else "Thinking..."
            except Exception:
                loading_msg = "Đang suy nghĩ..." if language == 'vi' else "Thinking..."

            def _task():
                try:
                    from PySide6.QtCore import QTimer

                    def _progress(step: str):
                        def _append():
                            try:
                                self.chat_widget.set_status(str(step))
                            except Exception:
                                pass

                        try:
                            QTimer.singleShot(0, self, _append)
                        except Exception:
                            QTimer.singleShot(0, _append)

                    raw = ""
                    try:
                        raw = str(content or "")
                    except Exception:
                        raw = ""

                    file_path = ""
                    body = raw
                    try:
                        s = raw.strip()
                        if s.startswith("[Uploaded file]"):
                            lines = raw.splitlines()
                            if len(lines) >= 2:
                                file_path = str(lines[1] or "").strip()
                                body = "\n".join(lines[2:]).strip()
                    except Exception:
                        file_path = ""
                        body = raw

                    extracted_rel = []
                    if file_path:
                        try:
                            from pathlib import Path
                            import re
                            p = Path(file_path)
                            if p.exists() and p.is_file():
                                proj = self.project_manager.get_current_project()
                                if isinstance(proj, dict) and str(proj.get("id") or "").strip():
                                    pid = str(proj.get("id") or "").strip()
                                    media_dir = self.project_manager.projects_dir / pid / "media"
                                    try:
                                        media_dir.mkdir(parents=True, exist_ok=True)
                                    except Exception:
                                        pass

                                    lower = p.name.lower()
                                    is_pdf = lower.endswith(".pdf")
                                    is_docx = lower.endswith(".docx")
                                    is_img = lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"))

                                    def _unique_name(base: str, ext: str) -> str:
                                        safe_base = re.sub(r"[^a-zA-Z0-9]+", "_", base).strip("_")[:48] or "img"
                                        safe_ext = ext if ext.startswith(".") else ("." + ext if ext else ".png")
                                        cand = media_dir / f"{safe_base}{safe_ext}"
                                        idx = 1
                                        while cand.exists():
                                            cand = media_dir / f"{safe_base}_{idx}{safe_ext}"
                                            idx += 1
                                        return cand.name

                                    if is_img:
                                        name = _unique_name(p.stem, p.suffix or ".png")
                                        dest = media_dir / name
                                        try:
                                            dest.write_bytes(p.read_bytes())
                                            extracted_rel.append(f"media/{dest.name}")
                                        except Exception:
                                            pass
                                    elif is_pdf:
                                        try:
                                            import fitz
                                            doc = fitz.open(str(p))
                                            saved = 0
                                            for page_index in range(min(6, int(getattr(doc, "page_count", 0) or 0))):
                                                try:
                                                    page = doc.load_page(page_index)
                                                except Exception:
                                                    continue
                                                try:
                                                    imgs = page.get_images(full=True) or []
                                                except Exception:
                                                    imgs = []
                                                for img_i, img in enumerate(imgs):
                                                    if saved >= 8:
                                                        break
                                                    try:
                                                        xref = img[0]
                                                        data = doc.extract_image(xref) or {}
                                                        blob = data.get("image") or b""
                                                        ext = data.get("ext") or "png"
                                                        if not blob:
                                                            continue
                                                        name = _unique_name(f"{p.stem}_p{page_index+1}_{img_i+1}", "." + str(ext).lstrip("."))
                                                        dest = media_dir / name
                                                        dest.write_bytes(blob)
                                                        extracted_rel.append(f"media/{dest.name}")
                                                        saved += 1
                                                    except Exception:
                                                        continue
                                                if saved >= 8:
                                                    break
                                        except Exception:
                                            pass
                                    elif is_docx:
                                        try:
                                            from docx import Document
                                            doc = Document(str(p))
                                            rels = getattr(getattr(doc, "part", None), "rels", {}) or {}
                                            saved = 0
                                            for r in rels.values():
                                                if saved >= 8:
                                                    break
                                                try:
                                                    reltype = str(getattr(r, "reltype", "") or "")
                                                    if "image" not in reltype:
                                                        continue
                                                    part = getattr(r, "target_part", None)
                                                    blob = getattr(part, "blob", None)
                                                    if not blob:
                                                        continue
                                                    ref = str(getattr(r, "target_ref", "") or "")
                                                    ext = ""
                                                    try:
                                                        ext = Path(ref).suffix
                                                    except Exception:
                                                        ext = ""
                                                    name = _unique_name(f"{p.stem}_{saved+1}", ext or ".png")
                                                    dest = media_dir / name
                                                    dest.write_bytes(blob)
                                                    extracted_rel.append(f"media/{dest.name}")
                                                    saved += 1
                                                except Exception:
                                                    continue
                                        except Exception:
                                            pass
                        except Exception:
                            extracted_rel = []

                    suffix = ""
                    if extracted_rel:
                        suffix = "\n\nẢnh đã trích xuất và lưu trong media:\n" + "\n".join(extracted_rel[:8])

                    if language == "vi":
                        local_prompt = f"Phân tích nội dung tệp sau và đưa ra gợi ý cải thiện:\n\n{body or raw}{suffix}"
                    else:
                        local_prompt = f"Analyze the following file content and give improvement suggestions:\n\n{body or raw}{suffix}"
                    return self.ai_service.chat_with_ai(local_prompt, language)
                except TypeError:
                    if language == "vi":
                        local_prompt = f"Phân tích nội dung tệp sau và đưa ra gợi ý cải thiện:\n\n{content}"
                    else:
                        local_prompt = f"Analyze the following file content and give improvement suggestions:\n\n{content}"
                    return self.ai_service.chat_with_ai(local_prompt, language)

            def _done(result, error):
                try:
                    self.chat_widget.clear_status()
                except Exception:
                    pass
                if error:
                    self.chat_widget.set_ai_response(str(error))
                else:
                    self.chat_widget.set_ai_response(result)

            self._run_background_task(_task, _done, loading_msg, "default")
        except Exception as e:
            self.chat_widget.set_ai_response(str(e))

    
    def closeEvent(self, event):
        try:
            from PySide6.QtGui import QPixmap
            from PySide6.QtWidgets import QSplashScreen
            from PySide6.QtCore import Qt, QTimer
            import os
            icon_path = os.path.join(os.path.dirname(__file__), "../resources/icons/icon.png")
            icon_path = os.path.abspath(icon_path)
            if os.path.exists(icon_path):
                pixmap = QPixmap(icon_path).scaled(180, 180, Qt.KeepAspectRatio, Qt.SmoothTransformation)
                splash = QSplashScreen(pixmap)
                try:
                    from eduplay.core.i18n import I18n
                    lang = self.settings_manager.get_language() if self.settings_manager else 'en'
                    text = I18n.t('app.goodbye', lang)
                except Exception:
                    text = "Goodbye!"
                splash.showMessage(text, Qt.AlignBottom | Qt.AlignHCenter, Qt.white)
                splash.show()
                QTimer.singleShot(700, lambda: splash.close())
        except Exception:
            pass
        event.accept()

    def _show_loading(self, message: str, icon_variant: str = "default"):
        try:
            from PySide6.QtWidgets import QWidget, QVBoxLayout, QApplication
            try:
                theme = self.settings_manager.get_theme() if self.settings_manager else None
                theme = theme or "dark"
            except Exception:
                theme = "dark"
            if theme == "dark":
                overlay_bg = "rgba(232, 235, 240, 0.70)"
                pill_bg = "transparent"
                pill_fg = "#0F1728"
                pill_border = "transparent"
            else:
                overlay_bg = "rgba(232, 235, 240, 0.72)"
                pill_bg = "transparent"
                pill_fg = "#0F1728"
                pill_border = "transparent"
            if not getattr(self, "_loading_overlay", None):
                overlay = QWidget(self)
                overlay.setObjectName("global-loading-overlay")
                overlay.setStyleSheet(f"background-color: {overlay_bg};")
                try:
                    overlay.setAttribute(Qt.WA_TransparentForMouseEvents, False)
                except Exception:
                    pass
                layout = QVBoxLayout(overlay)
                layout.setContentsMargins(0, 0, 0, 0)
                layout.setSpacing(16)
                layout.setAlignment(Qt.AlignCenter)
                icon_label = QLabel()
                icon_label.setAlignment(Qt.AlignCenter)
                icon_label.setStyleSheet("background: transparent;")
                try:
                    icon_label.setFixedSize(260, 260)
                except Exception:
                    pass
                layout.addWidget(icon_label, 0, Qt.AlignHCenter)
                text_label = QLabel(message)
                text_label.setAlignment(Qt.AlignCenter)
                text_label.setWordWrap(True)
                text_label.setStyleSheet(
                    f"""
                    QLabel {{
                        color: {pill_fg};
                        font-size: 16px;
                        font-weight: 600;
                        padding: 8px 24px;
                        background-color: {pill_bg};
                        border-radius: 999px;
                        border: 1px solid {pill_border};
                        min-height: 38px;
                    }}
                """
                )
                layout.addWidget(text_label, 0, Qt.AlignHCenter)
                self._loading_overlay = overlay
                self._loading_label = text_label
                self._loading_icon_label = icon_label
            else:
                if getattr(self, "_loading_label", None):
                    try:
                        self._loading_label.setText(message)
                    except Exception:
                        pass
                try:
                    if getattr(self, "_loading_overlay", None):
                        self._loading_overlay.setStyleSheet(f"background-color: {overlay_bg};")
                        try:
                            self._loading_overlay.setAttribute(Qt.WA_TransparentForMouseEvents, False)
                        except Exception:
                            pass
                    if getattr(self, "_loading_label", None):
                        self._loading_label.setStyleSheet(
                            f"""
                            QLabel {{
                                color: {pill_fg};
                                font-size: 16px;
                                font-weight: 600;
                                padding: 8px 24px;
                                background-color: {pill_bg};
                                border-radius: 999px;
                                border: 1px solid {pill_border};
                            }}
                        """
                        )
                except Exception:
                    pass

            try:
                base_dir = os.path.join(os.path.dirname(__file__), "../resources/icons")
                if icon_variant == "settings":
                    icon_file = "loading-setting.gif"
                elif icon_variant == "search":
                    icon_file = "loading-search.gif"
                elif icon_variant == "web":
                    icon_file = "loading-web.gif"
                else:
                    icon_file = "loading.gif"
                icon_path = os.path.abspath(os.path.join(base_dir, icon_file))
                if os.path.exists(icon_path) and getattr(self, "_loading_icon_label", None):
                    try:
                        if getattr(self, "_loading_movie", None):
                            self._loading_movie.stop()
                            self._loading_movie = None
                    except Exception:
                        pass
                    movie = QMovie(icon_path)
                    try:
                        movie.setScaledSize(QSize(240, 240))
                    except Exception:
                        pass
                    self._loading_icon_label.setMovie(movie)
                    movie.start()
                    self._loading_movie = movie
            except Exception:
                pass

            if self._loading_overlay:
                try:
                    self._loading_overlay.setGeometry(self.rect())
                except Exception:
                    pass
                try:
                    try:
                        if getattr(self, "central_widget", None):
                            if getattr(self, "_loading_blur", None) is None:
                                blur = QGraphicsBlurEffect(self.central_widget)
                                blur.setBlurRadius(8)
                                self.central_widget.setGraphicsEffect(blur)
                                self._loading_blur = blur
                            self.central_widget.setEnabled(False)
                    except Exception:
                        pass
                    try:
                        if getattr(self, "_chat_toggle_btn", None):
                            self._chat_toggle_btn.setEnabled(False)
                    except Exception:
                        pass
                    try:
                        if getattr(self, "chat_widget", None):
                            self.chat_widget.setEnabled(False)
                    except Exception:
                        pass
                    self._loading_overlay.show()
                    self._loading_overlay.raise_()
                    try:
                        self._loading_overlay.setFocus(Qt.ActiveWindowFocusReason)
                    except Exception:
                        pass
                except Exception:
                    pass
                try:
                    QApplication.processEvents()
                except Exception:
                    pass
            # Watchdog: nếu overlay tồn tại quá lâu mà không được ẩn, tự gỡ để tránh kẹt UI
            try:
                from PySide6.QtCore import QTimer
                if getattr(self, "_loading_watchdog", None):
                    try:
                        self._loading_watchdog.stop()
                    except Exception:
                        pass
                self._loading_watchdog = QTimer(self)
                self._loading_watchdog.setSingleShot(True)
                def _wd():
                    try:
                        if getattr(self, "_loading_overlay", None) and self._loading_overlay.isVisible():
                            self._hide_loading()
                            self.setEnabled(True)
                    except Exception:
                        pass
                self._loading_watchdog.timeout.connect(_wd)
                watchdog_ms = 60000 if icon_variant == "web" else 12000
                self._loading_watchdog.start(watchdog_ms)
            except Exception:
                pass
        except Exception:
            pass

    def _hide_loading(self):
        try:
            if getattr(self, "_loading_overlay", None):
                try:
                    self._loading_overlay.hide()
                except Exception:
                    pass
            if getattr(self, "_loading_blur", None):
                try:
                    self.central_widget.setGraphicsEffect(None)
                except Exception:
                    pass
                self._loading_blur = None
            try:
                if getattr(self, "central_widget", None):
                    self.central_widget.setEnabled(True)
            except Exception:
                pass
            try:
                if getattr(self, "_chat_toggle_btn", None):
                    self._chat_toggle_btn.setEnabled(True)
            except Exception:
                pass
            try:
                if getattr(self, "chat_widget", None):
                    self.chat_widget.setEnabled(True)
            except Exception:
                pass
            try:
                if getattr(self, "_loading_movie", None):
                    self._loading_movie.stop()
                    self._loading_movie = None
            except Exception:
                pass
            try:
                self.setEnabled(True)
            except Exception:
                pass
        except Exception:
            pass
    
    def show_system_notification(self, title: str, message: str, level: str = "info"):
        try:
            from eduplay.core.settings_manager import SettingsManager
            sm = self.settings_manager or SettingsManager()
            enabled = bool(sm.get("notifications.system_enabled", True))
            only_bg = bool(sm.get("notifications.only_when_background", True))
        except Exception:
            enabled = True
            only_bg = True
        if not enabled:
            return
        try:
            if only_bg:
                try:
                    if self.isActiveWindow():
                        return
                except Exception:
                    pass
            if not getattr(self, "_tray_icon", None):
                if QSystemTrayIcon.isSystemTrayAvailable():
                    self._tray_icon = QSystemTrayIcon(self)
                    self._tray_icon.setIcon(self.windowIcon())
                    self._tray_icon.setToolTip("EduPlay Studio")
                    self._tray_icon.show()
            tray = getattr(self, "_tray_icon", None)
            if not tray:
                return
            level_l = (level or "").lower()
            if level_l == "warning":
                icon = QSystemTrayIcon.Warning
            elif level_l == "error":
                icon = QSystemTrayIcon.Critical
            else:
                icon = QSystemTrayIcon.Information
            tray.showMessage(title or "EduPlay Studio", message or "", icon, 6000)
        except Exception:
            pass
    
    def _show_message(self, title: str, message: str, level: str = "info"):
        from PySide6.QtCore import QTimer
        box = QMessageBox(self)
        try:
            theme = (self.settings_manager.get_theme() if self.settings_manager else "dark") or "dark"
        except Exception:
            theme = "dark"
        if str(theme).lower() == "dark":
            box.setStyleSheet(
                """
                QMessageBox { background-color: #0F172A; color: #E5E7EB; }
                QMessageBox QLabel { color: #E5E7EB; background: transparent; }
                QMessageBox QPushButton {
                    background-color: #111827;
                    color: #E5E7EB;
                    border: 1px solid #374151;
                    border-radius: 10px;
                    padding: 8px 14px;
                    min-width: 96px;
                }
                QMessageBox QPushButton:hover { background-color: #1F2937; color: #FFFFFF; border-color: #4B5563; }
                QMessageBox QPushButton:default { background-color: #7F56D9; border-color: #7F56D9; color: #FFFFFF; }
                """
            )
        else:
            box.setStyleSheet(
                """
                QMessageBox { background-color: #FFFFFF; color: #111827; }
                QMessageBox QLabel { color: #1F2937; background: transparent; }
                QMessageBox QPushButton {
                    background-color: #F3F4F6;
                    color: #111827;
                    border: 1px solid #D0D5DD;
                    border-radius: 10px;
                    padding: 8px 14px;
                    min-width: 96px;
                }
                QMessageBox QPushButton:hover { background-color: #E5E7EB; color: #111827; border-color: #9CA3AF; }
                QMessageBox QPushButton:default { background-color: #7F56D9; border-color: #7F56D9; color: #FFFFFF; }
                """
            )
        box.setWindowTitle(title)
        box.setText(message)
        severity = (level or "").lower()
        if severity == "error":
            box.setIcon(QMessageBox.Critical)
        elif severity == "warning":
            box.setIcon(QMessageBox.Warning)
        elif severity == "success":
            box.setIcon(QMessageBox.Information)
        else:
            box.setIcon(QMessageBox.Information)
        try:
            box.setWindowModality(Qt.NonModal)
            box.show()
            # Auto-close success/info after 2.5s to avoid user having to click
            if severity in ("success", "info", ""):
                QTimer.singleShot(2500, box.close)
        except Exception:
            try:
                box.exec()
            except Exception:
                pass
    
    def handle_export(self, export_type: str):
        """Handle export requests from editor screen"""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from eduplay.core.export_service import ExportService

        try:
            dirty = hasattr(self.editor_screen, "_unsaved_changes") and bool(getattr(self.editor_screen, "_unsaved_changes"))
        except Exception:
            dirty = False
        if dirty:
            try:
                from eduplay.core.i18n import I18n
                lang = self.settings_manager.get_language() if self.settings_manager else 'en'
                title = I18n.t('editor.unsaved_title', lang) if hasattr(I18n, 't') else "Thay đổi chưa lưu"
                msg = I18n.t('editor.unsaved_msg', lang) if hasattr(I18n, 't') else "Bạn có thay đổi chưa lưu. Lưu trước khi xuất?"
            except Exception:
                title = "Thay đổi chưa lưu"
                msg = "Bạn có thay đổi chưa lưu. Lưu trước khi xuất?"
            box = QMessageBox(self)
            box.setIcon(QMessageBox.Warning)
            box.setWindowTitle(title)
            box.setText(msg)
            save_btn = box.addButton("Lưu", QMessageBox.AcceptRole)
            cont_btn = box.addButton("Tiếp tục", QMessageBox.DestructiveRole)
            cancel_btn = box.addButton("Hủy", QMessageBox.RejectRole)
            box.exec()
            if box.clickedButton() is cancel_btn:
                return
            if box.clickedButton() is save_btn:
                try:
                    self.editor_screen.save_project()
                except Exception:
                    pass
        
        # Get current project from the active editor screen
        try:
            current_widget = self.stacked_widget.currentWidget()
        except Exception:
            current_widget = None
        try:
            if current_widget and hasattr(current_widget, 'get_current_project'):
                current_project = current_widget.get_current_project()
            else:
                current_project = self.editor_screen.get_current_project()
        except Exception:
            current_project = self.editor_screen.get_current_project()
        try:
            from eduplay.core.i18n import I18n
            lang = self.settings_manager.get_language() if self.settings_manager else 'en'
        except Exception:
            I18n = None  # type: ignore
            lang = 'en'
        if not current_project:
            if I18n:
                title = I18n.t('export.error_title', lang)
                msg = I18n.t('export.error_no_project', lang)
            else:
                title = "Export Error"
                msg = "No project loaded for export."
            self._show_message(title, msg, "warning")
            return
        
        if export_type == 'html':
            suggested = (self.editor_screen.get_current_project() or {}).get('name', 'EduPlay_Game')
            if I18n:
                title = I18n.t('export.save_html_dialog_title', lang)
            else:
                title = "Export HTML - Choose file name"
            from PySide6.QtWidgets import QFileDialog
            from PySide6.QtCore import Qt, QUrl
            try:
                from PySide6.QtWidgets import QApplication, QStyleFactory
                app = QApplication.instance()
                _old_ss = app.styleSheet() if app else ""
                _old_style = (app.style().objectName() if app else "")
                if app:
                    try:
                        app.setStyleSheet("")
                        if QStyleFactory.keys():
                            app.setStyle(QStyleFactory.create("WindowsVista") or QStyleFactory.create("Windows"))
                    except Exception:
                        pass
            except Exception:
                app = None
                _old_ss = ""
                _old_style = ""
            dlg = QFileDialog(self, title, suggested + ".html", "HTML Files (*.html)")
            dlg.setAcceptMode(QFileDialog.AcceptSave)
            dlg.setFileMode(QFileDialog.AnyFile)
            try:
                dlg.setOption(QFileDialog.DontUseNativeDialog, True)
            except Exception:
                pass
            try:
                dlg.setWindowModality(Qt.ApplicationModal)
            except Exception:
                pass
            try:
                dlg.setStyleSheet("""
                    QFileDialog, QDialog, QWidget {
                        background-color: #FFFFFF;
                        color: #101828;
                    }
                    QLabel { color: #101828; }
                    QLineEdit, QListView, QTreeView {
                        background-color: #FFFFFF;
                        color: #101828;
                        selection-background-color: #7F56D9;
                        selection-color: #FFFFFF;
                    }
                    QPushButton {
                        background-color: #F2F4F7;
                        color: #101828;
                        border: 1px solid #D0D5DD;
                        border-radius: 6px;
                        padding: 6px 10px;
                    }
                    QPushButton:hover { background-color: #EAECF0; }
                """)
            except Exception:
                pass
            # Add common sidebar locations similar to Windows Quick Access
            try:
                import os
                sidebar = []
                home = os.path.expanduser("~")
                def add(path):
                    try:
                        if path and os.path.exists(path):
                            sidebar.append(QUrl.fromLocalFile(path))
                    except Exception:
                        pass
                add(home)
                add(os.path.join(home, "Desktop"))
                add(os.path.join(home, "Documents"))
                add(os.path.join(home, "Downloads"))
                add(os.path.join(home, "Pictures"))
                add(os.path.join(home, "Music"))
                add(os.path.join(home, "Videos"))
                # common drives
                for drv in ["C:\\", "D:\\", "E:\\"]:
                    add(drv)
                if sidebar:
                    dlg.setSidebarUrls(sidebar)
            except Exception:
                pass
            if dlg.exec():
                sel = dlg.selectedFiles()
                file_path = sel[0] if sel else ""
            else:
                file_path = ""
            try:
                if app:
                    if _old_style:
                        from PySide6.QtWidgets import QStyleFactory as _QSF
                        try:
                            app.setStyle(_QSF.create(_old_style) or app.style())
                        except Exception:
                            pass
                    app.setStyleSheet(_old_ss or "")
            except Exception:
                pass
            if not file_path:
                return
            from pathlib import Path
            p = Path(file_path)
            if p.suffix.lower() != ".html":
                p = p.with_suffix(".html")
                file_path = str(p)
            export_dir = str(p.parent)
            chosen_name = p.stem
            selected_html_path = file_path
        else:
            if I18n:
                dialog_title = I18n.t('export.select_output_dir', lang)
            else:
                dialog_title = f"Export {export_type.upper()} - Select Output Directory"
            from PySide6.QtWidgets import QFileDialog
            from PySide6.QtCore import Qt
            dlg = QFileDialog(self, dialog_title)
            dlg.setFileMode(QFileDialog.Directory)
            try:
                dlg.setOption(QFileDialog.ShowDirsOnly, True)
            except Exception:
                pass
            try:
                dlg.setOption(QFileDialog.DontUseNativeDialog, True)
            except Exception:
                pass
            try:
                dlg.setWindowModality(Qt.ApplicationModal)
            except Exception:
                pass
            try:
                dlg.setStyleSheet("""
                    QFileDialog, QDialog, QWidget {
                        background-color: #FFFFFF;
                        color: #101828;
                    }
                    QLabel { color: #101828; }
                    QLineEdit, QListView, QTreeView {
                        background-color: #FFFFFF;
                        color: #101828;
                        selection-background-color: #7F56D9;
                        selection-color: #FFFFFF;
                    }
                    QPushButton {
                        background-color: #F2F4F7;
                        color: #101828;
                        border: 1px solid #D0D5DD;
                        border-radius: 6px;
                        padding: 6px 10px;
                    }
                    QPushButton:hover { background-color: #EAECF0; }
                """)
            except Exception:
                pass
            if dlg.exec():
                sel = dlg.selectedFiles()
                export_dir = sel[0] if sel else ""
            else:
                export_dir = ""
            if not export_dir:
                return
            chosen_name = None
            selected_html_path = ""
        
        export_service = ExportService()
        try:
            try:
                if I18n:
                    loading_msg = I18n.t('export.loading', lang)
                else:
                    loading_msg = "Exporting..."
            except Exception:
                loading_msg = "Exporting..."

            def _task():
                result = {"export_type": export_type, "success": False, "extra": {}}
                try:
                    try:
                        print(f"[Export] Begin export: type={export_type}")
                    except Exception:
                        pass
                    if export_type == 'html':
                        name = chosen_name or current_project.get('name', 'EduPlay_Game')
                        try:
                            print(f"[Export] Output dir: {export_dir}, name: {name}")
                        except Exception:
                            pass
                        success = export_service.export_to_html(current_project, export_dir, bundle_resources=True, single_file=True, output_filename=name)
                        result["success"] = bool(success)
                        result["extra"]["name"] = name
                        result["extra"]["selected_path"] = selected_html_path
                    elif export_type == 'native':
                        success = export_service.export_to_native(current_project, export_dir)
                        result["success"] = bool(success)
                    elif export_type == 'exe':
                        success = export_service.export_to_exe(current_project, export_dir)
                        result["success"] = bool(success)
                    try:
                        print(f"[Export] Result success={result['success']}")
                    except Exception:
                        pass
                    return result
                except Exception as e:
                    result["error"] = str(e)
                    return result

            def _done(result, error):
                try:
                    try:
                        print("[Export] Done callback invoked. Error:", error, "Result:", result)
                    except Exception:
                        pass
                    if error:
                        raise error
                    if not isinstance(result, dict):
                        raise Exception("Invalid export result")
                    ok = bool(result.get("success"))
                    kind = result.get("export_type")
                    extra = result.get("extra") or {}
                    if kind == 'html':
                        name = extra.get("name") or (chosen_name or current_project.get('name', 'EduPlay_Game'))
                        if ok:
                            from pathlib import Path
                            safe = ''.join([c for c in name if c.isalnum() or c in (' ', '-', '_')]).strip().replace(' ', '_') or 'EduPlay_Game'
                            out = Path(export_dir) / (safe + '.html')
                            try:
                                sel = extra.get("selected_path") or ""
                                if sel:
                                    sp = Path(sel)
                                    if sp.parent.exists():
                                        if out.exists() and out.resolve() != sp.resolve():
                                            import shutil
                                            shutil.copyfile(str(out), str(sp))
                                            out = sp
                                        else:
                                            out = sp if sp.exists() else out
                            except Exception:
                                pass
                            try:
                                if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                                    if I18n:
                                        msg = I18n.t('export.success_html_single', lang, path=str(out))
                                    else:
                                        msg = f"Xuất thành công:\n\n{out}\n\nMở file này để chơi."
                                    self.editor_screen.show_toast(msg, kind="success", duration_ms=7000)
                                    try:
                                        title_n = I18n.t('export.success_title', lang) if I18n else "Export Success"
                                        self.show_system_notification(title_n, str(out), "success")
                                    except Exception:
                                        pass
                                else:
                                    # Fallback small info
                                    self._show_message("Export Success", f"Xuất thành công: {out}", "success")
                            except Exception:
                                pass
                        else:
                            try:
                                if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                                    msg = I18n.t('export.error_html_failed', lang) if I18n else "Xuất HTML thất bại"
                                    self.editor_screen.show_toast(msg, kind="error", duration_ms=7000)
                                    try:
                                        ttl = I18n.t('export.failed_title', lang) if I18n else "Export Failed"
                                        self.show_system_notification(ttl, msg, "error")
                                    except Exception:
                                        pass
                                else:
                                    self._show_message("Export Failed", "HTML export failed. Please check the console for details.", "error")
                            except Exception:
                                pass
                    elif kind == 'native':
                        if ok:
                            try:
                                if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                                    txt = I18n.t('export.success_native', lang, dir=str(export_dir)) if I18n else f"Xuất Native thành công:\n\n{export_dir}\n\nChạy run.bat để chơi."
                                    self.editor_screen.show_toast(txt, kind="success", duration_ms=7000)
                                    try:
                                        ttl = I18n.t('export.success_title', lang) if I18n else "Export Success"
                                        self.show_system_notification(ttl, txt, "success")
                                    except Exception:
                                        pass
                                else:
                                    self._show_message("Export Success", f"Native export completed.\nFolder: {export_dir}", "success")
                            except Exception:
                                pass
                        else:
                            try:
                                if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                                    txt = I18n.t('export.error_native_failed', lang) if I18n else "Xuất Native thất bại."
                                    self.editor_screen.show_toast(txt, kind="error", duration_ms=7000)
                                    try:
                                        ttl = I18n.t('export.failed_title', lang) if I18n else "Export Failed"
                                        self.show_system_notification(ttl, txt, "error")
                                    except Exception:
                                        pass
                                else:
                                    self._show_message("Export Failed", "Native export failed.", "error")
                            except Exception:
                                pass
                    elif kind == 'exe':
                        from pathlib import Path
                        exe_path = Path(export_dir) / 'dist' / 'EduPlayGame.exe'
                        if ok:
                            try:
                                if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                                    txt = I18n.t('export.success_exe', lang, dir=str(Path(export_dir)), exe=str(exe_path)) if I18n else f"Đã build EXE:\n\n{exe_path}"
                                    self.editor_screen.show_toast(txt, kind="success", duration_ms=8000)
                                    try:
                                        ttl = I18n.t('export.success_title', lang) if I18n else "Export Success"
                                        self.show_system_notification(ttl, str(exe_path), "success")
                                    except Exception:
                                        pass
                                else:
                                    self._show_message("Export Success", f"Built EXE: {exe_path}", "success")
                            except Exception:
                                pass
                        else:
                            try:
                                if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                                    txt = I18n.t('export.error_exe_failed', lang) if I18n else "Build EXE thất bại."
                                    self.editor_screen.show_toast(txt, kind="error", duration_ms=7000)
                                    try:
                                        ttl = I18n.t('export.failed_title', lang) if I18n else "Export Failed"
                                        self.show_system_notification(ttl, txt, "error")
                                    except Exception:
                                        pass
                                else:
                                    self._show_message("Export Failed", "Executable export failed.", "error")
                            except Exception:
                                pass
                except Exception as e:
                    try:
                        if hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                            txt = (I18n.t('export.error_generic', lang, error=str(e)) if I18n else f"Export failed: {str(e)}")
                            self.editor_screen.show_toast(txt, kind="error", duration_ms=8000)
                        try:
                            ttl = I18n.t('export.error_title', lang) if I18n else "Export Error"
                            self.show_system_notification(ttl, str(e), "error")
                        except Exception:
                            pass
                        else:
                            self._show_message("Export Error", f"Export failed: {str(e)}", "error")
                    except Exception:
                        pass
                    try:
                        print(f"Export error: {e}")
                    except Exception:
                        pass

            self._run_background_task(_task, _done, loading_msg, "default")
        except Exception as e:
            if I18n:
                title = I18n.t('export.error_title', lang)
                msg = I18n.t('export.error_generic', lang)
                try:
                    msg = msg.format(error=str(e))
                except Exception:
                    pass
            else:
                title = "Export Error"
                msg = f"Export failed with error:\n\n{str(e)}"
            self._show_message(title, msg, "error")

    def closeEvent(self, event):
        def _stop_bg_threads():
            try:
                threads = list(getattr(self, "_background_threads", []) or [])
            except Exception:
                threads = []
            for t in threads:
                try:
                    if t and t.isRunning():
                        t.quit()
                except Exception:
                    pass
            for t in threads:
                try:
                    if t:
                        t.wait(1500)
                except Exception:
                    pass

        try:
            dirty = hasattr(self, 'editor_screen') and hasattr(self.editor_screen, "_unsaved_changes") and bool(getattr(self.editor_screen, "_unsaved_changes"))
        except Exception:
            dirty = False
        if dirty:
            from PySide6.QtWidgets import QMessageBox
            try:
                from eduplay.core.i18n import I18n
                lang = self.settings_manager.get_language() if self.settings_manager else 'en'
                title = I18n.t('editor.unsaved_title', lang) if hasattr(I18n, 't') else "Thay đổi chưa lưu"
                msg = I18n.t('editor.unsaved_exit_msg', lang) if hasattr(I18n, 't') else "Bạn có thay đổi chưa lưu. Lưu trước khi thoát?"
            except Exception:
                title = "Thay đổi chưa lưu"
                msg = "Bạn có thay đổi chưa lưu. Lưu trước khi thoát?"
            box = QMessageBox(self)
            box.setIcon(QMessageBox.Warning)
            box.setWindowTitle(title)
            box.setText(msg)
            save_btn = box.addButton("Lưu", QMessageBox.AcceptRole)
            quit_btn = box.addButton("Thoát", QMessageBox.DestructiveRole)
            cancel_btn = box.addButton("Hủy", QMessageBox.RejectRole)
            box.exec()
            if box.clickedButton() is save_btn:
                try:
                    self.editor_screen.save_project()
                    _stop_bg_threads()
                    event.accept()
                    return
                except Exception:
                    event.ignore()
                    return
            if box.clickedButton() is cancel_btn:
                event.ignore()
                return
            _stop_bg_threads()
            event.accept()
        else:
            _stop_bg_threads()
            event.accept()

    def handle_publish_web(self):
        """Xuất HTML 1-file vào vùng tạm và đăng lên Firebase + trả link hosting"""
        from PySide6.QtWidgets import QMessageBox, QDialog, QVBoxLayout, QLabel, QLineEdit, QPushButton
        from pathlib import Path
        from eduplay.core.export_service import ExportService
        try:
            from eduplay.core.i18n import I18n
            lang = self.settings_manager.get_language() if self.settings_manager else 'en'
        except Exception:
            I18n = None  # type: ignore
            lang = 'en'
        try:
            current_widget = self.stacked_widget.currentWidget()
            if current_widget and hasattr(current_widget, 'get_current_project'):
                current_project = current_widget.get_current_project()
            else:
                current_project = self.editor_screen.get_current_project()
        except Exception:
            current_project = self.editor_screen.get_current_project()
        if not current_project:
            if I18n:
                title = I18n.t('publish.error_title', lang)
                msg = I18n.t('publish.error_no_project', lang)
            else:
                title = "Publish Error"
                msg = "No project available to publish."
            self._show_message(title, msg, "warning")
            return
        try:
            try:
                if I18n:
                    loading_msg = I18n.t('publish.loading', lang)
                else:
                    loading_msg = "Publishing..."
            except Exception:
                loading_msg = "Publishing..."

            name = (current_project or {}).get('name', 'EduPlay_Game')

            def _task():
                tmp_dir = PathResolver.resolve_publish_cache_dir()
                tmp_dir.mkdir(parents=True, exist_ok=True)
                export_service = ExportService()
                ok = export_service.export_to_html(current_project, str(tmp_dir), bundle_resources=True, single_file=True, output_filename=name)
                if not ok:
                    err_text = ""
                    try:
                        err_text = getattr(export_service, "_last_export_error", "") or ""
                    except Exception:
                        err_text = ""
                    try:
                        bundled = export_service._bundle_media_files(current_project or {})
                    except Exception:
                        bundled = current_project or {}
                    try:
                        html = export_service._generate_html_content(bundled)
                        html = export_service._inject_favicon(html)
                    except Exception:
                        return {
                            "success": False,
                            "tmp_dir": str(tmp_dir),
                            "name": name,
                            "links": {},
                            "stage": "export",
                            "error": err_text or "Fallback HTML generation failed",
                        }
                    safe_name = ''.join([c for c in name if c.isalnum() or c in (' ', '-', '_')]).strip().replace(' ', '_') or 'EduPlay_Game'
                    out_html = tmp_dir / (safe_name + '.html')
                    try:
                        out_html.write_text(html, encoding='utf-8')
                    except Exception:
                        return {
                            "success": False,
                            "tmp_dir": str(tmp_dir),
                            "name": name,
                            "links": {},
                            "stage": "export",
                            "error": "Could not write temporary HTML file",
                        }
                safe = ''.join([c for c in name if c.isalnum() or c in (' ', '-', '_')]).strip().replace(' ', '_') or 'EduPlay_Game'
                out = tmp_dir / (safe + '.html')
                proj_id = (current_project or {}).get('id', '')
                
                # Progress callback to update loading message
                def _progress(payload, total=None):
                    try:
                        if isinstance(payload, dict):
                            stage = str(payload.get("stage") or "").strip().lower()
                            current = int(payload.get("current") or 0)
                            total_chunks = int(payload.get("total") or 0)
                            percent = int(payload.get("percent") or 0)
                            if stage == "compressing":
                                msg = "Đang nén toàn bộ dữ liệu trước khi upload..."
                            elif stage == "uploading":
                                msg = f"Đang tải dữ liệu lên server: {current}/{total_chunks} phần ({percent}%)"
                            elif stage == "finalizing":
                                msg = "Đang tạo link chia sẻ..."
                            elif stage == "completed":
                                msg = "Đã tạo xong link, đang hoàn tất hiển thị..."
                            else:
                                msg = str(payload.get("message") or "").strip() or "Đang xuất lên web..."
                        else:
                            current = int(payload or 0)
                            total_chunks = int(total or 0)
                            percent = int((current / total_chunks) * 100) if total_chunks else 0
                            msg = f"Đang tải dữ liệu lên server: {current}/{total_chunks} phần ({percent}%)"
                        print(f"[PUBLISH] {msg}")
                        from PySide6.QtCore import QTimer

                        def _apply():
                            try:
                                self._show_loading(msg, "web")
                            except Exception:
                                pass

                        try:
                            QTimer.singleShot(0, self, _apply)
                        except Exception:
                            QTimer.singleShot(0, _apply)
                    except Exception:
                        pass
                
                pub = export_service.publish_to_firebase(
                    str(out),
                    name,
                    ""  # REDACTED,
                    project_id=proj_id,
                    progress_callback=_progress,
                    project_data=current_project,
                )
                export_service.cleanup_firebase_old(""  # REDACTED, days=15, max_items=300)
                if not isinstance(pub, dict) or not pub.get("ok") or not pub.get("play_link"):
                    err_text = ""
                    try:
                        err_text = str(pub.get("error") or "")
                    except Exception:
                        err_text = ""
                    return {
                        "success": False,
                        "tmp_dir": str(tmp_dir),
                        "name": name,
                        "links": {},
                        "stage": "publish",
                        "error": err_text or "Firebase publish failed",
                    }
                link_line = pub.get("play_link") or ""
                db_link = pub.get("db_link") or ""
                key_val = pub.get("key") or ""
                try:
                    if out.exists():
                        out.unlink()
                except Exception:
                    pass
                return {
                    "success": True,
                    "tmp_dir": str(tmp_dir),
                    "name": name,
                    "links": {"play": link_line, "db": db_link, "key": key_val},
                }

            def _done(result, error):
                if error:
                    if I18n:
                        title = I18n.t('publish.error_title', lang)
                        msg = I18n.t('publish.error_html_failed', lang)
                    else:
                        title = "Publish Error"
                        msg = "HTML export failed. Cannot publish to the web."
                    try:
                        extra = str(error)
                        if extra:
                            msg = f"{msg}\n\nDetails: {extra}"
                    except Exception:
                        pass
                    self._show_message(title, msg, "error")
                    try:
                        self.show_system_notification(title, msg, "error")
                    except Exception:
                        pass
                    return
                if not isinstance(result, dict) or not result.get("success"):
                    stage = ""
                    err_text = ""
                    try:
                        stage = str(result.get("stage") or "")
                        err_text = str(result.get("error") or "")
                    except Exception:
                        stage = ""
                        err_text = ""
                    try:
                        if stage or err_text:
                            print(f"Publish web failed - stage={stage}, details={err_text}")
                    except Exception:
                        pass
                    if I18n:
                        title = I18n.t('publish.error_title', lang)
                        msg = I18n.t('publish.error_html_failed', lang)
                    else:
                        title = "Publish Error"
                        msg = "HTML export failed. Cannot publish to the web."
                    self._show_message(title, msg, "error")
                    try:
                        self.show_system_notification(title, msg, "error")
                    except Exception:
                        pass
                    return
                links = result.get("links") or {}
                link_line = links.get("play") or ""
                db_line = links.get("db") or ""
                key_val = links.get("key") or ""
                if not link_line:
                    if I18n:
                        title = I18n.t('publish.error_title', lang)
                        msg = I18n.t('publish.error_html_failed', lang)
                    else:
                        title = "Publish Error"
                        msg = "Publishing failed. No link was generated."
                    self._show_message(title, msg, "error")
                    try:
                        self.show_system_notification(title, msg, "error")
                    except Exception:
                        pass
                    return
                dlg = QDialog(self)
                if I18n:
                    dlg.setWindowTitle(I18n.t('publish.dialog_title', lang))
                else:
                    dlg.setWindowTitle("Game Links")
                lay = QVBoxLayout(dlg)
                try:
                    from eduplay.core.settings_manager import SettingsManager
                    sm = self.settings_manager or SettingsManager()
                    theme = sm.get_theme() or "light"
                    brand = sm.get("brand_color", "#10B981")
                except Exception:
                    theme = "light"
                    brand = "#10B981"
                def _darker(hex_color: str, delta: int = 18) -> str:
                    try:
                        c = (hex_color or "").lstrip("#")
                        if len(c) != 6:
                            return hex_color
                        r = max(0, int(c[0:2], 16) - delta)
                        g = max(0, int(c[2:4], 16) - delta)
                        b = max(0, int(c[4:6], 16) - delta)
                        return f"#{r:02X}{g:02X}{b:02X}"
                    except Exception:
                        return hex_color
                brand_hover = _darker(brand)
                if theme == "dark":
                    dlg_bg = "#020617"
                    text_color = "#E5E7EB"
                    input_bg = "#020617"
                    input_border = "#1E293B"
                    input_text = "#E5E7EB"
                    badge_bg = brand
                    badge_text = "#0B1120"
                    footer_bg = "#020617"
                    primary_text = "#000000"
                    secondary_hover_bg = "rgba(148,163,184,0.14)"
                else:
                    dlg_bg = "#FFFFFF"
                    text_color = "#0F1728"
                    input_bg = "#F9FAFF"
                    input_border = brand
                    input_text = "#0F1728"
                    badge_bg = brand
                    badge_text = "#FFFFFF"
                    footer_bg = "#F1F5F9"
                    primary_text = "#000000"
                    secondary_hover_bg = "rgba(148,163,184,0.25)"
                dlg.setStyleSheet(
                    f"""
                    QDialog {{
                        background-color: {dlg_bg};
                        color: {text_color};
                        border-radius: 12px;
                    }}
                    QLabel {{
                        font-size: 13px;
                    }}
                    QLineEdit {{
                        background-color: {input_bg};
                        border: 1px solid {input_border};
                        border-radius: 8px;
                        padding: 6px 8px;
                        font-size: 13px;
                        color: {input_text};
                        selection-background-color: {brand};
                    }}
                    QPushButton {{
                        border-radius: 8px;
                        padding: 6px 14px;
                        font-weight: 600;
                    }}
                    QPushButton#primary-btn {{
                        background-color: {brand};
                        color: {primary_text};
                        border: 1px solid {brand};
                    }}
                    QPushButton#primary-btn:hover {{
                        background-color: {brand_hover};
                        border-color: {brand_hover};
                    }}
                    QPushButton#secondary-btn {{
                        background-color: transparent;
                        color: {text_color};
                        border: 1px solid {input_border};
                    }}
                    QPushButton#secondary-btn:hover {{
                        background-color: {secondary_hover_bg};
                    }}
                    """
                )
                if I18n:
                    lbl1_text = I18n.t('publish.play_link_label', lang)
                    btn_text = I18n.t('publish.copy_button', lang)
                    try:
                        open_text = I18n.t('publish.open_browser', lang)
                    except Exception:
                        open_text = "Open in browser"
                    try:
                        close_text = I18n.t('publish.close_button', lang)
                    except Exception:
                        close_text = "Close"
                else:
                    lbl1_text = "Play link:"
                    btn_text = "Copy"
                    open_text = "Open in browser"
                    close_text = "Close"
                lbl1 = QLabel(lbl1_text); txt1 = QLineEdit(link_line); btn1 = QPushButton(btn_text)
                if key_val:
                    txt1.setToolTip(key_val)
                txt1.setReadOnly(True)
                def _copy1():
                    from PySide6.QtWidgets import QApplication
                    QApplication.clipboard().setText(txt1.text())
                    try:
                        try:
                            if I18n:
                                msg = I18n.t('toast.copy_success', lang)
                            else:
                                msg = "Copied to clipboard"
                        except Exception:
                            msg = "Copied to clipboard"
                        target = None
                        try:
                            target = self.stacked_widget.currentWidget()
                        except Exception:
                            target = None
                        if target and hasattr(target, 'show_toast'):
                            target.show_toast(msg, kind="success")
                        elif hasattr(self, 'editor_screen') and hasattr(self.editor_screen, 'show_toast'):
                            self.editor_screen.show_toast(msg, kind="success")
                    except Exception:
                        pass
                btn1.clicked.connect(_copy1)
                lay.addWidget(lbl1); lay.addWidget(txt1); lay.addWidget(btn1)
                try:
                    if I18n:
                        hint = I18n.t('publish.link_hint', lang)
                    else:
                        hint = "Link is valid for about 15 days. Republishing the same project reuses this link if content does not change."
                    if hint:
                        hint_lbl = QLabel(hint)
                        lay.addWidget(hint_lbl)
                except Exception:
                    pass
                footer = QHBoxLayout()
                footer_widget = QWidget()
                footer_widget.setLayout(footer)
                footer_widget.setStyleSheet(f"background-color: {footer_bg}; border-top: 1px solid {input_border}; margin-top: 8px;")
                footer.addStretch()
                open_btn = QPushButton(open_text)
                open_btn.setObjectName("primary-btn")
                close_btn = QPushButton(close_text)
                close_btn.setObjectName("secondary-btn")
                def _open():
                    try:
                        from PySide6.QtGui import QDesktopServices
                        from PySide6.QtCore import QUrl
                        QDesktopServices.openUrl(QUrl(link_line))
                    except Exception:
                        pass
                open_btn.clicked.connect(_open)
                close_btn.clicked.connect(dlg.accept)
                footer.addWidget(open_btn)
                footer.addWidget(close_btn)
                lay.addWidget(footer_widget)
                dlg.exec()
                try:
                    ttl = I18n.t('publish.dialog_title', lang) if I18n else "Game Links"
                    self.show_system_notification(ttl, link_line, "success")
                except Exception:
                    pass

            self._run_background_task(_task, _done, loading_msg, "web")
        except Exception as e:
            if I18n:
                title = I18n.t('publish.error_title', lang)
                msg = I18n.t('publish.error_html_failed', lang)
            else:
                title = "Publish Error"
                msg = str(e)
            self._show_message(title, msg, "error")
            try:
                self.show_system_notification(title, msg, "error")
            except Exception:
                pass
            
"""
Nguyen-Thanh-Tan ¬_¬
"""
